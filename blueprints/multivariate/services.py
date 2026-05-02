from __future__ import annotations

import os
import math
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from sklearn.decomposition import PCA
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import StandardScaler
from sklearn.cluster import KMeans

import statsmodels.api as sm
from statsmodels.formula.api import ols

def build_multivariate_profile(df):
    import pandas as pd
    import numpy as np

    profile = {}

    n_rows, n_cols = df.shape
    profile["n_rows"] = int(n_rows)
    profile["n_cols"] = int(n_cols)

    num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    cat_cols = [c for c in df.columns if c not in num_cols]

    profile["num_cols"] = num_cols
    profile["cat_cols"] = cat_cols
    profile["n_num"] = len(num_cols)
    profile["n_cat"] = len(cat_cols)

    total_cells = n_rows * n_cols if n_rows and n_cols else 0
    missing_total = int(df.isna().sum().sum()) if total_cells else 0
    missing_pct = round((missing_total / total_cells) * 100, 2) if total_cells else 0.0
    profile["missing_pct"] = missing_pct

    # ===== numéricas utilizables =====
    usable_num = []
    zero_variance_cols = []

    for c in num_cols:
        s = df[c].dropna()
        if s.empty:
            continue
        if s.nunique() <= 1:
            zero_variance_cols.append(c)
            continue
        usable_num.append(c)

    profile["usable_num_cols"] = usable_num
    profile["zero_variance_cols"] = zero_variance_cols

    # ===== categóricas utilizables =====
    usable_cat = []
    for c in cat_cols:
        try:
            if df[c].dropna().nunique() >= 2:
                usable_cat.append(c)
        except Exception:
            pass

    profile["usable_cat_cols"] = usable_cat

    # ===== correlaciones altas =====
    high_corr_pairs = []
    if len(usable_num) >= 2:
        try:
            corr = df[usable_num].corr(numeric_only=True)
            for i in range(len(corr.columns)):
                for j in range(i + 1, len(corr.columns)):
                    val = corr.iloc[i, j]
                    if pd.notna(val) and abs(val) >= 0.70:
                        high_corr_pairs.append({
                            "col1": corr.columns[i],
                            "col2": corr.columns[j],
                            "corr": round(float(val), 4),
                        })
        except Exception:
            pass

    profile["high_corr_pairs"] = high_corr_pairs

    # ===== tipo de perfil =====
    if len(usable_num) >= 5 and len(usable_cat) >= 1:
        profile_type = "Mixto avanzado"
    elif len(usable_num) >= 5:
        profile_type = "Numérico multivariado"
    elif len(usable_num) >= 2 and len(usable_cat) >= 1:
        profile_type = "Mixto"
    elif len(usable_num) >= 2:
        profile_type = "Numérico básico"
    elif len(usable_cat) >= 2:
        profile_type = "Categórico"
    else:
        profile_type = "Simple"

    profile["profile_type"] = profile_type

    # ===== recomendaciones =====
    recommendations = []

    if len(usable_num) >= 3:
        recommendations.append({
            "method": "PCA",
            "reason": f"Se detectaron {len(usable_num)} variables numéricas utilizables; conviene reducir dimensionalidad y explorar estructura.",
            "priority": 1,
        })

    if len(usable_num) >= 2:
        recommendations.append({
            "method": "Heatmap de correlación",
            "reason": f"Hay {len(usable_num)} variables numéricas válidas; se pueden explorar relaciones entre variables.",
            "priority": 2,
        })

    if len(usable_num) >= 3 and profile["n_rows"] >= 20:
        recommendations.append({
            "method": "Clustering",
            "reason": "El dataset parece apto para segmentación de observaciones mediante K-Means.",
            "priority": 3,
        })

    if len(usable_num) >= 4 and profile["n_rows"] >= 20:
        recommendations.append({
            "method": "Regresión",
            "reason": "El dataset tiene suficientes variables numéricas para modelar una variable objetivo mediante regresión lineal múltiple.",
            "priority": 4,
        })

    if len(usable_num) >= 1 and len(usable_cat) >= 1 and profile["n_rows"] >= 12:
        recommendations.append({
            "method": "ANOVA",
            "reason": "Se detectaron variables numéricas y al menos una variable categórica de grupo, por lo que se pueden comparar medias entre grupos.",
            "priority": 5,
        })

    if len(usable_num) >= 2 and len(usable_cat) >= 1 and profile["n_rows"] >= 20:
        recommendations.append({
            "method": "MANOVA",
            "reason": "Hay múltiples variables numéricas y al menos una variable categórica de grupo, lo que permite comparar perfiles multivariados entre grupos.",
            "priority": 6,
        })

    if len(usable_num) >= 2 and len(usable_cat) >= 1 and profile["n_rows"] >= 20:
        recommendations.append({
            "method": "PERMANOVA",
            "reason": "El dataset permite contrastar diferencias multivariadas entre grupos mediante permutaciones.",
            "priority": 7,
        })

    if len(high_corr_pairs) >= 2 and len(usable_num) >= 4:
        recommendations.append({
            "method": "Análisis factorial",
            "reason": "Se observan correlaciones altas entre varias variables; podría existir estructura latente.",
            "priority": 8,
        })

    if len(cat_cols) >= 2 and len(cat_cols) > len(num_cols):
        recommendations.append({
            "method": "MCA",
            "reason": "Predominan variables categóricas; el análisis de correspondencias múltiples puede ser útil.",
            "priority": 9,
        })

    if len(usable_num) >= 1:
        recommendations.append({
            "method": "Histogramas y boxplots",
            "reason": "Conviene revisar distribución, dispersión y posibles outliers.",
            "priority": 10,
        })

    recommendations = sorted(recommendations, key=lambda x: x.get("priority", 999))
    profile["recommendations"] = recommendations

    # ===== recomendaciones rápidas =====
    quick_reco = []

    if len(usable_num) >= 3:
        quick_reco.append(f"PCA ({len(usable_num)} variables numéricas detectadas)")

    if len(usable_num) >= 2:
        quick_reco.append("Heatmap de correlación")

    if len(usable_num) >= 3 and profile["n_rows"] >= 20:
        quick_reco.append("Clustering (posible segmentación)")

    if len(usable_num) >= 4 and profile["n_rows"] >= 20:
        quick_reco.append("Regresión (modelo predictivo lineal)")

    if len(usable_num) >= 1 and len(usable_cat) >= 1 and profile["n_rows"] >= 12:
        quick_reco.append("ANOVA (comparación entre grupos)")

    if len(usable_num) >= 2 and len(usable_cat) >= 1 and profile["n_rows"] >= 20:
        quick_reco.append("MANOVA (comparación multivariada entre grupos)")
        quick_reco.append("PERMANOVA (comparación por permutaciones)")

    if len(high_corr_pairs) >= 2 and len(usable_num) >= 4:
        quick_reco.append("EFA (estructura latente)")

    profile["quick_recommendations"] = quick_reco

    return profile



def get_or_build_multivariate_analysis_cache(
    *,
    db,
    ds,
    df,
    dataset_id: int,
    plots_dir: str,
    research_type: str,
    regression_target_col: str | None = None,
    rf_target_col: str | None = None,
    logistic_target_col: str | None = None,
    rf_classification_target_col: str | None = None,
    anova_target_col: str | None = None,
    anova_group_col: str | None = None,
    manova_dependent_cols: list[str] | None = None,
    manova_group_col: str | None = None,
    permanova_group_col: str | None = None,
    force_rebuild: bool = False,
):
    """
    Devuelve analysis_results desde cache si existe.
    Si no existe o force_rebuild=True, recalcula y guarda.
    """

    if not force_rebuild and getattr(ds, "analysis_cache", None):
        return ds.analysis_cache, False

    analysis_results = run_full_sadi_analysis(
        df=df,
        dataset_id=dataset_id,
        plots_dir=plots_dir,
        research_type=research_type,
        regression_target_col=regression_target_col,
        rf_target_col=rf_target_col,
        logistic_target_col=logistic_target_col,
        rf_classification_target_col=rf_classification_target_col,
        anova_target_col=anova_target_col,
        anova_group_col=anova_group_col,
        manova_dependent_cols=manova_dependent_cols,
        manova_group_col=manova_group_col,
        permanova_group_col=permanova_group_col,
    )

    try:
        clean_cache = sanitize_for_json(analysis_results)

        ds.analysis_cache = clean_cache
        db.add(ds)
        db.commit()
        db.refresh(ds)
    except Exception:
        db.rollback()
        raise

    return analysis_results, True

def json_safe_number(val, digits: int | None = None):
    import math
    import numpy as np
    import pandas as pd

    if val is None:
        return None

    try:
        if pd.isna(val):
            return None
    except Exception:
        pass

    if isinstance(val, (np.integer,)):
        val = int(val)
    elif isinstance(val, (np.floating, float)):
        val = float(val)
        if math.isnan(val) or math.isinf(val):
            return None

    if digits is not None and isinstance(val, (int, float)):
        return round(float(val), digits)

    return val

def sanitize_for_json(value):
    import math
    import numpy as np
    import pandas as pd

    # None, bool, int, str
    if value is None or isinstance(value, (bool, int, str)):
        return value

    # float
    if isinstance(value, float):
        if math.isnan(value) or math.isinf(value):
            return None
        return value

    # numpy scalars
    if isinstance(value, (np.integer,)):
        return int(value)

    if isinstance(value, (np.floating,)):
        f = float(value)
        if math.isnan(f) or math.isinf(f):
            return None
        return f

    if isinstance(value, (np.bool_,)):
        return bool(value)

    # pandas missing
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass

    # dict
    if isinstance(value, dict):
        return {str(k): sanitize_for_json(v) for k, v in value.items()}

    # list / tuple / set
    if isinstance(value, (list, tuple, set)):
        return [sanitize_for_json(v) for v in value]

    # pandas objects
    if hasattr(value, "to_dict"):
        try:
            return sanitize_for_json(value.to_dict())
        except Exception:
            pass

    if hasattr(value, "tolist"):
        try:
            return sanitize_for_json(value.tolist())
        except Exception:
            pass

    # fallback final
    return str(value)

def _get_clean_numeric_df(df: pd.DataFrame) -> pd.DataFrame:
    num_df = df.select_dtypes(include=[np.number]).copy()
    if num_df.empty:
        return num_df

    keep_cols = []
    for c in num_df.columns:
        series = num_df[c]
        if series.notna().sum() == 0:
            continue
        if series.dropna().nunique() <= 1:
            continue
        keep_cols.append(c)

    if not keep_cols:
        return pd.DataFrame()

    return num_df[keep_cols].copy()


def _prepare_scaled_numeric_matrix(df: pd.DataFrame):
    num_df = _get_clean_numeric_df(df)

    if num_df.shape[1] < 2:
        raise ValueError("Se necesitan al menos 2 variables numéricas válidas.")

    imputer = SimpleImputer(strategy="median")
    X = imputer.fit_transform(num_df)

    scaler = StandardScaler()
    Xs = scaler.fit_transform(X)

    return num_df, Xs

def interpret_stat(value: float, metric: str, research_type: str) -> str:
    """
    Genera una interpretación automática simple según el tipo de investigación.
    """

    try:
        v = float(value)
    except Exception:
        return "No se pudo generar una interpretación automática."

    av = abs(v)

    if av >= 0.7:
        level = "alta"
    elif av >= 0.4:
        level = "moderada"
    elif av >= 0.2:
        level = "baja"
    else:
        level = "muy baja"

    if research_type == "biomedical":
        return (
            f"Se observa una asociación {level} ({v:.2f}), "
            "lo que podría indicar un hallazgo clínico relevante."
        )

    if research_type == "financial":
        return (
            f"Existe una relación {level} ({v:.2f}), "
            "lo que sugiere un efecto importante en el comportamiento financiero."
        )

    if research_type == "educational":
        return (
            f"Se evidencia una relación {level} ({v:.2f}), "
            "que podría influir en el rendimiento académico."
        )

    if research_type == "experimental":
        return (
            f"Se detecta un efecto {level} ({v:.2f}), "
            "lo que podría reflejar diferencias relevantes entre grupos."
        )

    if research_type == "textual":
        return (
            f"La relación encontrada es {level} ({v:.2f}), "
            "lo que podría reflejar patrones importantes en el contenido analizado."
        )

    if research_type == "social":
        return (
            f"Se observa una relación {level} ({v:.2f}), "
            "que podría estar asociada a dinámicas sociales relevantes."
        )

    if research_type == "market_research":
        return (
            f"Existe una relación {level} ({v:.2f}), "
            "que podría influir en el comportamiento del consumidor."
        )

    return f"Se observa una relación {level} entre las variables ({v:.2f})."

def interpret_pca_result(explained_variance_ratio, research_type: str) -> str:
    try:
        if explained_variance_ratio is None or len(explained_variance_ratio) == 0:
            return "No se pudo generar una interpretación automática del PCA."

        total_2 = float(sum(explained_variance_ratio[:2])) * 100.0
    except Exception:
        return "No se pudo generar una interpretación automática del PCA."

    if total_2 >= 80:
        level = "muy alta"
    elif total_2 >= 60:
        level = "alta"
    elif total_2 >= 40:
        level = "moderada"
    else:
        level = "limitada"

    if research_type == "biomedical":
        return (
            f"Los dos primeros componentes explican una proporción {level} de la variabilidad "
            f"({total_2:.1f}%), lo que sugiere patrones clínicos bien definidos."
        )

    if research_type == "financial":
        return (
            f"Los dos primeros componentes explican una proporción {level} de la variabilidad "
            f"({total_2:.1f}%), lo que sugiere estructuras financieras relevantes en los datos."
        )

    if research_type == "educational":
        return (
            f"Los dos primeros componentes explican una proporción {level} de la variabilidad "
            f"({total_2:.1f}%), lo que podría reflejar dimensiones académicas importantes."
        )

    if research_type == "experimental":
        return (
            f"Los dos primeros componentes explican una proporción {level} de la variabilidad "
            f"({total_2:.1f}%), lo que sugiere patrones experimentales relevantes."
        )

    if research_type == "textual":
        return (
            f"Los dos primeros componentes explican una proporción {level} de la variabilidad "
            f"({total_2:.1f}%), lo que podría reflejar estructuras temáticas importantes."
        )

    if research_type == "social":
        return (
            f"Los dos primeros componentes explican una proporción {level} de la variabilidad "
            f"({total_2:.1f}%), lo que podría reflejar patrones sociales relevantes."
        )

    if research_type == "market_research":
        return (
            f"Los dos primeros componentes explican una proporción {level} de la variabilidad "
            f"({total_2:.1f}%), lo que podría reflejar segmentos o comportamientos de mercado."
        )

    return (
        f"Los dos primeros componentes explican una proporción {level} de la variabilidad "
        f"({total_2:.1f}%), lo que resume parte importante de la estructura del dataset."
    )

def interpret_cluster_result(best_k: int, research_type: str) -> str:
    try:
        k = int(best_k)
    except Exception:
        return "No se pudo generar una interpretación automática del clustering."

    if research_type == "biomedical":
        return (
            f"SADI identificó {k} grupos principales, lo que podría reflejar "
            "perfiles clínicos diferenciados entre los pacientes."
        )

    if research_type == "financial":
        return (
            f"SADI identificó {k} grupos principales, lo que podría reflejar "
            "segmentos financieros con comportamientos económicos distintos."
        )

    if research_type == "educational":
        return (
            f"SADI identificó {k} grupos principales, lo que podría reflejar "
            "perfiles académicos diferenciados entre los estudiantes o cohortes."
        )

    if research_type == "experimental":
        return (
            f"SADI identificó {k} grupos principales, lo que podría indicar "
            "patrones diferenciados en los resultados experimentales."
        )

    if research_type == "textual":
        return (
            f"SADI identificó {k} grupos principales, lo que podría representar "
            "temas o patrones de contenido diferenciados."
        )

    if research_type == "social":
        return (
            f"SADI identificó {k} grupos principales, lo que podría reflejar "
            "subgrupos sociales con características diferenciadas."
        )

    if research_type == "market_research":
        return (
            f"SADI identificó {k} grupos principales, lo que podría reflejar "
            "segmentos de consumidores o mercados diferenciados."
        )

    return (
        f"SADI identificó {k} grupos principales, lo que sugiere "
        "la presencia de patrones diferenciados dentro del dataset."
    )

def interpret_regression_result(r2: float, research_type: str) -> str:
    try:
        r2 = float(r2)
    except Exception:
        return "No se pudo generar interpretación automática del modelo de regresión."

    if r2 >= 0.75:
        level = "muy alto"
    elif r2 >= 0.50:
        level = "alto"
    elif r2 >= 0.30:
        level = "moderado"
    elif r2 >= 0.10:
        level = "bajo"
    else:
        level = "muy bajo"

    if research_type == "biomedical":
        return (
            f"El modelo presenta un nivel de explicación {level} (R²={r2:.2f}), "
            "lo que sugiere capacidad relevante para explicar variación clínica."
        )
    if research_type == "financial":
        return (
            f"El modelo presenta un nivel de explicación {level} (R²={r2:.2f}), "
            "lo que indica capacidad importante para modelar comportamiento financiero."
        )
    if research_type == "educational":
        return (
            f"El modelo presenta un nivel de explicación {level} (R²={r2:.2f}), "
            "lo que sugiere influencia significativa en variables académicas."
        )
    if research_type == "experimental":
        return (
            f"El modelo presenta un nivel de explicación {level} (R²={r2:.2f}), "
            "lo que indica efecto relevante de las variables experimentales."
        )
    if research_type == "social":
        return (
            f"El modelo presenta un nivel de explicación {level} (R²={r2:.2f}), "
            "lo que refleja patrones importantes en el comportamiento social."
        )
    return f"El modelo presenta un nivel de explicación {level} (R²={r2:.2f})."


def interpret_rf_regression_result(r2: float, research_type: str) -> str:
    try:
        r2 = float(r2)
    except Exception:
        return "No se pudo generar interpretación automática del modelo Random Forest."

    if r2 >= 0.75:
        level = "muy alto"
    elif r2 >= 0.50:
        level = "alto"
    elif r2 >= 0.30:
        level = "moderado"
    elif r2 >= 0.10:
        level = "bajo"
    else:
        level = "muy bajo"

    if research_type == "biomedical":
        return (
            f"Random Forest presenta un ajuste {level} (R²={r2:.2f}), "
            "lo que podría captar patrones clínicos no lineales relevantes."
        )
    if research_type == "financial":
        return (
            f"Random Forest presenta un ajuste {level} (R²={r2:.2f}), "
            "lo que sugiere capacidad para capturar comportamientos financieros complejos."
        )
    if research_type == "educational":
        return (
            f"Random Forest presenta un ajuste {level} (R²={r2:.2f}), "
            "lo que podría captar patrones académicos complejos."
        )
    return f"Random Forest presenta un ajuste {level} (R²={r2:.2f})."


def interpret_anova_result(p_value: float | None, target_col: str, group_col: str, research_type: str) -> str:
    if p_value is None:
        return "No se pudo generar una interpretación automática del ANOVA."

    significant = p_value < 0.05
    if research_type == "biomedical":
        if significant:
            return (
                f"Se observaron diferencias estadísticamente significativas entre grupos de '{group_col}' "
                f"sobre '{target_col}' (p={p_value:.6f}), lo que podría reflejar perfiles clínicos distintos."
            )
        return (
            f"No se observaron diferencias estadísticamente significativas entre grupos de '{group_col}' "
            f"sobre '{target_col}' (p={p_value:.6f})."
        )

    if research_type == "financial":
        if significant:
            return (
                f"Se observaron diferencias estadísticamente significativas entre grupos de '{group_col}' "
                f"sobre '{target_col}' (p={p_value:.6f}), lo que podría indicar comportamientos financieros diferenciados."
            )
        return (
            f"No se observaron diferencias estadísticamente significativas entre grupos de '{group_col}' "
            f"sobre '{target_col}' (p={p_value:.6f})."
        )

    if significant:
        return (
            f"Se observaron diferencias estadísticamente significativas entre grupos de '{group_col}' "
            f"sobre '{target_col}' (p={p_value:.6f})."
        )
    return (
        f"No se observaron diferencias estadísticamente significativas entre grupos de '{group_col}' "
        f"sobre '{target_col}' (p={p_value:.6f})."
    )


def interpret_manova_result(p_value: float | None, group_col: str, research_type: str) -> str:
    if p_value is None:
        return "No se pudo generar una interpretación automática de la MANOVA."

    if p_value < 0.05:
        if research_type == "biomedical":
            return (
                f"MANOVA sugiere diferencias multivariadas significativas entre grupos de '{group_col}' "
                f"(p={p_value:.6f}), lo que podría reflejar perfiles clínicos diferenciados."
            )
        if research_type == "financial":
            return (
                f"MANOVA sugiere diferencias multivariadas significativas entre grupos de '{group_col}' "
                f"(p={p_value:.6f}), lo que podría reflejar estructuras financieras diferenciadas."
            )
        return (
            f"MANOVA sugiere diferencias multivariadas significativas entre grupos de '{group_col}' "
            f"(p={p_value:.6f})."
        )

    return (
        f"MANOVA no detectó diferencias multivariadas significativas entre grupos de '{group_col}' "
        f"(p={p_value:.6f})."
    )


def interpret_permanova_result(p_value: float | None, group_col: str, research_type: str) -> str:
    if p_value is None:
        return "No se pudo generar una interpretación automática de la PERMANOVA."

    if p_value < 0.05:
        if research_type == "biomedical":
            return (
                f"PERMANOVA detectó diferencias multivariadas significativas entre grupos de '{group_col}' "
                f"(p={p_value:.6f}), lo que podría reflejar subpoblaciones clínicas distintas."
            )
        if research_type == "financial":
            return (
                f"PERMANOVA detectó diferencias multivariadas significativas entre grupos de '{group_col}' "
                f"(p={p_value:.6f}), lo que podría reflejar segmentos financieros distintos."
            )
        return (
            f"PERMANOVA detectó diferencias multivariadas significativas entre grupos de '{group_col}' "
            f"(p={p_value:.6f})."
        )

    return (
        f"PERMANOVA no detectó diferencias multivariadas significativas entre grupos de '{group_col}' "
        f"(p={p_value:.6f})."
    )

def get_representative_corr_value(df):
    import numpy as np

    numeric_df = df.select_dtypes(include=[np.number])

    if numeric_df.shape[1] < 2:
        return None

    corr = numeric_df.corr(numeric_only=True)

    values = []
    for i in range(len(corr.columns)):
        for j in range(i + 1, len(corr.columns)):
            val = corr.iloc[i, j]
            if pd.notna(val):
                values.append(abs(float(val)))

    if not values:
        return None

    return sum(values) / len(values)

def generate_correlation_heatmap(*, df, dataset_id, plots_dir, research_type="general"):
    import os
    import numpy as np
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    num_df = _get_clean_numeric_df(df)

    if num_df.shape[1] < 2:
        raise ValueError("Se necesitan al menos 2 variables numéricas válidas para generar correlación.")

    corr = num_df.corr(numeric_only=True)

    # ===== correlación representativa promedio (sin diagonal) =====
    corr_values = []
    for i in range(len(corr.columns)):
        for j in range(i + 1, len(corr.columns)):
            val = corr.iloc[i, j]
            if pd.notna(val):
                corr_values.append(abs(float(val)))

    if corr_values:
        corr_value = sum(corr_values) / len(corr_values)
    else:
        corr_value = 0.0

    # ===== interpretación automática =====
    if corr_value >= 0.7:
        interpretacion = "Las variables presentan relaciones fuertes, lo que sugiere alta dependencia estructural."
    elif corr_value >= 0.4:
        interpretacion = "Se observan relaciones moderadas entre variables."
    else:
        interpretacion = "Las variables muestran baja correlación, indicando independencia relativa."

    # ===== guardar imagen =====
    os.makedirs(plots_dir, exist_ok=True)

    out_name = f"ds{dataset_id}_multivariate_corr.png"
    out_path = os.path.join(plots_dir, out_name)

    fig_w = max(6, min(10, 0.45 * len(corr.columns)))
    fig_h = max(5, min(8, 0.45 * len(corr.columns)))

    plt.figure(figsize=(fig_w, fig_h))
    plt.imshow(corr, aspect="auto")
    plt.colorbar()
    plt.xticks(range(len(corr.columns)), corr.columns, rotation=90)
    plt.yticks(range(len(corr.columns)), corr.columns)
    plt.title("Heatmap de correlación")
    plt.tight_layout()
    plt.savefig(out_path, dpi=180, bbox_inches="tight")
    plt.close()

    # ===== matriz segura para JSON/cache =====
    corr_matrix = {}
    for col in corr.columns:
        corr_matrix[str(col)] = {}
        for idx in corr.index:
            corr_matrix[str(col)][str(idx)] = json_safe_number(corr.loc[idx, col], 6)

    return {
        "image_filename": out_name,
        "columns_used": list(num_df.columns),
        "n_variables": int(num_df.shape[1]),
        "corr_value": json_safe_number(corr_value, 6),
        "corr_matrix": corr_matrix,
        "interpretacion": interpretacion,
    }


def run_pca_analysis(
    *,
    df: pd.DataFrame,
    dataset_id: int,
    plots_dir: str,
    research_type: str = "general",
) -> dict:
    num_df, Xs = _prepare_scaled_numeric_matrix(df)

    if num_df.shape[1] < 3:
        raise ValueError("Se necesitan al menos 3 variables numéricas válidas para ejecutar PCA.")

    n_components = min(num_df.shape[1], 5)
    pca = PCA(n_components=n_components)
    components = pca.fit_transform(Xs)

    interpretacion = interpret_pca_result(
        pca.explained_variance_ratio_,
        research_type
    )

    explained = pca.explained_variance_ratio_
    cumulative = np.cumsum(explained)

    os.makedirs(plots_dir, exist_ok=True)

    scree_name = f"ds{dataset_id}_pca_scree.png"
    scree_path = os.path.join(plots_dir, scree_name)

    plt.figure(figsize=(6.5, 5))
    plt.plot(range(1, len(explained) + 1), explained, marker="o")
    plt.xticks(range(1, len(explained) + 1))
    plt.xlabel("Componente principal")
    plt.ylabel("Varianza explicada")
    plt.title("Scree Plot - PCA")
    plt.tight_layout()
    plt.savefig(scree_path, dpi=180, bbox_inches="tight")
    plt.close()

    scatter_name = f"ds{dataset_id}_pca_scatter.png"
    scatter_path = os.path.join(plots_dir, scatter_name)

    if components.shape[1] >= 2:
        plt.figure(figsize=(6.5, 5))
        plt.scatter(components[:, 0], components[:, 1], alpha=0.7)
        plt.xlabel("PC1")
        plt.ylabel("PC2")
        plt.title("PCA - Proyección PC1 vs PC2")
        plt.tight_layout()
        plt.savefig(scatter_path, dpi=180, bbox_inches="tight")
        plt.close()
    else:
        scatter_name = None

    summary_rows = []
    for i, val in enumerate(explained, start=1):
        summary_rows.append({
            "component": f"PC{i}",
            "explained_variance_ratio": json_safe_number(val, 6),
            "explained_pct": json_safe_number(float(val) * 100, 2),
            "cumulative_variance": json_safe_number(cumulative[i - 1], 6),
            "cumulative_pct": json_safe_number(float(cumulative[i - 1]) * 100, 2),
        })

    loadings_df = pd.DataFrame(
        pca.components_.T,
        index=num_df.columns,
        columns=[f"PC{i}" for i in range(1, len(explained) + 1)]
    ).round(4)

    loadings_table = (
        loadings_df.reset_index()
        .rename(columns={"index": "variable"})
        .to_dict(orient="records")
    )

    return {
        "n_components": int(n_components),
        "explained_variance_ratio": [json_safe_number(x, 6) for x in pca.explained_variance_ratio_.tolist()],
        "total_variance_explained": json_safe_number(sum(pca.explained_variance_ratio_), 6),
        "n_variables": int(num_df.shape[1]),
        "n_samples": int(Xs.shape[0]),
        "columns_used": list(num_df.columns),
        "summary": summary_rows,
        "summary_rows": summary_rows,
        "loadings": loadings_table,
        "loadings_table": loadings_table,
        "scree_image": scree_name,
        "scatter_image": scatter_name,
        "interpretacion": interpretacion,
    }


def run_kmeans_analysis(
    *,
    df: pd.DataFrame,
    dataset_id: int,
    plots_dir: str,
    research_type: str = "general",
) -> dict:
    num_df, Xs = _prepare_scaled_numeric_matrix(df)

    if num_df.shape[1] < 2:
        raise ValueError("Se necesitan al menos 2 variables numéricas válidas para clustering.")

    n_samples = Xs.shape[0]
    if n_samples < 8:
        raise ValueError("Se requieren al menos 8 filas para realizar clustering de forma razonable.")

    max_k = min(8, max(3, n_samples // 5))
    k_values = list(range(2, max_k + 1))

    inertias = []
    for k in k_values:
        km = KMeans(n_clusters=k, random_state=42, n_init=10)
        km.fit(Xs)
        inertias.append(float(km.inertia_))

    # heurística simple del codo
    best_k = 3

    if len(inertias) >= 3:
        deltas = []
        for i in range(1, len(inertias)):
            deltas.append(inertias[i - 1] - inertias[i])

        ratios = []
        for i in range(1, len(deltas)):
            prev_drop = deltas[i - 1] if deltas[i - 1] != 0 else 1e-9
            ratios.append(deltas[i] / prev_drop)

        elbow_idx = 0
        min_ratio = math.inf
        for i, r in enumerate(ratios):
            if r < min_ratio:
                min_ratio = r
                elbow_idx = i + 1

        best_k = k_values[min(elbow_idx, len(k_values) - 1)]
    elif len(k_values) > 0:
        best_k = k_values[0]

    final_km = KMeans(n_clusters=best_k, random_state=42, n_init=10)
    labels = final_km.fit_predict(Xs)

    pca2 = PCA(n_components=2)
    coords = pca2.fit_transform(Xs)

    os.makedirs(plots_dir, exist_ok=True)

    elbow_name = f"ds{dataset_id}_kmeans_elbow.png"
    elbow_path = os.path.join(plots_dir, elbow_name)

    plt.figure(figsize=(6.8, 4.2))
    plt.plot(k_values, inertias, marker="o")
    plt.xticks(k_values)
    plt.xlabel("Número de clusters (k)")
    plt.ylabel("Inercia")
    plt.title("Elbow Plot - KMeans")
    plt.tight_layout()
    plt.savefig(elbow_path, dpi=180, bbox_inches="tight")
    plt.close()

    cluster_name = f"ds{dataset_id}_kmeans_clusters.png"
    cluster_path = os.path.join(plots_dir, cluster_name)

    plt.figure(figsize=(6.8, 4.2))
    plt.scatter(coords[:, 0], coords[:, 1], c=labels, alpha=0.8)
    plt.xlabel("PC1")
    plt.ylabel("PC2")
    plt.title(f"K-Means Clustering (k={best_k})")
    plt.tight_layout()
    plt.savefig(cluster_path, dpi=180, bbox_inches="tight")
    plt.close()

    cluster_counts = (
        pd.Series(labels)
        .value_counts()
        .sort_index()
        .rename_axis("cluster")
        .reset_index(name="n")
    )
    cluster_counts["cluster"] = cluster_counts["cluster"].apply(lambda x: f"Cluster {x}")

    interpretacion = interpret_cluster_result(best_k, research_type)

    return {
    "best_k": int(best_k),
    "elbow_image": elbow_name,
    "cluster_image": cluster_name,
    "cluster_counts": [
        {
            "cluster": row["cluster"],
            "n": int(row["n"]) if row["n"] is not None else None,
        }
        for _, row in cluster_counts.iterrows()
    ],
    "n_variables": int(num_df.shape[1]),
    "n_samples": int(n_samples),
    "columns_used": list(num_df.columns),
    "interpretacion": interpretacion,
}
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Image as RLImage,
    Table,
    TableStyle,
    PageBreak,
    KeepTogether
)
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import cm


def generate_multivariate_report_pdf(
    *,
    dataset_title: str,
    profile: dict,
    plots_dir: str,
    dataset_id: int,
    output_path: str,
    pca_summary=None,
    pca_loadings=None,
    cluster_counts=None,
    cluster_best_k=None,
    efa_kmo=None,
    efa_bartlett_p=None,
    efa_n_factors=None,
    efa_loadings=None,
):
    styles = getSampleStyleSheet()
    story = []

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    story.append(Paragraph("Reporte Multivariado - SADI", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"Dataset: {dataset_title}", styles["Heading2"]))
    story.append(Spacer(1, 12))

    summary_data = [
        ["Filas", str(profile.get("n_rows", ""))],
        ["Columnas", str(profile.get("n_cols", ""))],
        ["Numéricas", str(profile.get("n_num", ""))],
        ["Categóricas", str(profile.get("n_cat", ""))],
        ["Missing (%)", str(profile.get("missing_pct", ""))],
        ["Perfil", str(profile.get("profile_type", ""))],
    ]
    t = Table(summary_data, colWidths=[6 * cm, 8 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))

    story.append(Paragraph("Recomendaciones SADI", styles["Heading2"]))
    for rec in profile.get("quick_recommendations", []):
        story.append(Paragraph(f"- {rec}", styles["BodyText"]))
    story.append(Spacer(1, 12))

    image_candidates = [
        f"ds{dataset_id}_multivariate_corr.png",
        f"ds{dataset_id}_pca_scree.png",
        f"ds{dataset_id}_pca_scatter.png",
        f"ds{dataset_id}_kmeans_elbow.png",
        f"ds{dataset_id}_kmeans_clusters.png",
    ]

    for img_name in image_candidates:
        img_path = os.path.join(plots_dir, img_name)
        if os.path.exists(img_path):
            story.append(Spacer(1, 8))
            story.append(RLImage(img_path, width=16 * cm, height=10 * cm))
            story.append(Spacer(1, 8))

    doc.build(story)

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Image as RLImage,
    Table,
    TableStyle,
    PageBreak,
)
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import cm


def generate_multivariate_report_pdf(
    *,
    dataset_title: str,
    profile: dict,
    plots_dir: str,
    dataset_id: int,
    output_path: str,
    pca_summary=None,
    pca_loadings=None,
    cluster_counts=None,
    cluster_best_k=None,
    efa_kmo=None,
    efa_bartlett_p=None,
    efa_n_factors=None,
    efa_loadings=None,
):
    styles = getSampleStyleSheet()

    title_style = styles["Title"]
    heading_style = styles["Heading2"]
    normal_style = styles["BodyText"]

    small_style = ParagraphStyle(
        "SmallCustom",
        parent=styles["BodyText"],
        fontSize=9,
        leading=12,
        spaceAfter=6,
    )

    story = []

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=1.8 * cm,
        leftMargin=1.8 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
    )

    # Portada simple
    story.append(Paragraph("Reporte Multivariado - SADI", title_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"Dataset: {dataset_title}", heading_style))
    story.append(Spacer(1, 14))

    # Resumen general
    story.append(Paragraph("1. Perfil del dataset", heading_style))
    summary_data = [
        ["Indicador", "Valor"],
        ["Filas", str(profile.get("n_rows", ""))],
        ["Columnas", str(profile.get("n_cols", ""))],
        ["Variables numéricas", str(profile.get("n_num", ""))],
        ["Variables categóricas", str(profile.get("n_cat", ""))],
        ["Missing (%)", str(profile.get("missing_pct", ""))],
        ["Perfil detectado", str(profile.get("profile_type", ""))],
    ]
    t = Table(summary_data, colWidths=[6.2 * cm, 9.0 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9e8fb")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("PADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))

    # Recomendaciones rápidas
    story.append(Paragraph("2. Recomendaciones automáticas de SADI", heading_style))
    quick = profile.get("quick_recommendations", []) or []
    if quick:
        for item in quick:
            story.append(Paragraph(f"• {item}", normal_style))
    else:
        story.append(Paragraph("No se generaron recomendaciones rápidas.", normal_style))
    story.append(Spacer(1, 10))

    # Detalles del perfil
    story.append(Paragraph("3. Interpretación inicial", heading_style))

    usable_num = profile.get("usable_num_cols", []) or []
    high_corr_pairs = profile.get("high_corr_pairs", []) or []

    interpretation_lines = []
    if len(usable_num) >= 3:
        interpretation_lines.append(
            f"El dataset presenta {len(usable_num)} variables numéricas utilizables, lo que lo hace apto para PCA y clustering."
        )
    elif len(usable_num) >= 2:
        interpretation_lines.append(
            f"El dataset presenta {len(usable_num)} variables numéricas utilizables, suficientes para análisis de correlación."
        )
    else:
        interpretation_lines.append(
            "El dataset tiene pocas variables numéricas utilizables para análisis multivariado avanzado."
        )

    if high_corr_pairs:
        interpretation_lines.append(
            f"Se detectaron {len(high_corr_pairs)} pares de variables con correlación alta, lo que sugiere posible redundancia o estructura latente."
        )
    else:
        interpretation_lines.append(
            "No se detectaron correlaciones altas destacadas entre variables numéricas."
        )

    if profile.get("missing_pct", 0) and float(profile.get("missing_pct", 0)) > 10:
        interpretation_lines.append(
            "El porcentaje de valores faltantes es relevante; conviene revisar limpieza de datos antes de análisis más exigentes."
        )
    else:
        interpretation_lines.append(
            "El nivel de datos faltantes parece manejable para análisis exploratorio."
        )

    for line in interpretation_lines:
        story.append(Paragraph(f"• {line}", normal_style))
    story.append(Spacer(1, 10))

    # Correlaciones altas
    story.append(Paragraph("4. Correlaciones altas detectadas", heading_style))
    if high_corr_pairs:
        corr_table_data = [["Variable 1", "Variable 2", "Correlación"]]
        for item in high_corr_pairs[:20]:
            corr_table_data.append([
                str(item.get("col1", "")),
                str(item.get("col2", "")),
                str(item.get("corr", "")),
            ])

        corr_table = Table(corr_table_data, colWidths=[5.5 * cm, 5.5 * cm, 3.0 * cm])
        corr_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef7")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(corr_table)
    else:
        story.append(Paragraph("No se detectaron correlaciones altas.", normal_style))
    story.append(Spacer(1, 12))

    # Función auxiliar para insertar imagen
    def add_image_block(title: str, explanation: str, filename: str):
        img_path = os.path.join(plots_dir, filename)
        if not os.path.exists(img_path):
            return

        block = []

        block.append(Paragraph(title, heading_style))
        block.append(Paragraph(explanation, small_style))
        block.append(Spacer(1, 6))

        img = RLImage(img_path, width=16.2 * cm, height=9.2 * cm)
        block.append(img)

        block.append(Spacer(1, 12))

        story.append(KeepTogether(block))

    # Imágenes
    add_image_block(
        "5. Heatmap de correlación",
        "El heatmap resume visualmente la intensidad de asociación entre las variables numéricas del dataset.",
        f"ds{dataset_id}_multivariate_corr.png",
    )

    add_image_block(
        "6. Scree Plot del PCA",
        "Este gráfico muestra cuánto aporta cada componente principal a la varianza explicada total.",
        f"ds{dataset_id}_pca_scree.png",
    )

    add_image_block(
        "7. Proyección PCA (PC1 vs PC2)",
        "Esta proyección permite observar la estructura general del dataset en dos dimensiones reducidas.",
        f"ds{dataset_id}_pca_scatter.png",
    )

    add_image_block(
        "8. Elbow Plot de K-Means",
        "Este gráfico ayuda a sugerir un número razonable de clusters para segmentación.",
        f"ds{dataset_id}_kmeans_elbow.png",
    )

    add_image_block(
        "9. Clusters proyectados en PC1 y PC2",
        "Esta visualización muestra la segmentación K-Means sobre una proyección bidimensional del dataset.",
        f"ds{dataset_id}_kmeans_clusters.png",
    )
    add_image_block(
        "10. Scree Plot del análisis factorial exploratorio",
        "Este gráfico ayuda a identificar el número de factores con autovalores relevantes en el análisis factorial exploratorio.",
        f"ds{dataset_id}_efa_scree.png",
    )
        # Tabla de varianza explicada del PCA
    if pca_summary:
        story.append(Paragraph("10. Tabla de varianza explicada del PCA", heading_style))
        story.append(Paragraph(
            "La siguiente tabla resume cuánto aporta cada componente principal a la explicación de la variabilidad total del dataset.",
            small_style
        ))
        story.append(Spacer(1, 6))

        pca_table_data = [[
            "Componente",
            "Var. explicada",
            "% explicada",
            "Var. acumulada",
            "% acumulado",
        ]]

        for row in pca_summary:
            pca_table_data.append([
                str(row.get("component", "")),
                str(row.get("explained_variance_ratio", "")),
                f"{row.get('explained_pct', '')}%",
                str(row.get("cumulative_variance", "")),
                f"{row.get('cumulative_pct', '')}%",
            ])

        pca_table = Table(
            pca_table_data,
            colWidths=[2.5 * cm, 3.0 * cm, 3.0 * cm, 3.2 * cm, 3.0 * cm]
        )
        pca_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef7")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 5),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(pca_table)
        story.append(Spacer(1, 12))

    # Tabla de loadings del PCA
    if pca_loadings:
        story.append(Paragraph("11. Loadings del PCA", heading_style))
        story.append(Paragraph(
            "Los loadings indican el peso de cada variable original dentro de los componentes principales y ayudan a interpretar cada dimensión reducida.",
            small_style
        ))
        story.append(Spacer(1, 6))

        # limitar columnas para que entre en A4
        header_keys = list(pca_loadings[0].keys()) if pca_loadings else []
        preferred_keys = ["variable", "PC1", "PC2", "PC3", "PC4", "PC5"]
        final_keys = [k for k in preferred_keys if k in header_keys]

        loadings_table_data = [final_keys]

        for row in pca_loadings:
            loadings_table_data.append([str(row.get(k, "")) for k in final_keys])

        loadings_col_widths = []
        for key in final_keys:
            if key == "variable":
                loadings_col_widths.append(5.0 * cm)
            else:
                loadings_col_widths.append(2.2 * cm)

        loadings_table = Table(loadings_table_data, colWidths=loadings_col_widths, repeatRows=1)
        loadings_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef7")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
        ]))
        story.append(loadings_table)
        story.append(Spacer(1, 12))

    # Tabla de clusters
    if cluster_counts:
        story.append(Paragraph("12. Distribución por cluster", heading_style))
        if cluster_best_k is not None:
            story.append(Paragraph(
                f"SADI sugiere una segmentación con k = {cluster_best_k}. La tabla resume cuántas observaciones fueron asignadas a cada grupo.",
                small_style
            ))
        else:
            story.append(Paragraph(
                "La siguiente tabla resume cuántas observaciones fueron asignadas a cada grupo detectado por K-Means.",
                small_style
            ))
        story.append(Spacer(1, 6))

        cluster_table_data = [["Cluster", "Cantidad"]]
        for row in cluster_counts:
            cluster_table_data.append([
                str(row.get("cluster", "")),
                str(row.get("n", "")),
            ])

        cluster_table = Table(cluster_table_data, colWidths=[6.0 * cm, 4.0 * cm])
        cluster_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef7")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(cluster_table)
        story.append(Spacer(1, 12))

    # Tabla EFA
    if efa_loadings:
        story.append(Paragraph("13. Resultados del análisis factorial exploratorio (EFA)", heading_style))

        efa_intro = f"KMO = {efa_kmo} | Bartlett p = {efa_bartlett_p}"
        if efa_n_factors is not None:
            efa_intro += f" | Factores sugeridos = {efa_n_factors}"

        story.append(Paragraph(efa_intro, small_style))
        story.append(Paragraph(
            "La siguiente tabla resume las cargas factoriales obtenidas para cada variable en los factores extraídos.",
            small_style
        ))
        story.append(Spacer(1, 6))

        header_keys = list(efa_loadings[0].keys()) if efa_loadings else []
        efa_table_data = [header_keys]

        for row in efa_loadings:
            efa_table_data.append([str(row.get(k, "")) for k in header_keys])

        efa_col_widths = []
        for key in header_keys:
            if str(key).lower() == "variable":
                efa_col_widths.append(5.0 * cm)
            else:
                efa_col_widths.append(2.2 * cm)

        efa_table = Table(efa_table_data, colWidths=efa_col_widths, repeatRows=1)
        efa_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8eef7")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
        ]))
        story.append(efa_table)
        story.append(Spacer(1, 12))
    # Cierre
    story.append(Paragraph("14. Conclusión general", heading_style))
    story.append(Paragraph(
        "Este reporte resume el perfil multivariado del dataset y ofrece una base para continuar con análisis más profundos como reducción de dimensionalidad, segmentación, análisis factorial o modelado predictivo.",
        normal_style
    ))

    doc.build(story)

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def generate_multivariate_report_docx(
    *,
    dataset_title: str,
    profile: dict,
    plots_dir: str,
    dataset_id: int,
    output_path: str,
    figure_catalog=None,
    pca_summary=None,
    pca_loadings=None,
    cluster_counts=None,
    cluster_best_k=None,
    efa_kmo=None,
    efa_bartlett_p=None,
    efa_n_factors=None,
    efa_loadings=None,
    regression_result=None,
    rf_result=None,
    logistic_result=None,
    rf_classification_result=None,
    model_comparison=None,
    anova_result=None,
    manova_result=None,
    permanova_result=None,
    group_visuals=None,
    corr_interpretacion=None,
    pca_interpretacion=None,
    cluster_interpretacion=None,
    conclusion=None,
    abstract_text=None,
    keywords=None,
    insights_ranking=None,

    analysis_meta=None,
    insights=None,
    insights_text=None,
    quick_recommendations=None,
    suggested_plan=None,
    plot_summary=None,
    general_figure_catalog=None,
):
    import os
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # =========================
    # Configuración general
    # =========================
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    # =========================
    # Helpers
    # =========================
    def safe_str(val, default=""):
        if val is None:
            return default
        return str(val)

    def fmt_num(val, digits=4, default="—"):
        try:
            if val is None or val == "":
                return default
            return f"{float(val):.{digits}f}"
        except Exception:
            return safe_str(val, default)

    def add_bullets(items, empty_text=None):
        if not items:
            if empty_text:
                doc.add_paragraph(empty_text)
            return
        for item in items:
            doc.add_paragraph(safe_str(item), style="List Bullet")

    def add_interpretacion_paragraph(text, prefix="Interpretación automática de SADI:"):
        if not text:
            return
        p = doc.add_paragraph()
        r1 = p.add_run(prefix + " ")
        r1.bold = True
        p.add_run(safe_str(text))

    def add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = safe_str(h)

        for row_data in rows:
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = safe_str(val)
        return table

    def add_image_section(title: str, explanation: str, filename: str, width=5.7):
        img_path = os.path.join(plots_dir, filename)
        if not os.path.exists(img_path):
            return False

        p_title = doc.add_paragraph()
        p_title.style = "Heading 2"
        p_title.paragraph_format.keep_with_next = True
        p_title.add_run(title)

        p_exp = doc.add_paragraph(explanation)
        p_exp.paragraph_format.keep_with_next = True

        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        cell = tbl.cell(0, 0)
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(img_path, width=Inches(width))

        doc.add_paragraph("")
        return True

    # =========================
    # Variables derivadas
    # =========================
    n_rows = profile.get("n_rows", "")
    n_cols = profile.get("n_cols", "")
    n_num = profile.get("n_num", "")
    n_cat = profile.get("n_cat", "")
    profile_type = profile.get("profile_type", "")

    usable_num = profile.get("usable_num_cols", []) or profile.get("num_cols", []) or []
    usable_cat = profile.get("usable_cat_cols", []) or profile.get("cat_cols", []) or []
    high_corr_pairs = profile.get("high_corr_pairs", []) or []
    quick = profile.get("quick_recommendations", []) or []

    has_exploratory = any([
        os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_multivariate_corr.png")),
        os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_pca_scree.png")),
        os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_kmeans_clusters.png")),
        os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_efa_scree.png")),
        os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_outliers_pca.png")),
    ])

    has_group_analysis = any([anova_result, manova_result, permanova_result, group_visuals])
    has_regression_models = any([regression_result, rf_result, model_comparison])
    has_classification_models = any([logistic_result, rf_classification_result])

    logistic_conf_exists = os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_logistic_confusion.png"))
    logistic_roc_exists = os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_logistic_roc.png"))
    rf_cls_conf_exists = os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_rf_classification_confusion.png"))
    rf_cls_roc_exists = os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_rf_classification_roc.png"))
    rf_cls_imp_exists = os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_rf_classification_importance.png"))

    # =========================
    # Portada
    # =========================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Reporte Multivariado - SADI")
    r.bold = True
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Dataset: {dataset_title}")
    r.italic = True
    r.font.size = Pt(12)

    doc.add_paragraph("")

    # =========================
    # Resumen / Abstract
    # =========================
    doc.add_heading(dataset_title, level=0)

    # ABSTRACT
    if abstract_text:
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(abstract_text)

    # KEYWORDS
    if keywords:
        doc.add_paragraph("Keywords: " + ", ".join(keywords))

    # =========================
    # 1. Resumen ejecutivo
    # =========================
    doc.add_heading("1. Resumen ejecutivo", level=1)

    executive_text = (
        f"El dataset '{dataset_title}' contiene {n_rows} filas y {n_cols} columnas. "
        f"SADI detectó {n_num} variables numéricas y {n_cat} variables categóricas, "
        f"clasificando el perfil general como '{profile_type}'. "
    )

    if has_exploratory:
        executive_text += "Se ejecutaron análisis exploratorios multivariados para examinar correlaciones, estructura latente, segmentación y posibles observaciones atípicas. "
    if has_group_analysis:
        executive_text += "Se desarrollaron análisis inferenciales por grupos para contrastar diferencias entre categorías. "
    if has_regression_models:
        executive_text += "Se evaluaron modelos predictivos de regresión y se comparó su rendimiento. "
    if has_classification_models or logistic_conf_exists or logistic_roc_exists:
        executive_text += "También se incorporaron modelos de clasificación binaria para evaluar capacidad predictiva sobre variables objetivo categóricas. "

    doc.add_paragraph(executive_text)

    if quick:
        doc.add_paragraph("Principales recomendaciones automáticas de SADI:")
        add_bullets(quick)

    # =========================
    # 2. Perfil del dataset
    # =========================
    doc.add_heading("2. Perfil del dataset", level=1)

    add_table(
        ["Indicador", "Valor"],
        [
            ("Filas", n_rows),
            ("Columnas", n_cols),
            ("Variables numéricas", n_num),
            ("Variables categóricas", n_cat),
            ("Missing (%)", profile.get("missing_pct", "")),
            ("Perfil detectado", profile_type),
        ]
    )
    doc.add_paragraph("")

    # =========================
    # 3. Interpretación inicial
    # =========================
    doc.add_heading("3. Interpretación inicial del perfil", level=1)
    interpretation_lines = []

    if len(usable_num) >= 3:
        interpretation_lines.append(
            f"Se identificaron {len(usable_num)} variables numéricas utilizables, lo que favorece análisis como PCA, clustering y modelado predictivo."
        )
    elif len(usable_num) >= 2:
        interpretation_lines.append(
            f"Se identificaron {len(usable_num)} variables numéricas utilizables, suficientes para correlación y análisis exploratorio básico."
        )
    else:
        interpretation_lines.append(
            "El número de variables numéricas utilizables es limitado para análisis multivariado avanzado."
        )

    if usable_cat:
        interpretation_lines.append(
            f"Se detectaron {len(usable_cat)} variables categóricas utilizables, lo que permite contrastes por grupos como ANOVA, MANOVA o PERMANOVA."
        )
    else:
        interpretation_lines.append(
            "No se detectaron variables categóricas suficientes para análisis comparativos entre grupos."
        )

    if high_corr_pairs:
        interpretation_lines.append(
            f"Se detectaron {len(high_corr_pairs)} pares de variables con correlación alta, lo que sugiere redundancia o posible estructura latente."
        )
    else:
        interpretation_lines.append(
            "No se observaron correlaciones altas especialmente destacables entre variables numéricas."
        )

    try:
        missing_pct = float(profile.get("missing_pct", 0) or 0)
    except Exception:
        missing_pct = 0.0

    if missing_pct > 10:
        interpretation_lines.append(
            "El porcentaje de valores faltantes es relativamente elevado y conviene considerarlo en la interpretación de resultados."
        )
    else:
        interpretation_lines.append(
            "El nivel de valores faltantes parece manejable para los análisis ejecutados."
        )

    add_bullets(interpretation_lines)


    # =========================
    # SADI: Insights automáticos
    # =========================
    if insights:
        doc.add_heading("Hallazgos automáticos de SADI", level=1)
        add_bullets(insights)

    if insights_text:
        doc.add_heading("Conclusión narrativa inicial", level=1)
        doc.add_paragraph(safe_str(insights_text))

    if quick_recommendations:
        doc.add_heading("Recomendaciones rápidas de SADI", level=1)
        add_bullets(quick_recommendations)

    if suggested_plan:
        doc.add_heading("Plan sugerido por SADI", level=1)

        recommended_analysis = suggested_plan.get("recommended_analysis", []) or []
        recommended_plots = suggested_plan.get("recommended_plots", []) or []
        narrative_focus = suggested_plan.get("narrative_focus")
        warnings = suggested_plan.get("warnings", []) or []

        if recommended_analysis:
            doc.add_paragraph("Análisis sugeridos:")
            add_bullets(recommended_analysis)

        if recommended_plots:
            doc.add_paragraph("Gráficos recomendados:")
            add_bullets(recommended_plots)

        if narrative_focus:
            p = doc.add_paragraph()
            r = p.add_run("Enfoque sugerido: ")
            r.bold = True
            p.add_run(safe_str(narrative_focus))

        if warnings:
            doc.add_paragraph("Advertencias:")
            add_bullets(warnings)

    # =========================
    # SADI INTELIGENTE (COMPLETO)
    # =========================

    if analysis_meta:
        doc.add_heading("Análisis inteligente SADI", level=1)

        # Tipo y área
        tipo = analysis_meta.get("dataset_type", "")
        area = analysis_meta.get("research_area", "")

        p = doc.add_paragraph()
        r = p.add_run("Tipo de dataset: ")
        r.bold = True
        p.add_run(str(tipo))

        p = doc.add_paragraph()
        r = p.add_run("Área de investigación detectada: ")
        r.bold = True
        p.add_run(str(area))


    # =========================
    # INSIGHTS AUTOMÁTICOS (MEJORADOS)
    # =========================
    if insights:
        doc.add_heading("Insights automáticos de SADI", level=1)

        for item in insights:
            if isinstance(item, dict):
                txt = item.get("text", "")
                cat = item.get("category", "")
                score = item.get("score", "")

                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"[{cat}] ")
                run.bold = True
                p.add_run(f"{txt} (impacto: {score})")
            else:
                doc.add_paragraph(str(item), style="List Bullet")


    # =========================
    # NARRATIVA INTELIGENTE
    # =========================
    if insights_text:
        doc.add_heading("Interpretación global automática", level=1)
        doc.add_paragraph(insights_text)


    # =========================
    # RECOMENDACIONES
    # =========================
    if quick_recommendations:
        doc.add_heading("Recomendaciones estratégicas", level=1)
        add_bullets(quick_recommendations)


    # =========================
    # PLAN SUGERIDO
    # =========================
    if suggested_plan:
        doc.add_heading("Plan de análisis sugerido por SADI", level=1)

        ra = suggested_plan.get("recommended_analysis", [])
        rp = suggested_plan.get("recommended_plots", [])
        nf = suggested_plan.get("narrative_focus")
        warn = suggested_plan.get("warnings", [])

        if ra:
            doc.add_paragraph("Análisis recomendados:")
            add_bullets(ra)

        if rp:
            doc.add_paragraph("Visualizaciones recomendadas:")
            add_bullets(rp)

        if nf:
            p = doc.add_paragraph()
            r = p.add_run("Enfoque narrativo sugerido: ")
            r.bold = True
            p.add_run(str(nf))

        if warn:
            doc.add_paragraph("Advertencias metodológicas:")
            add_bullets(warn)


    # =========================
    # PRIORIDAD DE INSIGHTS
    # =========================
    priority = analysis_meta.get("priority_order", [])
    if priority:
        doc.add_heading("Orden de prioridad analítica", level=1)
        for i, item in enumerate(priority, 1):
            doc.add_paragraph(f"{i}. {item}")

    # =========================
    # 4. Correlaciones altas
    # =========================
    doc.add_heading("4. Correlaciones altas detectadas", level=1)
    if high_corr_pairs:
        rows = []
        for item in high_corr_pairs[:20]:
            rows.append([
                item.get("col1", ""),
                item.get("col2", ""),
                item.get("corr", ""),
            ])
        add_table(["Variable 1", "Variable 2", "Correlación"], rows)
    else:
        doc.add_paragraph("No se detectaron correlaciones altas.")
    doc.add_paragraph("")

    # =========================
    # Resumen visual del dataset
    # =========================
    if plot_summary:
        doc.add_heading("Resumen de gráficos generados", level=1)
        rows = []
        for item in plot_summary:
            if isinstance(item, dict):
                rows.append([
                    item.get("tag", ""),
                    item.get("count", ""),
                ])
        if rows:
            add_table(["Tipo de gráfico", "Cantidad"], rows)

    if general_figure_catalog:
        doc.add_heading("Visualizaciones generales del dataset", level=1)
        for fig in general_figure_catalog:
            add_image_section(
                fig.get("title", "Figura"),
                fig.get("caption", ""),
                fig.get("basename", ""),
                width=5.7,
            )

    # =========================
    # 5. Catálogo de figuras
    # =========================
    doc.add_heading("5. Resultados gráficos", level=1)

    if figure_catalog:
        section_idx = 5
        for fig in figure_catalog:
            add_image_section(
                f"{section_idx}. {fig.get('title', 'Figura')}",
                fig.get("caption", ""),
                os.path.basename(fig.get("filename", "")),
            )
            section_idx += 1
    else:
        doc.add_paragraph("No se detectaron figuras disponibles para exportación.")
        if corr_interpretacion:
            add_interpretacion_paragraph(corr_interpretacion)
        add_image_section(
            "5.1 Heatmap de correlación",
            "El heatmap resume visualmente la intensidad de asociación entre las variables numéricas del dataset.",
            f"ds{dataset_id}_multivariate_corr.png",
        )
        add_image_section(
            "5.2 Scree Plot del PCA",
            "Este gráfico muestra cuánto aporta cada componente principal a la varianza explicada total.",
            f"ds{dataset_id}_pca_scree.png",
        )

    if pca_summary:
        doc.add_heading("6. Tabla de varianza explicada del PCA", level=1)
        doc.add_paragraph(
            "La siguiente tabla resume cuánto aporta cada componente principal a la explicación de la variabilidad total del dataset."
        )

        if pca_interpretacion:
            add_interpretacion_paragraph(pca_interpretacion)

        rows = []
        for idx, row_data in enumerate(pca_summary):
            comp_name = row_data.get("component", "")
            explained_pct = row_data.get("explained_pct", "")
            cumulative_pct = row_data.get("cumulative_pct", "")

            interp = ""
            try:
                exp_val = float(explained_pct)
                cum_val = float(cumulative_pct)

                if idx == 0:
                    if exp_val >= 40:
                        interp = "Componente dominante"
                    elif exp_val >= 20:
                        interp = "Componente relevante"
                    else:
                        interp = "Aporte limitado"
                else:
                    if cum_val >= 70:
                        interp = "Consolida buena explicación acumulada"
                    elif exp_val >= 15:
                        interp = "Aporta varianza complementaria"
                    else:
                        interp = "Aporte secundario"
            except Exception:
                interp = ""

            rows.append([
                comp_name,
                row_data.get("explained_variance_ratio", ""),
                f"{explained_pct}%",
                row_data.get("cumulative_variance", ""),
                f"{cumulative_pct}%",
                interp,
            ])

        add_table(
            ["Componente", "Var. explicada", "% explicada", "Var. acumulada", "% acumulado", "Interpretación"],
            rows
        )
        doc.add_paragraph("")

    if pca_loadings:
        doc.add_heading("7. Loadings del PCA", level=1)
        doc.add_paragraph(
            "Los loadings indican el peso de cada variable original dentro de los componentes principales."
        )
        header_keys = list(pca_loadings[0].keys())
        rows = []
        for item in pca_loadings:
            rows.append([item.get(k, "") for k in header_keys])
        add_table(header_keys, rows)
        doc.add_paragraph("")

    if cluster_counts:
        doc.add_heading("8. Distribución por cluster", level=1)
        if cluster_best_k is not None:
            doc.add_paragraph(
                f"SADI sugiere una segmentación con k = {cluster_best_k}. La siguiente tabla resume la distribución por cluster."
            )

        if cluster_interpretacion:
            add_interpretacion_paragraph(cluster_interpretacion)

        total_cluster_n = 0
        for item in cluster_counts:
            try:
                total_cluster_n += int(item.get("n", 0) or 0)
            except Exception:
                pass

        rows = []
        for item in cluster_counts:
            cluster_name = item.get("cluster", "")
            cluster_n = item.get("n", "")

            interp = ""
            try:
                n_val = int(cluster_n)
                if total_cluster_n > 0:
                    ratio = n_val / total_cluster_n
                    if ratio >= 0.5:
                        interp = "Grupo predominante"
                    elif ratio >= 0.2:
                        interp = "Grupo de tamaño intermedio"
                    else:
                        interp = "Grupo pequeño o específico"
            except Exception:
                interp = ""

            rows.append([cluster_name, cluster_n, interp])

        add_table(["Cluster", "Cantidad", "Interpretación"], rows)
        doc.add_paragraph("")

    if efa_loadings:
        doc.add_heading("9. Resultados del análisis factorial exploratorio (EFA)", level=1)

        efa_text = f"KMO = {fmt_num(efa_kmo)} | Bartlett p = {fmt_num(efa_bartlett_p)}"
        if efa_n_factors is not None:
            efa_text += f" | Factores sugeridos = {safe_str(efa_n_factors)}"

        doc.add_paragraph(efa_text)
        doc.add_paragraph(
            "La siguiente tabla muestra las cargas factoriales de cada variable en los factores extraídos."
        )

        header_keys = list(efa_loadings[0].keys())
        rows = []
        for item in efa_loadings:
            rows.append([item.get(k, "") for k in header_keys])
        add_table(header_keys, rows)
        doc.add_paragraph("")
    else:
        if os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_efa_scree.png")):
            doc.add_heading("9. Resultados del análisis factorial exploratorio (EFA)", level=1)
            doc.add_paragraph(
                "SADI generó el gráfico exploratorio del EFA, pero la estimación tabular completa no pudo reconstruirse de forma estable para este dataset. Esto suele ocurrir cuando existen colinealidades extremas, varianza casi nula o problemas numéricos en la matriz."
            )
            doc.add_paragraph("")

    # =========================
    # 10. Análisis por grupos
    # =========================
    if anova_result or manova_result or permanova_result or group_visuals:
        doc.add_heading("10. Análisis inferencial por grupos", level=1)

    if anova_result:
        doc.add_heading("10.1 ANOVA", level=2)
        doc.add_paragraph(
            f"Variable dependiente: {anova_result.get('target_col', '')} | "
            f"Grupo: {anova_result.get('group_col', '')}"
        )

        if anova_result.get("interpretacion"):
            add_interpretacion_paragraph(anova_result.get("interpretacion"))

        if anova_result.get("interpretation"):
            add_bullets(anova_result.get("interpretation", []))

        anova_table = anova_result.get("anova_table", []) or []
        if anova_table:
            rows = []
            for row in anova_table:
                rows.append([
                    row.get("source", ""),
                    row.get("sum_sq", ""),
                    row.get("df", ""),
                    row.get("f_value", ""),
                    row.get("p_value", ""),
                ])
            add_table(["Fuente", "Suma cuadrados", "gl", "F", "p-value"], rows)

        group_summary = anova_result.get("group_summary", []) or []
        if group_summary:
            rows = []
            group_col = anova_result.get("group_col", "Grupo")
            for row in group_summary:
                rows.append([
                    row.get(group_col, ""),
                    row.get("n", ""),
                    row.get("mean", ""),
                    row.get("std", ""),
                ])
            add_table([group_col, "n", "Media", "Std"], rows)

        add_image_section(
            "10.1.1 Boxplot por grupo",
            "Este gráfico compara visualmente la distribución de la variable dependiente entre grupos.",
            f"ds{dataset_id}_anova_boxplot.png",
        )

    if manova_result:
        doc.add_heading("10.2 MANOVA", level=2)
        doc.add_paragraph(
            f"Grupo: {manova_result.get('group_col', '')} | "
            f"Dependientes: {', '.join(manova_result.get('dependent_cols', []) or [])}"
        )

        if manova_result.get("interpretacion"):
            add_interpretacion_paragraph(manova_result.get("interpretacion"))

        if manova_result.get("interpretation"):
            add_bullets(manova_result.get("interpretation", []))

        doc.add_paragraph(manova_result.get("manova_text", "Sin salida textual de MANOVA."))
        doc.add_paragraph("")

    if permanova_result:
        doc.add_heading("10.3 PERMANOVA", level=2)
        doc.add_paragraph(
            f"Grupo: {permanova_result.get('group_col', '')} | "
            f"Variables numéricas: {permanova_result.get('n_variables', '')} | "
            f"Filas: {permanova_result.get('n_rows', '')}"
        )

        add_table(
            ["Indicador", "Valor"],
            [
                ("Pseudo-F", permanova_result.get("pseudo_f", "")),
                ("p-value", permanova_result.get("p_value", "")),
                ("Permutaciones", permanova_result.get("permutations", "")),
            ]
        )

        if permanova_result.get("interpretacion"):
            add_interpretacion_paragraph(permanova_result.get("interpretacion"))

        if permanova_result.get("interpretation"):
            add_bullets(permanova_result.get("interpretation", []))

    if group_visuals:
        doc.add_heading("10.4 Visualización por grupos", level=2)

        if group_visuals.get("visual_interpretation"):
            doc.add_paragraph(
                f"Lectura visual automática: {group_visuals.get('visual_interpretation', '')}"
            )
            if group_visuals.get("separation_score") is not None:
                doc.add_paragraph(
                    f"Índice de separación: {group_visuals.get('separation_score')}"
                )

        for title, explanation, key in [
            ("10.4.1 PCA coloreado por grupo", "Proyección de observaciones en PC1 y PC2 coloreadas por grupo.", "group_pca"),
            ("10.4.2 Centroides por grupo", "Ubicación promedio de cada grupo en el espacio reducido.", "group_centroids"),
            ("10.4.3 Distancia entre grupos", "Matriz visual de distancias entre centroides de grupos.", "group_distances"),
        ]:
            rel = group_visuals.get(key)
            if rel:
                filename = os.path.basename(rel)
                add_image_section(title, explanation, filename)

    # =========================
    # 11. Modelos predictivos
    # =========================
    if regression_result or rf_result or model_comparison or logistic_result or rf_classification_result or logistic_conf_exists or logistic_roc_exists:
        doc.add_heading("11. Modelos predictivos", level=1)

    if regression_result:
        doc.add_heading("11.1 Regresión lineal múltiple", level=2)
        doc.add_paragraph(
            f"Variable objetivo: {regression_result.get('target_col', '')} | "
            f"Predictores: {regression_result.get('n_predictors', '')} | "
            f"Filas usadas: {regression_result.get('n_rows', '')}"
        )

        add_table(
            ["Indicador", "Valor"],
            [
                ("Intercepto", regression_result.get("intercept", "")),
                ("R²", regression_result.get("r2", "")),
                ("R² ajustado", regression_result.get("adj_r2", "")),
                ("MAE", regression_result.get("mae", "")),
                ("RMSE", regression_result.get("rmse", "")),
            ]
        )

        if regression_result.get("interpretacion"):
            add_interpretacion_paragraph(regression_result.get("interpretacion"))

        if regression_result.get("interpretation"):
            add_bullets(regression_result.get("interpretation", []))

        coeff_table = regression_result.get("coeff_table", []) or []
        if coeff_table:
            rows = []
            for row in coeff_table:
                rows.append([
                    row.get("variable", ""),
                    row.get("coefficient", ""),
                    row.get("std_error", ""),
                    row.get("t_value", ""),
                    row.get("p_value", ""),
                    row.get("beta_std", ""),
                    row.get("abs_beta_std", ""),
                ])
            add_table(
                ["Variable", "Coeficiente", "Error estándar", "t", "p-value", "Beta est.", "|Beta est.|"],
                rows
            )

    if rf_result:
        doc.add_heading("11.2 Random Forest Regressor", level=2)
        doc.add_paragraph(
            f"Variable objetivo: {rf_result.get('target_col', '')} | "
            f"Predictores: {rf_result.get('n_predictors', '')} | "
            f"Filas usadas: {rf_result.get('n_rows', '')}"
        )

        add_table(
            ["Indicador", "Valor"],
            [
                ("R²", rf_result.get("r2", "")),
                ("MAE", rf_result.get("mae", "")),
                ("RMSE", rf_result.get("rmse", "")),
            ]
        )

        if rf_result.get("interpretacion"):
            add_interpretacion_paragraph(rf_result.get("interpretacion"))

        if rf_result.get("interpretation"):
            add_bullets(rf_result.get("interpretation", []))

        importance_table = rf_result.get("feature_importance_table", []) or []
        if importance_table:
            rows = []
            for row in importance_table:
                rows.append([
                    row.get("variable", ""),
                    row.get("importance", ""),
                ])
            add_table(["Variable", "Importancia"], rows)

    if model_comparison and model_comparison.get("available"):
        doc.add_heading("11.3 Comparación entre modelos de regresión", level=2)
        doc.add_paragraph(
            f"Variable objetivo comparada: {model_comparison.get('target_col', '')}"
        )

        linear = model_comparison.get("linear") or {}
        rf = model_comparison.get("random_forest") or {}

        def model_interp(model_name, r2):
            try:
                r2v = float(r2)
            except Exception:
                return ""

            if r2v >= 0.75:
                base = "Ajuste muy alto"
            elif r2v >= 0.50:
                base = "Ajuste alto"
            elif r2v >= 0.30:
                base = "Ajuste moderado"
            else:
                base = "Ajuste limitado"

            if model_name == "Random Forest":
                return f"{base}; buen candidato si hay relaciones no lineales."
            return f"{base}; útil como modelo base interpretable."

        add_table(
            ["Modelo", "R²", "MAE", "RMSE", "Interpretación"],
            [
                [
                    "Regresión lineal",
                    linear.get("r2", "—"),
                    linear.get("mae", "—"),
                    linear.get("rmse", "—"),
                    model_interp("Regresión lineal", linear.get("r2", None)),
                ],
                [
                    "Random Forest",
                    rf.get("r2", "—"),
                    rf.get("mae", "—"),
                    rf.get("rmse", "—"),
                    model_interp("Random Forest", rf.get("r2", None)),
                ],
            ]
        )

        winner = model_comparison.get("winner")
        winner_text = {
            "random_forest": "Random Forest",
            "linear_regression": "Regresión lineal",
            "tie": "Rendimiento similar",
        }.get(winner, "No determinado")

        doc.add_paragraph(f"Modelo recomendado por SADI: {winner_text}")

        if model_comparison.get("summary"):
            add_bullets(model_comparison.get("summary", []))

    if logistic_result or logistic_conf_exists or logistic_roc_exists:
        doc.add_heading("11.4 Logistic Regression", level=2)

        if logistic_result:
            doc.add_paragraph(
                f"Variable objetivo binaria: {logistic_result.get('target_col', '')}"
            )

            add_table(
                ["Indicador", "Valor"],
                [
                    ("Accuracy", logistic_result.get("accuracy", "")),
                    ("Precision", logistic_result.get("precision", "")),
                    ("Recall", logistic_result.get("recall", "")),
                    ("F1", logistic_result.get("f1", "")),
                    ("AUC", logistic_result.get("roc_auc", "")),
                ]
            )

            if logistic_result.get("interpretacion"):
                add_interpretacion_paragraph(logistic_result.get("interpretacion"))

            if logistic_result.get("interpretation"):
                add_bullets(logistic_result.get("interpretation", []))

            coef_table = logistic_result.get("coef_table", []) or []
            if coef_table:
                rows = []
                for row in coef_table:
                    rows.append([
                        row.get("variable", ""),
                        row.get("coef", ""),
                    ])
                add_table(["Variable", "Coeficiente"], rows)
        else:
            doc.add_paragraph(
                "Los gráficos de Logistic Regression fueron detectados en el sistema, aunque la reconstrucción tabular completa del modelo no estuvo disponible en esta exportación."
            )

        add_image_section(
            "11.4.1 Matriz de confusión",
            "La matriz de confusión resume los aciertos y errores de clasificación del modelo.",
            f"ds{dataset_id}_logistic_confusion.png",
        )

        add_image_section(
            "11.4.2 Curva ROC",
            "La curva ROC muestra la capacidad del modelo para discriminar entre clases.",
            f"ds{dataset_id}_logistic_roc.png",
        )

    if rf_classification_result or rf_cls_conf_exists or rf_cls_roc_exists or rf_cls_imp_exists:
        doc.add_heading("11.5 Random Forest Classification", level=2)

        if rf_classification_result:
            doc.add_paragraph(
                f"Variable objetivo binaria: {rf_classification_result.get('target_col', '')}"
            )

            add_table(
                ["Indicador", "Valor"],
                [
                    ("Accuracy", rf_classification_result.get("accuracy", "")),
                    ("Precision", rf_classification_result.get("precision", "")),
                    ("Recall", rf_classification_result.get("recall", "")),
                    ("F1", rf_classification_result.get("f1", "")),
                    ("AUC", rf_classification_result.get("roc_auc", "")),
                ]
            )

            if rf_classification_result.get("interpretacion"):
                add_interpretacion_paragraph(rf_classification_result.get("interpretacion"))

            if rf_classification_result.get("interpretation"):
                add_bullets(rf_classification_result.get("interpretation", []))

            importance_table = rf_classification_result.get("feature_importance_table", []) or []
            if importance_table:
                rows = []
                for row in importance_table:
                    rows.append([
                        row.get("variable", ""),
                        row.get("importance", ""),
                    ])
                add_table(["Variable", "Importancia"], rows)
        else:
            doc.add_paragraph(
                "Se detectaron salidas gráficas asociadas a Random Forest de clasificación, aunque la reconstrucción tabular completa no estuvo disponible en esta exportación."
            )

        add_image_section(
            "11.5.1 Matriz de confusión (Random Forest clasificación)",
            "Resume los aciertos y errores del clasificador Random Forest.",
            f"ds{dataset_id}_rf_classification_confusion.png",
        )

        add_image_section(
            "11.5.2 Curva ROC (Random Forest clasificación)",
            "Evalúa la capacidad del modelo para discriminar entre clases.",
            f"ds{dataset_id}_rf_classification_roc.png",
        )

        add_image_section(
            "11.5.3 Importancia de variables (Random Forest clasificación)",
            "Muestra qué variables aportan más al modelo de clasificación.",
            f"ds{dataset_id}_rf_classification_importance.png",
        )
    if insights_ranking:
        doc.add_heading("Hallazgos principales priorizados por SADI", level=1)
        for item in insights_ranking:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(f"[{item.get('category', 'Insight')}] ")
            run.bold = True
            p.add_run(item.get("text", ""))
    # =========================
    # 12. Conclusión general
    # =========================
    
    if conclusion:
        doc.add_heading("Conclusión general", level=1)
        for bloque in str(conclusion).split("\n\n"):
            doc.add_paragraph(bloque)

    doc.save(output_path)


def generate_multivariate_article_docx(
    *,
    dataset_title: str,
    profile: dict,
    plots_dir: str,
    dataset_id: int,
    output_path: str,
    figure_catalog=None,
    pca_summary=None,
    pca_loadings=None,
    cluster_counts=None,
    cluster_best_k=None,
    efa_kmo=None,
    efa_bartlett_p=None,
    efa_n_factors=None,
    efa_loadings=None,
    regression_result=None,
    rf_result=None,
    logistic_result=None,
    rf_classification_result=None,
    model_comparison=None,
    anova_result=None,
    manova_result=None,
    permanova_result=None,
    group_visuals=None,
    interpretacion=None,
    corr_interpretacion=None,
    pca_interpretacion=None,
    cluster_interpretacion=None,
    abstract_text=None,
    keywords=None,
    conclusion=None,
    limitations=None,
    future_work=None,
    objective_text=None,
    introduction_text=None,
    methodology_text=None,
    insights_ranking=None,

    analysis_meta=None,
    insights=None,
    insights_text=None,
    quick_recommendations=None,
    suggested_plan=None,
    plot_summary=None,
    general_figure_catalog=None,
):
    import os
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # =========================
    # Configuración general
    # =========================
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    def safe_str(val, default=""):
        if val is None:
            return default
        return str(val)

    def fmt_num(val, digits=4, default="—"):
        try:
            if val is None or val == "":
                return default
            return f"{float(val):.{digits}f}"
        except Exception:
            return safe_str(val, default)

    def add_table(headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = safe_str(h)

        for row_data in rows:
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = safe_str(val)
        return table

    def add_bullets(items):
        if not items:
            return
        if isinstance(items, str):
            doc.add_paragraph(items, style="List Bullet")
            return
        for item in items:
            doc.add_paragraph(safe_str(item), style="List Bullet")

    fig_counter = {"n": 1}

    def add_fig(explanation, filename, short_title=None, width=5.8):
        path = os.path.join(plots_dir, filename)
        if not os.path.exists(path):
            return False

        n = fig_counter["n"]
        title = f"Figura {n}. {short_title}" if short_title else f"Figura {n}"

        p_title = doc.add_paragraph()
        p_title.paragraph_format.keep_with_next = True
        run = p_title.add_run(title)
        run.bold = True

        p_exp = doc.add_paragraph(explanation)
        p_exp.paragraph_format.keep_with_next = True

        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"

        cell = tbl.cell(0, 0)
        p_img = cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_img.add_run()
        run.add_picture(path, width=Inches(width))

        doc.add_paragraph("")
        fig_counter["n"] += 1
        return True

    # =========================
    # Variables derivadas
    # =========================
    n_rows = profile.get("n_rows", 0)
    n_cols = profile.get("n_cols", 0)
    n_num = profile.get("n_num", 0)
    n_cat = profile.get("n_cat", 0)
    missing_pct = profile.get("missing_pct", 0)
    profile_type = profile.get("profile_type", "")

    has_efa = bool(efa_loadings) or os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_efa_scree.png"))
    has_regression = bool(regression_result or rf_result or model_comparison)
    has_classification = bool(logistic_result or rf_classification_result) or \
        os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_logistic_confusion.png")) or \
        os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_logistic_roc.png")) or \
        os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_rf_classification_confusion.png")) or \
        os.path.exists(os.path.join(plots_dir, f"ds{dataset_id}_rf_classification_roc.png"))
    has_group_analysis = bool(anova_result or manova_result or permanova_result or group_visuals)

    methods = [
        "análisis de correlación",
        "análisis de componentes principales (PCA)",
        "clustering K-Means",
    ]
    if has_efa:
        methods.append("análisis factorial exploratorio (EFA)")
    if has_regression:
        methods.append("modelos predictivos de regresión")
    if has_classification:
        methods.append("modelos de clasificación")
    if has_group_analysis:
        methods.append("análisis inferencial por grupos")

    abstract_auto = (
        f"Este estudio presenta un análisis multivariado aplicado al dataset "
        f"'{dataset_title}', el cual contiene {n_rows} observaciones y {n_cols} variables. "
        f"Se identificaron {n_num} variables numéricas y {n_cat} variables categóricas, "
        f"con un porcentaje de valores faltantes de {missing_pct}%. "
        f"El perfil detectado por SADI fue '{profile_type}'. "
        f"Se aplicaron técnicas como {', '.join(methods)} para explorar la estructura de los datos, "
        f"detectar patrones latentes, evaluar agrupamientos naturales y analizar capacidad predictiva. "
        f"Los resultados permiten identificar relaciones entre variables, posibles dimensiones subyacentes, "
        f"diferencias entre grupos y desempeño relativo de modelos según el contexto del dataset."
    )

    conclusions_auto = (
        "El análisis multivariado realizado mediante SADI permitió obtener una visión amplia de la estructura del dataset. "
        "Las técnicas aplicadas resultaron útiles para identificar relaciones entre variables, reducir la dimensionalidad de los datos y detectar agrupamientos potenciales."
    )

    if has_regression:
        conclusions_auto += (
            " Asimismo, la evaluación de modelos predictivos de regresión permitió valorar el desempeño de enfoques lineales y no lineales."
        )
    if has_classification:
        conclusions_auto += (
            " En presencia de variables objetivo binarias, los modelos de clasificación ampliaron la capacidad interpretativa y predictiva del análisis."
        )
    if has_group_analysis:
        conclusions_auto += (
            " Los análisis inferenciales por grupos complementaron la exploración al mostrar diferencias relevantes entre categorías."
        )
    if has_efa:
        conclusions_auto += (
            " La incorporación del análisis factorial exploratorio reforzó la interpretación de posibles dimensiones latentes dentro del conjunto de datos."
        )

    conclusions_auto += (
        " En conjunto, estos resultados ofrecen una base sólida para estudios posteriores, validaciones estadísticas o construcción de modelos más avanzados."
    )

    # =========================
    # Título
    # =========================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Análisis multivariado del dataset '{dataset_title}' mediante SADI")
    r.bold = True
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento preliminar para revisión y edición académica")
    r.italic = True
    r.font.size = Pt(11)

    doc.add_paragraph("")

    # =========================
    # Resumen
    # =========================
    doc.add_heading("Resumen", level=1)
    doc.add_paragraph(abstract_text if abstract_text else abstract_auto)

    if keywords:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("Palabras clave: ")
        run.bold = True
        p.add_run(", ".join(keywords) if isinstance(keywords, list) else str(keywords))

    # =========================
    # 1. Introducción
    # =========================
    doc.add_heading("1. Introducción", level=1)

    if objective_text:
        p = doc.add_paragraph()
        r = p.add_run("Objetivo del estudio. ")
        r.bold = True
        p.add_run(objective_text)

    if introduction_text:
        doc.add_paragraph(introduction_text)
    else:
        doc.add_paragraph(
            "El presente estudio aborda el análisis multivariado del dataset seleccionado, "
            "con el propósito de identificar patrones, relaciones entre variables y estructuras internas "
            "que permitan comprender el fenómeno analizado."
        )

    # =========================
    # SADI INTELIGENTE (ARTÍCULO)
    # =========================
    if analysis_meta:
        doc.add_heading("1. Contexto analítico del dataset", level=1)

        tipo = analysis_meta.get("dataset_type", "")
        area = analysis_meta.get("research_area", "")
        suggested_area = analysis_meta.get("research_area_suggested")

        p = doc.add_paragraph()
        r = p.add_run("Tipo de dataset: ")
        r.bold = True
        p.add_run(str(tipo))

        p = doc.add_paragraph()
        r = p.add_run("Área de investigación: ")
        r.bold = True
        p.add_run(str(area))

        if suggested_area and suggested_area != area:
            p = doc.add_paragraph()
            r = p.add_run("Área sugerida por SADI: ")
            r.bold = True
            p.add_run(str(suggested_area))

    # =========================
    # INSIGHTS AUTOMÁTICOS
    # =========================
    if insights:
        doc.add_heading("2. Hallazgos automáticos de SADI", level=1)
        doc.add_paragraph(
            "A partir de la estructura, composición y características analíticas del dataset, "
            "SADI identificó los siguientes hallazgos preliminares:"
        )

        for item in insights:
            if isinstance(item, dict):
                txt = item.get("text", "")
                cat = item.get("category", "")
                score = item.get("score", "")

                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"[{cat}] ")
                run.bold = True
                p.add_run(f"{txt} (impacto: {score})")
            else:
                doc.add_paragraph(str(item), style="List Bullet")

    # =========================
    # INTERPRETACIÓN GLOBAL
    # =========================
    if insights_text:
        doc.add_heading("3. Interpretación global automática", level=1)
        doc.add_paragraph(insights_text)

    # =========================
    # RECOMENDACIONES
    # =========================
    if quick_recommendations:
        doc.add_heading("4. Recomendaciones metodológicas iniciales", level=1)
        doc.add_paragraph(
            "Con base en el perfil detectado, SADI propone las siguientes recomendaciones "
            "para orientar el desarrollo analítico y la interpretación de resultados:"
        )
        for item in quick_recommendations:
            doc.add_paragraph(str(item), style="List Bullet")

    # =========================
    # PLAN SUGERIDO
    # =========================
    if suggested_plan:
        doc.add_heading("5. Plan de análisis sugerido por SADI", level=1)

        recommended_analysis = suggested_plan.get("recommended_analysis", []) or []
        recommended_plots = suggested_plan.get("recommended_plots", []) or []
        narrative_focus = suggested_plan.get("narrative_focus")
        warnings = suggested_plan.get("warnings", []) or []

        if recommended_analysis:
            doc.add_paragraph(
                "Los análisis recomendados para este dataset son los siguientes:"
            )
            for item in recommended_analysis:
                doc.add_paragraph(str(item), style="List Bullet")

        if recommended_plots:
            doc.add_paragraph(
                "Las visualizaciones sugeridas para complementar la interpretación son:"
            )
            for item in recommended_plots:
                doc.add_paragraph(str(item), style="List Bullet")

        if narrative_focus:
            p = doc.add_paragraph()
            r = p.add_run("Enfoque interpretativo sugerido: ")
            r.bold = True
            p.add_run(str(narrative_focus))

        if warnings:
            doc.add_paragraph("Advertencias metodológicas:")
            for item in warnings:
                doc.add_paragraph(str(item), style="List Bullet")

    # =========================
    # RESUMEN DE GRÁFICOS
    # =========================
    if plot_summary:
        doc.add_heading("6. Resumen de visualizaciones generadas", level=1)
        rows = []
        for item in plot_summary:
            if isinstance(item, dict):
                rows.append([
                    item.get("tag", ""),
                    item.get("count", ""),
                ])

        if rows:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"

            hdr = table.rows[0].cells
            hdr[0].text = "Tipo de gráfico"
            hdr[1].text = "Cantidad"

            for row_data in rows:
                row = table.add_row().cells
                row[0].text = str(row_data[0])
                row[1].text = str(row_data[1])

    # =========================
    # FIGURAS GENERALES SADI
    # =========================
    if general_figure_catalog:
        doc.add_heading("7. Visualizaciones generales del dataset", level=1)

        for fig in general_figure_catalog:
            title = fig.get("title", "Figura")
            caption = fig.get("caption", "")
            filename = fig.get("basename", "")

            img_path = os.path.join(plots_dir, filename)
            if os.path.exists(img_path):
                doc.add_heading(title, level=2)
                if caption:
                    doc.add_paragraph(caption)

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.7))

                doc.add_paragraph("")

    # =========================
    # INSIGHTS PRIORIZADOS
    # =========================
    if insights_ranking:
        doc.add_heading("8. Hallazgos principales priorizados por SADI", level=1)

        for item in insights_ranking:
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(f"[{item.get('category', 'Insight')}] ")
            run.bold = True
            p.add_run(item.get("text", ""))

    # =========================
    # ORDEN DE PRIORIDAD ANALÍTICA
    # =========================
    priority = analysis_meta.get("priority_order", []) if analysis_meta else []
    if priority:
        doc.add_heading("9. Prioridad analítica sugerida", level=1)
        for i, item in enumerate(priority, 1):
            doc.add_paragraph(f"{i}. {item}")

    # =========================
    # 2. Metodología
    # =========================
    doc.add_heading("2. Metodología", level=1)

    methodology_body = methodology_text or (
        "Se aplicaron técnicas de análisis multivariado incluyendo correlación, PCA, clustering "
        "y modelos predictivos, con el fin de explorar la estructura del dataset y evaluar su comportamiento."
    )

    if has_efa:
        methodology_body += (
            " Cuando las condiciones del dataset lo permitieron, se incorporó un análisis factorial exploratorio "
            "(EFA) para detectar factores subyacentes."
        )
    if has_regression:
        methodology_body += (
            " Adicionalmente, se evaluaron modelos predictivos de regresión lineal y Random Forest para analizar "
            "desempeño comparativo."
        )
    if has_classification:
        methodology_body += (
            " En los casos con variable objetivo binaria, se aplicaron modelos de clasificación para estimar la "
            "capacidad discriminativa."
        )
    if has_group_analysis:
        methodology_body += (
            " Finalmente, cuando existieron variables categóricas de agrupación, se incorporaron contrastes "
            "como ANOVA, MANOVA y PERMANOVA, junto con visualizaciones por grupos."
        )

    doc.add_paragraph(methodology_body)

    # =========================
    # 3. Resultados
    # =========================
    doc.add_heading("3. Resultados", level=1)

    results = (
        "Los resultados del análisis multivariado muestran patrones relevantes en la estructura del dataset. "
        "El heatmap de correlación permitió identificar asociaciones entre variables, mientras que el análisis "
        "PCA evidenció componentes principales útiles para resumir la variabilidad de los datos. "
        "El análisis de clustering permitió observar posibles agrupamientos naturales entre las observaciones."
    )
    if has_efa and (efa_kmo is not None or efa_bartlett_p is not None):
        results += (
            f" Además, el análisis factorial exploratorio mostró un índice KMO de {fmt_num(efa_kmo)} "
            f"y una prueba de esfericidad de Bartlett con p = {fmt_num(efa_bartlett_p)}."
        )
        if efa_n_factors is not None:
            results += f" Se sugirió una estructura de {safe_str(efa_n_factors)} factores."
    if has_regression:
        results += (
            " Los modelos predictivos de regresión aportaron evidencia sobre la capacidad explicativa de las variables disponibles."
        )
    if has_classification:
        results += (
            " En el contexto de clasificación, los modelos permitieron evaluar la discriminación entre clases y el comportamiento de la predicción binaria."
        )
    if has_group_analysis:
        results += (
            " Los análisis por grupos mostraron diferencias entre categorías y reforzaron la interpretación multivariada del conjunto de datos."
        )
    doc.add_paragraph(results)

    # =========================
    # Figuras completas
    # =========================
    if figure_catalog:
        for fig in figure_catalog:
            add_fig(
                fig.get("caption", ""),
                fig.get("basename", os.path.basename(fig.get("filename", ""))),
                fig.get("title", "Figura"),
            )
    else:
        add_fig(
            "Esta figura muestra la intensidad de asociación entre las variables numéricas del dataset.",
            f"ds{dataset_id}_multivariate_corr.png",
            "Heatmap de correlación",
        )
        add_fig(
            "Esta figura resume la contribución de cada componente principal a la varianza explicada total.",
            f"ds{dataset_id}_pca_scree.png",
            "Scree plot del PCA",
        )

    # =========================
    # 4. Resultados específicos
    # =========================
    section_number = 4

    if efa_loadings:
        doc.add_heading(f"{section_number}. Cargas factoriales", level=1)
        intro_efa = f"KMO = {fmt_num(efa_kmo)} | Bartlett p = {fmt_num(efa_bartlett_p)}"
        if efa_n_factors is not None:
            intro_efa += f" | Factores sugeridos = {safe_str(efa_n_factors)}"
        doc.add_paragraph(intro_efa)
        doc.add_paragraph(
            "La siguiente tabla presenta las cargas factoriales estimadas para las variables del dataset."
        )

        header_keys = list(efa_loadings[0].keys())
        rows = []
        for item in efa_loadings:
            rows.append([item.get(k, "") for k in header_keys])
        add_table(header_keys, rows)
        doc.add_paragraph("")
        section_number += 1

    if model_comparison and model_comparison.get("available"):
        doc.add_heading(f"{section_number}. Comparación entre modelos predictivos", level=1)
        doc.add_paragraph(
            "La comparación entre modelos permite identificar el enfoque más conveniente según el desempeño observado sobre la variable objetivo."
        )

        linear = model_comparison.get("linear") or {}
        rf = model_comparison.get("random_forest") or {}

        add_table(
            ["Modelo", "R²", "MAE", "RMSE"],
            [
                ["Regresión lineal", linear.get("r2", "—"), linear.get("mae", "—"), linear.get("rmse", "—")],
                ["Random Forest", rf.get("r2", "—"), rf.get("mae", "—"), rf.get("rmse", "—")],
            ]
        )

        winner = model_comparison.get("winner")
        winner_text = {
            "random_forest": "Random Forest",
            "linear_regression": "Regresión lineal",
            "tie": "rendimiento similar",
        }.get(winner, "no determinado")

        doc.add_paragraph(f"De acuerdo con SADI, el modelo con mejor desempeño general fue: {winner_text}.")
        section_number += 1

    if anova_result or manova_result or permanova_result:
        doc.add_heading(f"{section_number}. Contrastes inferenciales por grupos", level=1)

        if anova_result:
            doc.add_paragraph(
                f"ANOVA: se analizó la variable dependiente '{anova_result.get('target_col', '')}' "
                f"según el factor de agrupación '{anova_result.get('group_col', '')}'."
            )

        if manova_result:
            doc.add_paragraph(
                f"MANOVA: se evaluaron diferencias multivariadas según el grupo "
                f"'{manova_result.get('group_col', '')}' en las variables dependientes consideradas."
            )

        if permanova_result:
            doc.add_paragraph(
                f"PERMANOVA: se examinó la separación entre grupos con un pseudo-F de "
                f"{safe_str(permanova_result.get('pseudo_f', '—'))} y p = {safe_str(permanova_result.get('p_value', '—'))}."
            )
        section_number += 1
    # =========================
    # Hallazgos principales SADI
    # =========================
    if insights_ranking is not None and len(insights_ranking) > 0:
        doc.add_heading(f"{section_number}. Hallazgos principales", level=1)

        for item in insights_ranking:
            p = doc.add_paragraph(style="List Number")

            r = p.add_run(f"{item.get('category', 'Insight')}: ")
            r.bold = True

            p.add_run(item.get("text", ""))

        section_number += 1
    
    # =========================
    # Discusión
    # =========================
    doc.add_heading(f"{section_number}. Discusión", level=1)

    discussion = (
        "Los resultados obtenidos sugieren la presencia de estructuras relevantes dentro del dataset analizado. "
        "La reducción de dimensionalidad mediante PCA facilita la interpretación de relaciones complejas entre variables, "
        "mientras que la segmentación obtenida mediante clustering puede ser útil para identificar grupos de observaciones con características similares."
    )

    if has_efa:
        discussion += (
            " El análisis factorial exploratorio aporta una visión complementaria al identificar factores subyacentes que pueden resumir la organización interna de las variables."
        )
    if has_regression:
        discussion += (
            " Los modelos predictivos muestran que el dataset contiene señal analítica suficiente para construir aproximaciones útiles, aunque el rendimiento depende de la estructura de la variable objetivo y del tipo de relación entre predictores."
        )
    if has_classification:
        discussion += (
            " En clasificación, las métricas y gráficos asociados sugieren que la discriminación entre clases puede analizarse de manera objetiva y comparativa."
        )
    if has_group_analysis:
        discussion += (
            " Los contrastes por grupos y sus visualizaciones aportan evidencia adicional para interpretar diferencias entre categorías y respaldar conclusiones sustantivas."
        )

    doc.add_paragraph(discussion)

    if corr_interpretacion or pca_interpretacion or cluster_interpretacion or interpretacion:
        doc.add_paragraph("")
        p = doc.add_paragraph()
        r = p.add_run("Interpretación integrada de resultados.")
        r.bold = True

        if corr_interpretacion:
            doc.add_paragraph(safe_str(corr_interpretacion), style="List Bullet")
        if pca_interpretacion:
            doc.add_paragraph(safe_str(pca_interpretacion), style="List Bullet")
        if cluster_interpretacion:
            doc.add_paragraph(safe_str(cluster_interpretacion), style="List Bullet")

        if isinstance(interpretacion, str):
            doc.add_paragraph(interpretacion, style="List Bullet")
        elif isinstance(interpretacion, (list, tuple)):
            add_bullets(interpretacion)

    section_number += 1

    # =========================
    # Conclusiones
    # =========================
    doc.add_heading(f"{section_number}. Conclusiones", level=1)

    if conclusion:
        for bloque in str(conclusion).split("\n\n"):
            doc.add_paragraph(bloque)
    else:
        doc.add_paragraph(conclusions_auto)

    section_number += 1

    # =========================
    # Limitaciones del estudio
    # =========================
    if limitations:
        doc.add_heading(f"{section_number}. Limitaciones del estudio", level=1)
        if isinstance(limitations, str):
            doc.add_paragraph(limitations)
        else:
            for item in limitations:
                doc.add_paragraph(str(item), style="List Bullet")
        section_number += 1

    # =========================
    # Líneas futuras de investigación
    # =========================
    if future_work:
        doc.add_heading(f"{section_number}. Líneas futuras de investigación", level=1)
        if isinstance(future_work, str):
            doc.add_paragraph(future_work)
        else:
            for item in future_work:
                doc.add_paragraph(str(item), style="List Bullet")
        section_number += 1

    # =========================
    # Referencias
    # =========================
    doc.add_heading(f"{section_number}. Referencias", level=1)
    doc.add_paragraph("Jolliffe, I. T. (2002). Principal Component Analysis. Springer.")
    doc.add_paragraph(
        "Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2014). Multivariate Data Analysis. Pearson."
    )

    if has_efa:
        doc.add_paragraph(
            "Fabrigar, L. R., & Wegener, D. T. (2011). Exploratory Factor Analysis. Oxford University Press."
        )

    if has_regression or has_classification:
        doc.add_paragraph(
            "James, G., Witten, D., Hastie, T., & Tibshirani, R. (2021). An Introduction to Statistical Learning. Springer."
        )

    doc.save(output_path)

def run_efa_analysis(*, df, dataset_id: int, plots_dir: str):
    import numpy as np
    import pandas as pd
    import os
    import matplotlib.pyplot as plt

    from factor_analyzer import FactorAnalyzer
    from factor_analyzer.factor_analyzer import (
        calculate_kmo,
        calculate_bartlett_sphericity,
    )
    from sklearn.preprocessing import StandardScaler

    # =========================
    # 1. Seleccionar variables numéricas y limpiar nulos
    # =========================
    work_df = df.select_dtypes(include=["number"]).copy().dropna()

    if work_df.shape[1] < 3:
        raise ValueError("Se necesitan al menos 3 variables numéricas para EFA.")

    # =========================
    # 2. Eliminar columnas constantes o casi constantes
    # =========================
    valid_cols = []
    for c in work_df.columns:
        s = work_df[c].dropna()
        if s.nunique() > 1:
            valid_cols.append(c)

    work_df = work_df[valid_cols].copy()

    if work_df.shape[1] < 3:
        raise ValueError("EFA requiere al menos 3 variables numéricas con variabilidad.")

    # =========================
    # 3. Eliminar variables casi duplicadas / extremadamente correlacionadas
    # =========================
    corr_abs = work_df.corr(numeric_only=True).abs()
    upper = corr_abs.where(np.triu(np.ones(corr_abs.shape), k=1).astype(bool))

    to_drop = [col for col in upper.columns if any(upper[col] > 0.98)]
    if to_drop:
        work_df = work_df.drop(columns=to_drop, errors="ignore")

    if work_df.shape[1] < 3:
        raise ValueError(
            "No se pudo ejecutar EFA porque, después de eliminar variables redundantes o muy similares, "
            "quedaron menos de 3 variables útiles. Esto suele ocurrir cuando el dataset tiene pocas "
            "variables independientes o demasiada colinealidad."
        )

    # =========================
    # 4. Escalado
    # =========================
    scaler = StandardScaler()
    X = scaler.fit_transform(work_df)

    # =========================
    # 5. KMO y Bartlett
    # =========================
    try:
        _kmo_all, kmo_model = calculate_kmo(X)
    except Exception as e:
        raise ValueError(f"No se pudo calcular KMO: {e}")

    try:
        bartlett_chi2, bartlett_p = calculate_bartlett_sphericity(X)
    except Exception as e:
        raise ValueError(f"No se pudo calcular la prueba de Bartlett: {e}")

    # =========================
    # 6. Eigenvalues iniciales
    # =========================
    try:
        fa0 = FactorAnalyzer(rotation=None)
        fa0.fit(X)
        ev, _v = fa0.get_eigenvalues()
    except Exception as e:
        msg = str(e)
        if "Singular matrix" in msg:
            raise ValueError(
                "La matriz de correlación es singular. Esto suele ocurrir cuando hay variables muy "
                "redundantes, duplicadas o con colinealidad extrema."
            )
        raise ValueError(f"No se pudo estimar la estructura factorial inicial: {e}")

    # =========================
    # 7. Scree plot
    # =========================
    os.makedirs(plots_dir, exist_ok=True)

    scree_file = f"ds{dataset_id}_efa_scree.png"
    scree_path = os.path.join(plots_dir, scree_file)

    plt.figure(figsize=(6, 4))
    plt.scatter(range(1, len(ev) + 1), ev)
    plt.plot(range(1, len(ev) + 1), ev)
    plt.title("EFA Scree Plot")
    plt.xlabel("Factor")
    plt.ylabel("Eigenvalue")
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(scree_path, dpi=180, bbox_inches="tight")
    plt.close()

    # =========================
    # 8. Número sugerido de factores
    # =========================
    n_factors = int(sum(ev > 1))
    if n_factors < 1:
        n_factors = 1

    n_factors = min(n_factors, work_df.shape[1] - 1) if work_df.shape[1] > 1 else 1
    if n_factors < 1:
        n_factors = 1

    # =========================
    # 9. Modelo final
    # =========================
    try:
        fa = FactorAnalyzer(n_factors=n_factors, rotation="varimax")
        fa.fit(X)
        loadings = fa.loadings_
    except Exception as e:
        msg = str(e)
        if "Singular matrix" in msg:
            raise ValueError(
                "No se pudo ajustar el modelo factorial final porque la matriz es singular. "
                "Conviene reducir variables redundantes o usar un dataset con más variabilidad."
            )
        raise ValueError(f"No se pudo ajustar el modelo factorial final: {e}")

    # =========================
    # 10. Tabla de loadings
    # =========================
    loadings_df = pd.DataFrame(
        loadings,
        columns=[f"Factor{i+1}" for i in range(n_factors)],
        index=work_df.columns,
    )

    loadings_table = (
        loadings_df.reset_index()
        .rename(columns={"index": "variable"})
        .round(3)
    )

    interpretation = []

    if kmo_model >= 0.80:
        interpretation.append("El KMO indica una adecuación muestral muy buena para análisis factorial.")
    elif kmo_model >= 0.70:
        interpretation.append("El KMO indica una adecuación muestral aceptable para análisis factorial.")
    elif kmo_model >= 0.60:
        interpretation.append("El KMO indica una adecuación muestral moderada.")
    else:
        interpretation.append("El KMO sugiere precaución: la adecuación muestral es limitada para EFA.")

    if bartlett_p < 0.05:
        interpretation.append("La prueba de Bartlett sugiere que sí existe estructura correlacional aprovechable.")
    else:
        interpretation.append("La prueba de Bartlett no muestra evidencia fuerte de estructura factorial.")

    interpretation.append(f"SADI sugiere {n_factors} factor(es) con criterio de eigenvalores mayores que 1.")

    return {
        "kmo": json_safe_number(kmo_model, 6),
        "bartlett_chi2": json_safe_number(bartlett_chi2, 6),
        "bartlett_p": json_safe_number(bartlett_p, 6),
        "n_factors": int(n_factors),
        "loadings_table": loadings_table.to_dict(orient="records"),
        "loadings": loadings_table.to_dict(orient="records"),
        "scree_image": scree_file,
        "interpretation": interpretation,
    }

def generate_multivariate_interpretation(
    pca_loadings=None,
    efa_loadings=None,
    cluster_counts=None,
):
    def infer_theme_from_variables(var_names):
        """
        Intenta inferir un tema semántico general a partir de nombres de variables.
        """
        joined = " ".join([str(v).lower() for v in var_names])

        themes = [
            (
                "condiciones climáticas",
                ["rain", "rainfall", "temp", "temperature", "climate", "humidity", "weather"]
            ),
            (
                "fertilidad y propiedades del suelo",
                ["soil", "ph", "nitrogen", "phosphorus", "potassium", "fertility", "moisture"]
            ),
            (
                "productividad agrícola y rendimiento",
                ["yield", "crop", "harvest", "production", "productivity"]
            ),
            (
                "marketing, inversión y conversión",
                ["ad", "spend", "sales", "revenue", "conversion", "order", "visits", "customer"]
            ),
            (
                "salud y biometría",
                ["bmi", "cholesterol", "blood", "pressure", "disease", "risk", "sleep", "stress", "age"]
            ),
            (
                "condiciones socioeconómicas",
                ["income", "education", "employment", "house", "city", "satisfaction"]
            ),
            (
                "manejo agronómico e insumos",
                ["fertilizer", "irrigation", "management", "input", "dose", "application"]
            ),
        ]

        best_theme = None
        best_score = 0

        for theme_name, keywords in themes:
            score = 0
            for kw in keywords:
                if kw in joined:
                    score += 1
            if score > best_score:
                best_score = score
                best_theme = theme_name

        if best_theme:
            return best_theme

        return "una dimensión latente del dataset"

    def extract_top_contributors(loadings_rows, prefix):
        comp_map = {}

        if not loadings_rows:
            return comp_map

        for row in loadings_rows:
            var = row.get("variable")

            for k, v in row.items():
                if not str(k).startswith(prefix):
                    continue
                if v is None or var is None:
                    continue

                try:
                    comp_map.setdefault(k, [])
                    comp_map[k].append((var, abs(float(v))))
                except Exception:
                    pass

        for comp in comp_map:
            comp_map[comp] = sorted(comp_map[comp], key=lambda x: x[1], reverse=True)

        return comp_map

    lines = []

    # PCA
    if pca_loadings:
        comp_map = extract_top_contributors(pca_loadings, "PC")

        for comp, vals in comp_map.items():
            top_vals = vals[:3]
            vars_top = [v[0] for v in top_vals]
            theme = infer_theme_from_variables(vars_top)

            txt = (
                f"El componente {comp} parece estar influenciado principalmente por las variables "
                f"{', '.join(vars_top)}. En conjunto, este patrón sugiere una dimensión asociada a {theme}."
            )
            lines.append(txt)

    # EFA
    if efa_loadings:
        factor_map = extract_top_contributors(efa_loadings, "Factor")

        for factor, vals in factor_map.items():
            top_vals = vals[:3]
            vars_top = [v[0] for v in top_vals]
            theme = infer_theme_from_variables(vars_top)

            txt = (
                f"El {factor} presenta mayores cargas en las variables {', '.join(vars_top)}, "
                f"por lo que podría interpretarse como un factor relacionado con {theme}."
            )
            lines.append(txt)

    # Clustering
    if cluster_counts:
        try:
            largest = max(cluster_counts, key=lambda x: x.get("n", 0))
            total = sum(int(x.get("n", 0) or 0) for x in cluster_counts)
            largest_n = int(largest.get("n", 0) or 0)

            pct = round((largest_n / total) * 100, 2) if total else 0

            txt = (
                f"El análisis de clustering muestra que el grupo más numeroso corresponde al "
                f"{largest.get('cluster')}, con aproximadamente {largest_n} observaciones "
                f"({pct}% del total). Esto sugiere la existencia de un perfil predominante en el dataset."
            )
            lines.append(txt)

            if len(cluster_counts) >= 2:
                txt2 = (
                    "La presencia de varios clusters respalda la hipótesis de que las observaciones "
                    "no son homogéneas y pueden organizarse en grupos con características diferenciadas."
                )
                lines.append(txt2)

        except Exception:
            pass

    return lines
def run_regression_analysis(*, df, dataset_id, plots_dir, target_col, research_type: str = "general"):
    import os
    import numpy as np
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    from sklearn.impute import SimpleImputer
    from sklearn.linear_model import LinearRegression
    from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
    from sklearn.model_selection import train_test_split
    from sklearn.preprocessing import StandardScaler

    num_df = df.select_dtypes(include=[np.number]).copy()

    if target_col not in num_df.columns:
        raise ValueError(f"La variable objetivo '{target_col}' no es numérica o no existe.")

    predictor_cols = [c for c in num_df.columns if c != target_col]

    clean_predictors = []
    for c in predictor_cols:
        s = num_df[c]
        if s.notna().sum() == 0:
            continue
        if s.dropna().nunique() <= 1:
            continue
        clean_predictors.append(c)

    if len(clean_predictors) < 2:
        raise ValueError("Se necesitan al menos 2 predictores numéricos válidos para regresión.")

    work_df = num_df[[target_col] + clean_predictors].copy().dropna(subset=[target_col])

    if len(work_df) < 20:
        raise ValueError("Se necesitan al menos 20 filas válidas para ejecutar una regresión razonable.")

    X = work_df[clean_predictors]
    y = work_df[target_col]

    imputer = SimpleImputer(strategy="median")
    X_imp = imputer.fit_transform(X)

    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X_imp)

    X_train, X_test, y_train, y_test = train_test_split(
        X_scaled, y, test_size=0.25, random_state=42
    )

    model = LinearRegression()
    model.fit(X_train, y_train)

    y_pred = model.predict(X_test)

    r2 = r2_score(y_test, y_pred)
    n = len(y_test)
    p = len(clean_predictors)
    adj_r2 = 1 - (1 - r2) * (n - 1) / max((n - p - 1), 1)

    mae = mean_absolute_error(y_test, y_pred)
    rmse = float(np.sqrt(mean_squared_error(y_test, y_pred)))

    coeffs_df = pd.DataFrame({
        "variable": clean_predictors,
        "coefficient": model.coef_,
        "abs_coefficient": np.abs(model.coef_),
    }).sort_values("abs_coefficient", ascending=False)

    top_predictors = coeffs_df.head(5)["variable"].tolist()

    os.makedirs(plots_dir, exist_ok=True)

    pred_name = f"ds{dataset_id}_regression_pred_vs_real.png"
    pred_path = os.path.join(plots_dir, pred_name)

    fig, ax = plt.subplots(figsize=(7.0, 5.2), facecolor="white")
    ax.scatter(y_test, y_pred, alpha=0.75)
    minv = min(float(np.min(y_test)), float(np.min(y_pred)))
    maxv = max(float(np.max(y_test)), float(np.max(y_pred)))
    ax.plot([minv, maxv], [minv, maxv], linestyle="--", linewidth=1.5)
    ax.set_xlim(minv, maxv)
    ax.set_ylim(minv, maxv)
    ax.set_xlabel("Valor real")
    ax.set_ylabel("Valor predicho")
    ax.set_title(f"Regresión: real vs predicho ({target_col})")
    ax.grid(alpha=0.25)
    fig.tight_layout()
    fig.savefig(pred_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    resid_name = f"ds{dataset_id}_regression_residuals.png"
    resid_path = os.path.join(plots_dir, resid_name)
    residuals = y_test - y_pred

    fig, ax = plt.subplots(figsize=(7.0, 5.2), facecolor="white")
    ax.scatter(y_pred, residuals, alpha=0.75)
    ax.axhline(0, linestyle="--", linewidth=1.5)
    ax.set_xlabel("Predicción")
    ax.set_ylabel("Residuo")
    ax.set_title("Regresión: residuos vs predicción")
    ax.grid(alpha=0.25)
    fig.tight_layout()
    fig.savefig(resid_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    coef_name = f"ds{dataset_id}_regression_coefficients.png"
    coef_path = os.path.join(plots_dir, coef_name)

    plot_df = coeffs_df.head(10).iloc[::-1]
    fig, ax = plt.subplots(figsize=(7.4, 5.4), facecolor="white")
    ax.barh(plot_df["variable"], plot_df["abs_coefficient"])
    ax.set_xlabel("Magnitud absoluta del coeficiente")
    ax.set_ylabel("Variable")
    ax.set_title("Regresión: importancia relativa de predictores")
    ax.grid(alpha=0.20, axis="x")
    fig.tight_layout()
    fig.savefig(coef_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    coeff_table = [
        {
            "variable": row["variable"],
            "coefficient": json_safe_number(row["coefficient"], 6),
            "abs_coefficient": json_safe_number(row["abs_coefficient"], 6),
        }
        for _, row in coeffs_df.iterrows()
    ]

    interpretation = []
    if top_predictors:
        interpretation.append(
            f"Las variables con mayor peso relativo en el modelo fueron {', '.join(top_predictors[:3])}."
        )
    interpretation.append(
        f"El modelo alcanzó un R² de {round(float(r2), 4)} y un R² ajustado de {round(float(adj_r2), 4)}."
    )
    if r2 >= 0.70:
        interpretation.append("El ajuste del modelo puede considerarse alto para un análisis exploratorio.")
    elif r2 >= 0.40:
        interpretation.append("El ajuste del modelo puede considerarse moderado para un análisis exploratorio.")
    else:
        interpretation.append("El ajuste del modelo es bajo y conviene revisar variables o probar otros enfoques.")

    interpretacion = interpret_regression_result(r2, research_type)

    return {
        "target_col": target_col,
        "predictor_cols": clean_predictors,
        "n_predictors": len(clean_predictors),
        "n_rows": int(len(work_df)),
        "intercept": json_safe_number(model.intercept_, 6),
        "r2": json_safe_number(r2, 6),
        "adj_r2": json_safe_number(adj_r2, 6),
        "mae": json_safe_number(mae, 6),
        "rmse": json_safe_number(rmse, 6),
        "coeff_table": coeff_table,
        "pred_vs_real_image": pred_name,
        "residuals_image": resid_name,
        "coefficients_image": coef_name,
        "interpretation": interpretation,
        "interpretacion": interpretacion,
    }

def run_rf_classification_analysis(
    *,
    df,
    dataset_id,
    plots_dir,
    target_col,
    research_type: str = "general",
):
    import os
    import numpy as np
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    from sklearn.ensemble import RandomForestClassifier
    from sklearn.metrics import (
        accuracy_score,
        precision_score,
        recall_score,
        f1_score,
        confusion_matrix,
        classification_report,
    )
    from sklearn.model_selection import train_test_split

    if target_col not in df.columns:
        raise ValueError(f"La variable objetivo '{target_col}' no existe.")

    num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    feature_cols = [c for c in num_cols if c != target_col]

    if len(feature_cols) < 1:
        raise ValueError("Se requiere al menos una variable numérica como predictor.")

    work_df = df[feature_cols + [target_col]].dropna().copy()

    if len(work_df) < 15:
        raise ValueError("No hay suficientes datos para clasificación.")

    X = work_df[feature_cols]
    y = work_df[target_col]

    if y.nunique() < 2:
        raise ValueError("La variable objetivo no tiene clases suficientes.")

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.25, random_state=42, stratify=y
    )

    model = RandomForestClassifier(
        n_estimators=300,
        random_state=42,
        n_jobs=-1,
    )
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)

    acc = float(accuracy_score(y_test, y_pred))
    precision = float(precision_score(y_test, y_pred, average="weighted", zero_division=0))
    recall = float(recall_score(y_test, y_pred, average="weighted", zero_division=0))
    f1 = float(f1_score(y_test, y_pred, average="weighted", zero_division=0))

    labels = sorted(pd.Series(y).dropna().astype(str).unique().tolist())
    cm = confusion_matrix(y_test, y_pred, labels=labels)

    os.makedirs(plots_dir, exist_ok=True)

    # =========================
    # MATRIZ DE CONFUSIÓN
    # =========================
    cm_plot_name = f"ds{dataset_id}_rf_confusion_matrix.png"
    cm_plot_path = os.path.join(plots_dir, cm_plot_name)

    fig, ax = plt.subplots(figsize=(6, 5), facecolor="white")
    im = ax.imshow(cm)
    ax.set_title("Matriz de confusión (RF)")
    ax.set_xlabel("Predicho")
    ax.set_ylabel("Real")
    ax.set_xticks(range(len(labels)))
    ax.set_yticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=25, ha="right")
    ax.set_yticklabels(labels)

    for i in range(cm.shape[0]):
        for j in range(cm.shape[1]):
            ax.text(j, i, str(cm[i, j]), ha="center", va="center")

    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(cm_plot_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # =========================
    # IMPORTANCIA DE VARIABLES
    # =========================
    importance = pd.DataFrame({
        "variable": feature_cols,
        "importance": model.feature_importances_
    }).sort_values("importance", ascending=False)

    imp_plot_name = f"ds{dataset_id}_rf_class_importance.png"
    imp_plot_path = os.path.join(plots_dir, imp_plot_name)

    fig, ax = plt.subplots(figsize=(8, max(4, 0.4 * len(importance))), facecolor="white")
    ax.barh(importance["variable"][::-1], importance["importance"][::-1])
    ax.set_title("Importancia de variables (Clasificación RF)")
    ax.set_xlabel("Importancia")
    ax.set_ylabel("Variable")
    ax.grid(alpha=0.20, axis="x")
    fig.tight_layout()
    fig.savefig(imp_plot_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    report = classification_report(
        y_test, y_pred, output_dict=True, zero_division=0
    )

    interpretation = []

    if acc >= 0.85:
        interpretation.append("El modelo presenta una precisión muy alta.")
    elif acc >= 0.70:
        interpretation.append("El modelo presenta una precisión buena.")
    elif acc >= 0.55:
        interpretation.append("El modelo presenta una precisión moderada.")
    else:
        interpretation.append("El modelo presenta una precisión limitada.")

    if f1 >= 0.80:
        interpretation.append("El equilibrio global entre precisión y recuperación es alto.")
    elif f1 >= 0.60:
        interpretation.append("El equilibrio global entre precisión y recuperación es aceptable.")
    else:
        interpretation.append("El equilibrio entre precisión y recuperación todavía es mejorable.")

    if len(importance) > 0:
        top_vars = importance.head(3)["variable"].tolist()
        interpretation.append(
            f"Las variables más influyentes son: {', '.join(top_vars)}."
        )

    return {
        "target_col": target_col,
        "accuracy": round(acc, 6),
        "precision": round(precision, 6),
        "recall": round(recall, 6),
        "f1_score": round(f1, 6),
        "n_rows": int(len(work_df)),
        "n_predictors": int(len(feature_cols)),
        "labels": labels,
        "confusion_matrix": cm.tolist(),
        "classification_report": report,
        "feature_importance_table": [
            {
                "variable": row["variable"],
                "importance": round(float(row["importance"]), 6),
            }
            for _, row in importance.iterrows()
        ],
        "interpretation": interpretation,
        "confusion_matrix_plot": f"plots/{cm_plot_name}",
        "importance_plot": f"plots/{imp_plot_name}",
    }

def run_outlier_analysis(*, df, dataset_id: int, plots_dir: str):
    import os
    import numpy as np
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    from scipy.stats import chi2
    from sklearn.decomposition import PCA
    from sklearn.preprocessing import StandardScaler

    num_df = df.select_dtypes(include=[np.number]).dropna().copy()

    if num_df.shape[1] < 2:
        raise ValueError("Se necesitan al menos 2 variables numéricas para detectar outliers.")

    if len(num_df) < 15:
        raise ValueError("Se requieren al menos 15 filas válidas para detección de outliers multivariados.")

    X = num_df.values
    mean_vec = np.mean(X, axis=0)
    cov = np.cov(X, rowvar=False)

    inv_cov = np.linalg.pinv(cov)
    distances = []
    for row in X:
        diff = row - mean_vec
        d = float(diff.T @ inv_cov @ diff)
        distances.append(d)

    distances = np.array(distances)
    threshold = chi2.ppf(0.975, df=X.shape[1])

    outlier_mask = distances > threshold
    outlier_count = int(np.sum(outlier_mask))

    result_df = num_df.copy()
    result_df["mahalanobis_d2"] = distances
    result_df["is_outlier"] = outlier_mask

    os.makedirs(plots_dir, exist_ok=True)

    pca = PCA(n_components=2)
    X_scaled = StandardScaler().fit_transform(num_df)
    coords = pca.fit_transform(X_scaled)

    scatter_name = f"ds{dataset_id}_outliers_pca.png"
    scatter_path = os.path.join(plots_dir, scatter_name)

    plt.figure(figsize=(6.8, 5.0))
    plt.scatter(coords[~outlier_mask, 0], coords[~outlier_mask, 1], alpha=0.7, label="Normal")
    if outlier_count > 0:
        plt.scatter(coords[outlier_mask, 0], coords[outlier_mask, 1], alpha=0.9, label="Outlier")
    plt.xlabel("PC1")
    plt.ylabel("PC2")
    plt.title("Outliers multivariados sobre proyección PCA")
    plt.legend()
    plt.grid(alpha=0.25)
    plt.tight_layout()
    plt.savefig(scatter_path, dpi=180, bbox_inches="tight")
    plt.close()

    dist_name = f"ds{dataset_id}_outliers_distance.png"
    dist_path = os.path.join(plots_dir, dist_name)

    sorted_d = np.sort(distances)
    plt.figure(figsize=(6.8, 4.4))
    plt.plot(range(1, len(sorted_d) + 1), sorted_d)
    plt.axhline(threshold, linestyle="--")
    plt.xlabel("Observaciones ordenadas")
    plt.ylabel("Distancia de Mahalanobis²")
    plt.title("Distancias multivariadas y umbral de outliers")
    plt.grid(alpha=0.25)
    plt.tight_layout()
    plt.savefig(dist_path, dpi=180, bbox_inches="tight")
    plt.close()

    top_outliers = (
        result_df[result_df["is_outlier"]]
        .sort_values("mahalanobis_d2", ascending=False)
        .head(15)
        .reset_index()
        .rename(columns={"index": "row_index"})
    )

    interpretation = []
    interpretation.append(
        f"SADI detectó {outlier_count} posibles outliers multivariados usando distancia de Mahalanobis."
    )
    if outlier_count == 0:
        interpretation.append("No se observaron observaciones extremadamente alejadas de la estructura general.")
    elif outlier_count <= 5:
        interpretation.append("La cantidad de observaciones atípicas es baja y puede revisarse caso por caso.")
    else:
        interpretation.append("La cantidad de observaciones atípicas es relevante y conviene evaluar su impacto en PCA, clustering o regresión.")

    return {
        "n_rows_used": int(len(num_df)),
        "n_variables": int(num_df.shape[1]),
        "threshold": json_safe_number(threshold, 6),
        "outlier_count": outlier_count,
        "outlier_pct": json_safe_number((outlier_count / len(num_df)) * 100, 6),
        "top_outliers": [
            {
                "row_index": int(row["row_index"]) if row["row_index"] is not None else None,
                "mahalanobis_d2": json_safe_number(row["mahalanobis_d2"], 6),
            }
            for _, row in top_outliers[["row_index", "mahalanobis_d2"]].iterrows()
        ],
        "scatter_image": scatter_name,
        "distance_image": dist_name,
        "interpretation": interpretation,
    }

def run_anova_analysis(*, df, dataset_id, plots_dir, target_col, group_col, research_type: str = "general"):
    import os
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import statsmodels.api as sm
    from statsmodels.formula.api import ols

    if target_col not in df.columns:
        raise ValueError(f"La variable objetivo '{target_col}' no existe.")
    if group_col not in df.columns:
        raise ValueError(f"La variable de grupo '{group_col}' no existe.")

    work_df = df[[target_col, group_col]].copy().dropna()
    work_df = work_df[work_df[group_col].astype(str).str.strip() != ""]

    if work_df.empty:
        raise ValueError("No hay datos válidos para ANOVA.")

    if not pd.api.types.is_numeric_dtype(work_df[target_col]):
        raise ValueError("La variable dependiente de ANOVA debe ser numérica.")

    n_groups = work_df[group_col].nunique()
    if n_groups < 2:
        raise ValueError("ANOVA requiere al menos 2 grupos.")
    if len(work_df) < 12:
        raise ValueError("Se requieren al menos 12 observaciones para ANOVA.")

    formula = f'Q("{target_col}") ~ C(Q("{group_col}"))'
    model = ols(formula, data=work_df).fit()
    anova_table = sm.stats.anova_lm(model, typ=2).reset_index()
    anova_table.columns = ["source", "sum_sq", "df", "f_value", "p_value"]

    anova_rows = []
    for _, row in anova_table.iterrows():
        anova_rows.append({
            "source": row.get("source"),
            "sum_sq": json_safe_number(row.get("sum_sq"), 6),
            "df": json_safe_number(row.get("df"), 6),
            "f_value": json_safe_number(row.get("f_value"), 6),
            "p_value": json_safe_number(row.get("p_value"), 6),
        })

    group_summary = (
        work_df.groupby(group_col)[target_col]
        .agg(["count", "mean", "std"])
        .reset_index()
        .rename(columns={"count": "n", "mean": "mean", "std": "std"})
    )

    group_summary_rows = []
    for _, row in group_summary.iterrows():
        group_summary_rows.append({
            group_col: row.get(group_col),
            "n": int(row.get("n")) if row.get("n") is not None else None,
            "mean": json_safe_number(row.get("mean"), 6),
            "std": json_safe_number(row.get("std"), 6),
        })

    os.makedirs(plots_dir, exist_ok=True)
    box_name = f"ds{dataset_id}_anova_boxplot.png"
    box_path = os.path.join(plots_dir, box_name)

    plt.figure(figsize=(7.0, 4.8))
    work_df.boxplot(column=target_col, by=group_col, grid=False)
    plt.title(f"ANOVA: {target_col} por {group_col}")
    plt.suptitle("")
    plt.xlabel(group_col)
    plt.ylabel(target_col)
    plt.tight_layout()
    plt.savefig(box_path, dpi=180, bbox_inches="tight")
    plt.close()

    p_value = None
    try:
        row_group = anova_table.iloc[0]
        p_value = float(row_group["p_value"])
    except Exception:
        p_value = None

    interpretation = []
    if p_value is not None:
        if p_value < 0.05:
            interpretation.append(
                f"Se observaron diferencias estadísticamente significativas entre grupos de '{group_col}' sobre '{target_col}' (p = {round(p_value, 6)})."
            )
        else:
            interpretation.append(
                f"No se observaron diferencias estadísticamente significativas entre grupos de '{group_col}' sobre '{target_col}' (p = {round(p_value, 6)})."
            )

    try:
        largest_group = group_summary.sort_values("mean", ascending=False).head(1)
        if not largest_group.empty:
            interpretation.append(
                f"El grupo con mayor media en '{target_col}' fue '{largest_group.iloc[0][group_col]}' con un promedio de {round(float(largest_group.iloc[0]['mean']), 4)}."
            )
    except Exception:
        pass

    interpretacion = interpret_anova_result(p_value, target_col, group_col, research_type)

    return {
        "target_col": target_col,
        "group_col": group_col,
        "anova_table": anova_rows,
        "group_summary": group_summary_rows,
        "p_value": json_safe_number(p_value, 6),
        "boxplot_image": box_name,
        "interpretation": interpretation,
        "interpretacion": interpretacion,
    }


def run_manova_analysis(*, df, dependent_cols: list[str], group_col: str, research_type: str = "general"):
    import pandas as pd
    from statsmodels.multivariate.manova import MANOVA

    if group_col not in df.columns:
        raise ValueError(f"La variable de grupo '{group_col}' no existe.")

    valid_deps = [c for c in dependent_cols if c in df.columns and pd.api.types.is_numeric_dtype(df[c])]
    if len(valid_deps) < 2:
        raise ValueError("MANOVA requiere al menos 2 variables dependientes numéricas.")

    work_df = df[valid_deps + [group_col]].copy().dropna()
    work_df = work_df[work_df[group_col].astype(str).str.strip() != ""]

    if work_df[group_col].nunique() < 2:
        raise ValueError("MANOVA requiere al menos 2 grupos.")
    if len(work_df) < 20:
        raise ValueError("Se requieren al menos 20 observaciones para MANOVA.")

    lhs = " + ".join([f'Q("{c}")' for c in valid_deps])
    formula = f"{lhs} ~ C(Q(\"{group_col}\"))"

    mv = MANOVA.from_formula(formula, data=work_df)
    res = mv.mv_test()

    txt = str(res)

    p_value = None
    try:
        stat_table = res.results[f'C(Q("{group_col}"))']["stat"]
        if "Pr > F" in stat_table.columns and "Wilks' lambda" in stat_table.index:
            raw_p = stat_table.loc["Wilks' lambda", "Pr > F"]
            if pd.notna(raw_p):
                p_value = float(raw_p)
    except Exception:
        p_value = None

    interpretation = [
        f"Se ejecutó MANOVA con {len(valid_deps)} variables dependientes y '{group_col}' como factor de agrupación.",
        "Revisa especialmente Wilks' Lambda, Pillai's Trace y sus p-values para decidir si existen diferencias globales entre grupos."
    ]

    interpretacion = interpret_manova_result(p_value, group_col, research_type)

    return {
        "group_col": group_col,
        "dependent_cols": valid_deps,
        "manova_text": txt,
        "p_value": json_safe_number(p_value, 6),
        "interpretation": interpretation,
        "interpretacion": interpretacion,
    }


def run_permanova_analysis(*, df, group_col, permutations=499, research_type: str = "general"):
    import numpy as np
    import pandas as pd
    from sklearn.preprocessing import StandardScaler
    from sklearn.metrics import pairwise_distances

    rng = np.random.default_rng(42)

    if group_col not in df.columns:
        raise ValueError(f"La variable de grupo '{group_col}' no existe.")

    num_df = df.select_dtypes(include=[np.number]).copy()
    if num_df.shape[1] < 2:
        raise ValueError("PERMANOVA requiere al menos 2 variables numéricas.")

    work_df = pd.concat([num_df, df[[group_col]]], axis=1).dropna()
    work_df = work_df[work_df[group_col].astype(str).str.strip() != ""]

    if work_df[group_col].nunique() < 2:
        raise ValueError("PERMANOVA requiere al menos 2 grupos.")
    if len(work_df) < 20:
        raise ValueError("Se requieren al menos 20 observaciones para PERMANOVA.")

    X = work_df[num_df.columns].values
    groups = work_df[group_col].astype(str).values

    Xs = StandardScaler().fit_transform(X)
    _dist_matrix = pairwise_distances(Xs, metric="euclidean")

    grand_centroid = Xs.mean(axis=0)
    unique_groups = np.unique(groups)
    k = len(unique_groups)
    n = len(groups)

    def pseudo_f(labels):
        ss_between = 0.0
        ss_within = 0.0

        for g in np.unique(labels):
            idx = np.where(labels == g)[0]
            group_data = Xs[idx]
            centroid = group_data.mean(axis=0)
            ss_between += len(idx) * np.sum((centroid - grand_centroid) ** 2)
            ss_within += np.sum((group_data - centroid) ** 2)

        df_between = max(k - 1, 1)
        df_within = max(n - k, 1)
        ms_between = ss_between / df_between
        ms_within = ss_within / df_within if ss_within != 0 else 1e-12
        return ms_between / ms_within

    observed_f = float(pseudo_f(groups))

    perm_f = []
    for _i in range(permutations):
        shuffled = rng.permutation(groups)
        perm_f.append(float(pseudo_f(shuffled)))

    perm_f = np.array(perm_f)
    p_value = float((np.sum(perm_f >= observed_f) + 1) / (len(perm_f) + 1))

    group_sizes = (
        pd.Series(groups)
        .value_counts()
        .sort_index()
        .rename_axis(group_col)
        .reset_index(name="n")
    )

    group_sizes_rows = []
    for _, row in group_sizes.iterrows():
        group_sizes_rows.append({
            group_col: row.get(group_col),
            "n": int(row.get("n")) if row.get("n") is not None else None,
        })

    interpretation = []
    if p_value < 0.05:
        interpretation.append(
            f"PERMANOVA sugiere diferencias multivariadas significativas entre grupos de '{group_col}' "
            f"(pseudo-F = {round(observed_f, 4)}, p = {round(p_value, 6)})."
        )
    else:
        interpretation.append(
            f"PERMANOVA no encontró evidencia suficiente de diferencias multivariadas entre grupos de '{group_col}' "
            f"(pseudo-F = {round(observed_f, 4)}, p = {round(p_value, 6)})."
        )

    interpretacion = interpret_permanova_result(p_value, group_col, research_type)

    return {
        "group_col": group_col,
        "n_variables": int(num_df.shape[1]),
        "n_rows": int(len(work_df)),
        "pseudo_f": json_safe_number(observed_f, 6),
        "p_value": json_safe_number(p_value, 6),
        "permutations": int(permutations),
        "group_sizes": group_sizes_rows,
        "interpretation": interpretation,
        "interpretacion": interpretacion,
    }

import numpy as np
import matplotlib.pyplot as plt
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler
from scipy.spatial.distance import cdist
import os


def generate_group_visualizations(df, group_col, dataset_id, plots_dir):
    import os
    import numpy as np
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib.patches import Ellipse
    from sklearn.decomposition import PCA
    from sklearn.preprocessing import StandardScaler
    from scipy.spatial.distance import cdist

    result = {}

    work_df = df.copy()
    work_df = work_df.dropna(subset=[group_col])

    num_cols = [c for c in work_df.columns if pd.api.types.is_numeric_dtype(work_df[c])]
    if len(num_cols) < 2:
        return result

    X_df = work_df[num_cols].dropna()
    if X_df.empty:
        return result

    work_df = work_df.loc[X_df.index].copy()
    groups = work_df[group_col].astype(str).values

    if len(np.unique(groups)) < 2:
        return result

    X = X_df.values
    X_scaled = StandardScaler().fit_transform(X)

    pca = PCA(n_components=2)
    X_pca = pca.fit_transform(X_scaled)

    unique_groups = np.unique(groups)

    def add_confidence_ellipse(ax, x, y, n_std=1.8, label=None):
        if len(x) < 3:
            return
        cov = np.cov(x, y)
        if np.any(~np.isfinite(cov)):
            return

        vals, vecs = np.linalg.eigh(cov)
        order = vals.argsort()[::-1]
        vals = vals[order]
        vecs = vecs[:, order]

        angle = np.degrees(np.arctan2(*vecs[:, 0][::-1]))
        width, height = 2 * n_std * np.sqrt(np.maximum(vals, 1e-12))

        mean_x = np.mean(x)
        mean_y = np.mean(y)

        ell = Ellipse(
            xy=(mean_x, mean_y),
            width=width,
            height=height,
            angle=angle,
            fill=False,
            linewidth=2,
            alpha=0.9,
        )
        ax.add_patch(ell)

    # =============================
    # 1) PCA coloreado por grupo + elipses
    # =============================
    fig, ax = plt.subplots(figsize=(7, 5))

    centroids = []
    labels = []
    within_spreads = []

    for g in unique_groups:
        idx = groups == g
        pts = X_pca[idx]

        ax.scatter(pts[:, 0], pts[:, 1], label=str(g), alpha=0.65)

        add_confidence_ellipse(ax, pts[:, 0], pts[:, 1])

        centroid = pts.mean(axis=0)
        centroids.append(centroid)
        labels.append(str(g))

        dists = np.sqrt(((pts - centroid) ** 2).sum(axis=1))
        within_spreads.append(float(np.mean(dists)) if len(dists) else 0.0)

    centroids = np.array(centroids)

    ax.set_xlabel("PC1")
    ax.set_ylabel("PC2")
    ax.set_title("PCA por grupos con elipses de confianza")
    ax.legend()

    pca_path = os.path.join(plots_dir, f"ds{dataset_id}_group_pca.png")
    fig.tight_layout()
    fig.savefig(pca_path, dpi=180, bbox_inches="tight")
    plt.close(fig)

    result["group_pca"] = f"plots/ds{dataset_id}_group_pca.png"

    # =============================
    # 2) Centroides por grupo
    # =============================
    fig, ax = plt.subplots(figsize=(7, 5))

    for g in unique_groups:
        idx = groups == g
        pts = X_pca[idx]
        ax.scatter(pts[:, 0], pts[:, 1], alpha=0.20)

    ax.scatter(centroids[:, 0], centroids[:, 1], marker="X", s=220)

    for i, label in enumerate(labels):
        ax.text(centroids[i, 0], centroids[i, 1], f" {label}", fontsize=10)

    ax.set_title("Centroides por grupo (espacio PCA)")
    ax.set_xlabel("PC1")
    ax.set_ylabel("PC2")

    cent_path = os.path.join(plots_dir, f"ds{dataset_id}_group_centroids.png")
    fig.tight_layout()
    fig.savefig(cent_path, dpi=180, bbox_inches="tight")
    plt.close(fig)

    result["group_centroids"] = f"plots/ds{dataset_id}_group_centroids.png"

    # =============================
    # 3) Distancias entre grupos
    # =============================
    dist_matrix = cdist(centroids, centroids)

    fig, ax = plt.subplots(figsize=(6, 5))
    im = ax.imshow(dist_matrix)
    fig.colorbar(im, ax=ax)

    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, rotation=45, ha="right")
    ax.set_yticks(range(len(labels)))
    ax.set_yticklabels(labels)
    ax.set_title("Distancia entre centroides de grupos")

    for i in range(len(labels)):
        for j in range(len(labels)):
            ax.text(j, i, f"{dist_matrix[i, j]:.2f}", ha="center", va="center", fontsize=8)

    dist_path = os.path.join(plots_dir, f"ds{dataset_id}_group_distances.png")
    fig.tight_layout()
    fig.savefig(dist_path, dpi=180, bbox_inches="tight")
    plt.close(fig)

    result["group_distances"] = f"plots/ds{dataset_id}_group_distances.png"

    # =============================
    # 4) Test visual automático
    # =============================
    separation_score = None
    visual_interpretation = "No se pudo determinar la separación visual."

    if len(centroids) >= 2 and len(within_spreads) >= 2:
        non_diag = dist_matrix[np.triu_indices_from(dist_matrix, k=1)]
        mean_between = float(np.mean(non_diag)) if len(non_diag) else 0.0
        mean_within = float(np.mean(within_spreads)) if len(within_spreads) else 0.0

        if mean_within > 0:
            separation_score = round(mean_between / mean_within, 4)

            if separation_score >= 3.0:
                visual_interpretation = "Los grupos aparecen claramente separados en el espacio PCA."
            elif separation_score >= 1.8:
                visual_interpretation = "Los grupos muestran una separación visual moderada."
            else:
                visual_interpretation = "Los grupos presentan solapamiento visual importante."

    result["group_labels"] = labels
    result["centroid_distances"] = dist_matrix.round(4).tolist()
    result["separation_score"] = separation_score
    result["visual_interpretation"] = visual_interpretation

    return result

def detect_prediction_context(df):
    import pandas as pd

    numeric_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    categorical_cols = [c for c in df.columns if not pd.api.types.is_numeric_dtype(df[c])]

    binary_target_cols = [
        c for c in df.columns
        if df[c].dropna().nunique() == 2
    ]

    numeric_binary_cols = [
        c for c in binary_target_cols
        if pd.api.types.is_numeric_dtype(df[c]) and set(df[c].dropna().unique()).issubset({0, 1})
    ]

    categorical_binary_cols = [
        c for c in binary_target_cols
        if not pd.api.types.is_numeric_dtype(df[c])
    ]

    target = None
    target_type = None

    if categorical_binary_cols:
        target = categorical_binary_cols[0]
        target_type = "classification"
    elif numeric_binary_cols:
        target = numeric_binary_cols[0]
        target_type = "classification"
    elif numeric_cols:
        target = numeric_cols[-1]
        target_type = "regression"

    return {
        "target": target,
        "target_type": target_type,
        "numeric_cols": numeric_cols,
        "categorical_cols": categorical_cols,
        "binary_target_cols": binary_target_cols,
        "categorical_binary_cols": categorical_binary_cols,
        "numeric_binary_cols": numeric_binary_cols,
    }

def run_rf_regression_analysis(*, df, dataset_id, plots_dir, target_col, research_type: str = "general"):
    import os
    import numpy as np
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    from sklearn.ensemble import RandomForestRegressor
    from sklearn.metrics import r2_score, mean_absolute_error, mean_squared_error
    from sklearn.model_selection import train_test_split

    if target_col not in df.columns:
        raise ValueError(f"La variable objetivo '{target_col}' no existe en el dataset.")

    num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    feature_cols = [c for c in num_cols if c != target_col]

    if len(feature_cols) < 1:
        raise ValueError("Random Forest requiere al menos una variable predictora numérica.")

    work_df = df[feature_cols + [target_col]].dropna().copy()
    if len(work_df) < 10:
        raise ValueError("No hay suficientes filas válidas para entrenar Random Forest.")

    X = work_df[feature_cols]
    y = work_df[target_col]

    if y.nunique() < 2:
        raise ValueError("La variable objetivo no tiene variabilidad suficiente.")

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.25, random_state=42
    )

    model = RandomForestRegressor(
        n_estimators=300,
        random_state=42,
        max_depth=None,
        min_samples_split=2,
        min_samples_leaf=1,
        n_jobs=-1,
    )
    model.fit(X_train, y_train)

    y_pred = model.predict(X_test)

    r2 = float(r2_score(y_test, y_pred))
    mae = float(mean_absolute_error(y_test, y_pred))
    rmse = float(np.sqrt(mean_squared_error(y_test, y_pred)))

    os.makedirs(plots_dir, exist_ok=True)

    # ===== gráfico real vs predicho =====
    pred_plot_name = f"ds{dataset_id}_rf_pred_vs_real.png"
    pred_plot_path = os.path.join(plots_dir, pred_plot_name)

    fig, ax = plt.subplots(figsize=(7.2, 5.4), facecolor="white")
    ax.scatter(y_test, y_pred, alpha=0.75)
    lims = [
        min(float(np.min(y_test)), float(np.min(y_pred))),
        max(float(np.max(y_test)), float(np.max(y_pred))),
    ]
    ax.plot(lims, lims, linestyle="--", linewidth=1.5)
    ax.set_xlim(lims)
    ax.set_ylim(lims)
    ax.set_xlabel("Valor real")
    ax.set_ylabel("Valor predicho")
    ax.set_title("Random Forest: real vs predicho")
    ax.grid(alpha=0.25)
    fig.tight_layout()
    fig.savefig(pred_plot_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    # ===== importancia de variables =====
    imp_plot_name = f"ds{dataset_id}_rf_feature_importance.png"
    imp_plot_path = os.path.join(plots_dir, imp_plot_name)

    importance = pd.DataFrame({
        "variable": feature_cols,
        "importance": model.feature_importances_
    }).sort_values("importance", ascending=False)

    fig, ax = plt.subplots(figsize=(8.2, max(4.5, 0.50 * len(importance))), facecolor="white")
    ax.barh(importance["variable"][::-1], importance["importance"][::-1])
    ax.set_xlabel("Importancia")
    ax.set_ylabel("Variable")
    ax.set_title("Random Forest: importancia de variables")
    ax.grid(alpha=0.20, axis="x")
    fig.tight_layout()
    fig.savefig(imp_plot_path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    interpretation = []

    if r2 >= 0.80:
        interpretation.append("El modelo Random Forest muestra un ajuste predictivo muy alto.")
    elif r2 >= 0.60:
        interpretation.append("El modelo Random Forest muestra un ajuste predictivo bueno.")
    elif r2 >= 0.40:
        interpretation.append("El modelo Random Forest muestra un ajuste predictivo moderado.")
    else:
        interpretation.append("El modelo Random Forest muestra un ajuste predictivo limitado.")

    if len(importance) > 0:
        top_vars = importance.head(3)["variable"].tolist()
        interpretation.append(
            f"Las variables más influyentes en la predicción son: {', '.join(top_vars)}."
        )

    interpretacion = interpret_rf_regression_result(r2, research_type)

    return {
        "target_col": target_col,
        "n_rows": int(len(work_df)),
        "n_predictors": int(len(feature_cols)),
        "r2": json_safe_number(r2, 6),
        "mae": json_safe_number(mae, 6),
        "rmse": json_safe_number(rmse, 6),
        "feature_importance_table": [
            {
                "variable": row["variable"],
                "importance": json_safe_number(row["importance"], 6),
            }
            for _, row in importance.iterrows()
        ],
        "interpretation": interpretation,
        "interpretacion": interpretacion,
        "pred_plot": f"plots/{pred_plot_name}",
        "importance_plot": f"plots/{imp_plot_name}",
    }

def compare_regression_models(linear_result=None, rf_result=None):
    comparison = {
        "available": False,
        "target_col": None,
        "linear": linear_result,
        "random_forest": rf_result,
        "winner": None,
        "summary": [],
    }

    if not linear_result and not rf_result:
        return comparison

    target_col = None
    if linear_result and linear_result.get("target_col"):
        target_col = linear_result.get("target_col")
    elif rf_result and rf_result.get("target_col"):
        target_col = rf_result.get("target_col")

    comparison["target_col"] = target_col

    if linear_result and rf_result:
        comparison["available"] = True

        lin_r2 = linear_result.get("r2")
        rf_r2 = rf_result.get("r2")

        lin_mae = linear_result.get("mae")
        rf_mae = rf_result.get("mae")

        lin_rmse = linear_result.get("rmse")
        rf_rmse = rf_result.get("rmse")

        score_linear = 0
        score_rf = 0

        if lin_r2 is not None and rf_r2 is not None:
            if rf_r2 > lin_r2:
                score_rf += 1
            elif lin_r2 > rf_r2:
                score_linear += 1

        if lin_mae is not None and rf_mae is not None:
            if rf_mae < lin_mae:
                score_rf += 1
            elif lin_mae < rf_mae:
                score_linear += 1

        if lin_rmse is not None and rf_rmse is not None:
            if rf_rmse < lin_rmse:
                score_rf += 1
            elif lin_rmse < rf_rmse:
                score_linear += 1

        if score_rf > score_linear:
            comparison["winner"] = "random_forest"
            comparison["summary"].append(
                "Random Forest presenta mejor desempeño global que la regresión lineal en las métricas comparadas."
            )
        elif score_linear > score_rf:
            comparison["winner"] = "linear_regression"
            comparison["summary"].append(
                "La regresión lineal presenta mejor desempeño global que Random Forest en las métricas comparadas."
            )
        else:
            comparison["winner"] = "tie"
            comparison["summary"].append(
                "Ambos modelos muestran un rendimiento similar en las métricas comparadas."
            )

        if lin_r2 is not None and rf_r2 is not None:
            comparison["summary"].append(
                f"R² lineal = {lin_r2} | R² Random Forest = {rf_r2}."
            )

        if lin_mae is not None and rf_mae is not None:
            comparison["summary"].append(
                f"MAE lineal = {lin_mae} | MAE Random Forest = {rf_mae}."
            )

        if lin_rmse is not None and rf_rmse is not None:
            comparison["summary"].append(
                f"RMSE lineal = {lin_rmse} | RMSE Random Forest = {rf_rmse}."
            )

    return comparison
def run_logistic_regression_analysis(*, df, dataset_id, plots_dir, target_col, research_type: str = "general"):
    import os
    import numpy as np
    import pandas as pd
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    from sklearn.linear_model import LogisticRegression
    from sklearn.model_selection import train_test_split
    from sklearn.metrics import (
        accuracy_score,
        precision_score,
        recall_score,
        f1_score,
        confusion_matrix,
        roc_curve,
        auc,
    )

    if target_col not in df.columns:
        raise ValueError(f"Target '{target_col}' no existe.")

    # validar target binario
    y_raw = df[target_col].astype("category")
    if y_raw.nunique() != 2:
        raise ValueError("Logistic Regression requiere variable binaria.")

    num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    feature_cols = [c for c in num_cols if c != target_col]

    if len(feature_cols) < 1:
        raise ValueError("Se requieren variables numéricas predictoras.")

    work_df = df[feature_cols + [target_col]].dropna().copy()
    if len(work_df) < 10:
        raise ValueError("No hay suficientes filas válidas para entrenar Logistic Regression.")

    X = work_df[feature_cols]
    y = work_df[target_col].astype("category").cat.codes

    if pd.Series(y).nunique() != 2:
        raise ValueError("La variable objetivo no tiene exactamente 2 clases utilizables.")

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.25, random_state=42
    )

    model = LogisticRegression(max_iter=1000)
    model.fit(X_train, y_train)

    y_pred = model.predict(X_test)
    y_prob = model.predict_proba(X_test)[:, 1]

    acc = accuracy_score(y_test, y_pred)
    prec = precision_score(y_test, y_pred, zero_division=0)
    rec = recall_score(y_test, y_pred, zero_division=0)
    f1 = f1_score(y_test, y_pred, zero_division=0)

    os.makedirs(plots_dir, exist_ok=True)

    # ===== matriz de confusión =====
    cm = confusion_matrix(y_test, y_pred)

    cm_plot_name = f"ds{dataset_id}_logistic_confusion.png"
    cm_plot_path = os.path.join(plots_dir, cm_plot_name)

    plt.figure()
    plt.imshow(cm)
    plt.title("Matriz de confusión")
    plt.xlabel("Predicho")
    plt.ylabel("Real")

    for i in range(len(cm)):
        for j in range(len(cm)):
            plt.text(j, i, cm[i, j], ha="center", va="center")

    plt.tight_layout()
    plt.savefig(cm_plot_path, dpi=180, bbox_inches="tight")
    plt.close()

    # ===== ROC =====
    fpr, tpr, _ = roc_curve(y_test, y_prob)
    roc_auc = auc(fpr, tpr)

    roc_plot_name = f"ds{dataset_id}_logistic_roc.png"
    roc_plot_path = os.path.join(plots_dir, roc_plot_name)

    plt.figure()
    plt.plot(fpr, tpr)
    plt.plot([0, 1], [0, 1])
    plt.xlabel("FPR")
    plt.ylabel("TPR")
    plt.title(f"ROC (AUC={roc_auc:.3f})")
    plt.tight_layout()
    plt.savefig(roc_plot_path, dpi=180, bbox_inches="tight")
    plt.close()

    # ===== coeficientes =====
    coef_table = []
    for i, var in enumerate(feature_cols):
        coef_table.append({
            "variable": var,
            "coef": json_safe_number(model.coef_[0][i], 6),
        })

    interpretation = []

    if acc > 0.80:
        interpretation.append("El modelo muestra una capacidad predictiva alta.")
    elif acc > 0.60:
        interpretation.append("El modelo muestra capacidad predictiva moderada.")
    else:
        interpretation.append("El modelo muestra baja capacidad predictiva.")

    if roc_auc >= 0.80:
        interpretation.append("La discriminación entre clases es alta según el AUC.")
    elif roc_auc >= 0.65:
        interpretation.append("La discriminación entre clases es moderada según el AUC.")
    else:
        interpretation.append("La discriminación entre clases es limitada según el AUC.")

    if len(coef_table) > 0:
        top_vars = sorted(
            coef_table,
            key=lambda x: abs(float(x["coef"])) if x["coef"] is not None else 0,
            reverse=True
        )[:3]
        interpretation.append(
            f"Las variables con mayor peso en la clasificación fueron: {', '.join([x['variable'] for x in top_vars])}."
        )

    # interpretación ejecutiva por research_type
    if research_type == "biomedical":
        interpretacion = (
            f"El modelo logístico alcanzó un desempeño de accuracy={acc:.2f} y AUC={roc_auc:.2f}, "
            "lo que sugiere capacidad para discriminar perfiles clínicos entre las dos clases."
        )
    elif research_type == "financial":
        interpretacion = (
            f"El modelo logístico alcanzó un desempeño de accuracy={acc:.2f} y AUC={roc_auc:.2f}, "
            "lo que sugiere capacidad para discriminar categorías financieras entre las dos clases."
        )
    elif research_type == "educational":
        interpretacion = (
            f"El modelo logístico alcanzó un desempeño de accuracy={acc:.2f} y AUC={roc_auc:.2f}, "
            "lo que sugiere capacidad para discriminar perfiles académicos entre las dos clases."
        )
    else:
        interpretacion = (
            f"El modelo logístico alcanzó un desempeño de accuracy={acc:.2f} y AUC={roc_auc:.2f}, "
            "lo que indica una capacidad determinada de clasificación binaria en el dataset."
        )

    return {
        "target_col": target_col,
        "accuracy": json_safe_number(acc, 6),
        "precision": json_safe_number(prec, 6),
        "recall": json_safe_number(rec, 6),
        "f1": json_safe_number(f1, 6),
        "roc_auc": json_safe_number(roc_auc, 6),
        "coef_table": coef_table,
        "interpretation": interpretation,
        "interpretacion": interpretacion,
        "confusion_plot": f"plots/{cm_plot_name}",
        "roc_plot": f"plots/{roc_plot_name}",
    }

def build_unified_multivariate_figure_catalog(
    *,
    dataset_id: int,
    plots_dir: str,
    group_visuals: dict | None = None,
) -> list[dict]:
    import os

    catalog = []

    def add_fig(filename: str, title: str, caption: str, section: str = ""):
        full_path = os.path.join(plots_dir, filename)
        if os.path.exists(full_path):
            catalog.append({
                "filename": f"plots/{filename}",   # para HTML con url_for('static', ...)
                "basename": filename,              # para DOCX/article con os.path.basename(...)
                "title": title,
                "caption": caption,
                "section": section,
            })

    # =========================
    # Exploratorio multivariado
    # =========================
    add_fig(
        f"ds{dataset_id}_multivariate_corr.png",
        "Heatmap de correlación",
        "Mapa de calor de correlaciones entre variables numéricas.",
        "exploratory",
    )
    add_fig(
        f"ds{dataset_id}_pca_scree.png",
        "Scree Plot del PCA",
        "Varianza explicada por los componentes principales.",
        "exploratory",
    )
    add_fig(
        f"ds{dataset_id}_kmeans_elbow.png",
        "Elbow Plot de K-Means",
        "Curva de inercia para sugerir el número de clusters.",
        "exploratory",
    )
    add_fig(
        f"ds{dataset_id}_kmeans_clusters.png",
        "Clusters K-Means",
        "Distribución de observaciones en el espacio reducido.",
        "exploratory",
    )
    add_fig(
        f"ds{dataset_id}_efa_scree.png",
        "Scree Plot del EFA",
        "Autovalores del análisis factorial exploratorio.",
        "exploratory",
    )
    add_fig(
        f"ds{dataset_id}_outliers_pca.png",
        "Outliers en espacio PCA",
        "Detección visual de observaciones atípicas sobre componentes principales.",
        "exploratory",
    )

    # =========================
    # Regresión
    # =========================
    add_fig(
        f"ds{dataset_id}_regression_pred_vs_real.png",
        "Regresión lineal: real vs predicho",
        "Comparación entre valores reales y estimados por el modelo lineal.",
        "regression",
    )
    add_fig(
        f"ds{dataset_id}_regression_residuals.png",
        "Regresión lineal: residuos",
        "Distribución de residuos del modelo lineal.",
        "regression",
    )
    add_fig(
        f"ds{dataset_id}_regression_coefficients.png",
        "Regresión lineal: coeficientes",
        "Magnitud relativa de los coeficientes del modelo lineal.",
        "regression",
    )

    # =========================
    # Random Forest regresión
    # =========================
    add_fig(
        f"ds{dataset_id}_rf_pred_vs_real.png",
        "Random Forest: real vs predicho",
        "Comparación entre valores reales y predichos por Random Forest.",
        "rf_regression",
    )
    add_fig(
        f"ds{dataset_id}_rf_feature_importance.png",
        "Random Forest: importancia de variables",
        "Importancia relativa de variables en Random Forest de regresión.",
        "rf_regression",
    )

    # =========================
    # Logistic Regression
    # =========================
    add_fig(
        f"ds{dataset_id}_logistic_confusion.png",
        "Logistic Regression: matriz de confusión",
        "Resumen de aciertos y errores del clasificador logístico.",
        "classification",
    )
    add_fig(
        f"ds{dataset_id}_logistic_roc.png",
        "Logistic Regression: curva ROC",
        "Capacidad de discriminación del modelo logístico.",
        "classification",
    )

    # =========================
    # RF clasificación
    # =========================
    add_fig(
        f"ds{dataset_id}_rf_classification_confusion.png",
        "Random Forest clasificación: matriz de confusión",
        "Resumen de aciertos y errores del clasificador Random Forest.",
        "classification",
    )
    add_fig(
        f"ds{dataset_id}_rf_classification_roc.png",
        "Random Forest clasificación: curva ROC",
        "Capacidad de discriminación del clasificador Random Forest.",
        "classification",
    )
    add_fig(
        f"ds{dataset_id}_rf_classification_importance.png",
        "Random Forest clasificación: importancia de variables",
        "Variables más influyentes en el clasificador Random Forest.",
        "classification",
    )

    # =========================
    # ANOVA / grupos
    # =========================
    add_fig(
        f"ds{dataset_id}_anova_boxplot.png",
        "ANOVA: boxplot por grupo",
        "Comparación visual de la variable dependiente entre grupos.",
        "group_analysis",
    )

    # =========================
    # Visualizaciones por grupos
    # =========================
    if group_visuals:
        gp = group_visuals.get("group_pca")
        if gp:
            add_fig(
                os.path.basename(gp),
                "PCA coloreado por grupo",
                "Proyección de observaciones en PC1 y PC2 coloreadas por grupo.",
                "group_analysis",
            )

        gc = group_visuals.get("group_centroids")
        if gc:
            add_fig(
                os.path.basename(gc),
                "Centroides por grupo",
                "Ubicación promedio de cada grupo en el espacio reducido.",
                "group_analysis",
            )

        gd = group_visuals.get("group_distances")
        if gd:
            add_fig(
                os.path.basename(gd),
                "Distancia entre grupos",
                "Matriz visual de distancias entre centroides de grupos.",
                "group_analysis",
            )

    return catalog

def get_multivariate_figure_catalog(dataset_id: int):
    return [
        {
            "key": "corr",
            "filename": f"ds{dataset_id}_multivariate_corr.png",
            "title": "Heatmap de correlación",
            "caption": "Esta figura muestra la intensidad de asociación entre las variables numéricas del dataset.",
        },
        {
            "key": "pca_scree",
            "filename": f"ds{dataset_id}_pca_scree.png",
            "title": "Scree plot del PCA",
            "caption": "Esta figura resume la contribución de cada componente principal a la varianza explicada total.",
        },
        {
            "key": "pca_scatter",
            "filename": f"ds{dataset_id}_pca_scatter.png",
            "title": "Proyección PCA",
            "caption": "Esta figura permite observar la estructura general del dataset en una proyección bidimensional.",
        },
        {
            "key": "kmeans_elbow",
            "filename": f"ds{dataset_id}_kmeans_elbow.png",
            "title": "Elbow plot de K-Means",
            "caption": "Esta figura ayuda a sugerir un número razonable de clusters para la segmentación del dataset.",
        },
        {
            "key": "kmeans_clusters",
            "filename": f"ds{dataset_id}_kmeans_clusters.png",
            "title": "Clustering K-Means",
            "caption": "Esta figura representa la segmentación K-Means proyectada en los dos primeros componentes principales.",
        },
        {
            "key": "efa_scree",
            "filename": f"ds{dataset_id}_efa_scree.png",
            "title": "Scree plot del análisis factorial exploratorio",
            "caption": "Esta figura ayuda a visualizar el número de factores con autovalores relevantes en el análisis factorial exploratorio.",
        },
        {
            "key": "outliers_pca",
            "filename": f"ds{dataset_id}_outliers_pca.png",
            "title": "Outliers multivariados sobre PCA",
            "caption": "Esta figura muestra observaciones potencialmente atípicas proyectadas en el espacio PCA.",
        },
        {
            "key": "outliers_distance",
            "filename": f"ds{dataset_id}_outliers_distance.png",
            "title": "Distancias de Mahalanobis",
            "caption": "Esta figura resume las distancias multivariadas utilizadas para detectar posibles observaciones atípicas.",
        },
        {
            "key": "anova_boxplot",
            "filename": f"ds{dataset_id}_anova_boxplot.png",
            "title": "Boxplot del ANOVA",
            "caption": "Esta figura compara visualmente la distribución de la variable dependiente entre grupos.",
        },
        {
            "key": "group_pca",
            "filename": f"ds{dataset_id}_group_pca.png",
            "title": "Visualización PCA por grupos",
            "caption": "Esta figura muestra la proyección de observaciones coloreadas por grupo, útil para evaluar separación visual entre categorías.",
        },
        {
            "key": "group_centroids",
            "filename": f"ds{dataset_id}_group_centroids.png",
            "title": "Centroides por grupo",
            "caption": "Esta figura resume la posición promedio de cada grupo en el espacio multivariado reducido.",
        },
        {
            "key": "group_distances",
            "filename": f"ds{dataset_id}_group_distances.png",
            "title": "Distancia entre grupos",
            "caption": "Esta figura muestra las distancias entre centroides de grupos y apoya la interpretación de separación multivariada.",
        },
        {
            "key": "reg_pred",
            "filename": f"ds{dataset_id}_regression_pred_vs_real.png",
            "title": "Regresión lineal: real vs predicho",
            "caption": "Esta figura compara los valores observados con los estimados por el modelo lineal.",
        },
        {
            "key": "reg_resid",
            "filename": f"ds{dataset_id}_regression_residuals.png",
            "title": "Regresión lineal: residuos",
            "caption": "Esta figura permite examinar el comportamiento de los errores del modelo lineal.",
        },
        {
            "key": "reg_coef",
            "filename": f"ds{dataset_id}_regression_coefficients.png",
            "title": "Regresión lineal: importancia de predictores",
            "caption": "Esta figura resume la magnitud relativa de los coeficientes del modelo lineal.",
        },
        {
            "key": "rf_pred",
            "filename": f"ds{dataset_id}_rf_pred_vs_real.png",
            "title": "Random Forest: real vs predicho",
            "caption": "Esta figura compara los valores observados con los estimados por el modelo Random Forest.",
        },
        {
            "key": "rf_importance",
            "filename": f"ds{dataset_id}_rf_feature_importance.png",
            "title": "Random Forest: importancia de variables",
            "caption": "Esta figura muestra la importancia relativa de las variables dentro del modelo Random Forest.",
        },
        {
            "key": "logistic_confusion",
            "filename": f"ds{dataset_id}_logistic_confusion.png",
            "title": "Logistic Regression: matriz de confusión",
            "caption": "Esta figura resume aciertos y errores de clasificación para el modelo logístico.",
        },
        {
            "key": "logistic_roc",
            "filename": f"ds{dataset_id}_logistic_roc.png",
            "title": "Logistic Regression: curva ROC",
            "caption": "Esta figura muestra la capacidad de discriminación del modelo logístico.",
        },
        {
            "key": "rf_cls_confusion",
            "filename": f"ds{dataset_id}_rf_classification_confusion.png",
            "title": "Random Forest clasificación: matriz de confusión",
            "caption": "Esta figura resume aciertos y errores del clasificador Random Forest.",
        },
        {
            "key": "rf_cls_roc",
            "filename": f"ds{dataset_id}_rf_classification_roc.png",
            "title": "Random Forest clasificación: curva ROC",
            "caption": "Esta figura muestra la capacidad de discriminación del clasificador Random Forest.",
        },
        {
            "key": "rf_cls_importance",
            "filename": f"ds{dataset_id}_rf_classification_importance.png",
            "title": "Random Forest clasificación: importancia de variables",
            "caption": "Esta figura muestra qué variables aportan más al modelo de clasificación Random Forest.",
        },
    ]
def run_full_sadi_analysis(
    *,
    df,
    dataset_id: int,
    plots_dir: str,
    research_type: str = "general",
    regression_target_col: str | None = None,
    rf_target_col: str | None = None,
    logistic_target_col: str | None = None,
    rf_classification_target_col: str | None = None,
    anova_target_col: str | None = None,
    anova_group_col: str | None = None,
    manova_dependent_cols: list[str] | None = None,
    manova_group_col: str | None = None,
    permanova_group_col: str | None = None,
):
    """
    Ejecuta en cascada los análisis principales de SADI.
    Si un análisis falla, lo registra y continúa con los demás.
    """

    results = {
        "corr": None,
        "pca": None,
        "cluster": None,
        "efa": None,
        "regression": None,
        "rf_regression": None,
        "logistic": None,
        "rf_classification": None,
        "anova": None,
        "manova": None,
        "permanova": None,
        "errors": {},
    }

    # =========================
    # Correlación
    # =========================
    try:
        results["corr"] = generate_correlation_heatmap(
            df=df,
            dataset_id=dataset_id,
            plots_dir=plots_dir,
            research_type=research_type,
        )
    except Exception as e:
        results["errors"]["corr"] = str(e)

    # =========================
    # PCA
    # =========================
    try:
        results["pca"] = run_pca_analysis(
            df=df,
            dataset_id=dataset_id,
            plots_dir=plots_dir,
            research_type=research_type,
        )
    except Exception as e:
        results["errors"]["pca"] = str(e)

    # =========================
    # Clustering
    # =========================
    try:
        results["cluster"] = run_kmeans_analysis(
            df=df,
            dataset_id=dataset_id,
            plots_dir=plots_dir,
            research_type=research_type,
        )
    except Exception as e:
        results["errors"]["cluster"] = str(e)

    # =========================
    # EFA
    # =========================
    try:
        results["efa"] = run_efa_analysis(
            df=df,
            dataset_id=dataset_id,
            plots_dir=plots_dir,
        )
    except Exception as e:
        results["errors"]["efa"] = str(e)

    # =========================
    # Regresión lineal
    # =========================
    if regression_target_col:
        try:
            results["regression"] = run_regression_analysis(
                df=df,
                dataset_id=dataset_id,
                plots_dir=plots_dir,
                target_col=regression_target_col,
                research_type=research_type,
            )
        except Exception as e:
            results["errors"]["regression"] = str(e)

    # =========================
    # Random Forest regresión
    # =========================
    if rf_target_col:
        try:
            results["rf_regression"] = run_rf_regression_analysis(
                df=df,
                dataset_id=dataset_id,
                plots_dir=plots_dir,
                target_col=rf_target_col,
                research_type=research_type,
            )
        except Exception as e:
            results["errors"]["rf_regression"] = str(e)

    # =========================
    # Regresión logística
    # =========================
    if logistic_target_col:
        try:
            results["logistic"] = run_logistic_regression_analysis(
                df=df,
                dataset_id=dataset_id,
                plots_dir=plots_dir,
                target_col=logistic_target_col,
                research_type=research_type,
            )
        except Exception as e:
            results["errors"]["logistic"] = str(e)

    # =========================
    # RF clasificación
    # =========================
    if rf_classification_target_col:
        try:
            results["rf_classification"] = run_rf_classification_analysis(
                df=df,
                dataset_id=dataset_id,
                plots_dir=plots_dir,
                target_col=rf_classification_target_col,
                research_type=research_type,
            )
        except Exception as e:
            results["errors"]["rf_classification"] = str(e)

    # =========================
    # ANOVA
    # =========================
    if anova_target_col and anova_group_col:
        try:
            results["anova"] = run_anova_analysis(
                df=df,
                dataset_id=dataset_id,
                plots_dir=plots_dir,
                target_col=anova_target_col,
                group_col=anova_group_col,
                research_type=research_type,
            )
        except Exception as e:
            results["errors"]["anova"] = str(e)

    # =========================
    # MANOVA
    # =========================
    if manova_dependent_cols and manova_group_col:
        try:
            results["manova"] = run_manova_analysis(
                df=df,
                dependent_cols=manova_dependent_cols,
                group_col=manova_group_col,
                research_type=research_type,
            )
        except Exception as e:
            results["errors"]["manova"] = str(e)

    # =========================
    # PERMANOVA
    # =========================
    if permanova_group_col:
        try:
            results["permanova"] = run_permanova_analysis(
                df=df,
                group_col=permanova_group_col,
                permutations=499,
                research_type=research_type,
            )
        except Exception as e:
            results["errors"]["permanova"] = str(e)

    return results

def generate_sadi_conclusion(
    *,
    profile=None,
    corr_result=None,
    pca_result=None,
    cluster_result=None,
    efa_result=None,
    regression_result=None,
    rf_result=None,
    logistic_result=None,
    anova_result=None,
    manova_result=None,
    permanova_result=None,
    research_type="general",
):
    parts = []

    # =========================
    # 1. OBJETIVO
    # =========================
    parts.append(
        "El presente estudio tuvo como objetivo analizar la estructura y relaciones internas del dataset mediante técnicas estadísticas y multivariadas avanzadas."
    )

    # =========================
    # 2. HALLAZGOS CLAVE
    # =========================
    findings = []

    if corr_result:
        findings.append("Se identificaron patrones de correlación relevantes entre variables.")

    if pca_result:
        findings.append("El análisis de componentes principales evidenció reducción de dimensionalidad significativa.")

    if cluster_result:
        k = cluster_result.get("best_k")
        findings.append(f"Se detectaron {k} grupos diferenciados dentro de los datos.")

    if efa_result:
        findings.append("El análisis factorial reveló estructuras latentes subyacentes en las variables.")

    if regression_result:
        r2 = regression_result.get("r2")
        if r2 is not None:
            findings.append(f"La regresión mostró capacidad explicativa (R²={r2}).")

    if rf_result:
        r2 = rf_result.get("r2")
        if r2 is not None:
            findings.append(f"El modelo Random Forest mostró capacidad predictiva (R²={r2}).")

    if logistic_result:
        acc = logistic_result.get("accuracy")
        if acc is not None:
            findings.append(f"La clasificación logística alcanzó una precisión de {acc}.")

    if anova_result:
        findings.append("Se encontraron diferencias significativas entre grupos (ANOVA).")

    if manova_result:
        findings.append("Se evidenciaron diferencias multivariadas entre grupos (MANOVA).")

    if permanova_result:
        findings.append("Se detectaron diferencias estructurales entre grupos (PERMANOVA).")

    if findings:
        parts.append("Principales hallazgos: " + " ".join(findings))

    # =========================
    # 3. IMPLICACIONES
    # =========================
    if research_type == "biomedical":
        parts.append(
            "Estos resultados sugieren posibles patrones clínicos o biomédicos relevantes que podrían ser útiles en procesos diagnósticos o de investigación aplicada."
        )
    elif research_type == "financial":
        parts.append(
            "Los resultados evidencian patrones económicos y financieros que pueden ser utilizados para la toma de decisiones estratégicas y modelado predictivo."
        )
    elif research_type == "educational":
        parts.append(
            "Los hallazgos permiten comprender mejor los factores educativos involucrados, aportando evidencia para la mejora de procesos formativos."
        )
    else:
        parts.append(
            "Los resultados evidencian patrones estructurales relevantes que aportan comprensión sobre el comportamiento del dataset analizado."
        )

    # =========================
    # 4. RECOMENDACIONES
    # =========================
    recommendations = []

    if regression_result or rf_result:
        recommendations.append("Profundizar en modelos predictivos con mayor volumen de datos.")

    if efa_result:
        recommendations.append("Validar la estructura factorial mediante análisis confirmatorio.")

    if cluster_result:
        recommendations.append("Explorar perfiles específicos dentro de cada cluster identificado.")

    if recommendations:
        parts.append("Se recomienda: " + " ".join(recommendations))

    # =========================
    # FINAL
    # =========================
    return "\n\n".join(parts)

def generate_sadi_abstract_and_keywords(
    *,
    profile=None,
    corr_result=None,
    pca_result=None,
    cluster_result=None,
    efa_result=None,
    regression_result=None,
    rf_result=None,
    logistic_result=None,
    anova_result=None,
    manova_result=None,
    permanova_result=None,
    research_type="general",
):
    import random

    n_rows = profile.get("n_rows", 0) if profile else 0
    n_cols = profile.get("n_cols", 0) if profile else 0

    # =========================
    # ABSTRACT
    # =========================
    sentences = []

    sentences.append(
        f"Este estudio analiza un dataset compuesto por {n_rows} observaciones y {n_cols} variables, utilizando técnicas estadísticas y multivariadas."
    )

    if corr_result:
        sentences.append("Se evaluaron las relaciones entre variables mediante análisis de correlación.")

    if pca_result:
        sentences.append("Se aplicó análisis de componentes principales para reducir la dimensionalidad.")

    if cluster_result:
        sentences.append("Se emplearon técnicas de clustering para identificar patrones estructurales.")

    if efa_result:
        sentences.append("El análisis factorial permitió identificar estructuras latentes en los datos.")

    if regression_result or rf_result:
        sentences.append("Se desarrollaron modelos predictivos para explicar el comportamiento de la variable objetivo.")

    if logistic_result:
        sentences.append("Se aplicaron modelos de clasificación para discriminar entre categorías.")

    if anova_result or manova_result or permanova_result:
        sentences.append("Se realizaron pruebas de hipótesis para evaluar diferencias entre grupos.")

    # cierre
    if research_type == "biomedical":
        sentences.append("Los resultados aportan evidencia relevante para la comprensión de fenómenos biomédicos.")
    elif research_type == "financial":
        sentences.append("Los hallazgos permiten comprender patrones económicos y mejorar la toma de decisiones financieras.")
    elif research_type == "educational":
        sentences.append("Los resultados contribuyen al análisis de variables educativas y procesos de aprendizaje.")
    else:
        sentences.append("Los resultados evidencian patrones relevantes en el comportamiento del dataset.")

    abstract = " ".join(sentences)

    # =========================
    # KEYWORDS
    # =========================
    keywords = set()

    keywords.add("multivariate analysis")

    if corr_result:
        keywords.add("correlation")

    if pca_result:
        keywords.add("principal component analysis")

    if cluster_result:
        keywords.add("clustering")

    if efa_result:
        keywords.add("factor analysis")

    if regression_result:
        keywords.add("linear regression")

    if rf_result:
        keywords.add("random forest")

    if logistic_result:
        keywords.add("logistic regression")

    if anova_result:
        keywords.add("anova")

    if manova_result:
        keywords.add("manova")

    if permanova_result:
        keywords.add("permanova")

    # research type
    if research_type != "general":
        keywords.add(research_type)

    keywords = sorted(list(keywords))

    return {
        "abstract": abstract,
        "keywords": keywords,
    }
def generate_sadi_limitations_and_future_work(
    *,
    profile=None,
    corr_result=None,
    pca_result=None,
    cluster_result=None,
    efa_result=None,
    regression_result=None,
    rf_result=None,
    logistic_result=None,
    anova_result=None,
    manova_result=None,
    permanova_result=None,
    research_type="general",
):
    limitations = []
    future_work = []

    n_rows = profile.get("n_rows", 0) if profile else 0
    n_cols = profile.get("n_cols", 0) if profile else 0
    missing_pct = profile.get("missing_pct", 0) if profile else 0

    # Limitaciones generales
    if n_rows and n_rows < 30:
        limitations.append(
            "El tamaño muestral es relativamente reducido, lo que puede limitar la estabilidad de algunos resultados multivariados y predictivos."
        )

    if missing_pct and float(missing_pct) > 10:
        limitations.append(
            "La presencia de valores faltantes podría afectar parcialmente la robustez de ciertos análisis y la generalización de los hallazgos."
        )

    if not efa_result:
        limitations.append(
            "No fue posible consolidar evidencia factorial robusta, por lo que la identificación de estructuras latentes debe interpretarse con cautela."
        )

    if regression_result:
        r2 = regression_result.get("r2")
        try:
            if r2 is not None and float(r2) < 0.40:
                limitations.append(
                    "El poder explicativo del modelo de regresión fue limitado, lo que sugiere que pueden existir variables relevantes no incluidas en el dataset."
                )
        except Exception:
            pass

    if logistic_result:
        acc = logistic_result.get("accuracy")
        try:
            if acc is not None and float(acc) < 0.70:
                limitations.append(
                    "El desempeño del modelo de clasificación fue moderado o bajo, por lo que sus resultados deben considerarse exploratorios."
                )
        except Exception:
            pass

    if not (anova_result or manova_result or permanova_result):
        limitations.append(
            "La evidencia inferencial por grupos fue limitada o no aplicable, lo que restringe comparaciones categóricas más profundas."
        )

    if not limitations:
        limitations.append(
            "Aunque los análisis realizados ofrecen resultados consistentes, toda interpretación debe considerarse dentro del carácter exploratorio del estudio."
        )

    # Líneas futuras
    future_work.append(
        "Ampliar el tamaño muestral para aumentar la estabilidad de los resultados y la capacidad de generalización."
    )

    if corr_result or pca_result:
        future_work.append(
            "Profundizar en la validación de la estructura multivariada mediante análisis confirmatorios y comparación con nuevos datasets."
        )

    if cluster_result:
        future_work.append(
            "Explorar con mayor detalle los perfiles específicos de los grupos identificados y evaluar su utilidad práctica en contextos aplicados."
        )

    if efa_result:
        future_work.append(
            "Complementar los hallazgos factoriales con análisis factorial confirmatorio o modelos de ecuaciones estructurales."
        )

    if regression_result or rf_result or logistic_result:
        future_work.append(
            "Evaluar modelos predictivos adicionales y comparar su desempeño con técnicas más avanzadas de aprendizaje automático."
        )

    if research_type == "biomedical":
        future_work.append(
            "Incorporar nuevas variables clínicas o biomédicas para fortalecer la utilidad aplicada de los modelos analíticos."
        )
    elif research_type == "financial":
        future_work.append(
            "Integrar variables temporales y contextuales para mejorar la capacidad explicativa y predictiva en escenarios financieros."
        )
    elif research_type == "educational":
        future_work.append(
            "Incluir variables pedagógicas, institucionales y contextuales que permitan comprender con mayor profundidad los procesos educativos."
        )
    else:
        future_work.append(
            "Replicar el análisis en contextos adicionales para contrastar la estabilidad de los patrones observados."
        )

    return {
        "limitations": limitations,
        "future_work": future_work,
    }
def generate_sadi_paper_sections(
    *,
    dataset_title: str,
    profile: dict,
    research_type: str = "general",
    has_efa: bool = False,
    has_regression: bool = False,
    has_classification: bool = False,
    has_group_analysis: bool = False,
):
    n_rows = profile.get("n_rows", 0)
    n_cols = profile.get("n_cols", 0)
    n_num = profile.get("n_num", 0)
    n_cat = profile.get("n_cat", 0)
    missing_pct = profile.get("missing_pct", 0)

    # =========================
    # Introducción adaptada
    # =========================
    if research_type == "biomedical":
        intro = (
            "El análisis multivariado constituye una herramienta de gran relevancia en estudios biomédicos, "
            "debido a la complejidad inherente de los datos clínicos, fisiológicos y experimentales. "
            "En este tipo de contextos, múltiples variables pueden interactuar simultáneamente, por lo que resulta "
            "necesario aplicar técnicas capaces de identificar patrones, reducir dimensionalidad y apoyar la "
            "interpretación de fenómenos biológicos o clínicos. En este trabajo se empleó SADI para desarrollar "
            "un análisis automatizado del dataset, integrando componentes exploratorios, inferenciales y predictivos."
        )
    elif research_type == "financial":
        intro = (
            "En investigación financiera, el análisis multivariado resulta fundamental para comprender relaciones "
            "entre indicadores económicos, variables de riesgo y comportamientos de mercado. La naturaleza dinámica "
            "y multidimensional de estos datos exige herramientas capaces de identificar estructuras internas, "
            "segmentaciones y señales predictivas. En este trabajo se utilizó SADI para desarrollar un análisis "
            "automatizado del dataset, integrando métodos exploratorios, inferenciales y predictivos."
        )
    elif research_type == "educational":
        intro = (
            "En el ámbito educativo, los datos suelen reflejar múltiples dimensiones del aprendizaje, desempeño, "
            "contexto institucional y características del estudiantado. El análisis multivariado permite estudiar "
            "estas relaciones de forma integrada, generando evidencia útil para la comprensión de procesos formativos. "
            "En este trabajo se utilizó SADI para analizar el dataset mediante una estrategia automatizada que combina "
            "componentes exploratorios, inferenciales y predictivos."
        )
    else:
        intro = (
            "El análisis multivariado constituye una herramienta fundamental en la exploración de datasets complejos. "
            "En diversos contextos de investigación, los datos suelen incluir múltiples variables interrelacionadas, "
            "por lo que se requieren técnicas estadísticas y computacionales capaces de resumir estructura, detectar "
            "patrones y apoyar la toma de decisiones analíticas. En este trabajo se utilizó SADI para desarrollar "
            "un análisis automatizado del dataset seleccionado, integrando componentes exploratorios, inferenciales y predictivos."
        )

    # =========================
    # Metodología adaptada
    # =========================
    methodology = (
        f"El estudio se desarrolló sobre un dataset compuesto por {n_rows} observaciones y {n_cols} variables, "
        f"de las cuales {n_num} fueron numéricas y {n_cat} categóricas. El porcentaje de valores faltantes fue de "
        f"{missing_pct}%. Inicialmente se realizó una caracterización general del conjunto de datos y posteriormente "
        f"se calculó la matriz de correlación para examinar relaciones entre variables."
    )

    methodology += (
        " Luego se aplicó Análisis de Componentes Principales (PCA) con el propósito de reducir la dimensionalidad "
        "y resumir la variabilidad del dataset. De manera complementaria, se utilizó clustering K-Means para "
        "identificar posibles agrupamientos naturales entre las observaciones."
    )

    if has_efa:
        methodology += (
            " Adicionalmente, se implementó Análisis Factorial Exploratorio (EFA) para investigar la posible "
            "presencia de dimensiones latentes subyacentes."
        )

    if has_regression:
        methodology += (
            " También se evaluaron modelos predictivos de regresión lineal y Random Forest con el fin de analizar "
            "la capacidad explicativa y predictiva de las variables disponibles."
        )

    if has_classification:
        methodology += (
            " En presencia de variables objetivo binarias, se incorporaron modelos de clasificación para examinar "
            "la capacidad discriminativa del dataset."
        )

    if has_group_analysis:
        methodology += (
            " Finalmente, cuando existieron variables categóricas relevantes de agrupación, se aplicaron pruebas "
            "como ANOVA, MANOVA y PERMANOVA, junto con visualizaciones específicas por grupos."
        )

    # =========================
    # Objetivo del estudio
    # =========================
    if research_type == "biomedical":
        objective = (
            "El objetivo del estudio fue identificar patrones estructurales, relaciones internas y posibles señales "
            "predictivas en datos biomédicos, aportando evidencia útil para la comprensión del fenómeno analizado."
        )
    elif research_type == "financial":
        objective = (
            "El objetivo del estudio fue identificar patrones estructurales, relaciones internas y posibles señales "
            "predictivas en variables financieras, aportando evidencia útil para la interpretación y toma de decisiones."
        )
    elif research_type == "educational":
        objective = (
            "El objetivo del estudio fue identificar relaciones, agrupamientos y estructuras relevantes en variables "
            "educativas, contribuyendo a una mejor comprensión del fenómeno formativo analizado."
        )
    else:
        objective = (
            "El objetivo del estudio fue identificar relaciones, estructuras latentes, agrupamientos y capacidad "
            "predictiva en el dataset analizado, a fin de obtener una visión integral del fenómeno bajo estudio."
        )

    return {
        "objective": objective,
        "introduction": intro,
        "methodology": methodology,
    }
def generate_insights_ranking(
    *,
    corr_result=None,
    pca_result=None,
    cluster_result=None,
    efa_result=None,
    regression_result=None,
    rf_result=None,
    logistic_result=None,
    rf_classification_result=None,
    anova_result=None,
    manova_result=None,
    permanova_result=None,
    research_type: str = "general",
    top_n: int = 8,
):
    insights = []

    def add_insight(score: float, category: str, text: str):
        insights.append({
            "score": round(float(score), 3),
            "category": category,
            "text": text,
        })

    # =========================
    # Correlación
    # =========================
    if corr_result:
        corr_value = corr_result.get("corr_value")
        try:
            corr_value = float(corr_value) if corr_value is not None else None
        except Exception:
            corr_value = None

        if corr_value is not None:
            if corr_value >= 0.80:
                add_insight(
                    9.5,
                    "Correlación",
                    f"Se observaron relaciones muy fuertes entre variables (correlación promedio ≈ {corr_value:.2f}), lo que sugiere alta dependencia estructural."
                )
            elif corr_value >= 0.60:
                add_insight(
                    8.0,
                    "Correlación",
                    f"Se identificaron relaciones fuertes o moderadas entre variables (correlación promedio ≈ {corr_value:.2f})."
                )
            elif corr_value >= 0.40:
                add_insight(
                    6.5,
                    "Correlación",
                    f"Las variables presentan asociaciones moderadas (correlación promedio ≈ {corr_value:.2f})."
                )

    # =========================
    # PCA
    # =========================
    if pca_result:
        total_var = pca_result.get("total_variance_explained")
        try:
            total_var = float(total_var) if total_var is not None else None
        except Exception:
            total_var = None

        if total_var is not None:
            if total_var >= 0.80:
                add_insight(
                    9.0,
                    "PCA",
                    f"El PCA explicó una proporción muy alta de la varianza total ({total_var*100:.1f}%), indicando una reducción dimensional altamente eficiente."
                )
            elif total_var >= 0.60:
                add_insight(
                    7.5,
                    "PCA",
                    f"El PCA explicó una parte importante de la varianza total ({total_var*100:.1f}%), lo que sugiere buena compresión estructural."
                )
            elif total_var >= 0.40:
                add_insight(
                    6.0,
                    "PCA",
                    f"El PCA explicó una proporción moderada de la varianza total ({total_var*100:.1f}%)."
                )

    # =========================
    # Clustering
    # =========================
    if cluster_result:
        best_k = cluster_result.get("best_k")
        cluster_counts = cluster_result.get("cluster_counts") or []

        if best_k:
            score = 7.0
            if best_k in [2, 3, 4]:
                score = 8.0

            add_insight(
                score,
                "Clustering",
                f"El análisis de clustering identificó {best_k} grupos diferenciados, sugiriendo segmentación natural dentro del dataset."
            )

            try:
                sizes = []
                for row in cluster_counts:
                    n = row.get("n")
                    if n is not None:
                        sizes.append(int(n))
                if sizes and len(sizes) >= 2:
                    ratio = max(sizes) / max(min(sizes), 1)
                    if ratio <= 2.5:
                        add_insight(
                            6.8,
                            "Clustering",
                            "Los clusters presentan tamaños relativamente equilibrados, lo que refuerza la consistencia de la segmentación."
                        )
            except Exception:
                pass

    # =========================
    # EFA
    # =========================
    if efa_result:
        kmo = efa_result.get("kmo")
        n_factors = efa_result.get("n_factors")

        try:
            kmo = float(kmo) if kmo is not None else None
        except Exception:
            kmo = None

        if kmo is not None:
            if kmo >= 0.80:
                add_insight(
                    8.8,
                    "EFA",
                    f"El índice KMO fue alto ({kmo:.2f}), lo que sugiere una base sólida para identificar factores latentes."
                )
            elif kmo >= 0.70:
                add_insight(
                    7.2,
                    "EFA",
                    f"El índice KMO fue aceptable ({kmo:.2f}), apoyando la viabilidad del análisis factorial."
                )
            elif kmo >= 0.60:
                add_insight(
                    6.0,
                    "EFA",
                    f"El índice KMO fue moderado ({kmo:.2f}), por lo que la estructura factorial debe interpretarse con cautela."
                )

        if n_factors:
            add_insight(
                6.5,
                "EFA",
                f"El análisis factorial sugirió una estructura de {n_factors} factores latentes."
            )

    # =========================
    # Regresión lineal
    # =========================
    if regression_result:
        r2 = regression_result.get("r2")
        try:
            r2 = float(r2) if r2 is not None else None
        except Exception:
            r2 = None

        if r2 is not None:
            if r2 >= 0.80:
                add_insight(
                    9.5,
                    "Regresión",
                    f"El modelo de regresión lineal mostró alta capacidad explicativa (R² = {r2:.3f})."
                )
            elif r2 >= 0.60:
                add_insight(
                    8.0,
                    "Regresión",
                    f"El modelo de regresión lineal mostró un buen ajuste (R² = {r2:.3f})."
                )
            elif r2 >= 0.40:
                add_insight(
                    6.5,
                    "Regresión",
                    f"El modelo de regresión lineal mostró capacidad explicativa moderada (R² = {r2:.3f})."
                )

    # =========================
    # RF regresión
    # =========================
    if rf_result:
        r2 = rf_result.get("r2")
        try:
            r2 = float(r2) if r2 is not None else None
        except Exception:
            r2 = None

        if r2 is not None:
            if r2 >= 0.80:
                add_insight(
                    9.7,
                    "Random Forest",
                    f"El modelo Random Forest mostró una capacidad predictiva muy alta (R² = {r2:.3f})."
                )
            elif r2 >= 0.60:
                add_insight(
                    8.3,
                    "Random Forest",
                    f"El modelo Random Forest mostró un buen rendimiento predictivo (R² = {r2:.3f})."
                )
            elif r2 >= 0.40:
                add_insight(
                    6.8,
                    "Random Forest",
                    f"El modelo Random Forest mostró un rendimiento predictivo moderado (R² = {r2:.3f})."
                )

        try:
            fi = rf_result.get("feature_importance_table") or []
            top_vars = [x["variable"] for x in fi[:3] if x.get("variable")]
            if top_vars:
                add_insight(
                    6.2,
                    "Random Forest",
                    f"Las variables con mayor influencia predictiva fueron: {', '.join(top_vars)}."
                )
        except Exception:
            pass

    # =========================
    # Logistic
    # =========================
    if logistic_result:
        acc = logistic_result.get("accuracy")
        auc = logistic_result.get("roc_auc")

        try:
            acc = float(acc) if acc is not None else None
        except Exception:
            acc = None
        try:
            auc = float(auc) if auc is not None else None
        except Exception:
            auc = None

        if acc is not None:
            if acc >= 0.85:
                add_insight(
                    8.8,
                    "Clasificación",
                    f"El modelo logístico alcanzó una precisión alta (accuracy = {acc:.3f})."
                )
            elif acc >= 0.70:
                add_insight(
                    7.0,
                    "Clasificación",
                    f"El modelo logístico alcanzó una precisión aceptable (accuracy = {acc:.3f})."
                )

        if auc is not None and auc >= 0.80:
            add_insight(
                8.2,
                "Clasificación",
                f"La curva ROC mostró buena capacidad de discriminación (AUC = {auc:.3f})."
            )

    # =========================
    # RF clasificación
    # =========================
    if rf_classification_result:
        acc = rf_classification_result.get("accuracy")
        auc = rf_classification_result.get("roc_auc")

        try:
            acc = float(acc) if acc is not None else None
        except Exception:
            acc = None
        try:
            auc = float(auc) if auc is not None else None
        except Exception:
            auc = None

        if acc is not None and acc >= 0.75:
            add_insight(
                8.4,
                "RF Clasificación",
                f"Random Forest clasificación mostró buen desempeño (accuracy = {acc:.3f})."
            )

        if auc is not None and auc >= 0.80:
            add_insight(
                8.0,
                "RF Clasificación",
                f"Random Forest clasificación mostró adecuada discriminación entre clases (AUC = {auc:.3f})."
            )

    # =========================
    # ANOVA
    # =========================
    if anova_result:
        p = anova_result.get("p_value")
        target = anova_result.get("target_col")
        group = anova_result.get("group_col")

        try:
            p = float(p) if p is not None else None
        except Exception:
            p = None

        if p is not None:
            if p < 0.01:
                add_insight(
                    8.5,
                    "ANOVA",
                    f"Se detectaron diferencias muy significativas entre grupos de '{group}' sobre '{target}' (p = {p:.4f})."
                )
            elif p < 0.05:
                add_insight(
                    7.3,
                    "ANOVA",
                    f"Se detectaron diferencias significativas entre grupos de '{group}' sobre '{target}' (p = {p:.4f})."
                )

    # =========================
    # MANOVA
    # =========================
    if manova_result:
        p = manova_result.get("p_value")
        group = manova_result.get("group_col")

        try:
            p = float(p) if p is not None else None
        except Exception:
            p = None

        if p is not None and p < 0.05:
            add_insight(
                8.0,
                "MANOVA",
                f"Se observaron diferencias multivariadas significativas entre grupos definidos por '{group}' (p = {p:.4f})."
            )

    # =========================
    # PERMANOVA
    # =========================
    if permanova_result:
        p = permanova_result.get("p_value")
        pseudo_f = permanova_result.get("pseudo_f")
        group = permanova_result.get("group_col")

        try:
            p = float(p) if p is not None else None
        except Exception:
            p = None
        try:
            pseudo_f = float(pseudo_f) if pseudo_f is not None else None
        except Exception:
            pseudo_f = None

        if p is not None and p < 0.05:
            txt = f"PERMANOVA confirmó diferencias multivariadas entre grupos de '{group}'"
            if pseudo_f is not None:
                txt += f" (pseudo-F = {pseudo_f:.3f}, p = {p:.4f})."
            else:
                txt += f" (p = {p:.4f})."

            add_insight(8.1, "PERMANOVA", txt)

    # =========================
    # Ajuste contextual por research_type
    # =========================
    context_boost = {
        "biomedical": 0.2,
        "financial": 0.15,
        "educational": 0.1,
        "general": 0.0,
    }.get(research_type, 0.0)

    for item in insights:
        item["score"] = round(item["score"] + context_boost, 3)

    # Ordenar
    insights = sorted(insights, key=lambda x: x["score"], reverse=True)

    # Ranking explícito
    ranked = []
    for i, item in enumerate(insights[:top_n], start=1):
        ranked.append({
            "rank": i,
            "score": item["score"],
            "category": item["category"],
            "text": item["text"],
        })

    return ranked