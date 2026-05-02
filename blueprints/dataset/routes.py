
from __future__ import annotations

import csv
import glob
import json
import os
import re
import shutil
import time
import zipfile
from datetime import datetime

from flask import (
    Blueprint,
    flash,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    url_for,
    current_app,
)
from flask_login import current_user, login_required
from sqlalchemy import select
from werkzeug.utils import secure_filename

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, inch
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from .constans import DATASET_TYPE_CHOICES, RESEARCH_AREA_CHOICES, normalize_research_area

from blueprints.dataset.analysis import (    
    analyze_general_dataset,
    generate_general_dataset_plots,
    analyze_dataset_with_recommendations,
    suggest_research_area,
    build_next_step_recommendation,
    build_model_comparison_summary,
    build_advanced_classification_interpretation
)
from blueprints.survey.analysis import read_dataframe, repair_broken_csv_file
from blueprints.multivariate.routes import multivariate_report_word, multivariate_article
from config import PLOTS_DIR, UPLOAD_DIR
from extensions import SessionLocal, engine
from models import Dataset
from utils.manifest import read_manifest_data, write_manifest, auto_run_sadi_model
from utils.plot_manager import normalize_plot_catalog, split_plots
from utils.plot_meta import (
    classify_plot_tag,
    describe_plot,
    prettify_plot_title,
    summarize_plot_tags,
    build_general_dataset_figure_catalog,
)

# =========================
# LIKERT
# =========================
from blueprints.likert.routes import build_likert_summary

# =========================
# SADI CORE / MULTIVARIANTE
# =========================
from blueprints.multivariate.services import build_multivariate_profile
from blueprints.multivariate.services import generate_insights_ranking
from blueprints.multivariate.services import generate_sadi_conclusion

from blueprints.multivariate.services import generate_correlation_heatmap
from blueprints.multivariate.services import run_regression_analysis
from blueprints.multivariate.services import run_rf_regression_analysis

dataset_bp = Blueprint("dataset", __name__)

def _read_dataset_with_auto_repair(path: str, delimiter: str | None = None):
    """
    Lee el dataset usando el lector robusto central.
    Si detecta una sola columna sospechosa, intenta reparación automática.
    """
    df = read_dataframe(path, delimiter)

    if df.shape[1] == 1:
        only_col = str(df.columns[0]) if len(df.columns) > 0 else ""
        if any(x in only_col for x in [",", ";", "\t", "|", ":"]) or df.shape[0] == 0:
            repair_info = repair_broken_csv_file(path, delimiter)

            if repair_info.get("ok") and repair_info.get("repaired"):
                current_app.logger.warning(
                    f"[AUTO-REPAIR][dataset] Archivo reparado: {path} backup={repair_info.get('backup_path')}"
                )
                df = read_dataframe(path, delimiter)
            elif not repair_info.get("ok"):
                current_app.logger.warning(
                    f"[AUTO-REPAIR][dataset] No se pudo reparar {path}: {repair_info.get('message')}"
                )

    return df
# =========================================================
# Helpers
# =========================================================

def _fmt_num(x, decimals=3):
    try:
        return f"{float(x):.{decimals}f}"
    except Exception:
        return "—"


def _fmt_pct(x):
    try:
        return f"{float(x):.1f}%"
    except Exception:
        return "—"


def _latex_escape(text: str) -> str:
    if text is None:
        return ""
    text = str(text)
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def filter_plots_for_dataset(dataset_id: int, plots: list[str]) -> list[str]:
    out = []
    seen = set()

    for p in plots or []:
        if not isinstance(p, str):
            continue
        rel = p.replace("\\", "/")
        fname = rel.split("/")[-1]

        if not fname.lower().endswith(".png"):
            continue
        if not re.match(rf"^ds{dataset_id}_.+\.png$", fname, re.IGNORECASE):
            continue

        rel_std = os.path.join("plots", fname).replace("\\", "/")
        full_path = os.path.join(PLOTS_DIR, fname)
        if os.path.exists(full_path) and rel_std not in seen:
            out.append(rel_std)
            seen.add(rel_std)

    return out

def _dataset_word_build_abstract(meta: dict, ds_title: str, research_area: str = "general") -> str:
    meta = meta or {}

    n_rows = meta.get("n_rows", "—")
    n_cols = meta.get("n_cols", "—")
    n_num = meta.get("n_num", "—")
    n_cat = meta.get("n_cat", "—")
    n_dt = meta.get("n_dt", "—")
    missing = meta.get("missing_global_pct", None)

    try:
        missing_txt = f"{float(missing):.2f}%"
    except Exception:
        missing_txt = "—"

    return (
        f"Este informe presenta un análisis exploratorio automatizado del dataset "
        f"\"{ds_title}\", clasificado por SADI como dataset general en el área "
        f"\"{research_area}\". El conjunto contiene {n_rows} filas y {n_cols} columnas, "
        f"con {n_num} variables numéricas, {n_cat} categóricas y {n_dt} variables de fecha. "
        f"El análisis resume estructura, calidad de datos, patrones exploratorios, posibles "
        f"objetivos predictivos y recomendaciones metodológicas. El porcentaje global de "
        f"datos faltantes fue de {missing_txt}."
    )


def _dataset_word_build_introduction(meta: dict, ds_title: str) -> str:
    meta = meta or {}

    target_candidate = meta.get("target_candidate")
    target_type = meta.get("target_type")
    model_suggestion = meta.get("model_suggestion")

    base = (
        f"El análisis del dataset \"{ds_title}\" fue realizado automáticamente por SADI "
        f"con el propósito de describir su perfil estructural, detectar patrones relevantes, "
        f"identificar problemas de calidad y sugerir los siguientes pasos analíticos más útiles. "
        f"Además del análisis exploratorio, SADI evalúa si existe una variable objetivo plausible "
        f"para tareas de regresión o clasificación."
    )

    if target_candidate:
        extra = f" En este caso, SADI detectó como posible variable objetivo a \"{target_candidate}\""
        if target_type:
            extra += f", asociada a un problema de tipo {target_type}"
        if model_suggestion:
            extra += f", con sugerencia metodológica inicial basada en {model_suggestion}"
        extra += "."
        base += extra

    return base


def _safe_add_heading_paragraph(doc, title, text):
    if text is None:
        return
    text = str(text).strip()
    if not text:
        return
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)


def _safe_add_bullets(doc, title, items, prefix="• "):
    if not items:
        return
    doc.add_heading(title, level=1)
    for item in items:
        doc.add_paragraph(f"{prefix}{item}")


def _safe_add_table(doc, title, headers, rows):
    if not rows:
        return
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)

    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = "" if val is None else str(val)


def _dataset_is_model_plot(p: str) -> bool:
    low = str(p).lower()
    model_tokens = [
        "rf_",
        "regression",
        "logistic",
        "pred_vs_real",
        "feature_importance",
        "residuals",
        "coefficients",
        "roc",
        "confusion",
        "classification",
    ]
    return any(tok in low for tok in model_tokens)


def _dataset_add_figure_block(doc, idx, rel_path, plots_dir):
    rel_path = str(rel_path).replace("\\", "/")
    fname = rel_path.split("/")[-1]
    img_path = os.path.join(plots_dir, fname)

    if not os.path.exists(img_path):
        return False

    nice_title = prettify_plot_title(rel_path)
    nice_desc = describe_plot(rel_path)
    nice_tag = classify_plot_tag(rel_path)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.keep_with_next = True
    r = p_title.add_run(f"Gráfico {idx}. {nice_title}")
    r.bold = True
    r.font.size = Pt(12)

    p_type = doc.add_paragraph()
    p_type.paragraph_format.keep_with_next = True
    r2 = p_type.add_run(f"Tipo: {nice_tag}")
    r2.italic = True

    if nice_desc:
        p_desc = doc.add_paragraph(nice_desc)
        p_desc.paragraph_format.keep_with_next = True

    inserted = safe_add_picture(doc, img_path, width_inches=6.2)
    if inserted:
        doc.add_paragraph("")
        return True

    return False

def build_advanced_model_interpretation(
    *,
    regression_result: dict | None = None,
    rf_result: dict | None = None,
) -> list[str]:
    notes: list[str] = []

    def _to_float(v):
        try:
            return float(v)
        except Exception:
            return None

    lr_r2 = _to_float((regression_result or {}).get("r2"))
    lr_rmse = _to_float((regression_result or {}).get("rmse"))
    rf_r2 = _to_float((rf_result or {}).get("r2"))
    rf_rmse = _to_float((rf_result or {}).get("rmse"))

    if lr_r2 is not None:
        if lr_r2 >= 0.80:
            notes.append("La regresión lineal presenta un ajuste alto, lo que sugiere que las relaciones entre predictores y variable objetivo son relativamente estables.")
        elif lr_r2 >= 0.50:
            notes.append("La regresión lineal presenta un ajuste moderado, suficiente para exploración inicial pero susceptible de mejora.")
        else:
            notes.append("La regresión lineal presenta un ajuste limitado, por lo que conviene revisar predictores, outliers o relaciones no lineales.")

    if rf_r2 is not None:
        if rf_r2 >= 0.80:
            notes.append("Random Forest muestra una capacidad predictiva alta, lo que sugiere que el dataset contiene patrones aprovechables por modelos no lineales.")
        elif rf_r2 >= 0.50:
            notes.append("Random Forest muestra un desempeño moderado, útil para exploración predictiva y comparación de variables influyentes.")
        else:
            notes.append("Random Forest muestra un ajuste limitado, lo que sugiere revisar calidad de datos, tamaño muestral o relevancia del target detectado.")

    if lr_r2 is not None and rf_r2 is not None:
        diff = rf_r2 - lr_r2
        if diff >= 0.10:
            notes.append("El mejor rendimiento de Random Forest frente a la regresión lineal sugiere posibles relaciones no lineales en el dataset.")
        elif diff <= -0.10:
            notes.append("La regresión lineal supera claramente a Random Forest, lo que sugiere una estructura más lineal y simple en los datos.")
        else:
            notes.append("Ambos modelos presentan rendimientos relativamente cercanos, lo que sugiere una estructura predictiva sin diferencias extremas entre enfoques lineales y no lineales.")

    if lr_rmse is not None and rf_rmse is not None:
        if rf_rmse < lr_rmse:
            notes.append("Random Forest presenta menor error de predicción que la regresión lineal.")
        elif lr_rmse < rf_rmse:
            notes.append("La regresión lineal presenta menor error de predicción que Random Forest.")

    if not notes:
        notes.append("No hay suficiente información para generar una interpretación avanzada del modelo.")

    # quitar duplicados preservando orden
    seen = set()
    deduped = []
    for x in notes:
        if x not in seen:
            deduped.append(x)
            seen.add(x)

    return deduped



def safe_add_picture(doc, img_path: str, width_inches: float = 6.2) -> bool:
    try:
        if not os.path.exists(img_path):
            return False
        doc.add_picture(img_path, width=Inches(width_inches))
        return True
    except Exception:
        return False


def allowed_file(filename: str) -> bool:
    if not filename or "." not in filename:
        return False
    ext = filename.rsplit(".", 1)[1].lower()
    return ext in {"csv", "txt", "xlsx", "xls", "sav", "json"}


def generate_scientific_article_dataset(meta: dict, dataset_title: str) -> str:
    meta = meta or {}

    # =========================
    # Variables base
    # =========================
    n_rows = meta.get("n_rows", "—")
    n_cols = meta.get("n_cols", "—")
    n_num = meta.get("n_num", "—")
    n_cat = meta.get("n_cat", "—")
    n_dt = meta.get("n_dt", "—")
    missing = meta.get("missing_global_pct", "—")

    insights = meta.get("insights", []) or []
    insights_text = meta.get("insights_text") or "No se generaron insights automáticos."

    quick_recommendations = meta.get("quick_recommendations", []) or []
    suggested_plan = meta.get("suggested_plan", {}) or {}

    target_candidate = meta.get("target_candidate")
    target_type = meta.get("target_type")
    model_suggestion = meta.get("model_suggestion")

    high_corr = meta.get("high_corr_pairs", []) or []
    variability = meta.get("top_numeric_by_variability", []) or []
    outliers = meta.get("outlier_summary", []) or []

    # =========================
    # ABSTRACT
    # =========================
    abstract = f"""
Este estudio presenta un análisis exploratorio automatizado del dataset "{dataset_title}" mediante el sistema SADI.
El conjunto de datos contiene {n_rows} observaciones y {n_cols} variables, incluyendo {n_num} variables numéricas,
{n_cat} categóricas y {n_dt} variables de tipo fecha. El análisis permitió evaluar la calidad de los datos,
identificar patrones estructurales y detectar posibles variables objetivo para modelado predictivo.
El porcentaje global de valores faltantes fue de {missing}%.
""".strip()

    # =========================
    # INTRODUCCIÓN
    # =========================
    intro = f"""
El análisis exploratorio de datos constituye una etapa fundamental en cualquier proceso de investigación
basado en datos. En este estudio, se analiza el dataset "{dataset_title}" utilizando el sistema SADI,
el cual permite automatizar la detección de patrones, problemas de calidad y oportunidades de modelado.

El objetivo es caracterizar la estructura del dataset, identificar relaciones entre variables
y sugerir posibles enfoques analíticos para su explotación.
""".strip()

    # =========================
    # METODOLOGÍA
    # =========================
    methodology = """
Se aplicaron técnicas de análisis exploratorio automatizado, incluyendo:
- Evaluación de estructura del dataset
- Detección de valores faltantes
- Análisis de correlación
- Identificación de variables relevantes
- Detección de outliers
- Evaluación preliminar de modelos predictivos

El sistema SADI integró estos análisis para generar recomendaciones automáticas.
""".strip()

    # =========================
    # RESULTADOS
    # =========================
    results = f"""
El dataset presenta {n_rows} registros y {n_cols} variables.
Se identificaron {n_num} variables numéricas y {n_cat} categóricas.

El análisis de correlación permitió detectar relaciones relevantes entre variables,
mientras que el análisis de variabilidad permitió identificar aquellas con mayor dispersión.
Asimismo, se detectaron posibles valores atípicos que podrían afectar el modelado.

"""

    if target_candidate:
        results += f"""
SADI identificó la variable "{target_candidate}" como posible objetivo para un problema de tipo {target_type}.
Se sugiere el uso de modelos como {model_suggestion}.
"""

    # =========================
    # INSIGHTS
    # =========================
    insights_block = "\n".join([f"- {i}" for i in insights]) if insights else "No disponibles."

    # =========================
    # DISCUSIÓN
    # =========================
    discussion = """
Los resultados evidencian que el dataset contiene estructura analítica suficiente para realizar estudios más avanzados.
Las relaciones detectadas entre variables sugieren posibles dependencias que pueden ser aprovechadas en modelos predictivos.

Sin embargo, la presencia de valores faltantes y outliers indica la necesidad de procesos de limpieza y validación previos.
""".strip()

    # =========================
    # CONCLUSIÓN
    # =========================
    conclusion = f"""
El sistema SADI permitió obtener una visión integral del dataset "{dataset_title}".
Los resultados destacan la importancia del análisis exploratorio como base para la toma de decisiones analíticas.

Los insights generados automáticamente proporcionan una guía clara para el desarrollo de modelos predictivos
y análisis más profundos.
""".strip()

    # =========================
    # PRÓXIMOS PASOS
    # =========================
    next_steps = "\n".join([f"- {r}" for r in quick_recommendations]) if quick_recommendations else "No definidos."

    # =========================
    # PLAN SUGERIDO
    # =========================
    plan_block = ""
    if suggested_plan:
        plan_block += "\nAnálisis sugeridos:\n"
        for a in suggested_plan.get("recommended_analysis", []) or []:
            plan_block += f"- {a}\n"

        plan_block += "\nVisualizaciones sugeridas:\n"
        for g in suggested_plan.get("recommended_plots", []) or []:
            plan_block += f"- {g}\n"

    # =========================
    # ARTÍCULO FINAL
    # =========================
    article = f"""
TÍTULO
Análisis exploratorio automatizado del dataset "{dataset_title}" mediante SADI

RESUMEN
{abstract}

1. INTRODUCCIÓN
{intro}

2. METODOLOGÍA
{methodology}

3. RESULTADOS
{results}

4. INSIGHTS AUTOMÁTICOS
{insights_block}

5. INTERPRETACIÓN AUTOMÁTICA
{insights_text}

6. RECOMENDACIONES
{next_steps}

7. PLAN DE ANÁLISIS SUGERIDO
{plan_block}

8. DISCUSIÓN
{discussion}

9. CONCLUSIÓN
{conclusion}

10. REFERENCIAS
- Jolliffe, I. T. (2002). Principal Component Analysis.
- James, G. et al. (2021). An Introduction to Statistical Learning.
""".strip()

    return article

def analyze_general_dataset(df: pd.DataFrame) -> dict:
    import pandas as pd
    import numpy as np

    meta = {}

    n_rows, n_cols = int(df.shape[0]), int(df.shape[1])
    meta["n_rows"] = n_rows
    meta["n_cols"] = n_cols

    total_cells = n_rows * n_cols
    missing_total = int(df.isna().sum().sum()) if total_cells else 0
    meta["missing_global_pct"] = round((missing_total / total_cells) * 100, 2) if total_cells else 0.0

    num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    dt_cols = [c for c in df.columns if pd.api.types.is_datetime64_any_dtype(df[c])]
    cat_cols = [c for c in df.columns if c not in num_cols and c not in dt_cols]

    meta["num_cols"] = num_cols
    meta["cat_cols"] = cat_cols
    meta["dt_cols"] = dt_cols
    meta["n_num"] = len(num_cols)
    meta["n_cat"] = len(cat_cols)
    meta["n_dt"] = len(dt_cols)

    # =========================
    # Correlaciones altas
    # =========================
    high_corr_pairs = []
    if len(num_cols) >= 2:
        try:
            corr = df[num_cols].corr(numeric_only=True)
            for i in range(len(corr.columns)):
                for j in range(i + 1, len(corr.columns)):
                    val = corr.iloc[i, j]
                    if pd.notna(val) and abs(val) >= 0.70:
                        high_corr_pairs.append({
                            "col1": corr.columns[i],
                            "col2": corr.columns[j],
                            "corr": round(float(val), 4),
                        })
        except Exception:
            pass
    meta["high_corr_pairs"] = high_corr_pairs[:12]

    # =========================
    # Variables más variables
    # =========================
    top_numeric_by_variability = []
    for c in num_cols:
        s = df[c].dropna()
        if s.empty:
            continue
        std = float(s.std()) if s.shape[0] > 1 else 0.0
        rng = float(s.max() - s.min()) if s.shape[0] > 0 else 0.0
        top_numeric_by_variability.append({
            "column": c,
            "std": round(std, 4),
            "range": round(rng, 4),
        })

    top_numeric_by_variability.sort(key=lambda x: x["std"], reverse=True)
    meta["top_numeric_by_variability"] = top_numeric_by_variability[:8]

    # =========================
    # Outliers por IQR
    # =========================
    outlier_summary = []
    for c in num_cols:
        s = df[c].dropna()
        if s.shape[0] < 8:
            continue
        q1 = s.quantile(0.25)
        q3 = s.quantile(0.75)
        iqr = q3 - q1
        if iqr == 0:
            continue
        low = q1 - 1.5 * iqr
        high = q3 + 1.5 * iqr
        n_out = int(((s < low) | (s > high)).sum())
        pct_out = round((n_out / len(s)) * 100, 2) if len(s) else 0.0
        if n_out > 0:
            outlier_summary.append({
                "column": c,
                "n_outliers": n_out,
                "pct_outliers": pct_out,
            })

    outlier_summary.sort(key=lambda x: x["n_outliers"], reverse=True)
    meta["outlier_summary"] = outlier_summary[:8]

    # =========================
    # Patrones rápidos
    # =========================
    pattern_notes = []

    if len(high_corr_pairs) >= 3:
        pattern_notes.append("Se observan múltiples relaciones fuertes entre variables numéricas, lo que sugiere estructura interna relevante.")

    if len(outlier_summary) >= 2:
        pattern_notes.append("Se detectan valores atípicos en varias variables, por lo que conviene revisar robustez y limpieza antes de modelar.")

    if len(top_numeric_by_variability) >= 1:
        pattern_notes.append(
            f"La variable con mayor dispersión observada fue '{top_numeric_by_variability[0]['column']}'."
        )

    if meta["missing_global_pct"] >= 10:
        pattern_notes.append("El porcentaje de datos faltantes podría afectar la estabilidad de análisis más avanzados.")

    meta["pattern_notes"] = pattern_notes

    # =========================
    # Detección automática de target
    # =========================
    target_candidate = None
    target_type = None
    target_reason = None

    # numéricas utilizables
    usable_num = []
    for c in num_cols:
        s = df[c].dropna()
        if s.empty:
            continue
        if s.nunique() <= 1:
            continue
        usable_num.append(c)

    # categóricas utilizables
    usable_cat = []
    for c in cat_cols:
        s = df[c].dropna()
        if s.empty:
            continue
        nun = s.nunique()
        if 2 <= nun <= 20:
            usable_cat.append(c)

    # 1. Intentar target binario / clasificación
    binary_candidates = []
    for c in usable_cat:
        s = df[c].dropna()
        nun = s.nunique()
        if nun == 2:
            binary_candidates.append(c)

    if binary_candidates:
        target_candidate = binary_candidates[0]
        target_type = "classification"
        target_reason = "Se detectó una variable categórica binaria apta como objetivo de clasificación."

    # 2. Si no hay binaria, intentar regresión con una numérica útil
    elif len(usable_num) >= 2:
        ranked = []
        for c in usable_num:
            s = df[c].dropna()
            std = float(s.std()) if s.shape[0] > 1 else 0.0
            missing_pct = round(df[c].isna().mean() * 100, 2)
            score = std * (1 - (missing_pct / 100))
            ranked.append((c, score))

        ranked.sort(key=lambda x: x[1], reverse=True)
        if ranked:
            target_candidate = ranked[0][0]
            target_type = "regression"
            target_reason = "Se detectó una variable numérica con suficiente variabilidad como posible objetivo de regresión."

    meta["target_candidate"] = target_candidate
    meta["target_type"] = target_type
    meta["target_reason"] = target_reason
    meta["binary_target_candidates"] = binary_candidates

    return meta


# =========================================================
# Dashboard / detail / stats
# =========================================================

@dataset_bp.get("/dashboard")
@login_required
def dashboard():
    from collections import Counter
    from datetime import datetime
    from sqlalchemy import text, select

    from .constans import (
        DATASET_TYPE_CHOICES,
        RESEARCH_AREA_CHOICES,
        DATASET_TYPE_LABELS,
        RESEARCH_AREA_LABELS,
    )


    print("ENGINE URL:", engine.url)

    with SessionLocal() as db:
        #test_rows = db.execute(text("SELECT current_database(), current_schema()")).fetchall()
        #print("DB ACTUAL:", test_rows)

        stmt = select(Dataset).where(Dataset.user_id == current_user.id)
        datasets = db.execute(stmt).scalars().all()

    print("TOTAL DATASETS LEIDOS:", len(datasets))

    def _dt_key(d):
        dt = getattr(d, "uploaded_at", None) or getattr(d, "created_at", None)
        if isinstance(dt, datetime):
            return dt
        return datetime.min

    try:
        datasets.sort(key=_dt_key, reverse=True)
    except Exception:
        pass

    total_datasets = len(datasets)
    total_rows = sum(int(getattr(d, "n_rows", 0) or 0) for d in datasets)
    total_cols = sum(int(getattr(d, "n_cols", 0) or 0) for d in datasets)

    dataset_type_counter = Counter(
        (getattr(d, "dataset_type", None) or "dataset")
        for d in datasets
    )

    research_area_counter = Counter(
        (getattr(d, "research_area", None) or "general")
        for d in datasets
    )

    dataset_type_counter_labeled = {
        DATASET_TYPE_LABELS.get(k, k): v
        for k, v in dataset_type_counter.items()
    }

    research_area_counter_labeled = {
        RESEARCH_AREA_LABELS.get(k, k): v
        for k, v in research_area_counter.items()
    }

    total_surveys = (
        dataset_type_counter.get("survey_normal", 0)
        + dataset_type_counter.get("survey_likert_5", 0)
        + dataset_type_counter.get("survey_likert_7", 0)
    )
    total_general = dataset_type_counter.get("dataset", 0)
    total_multivariate = dataset_type_counter.get("multivariate", 0)
    total_undefined = dataset_type_counter.get("no_definido", 0)

    recent_datasets = datasets[:5]

    largest_dataset = None
    if datasets:
        try:
            largest_dataset = max(
                datasets,
                key=lambda d: (
                    int(getattr(d, "n_rows", 0) or 0),
                    int(getattr(d, "n_cols", 0) or 0),
                )
            )
        except Exception:
            largest_dataset = None

    dashboard_stats = {
    "total_datasets": total_datasets,
    "total_rows": total_rows,
    "total_cols": total_cols,
    "total_surveys": total_surveys,
    "total_general": total_general,
    "total_multivariate": total_multivariate,
    "total_undefined": total_undefined,
    "recent_datasets": recent_datasets,
    "largest_dataset": largest_dataset,

    "dataset_type_counter": dict(dataset_type_counter),
    "research_area_counter": dict(research_area_counter),

    "dataset_type_counter_labeled": dataset_type_counter_labeled,
    "research_area_counter_labeled": research_area_counter_labeled,

    "total_area_general": research_area_counter.get("general", 0),
    "total_biomedicina": research_area_counter.get("biomedicina", 0),
    "total_educacion": research_area_counter.get("educacion", 0),
    "total_finanzas": research_area_counter.get("finanzas", 0),
    "total_marketing": research_area_counter.get("marketing", 0),
    "total_agronomia": research_area_counter.get("agronomia", 0),
    "total_social": research_area_counter.get("social", 0),
    "total_ingenieria": research_area_counter.get("ingenieria", 0),
    "total_medio_ambiente": research_area_counter.get("medio_ambiente", 0),
    "total_legal": research_area_counter.get("legal", 0),
    "total_psicologia": research_area_counter.get("psicologia", 0),
}

    return render_template(
        "dashboard.html",
        datasets=datasets,
        dashboard_stats=dashboard_stats,
        DATASET_TYPE_CHOICES=DATASET_TYPE_CHOICES,
        RESEARCH_AREA_CHOICES=RESEARCH_AREA_CHOICES,
        DATASET_TYPE_LABELS=DATASET_TYPE_LABELS,
        RESEARCH_AREA_LABELS=RESEARCH_AREA_LABELS,
    )


@dataset_bp.get("/datasets/<int:dataset_id>/run_suggested_model", endpoint="dataset_run_suggested_model")
@login_required
def dataset_run_suggested_model(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    file_path = os.path.join(UPLOAD_DIR, ds.filename)
    df = _read_dataset_with_auto_repair(file_path, ds.delimiter)

    kind = (
        getattr(ds, "dataset_kind", None)
        or getattr(ds, "dataset_type", None)
        or getattr(ds, "kind", None)
        or getattr(ds, "type", None)
        or "dataset"
    )
    kind = (str(kind).strip() or "dataset").lower()

    research_area = (
        getattr(ds, "research_area", None)
        or "general"
    )
    research_area = (str(research_area).strip() or "general")

    analysis_meta = analyze_dataset_with_recommendations(
        df,
        dataset_type=kind,
        research_area=research_area,
    ) or {}

    target_candidate = analysis_meta.get("target_candidate")
    target_type = analysis_meta.get("target_type")

    if not target_candidate or not target_type:
        flash("SADI no pudo detectar una variable objetivo adecuada para modelado.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    if target_type == "regression":
        return redirect(url_for("multivariate.multivariate_rf_regression_auto", dataset_id=dataset_id, target_col=target_candidate))

    if target_type == "classification":
        return redirect(url_for("multivariate.multivariate_rf_classification_auto", dataset_id=dataset_id, target_col=target_candidate))

    flash("No se pudo determinar el tipo de modelo sugerido.", "warning")
    return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

@dataset_bp.get("/datasets/new-manual", endpoint="dataset_new_manual")
@login_required
def dataset_new_manual():
    """
    Muestra la pantalla para crear un dataset manual tipo planilla.
    """


    example_grid = {
        "header": ["Edad", "Genero", "Satisfaccion"],
        "rows": [
            ["22", "Femenino", "5"],
            ["25", "Masculino", "4"],
            ["21", "Femenino", "5"],
        ],
    }

    return render_template(
        "dataset_new_manual.html",
        page_title="Crear dataset manual",
        allowed_dataset_types=DATASET_TYPE_CHOICES,
        allowed_research_areas=RESEARCH_AREA_CHOICES,
        default_dataset_type="dataset",
        default_research_area="general",
        default_title="",
        example_grid=example_grid,
    )


@dataset_bp.get("/datasets/<int:dataset_id>")
@login_required
def dataset_detail(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

        preview = None
        analysis_meta = {}
        suggested_plan = {}
        suggested_area = None
        advanced_model_interpretation = []
        next_step_recommendation = None
        regression_result = None
        rf_result = None
        model_comparison = {"available": False}
        best_model_selection = {}

        try:
            file_path = os.path.join(UPLOAD_DIR, ds.filename)
            df = _read_dataset_with_auto_repair(file_path, ds.delimiter)

            if df.shape[1] == 1:
                only_col = str(df.columns[0]) if len(df.columns) > 0 else ""
                if any(x in only_col for x in [",", ";", "\t", "|", ":"]) or df.shape[0] == 0:
                    current_app.logger.warning(
                        f"[dataset_detail] ds{dataset_id} quedó con 1 columna tras lectura/reparación. columns={list(df.columns)}"
                    )

            preview = df.head(25)

            ds.n_rows = int(df.shape[0])
            ds.n_cols = int(df.shape[1])

            kind = (
                getattr(ds, "dataset_kind", None)
                or getattr(ds, "dataset_type", None)
                or getattr(ds, "kind", None)
                or getattr(ds, "type", None)
                or "dataset"
            )
            kind = (str(kind).strip() or "dataset").lower()

            research_area = (
                getattr(ds, "research_area", None)
                or "general"
            )
            research_area = (str(research_area).strip() or "general")

            # =========================
            # ANÁLISIS INTELIGENTE
            # =========================
            try:
                if kind in ("survey_likert_5", "survey_likert_7"):
                    summary = build_likert_summary(dataset_id)

                    analysis_meta = summary.get("analysis_meta", {}) or {}
                    if not isinstance(analysis_meta, dict):
                        analysis_meta = {}

                    suggested_plan = summary.get("suggested_plan", {}) or {}
                    suggested_area = summary.get("research_area_suggested")

                    analysis_meta["dimension_summary"] = summary.get("dimension_summary")
                    analysis_meta["dimension_summary_plot"] = summary.get("dimension_summary_plot")
                    analysis_meta["dimension_radar_plot"] = summary.get("dimension_radar_plot")
                    analysis_meta["plots"] = summary.get("plots", [])
                    analysis_meta["plot_summary"] = summary.get("plot_summary", [])
                    analysis_meta["insights"] = summary.get("insights", [])
                    analysis_meta["insights_text"] = summary.get("insights_text", "")
                    analysis_meta["quick_recommendations"] = summary.get("quick_recommendations", [])
                    analysis_meta["suggested_plan"] = summary.get("suggested_plan", {})
                    analysis_meta["priority_order"] = summary.get("priority_order", [])
                    analysis_meta["research_area_suggested"] = summary.get("research_area_suggested")

                else:
                    # 1) análisis inteligente base en vivo
                    analysis_meta = analyze_dataset_with_recommendations(
                        df,
                        dataset_type=kind,
                        research_area=research_area,
                    ) or {}

                    if not isinstance(analysis_meta, dict):
                        analysis_meta = {}

                    # 2) fusionar con meta persistido del manifest
                    persisted_manifest = read_manifest_data(dataset_id) or {}
                    if not isinstance(persisted_manifest, dict):
                        persisted_manifest = {}

                    persisted_meta = persisted_manifest.get("meta") or {}
                    if not isinstance(persisted_meta, dict):
                        persisted_meta = {}

                    # el meta persistido debe prevalecer para survey
                    if kind == "survey_normal":
                        merged_meta = dict(analysis_meta)
                        merged_meta.update(persisted_meta)
                        analysis_meta = merged_meta

                        # aliases consistentes para pantalla, Word y artículo
                        analysis_meta["sadi_insights"] = (
                            analysis_meta.get("sadi_insights")
                            or analysis_meta.get("insights_text")
                            or analysis_meta.get("survey_insights")
                            or ""
                        )
                        analysis_meta["sadi_recommendations"] = (
                            analysis_meta.get("sadi_recommendations")
                            or analysis_meta.get("quick_recommendations")
                            or []
                        )
                        analysis_meta["sadi_plan"] = (
                            analysis_meta.get("sadi_plan")
                            or analysis_meta.get("suggested_plan")
                            or {}
                        )
                        analysis_meta["sadi_priority"] = (
                            analysis_meta.get("sadi_priority")
                            or analysis_meta.get("priority_order")
                            or []
                        )

                        analysis_meta["insights"] = (
                            analysis_meta.get("insights")
                            or analysis_meta.get("group_findings")
                            or analysis_meta.get("survey_key_findings")
                            or []
                        )
                        analysis_meta["insights_text"] = (
                            analysis_meta.get("insights_text")
                            or analysis_meta.get("survey_insights")
                            or analysis_meta.get("results_text")
                            or ""
                        )
                        analysis_meta["quick_recommendations"] = (
                            analysis_meta.get("quick_recommendations")
                            or analysis_meta.get("sadi_recommendations")
                            or []
                        )
                        analysis_meta["suggested_plan"] = (
                            analysis_meta.get("suggested_plan")
                            or analysis_meta.get("sadi_plan")
                            or {}
                        )
                        analysis_meta["priority_order"] = (
                            analysis_meta.get("priority_order")
                            or analysis_meta.get("sadi_priority")
                            or []
                        )
                        analysis_meta["plot_summary"] = analysis_meta.get("plot_summary") or []

                    suggested_plan = analysis_meta.get("suggested_plan") or {}
                    suggested_area = analysis_meta.get("research_area_suggested")

                if (not getattr(ds, "research_area", None)) and suggested_area:
                    ds.research_area = suggested_area

            except Exception as e:
                current_app.logger.warning(
                    f"[dataset_detail] ds{dataset_id}: no se pudo generar análisis inteligente: {e}"
                )
                analysis_meta = {}
                suggested_plan = {}
                suggested_area = None

            db.commit()

        except Exception as e:
            current_app.logger.warning(
                f"[dataset_detail] ds{dataset_id}: no se pudo generar preview/análisis: {e}"
            )

            kind = (
                getattr(ds, "dataset_kind", None)
                or getattr(ds, "dataset_type", None)
                or getattr(ds, "kind", None)
                or getattr(ds, "type", None)
                or "dataset"
            )
            kind = (str(kind).strip() or "dataset").lower()

            research_area = (
                getattr(ds, "research_area", None)
                or "general"
            )
            research_area = (str(research_area).strip() or "general")

            df = None

    manifest_data = read_manifest_data(dataset_id) or {}
    if not isinstance(manifest_data, dict):
        manifest_data = {}

    best_model_selection = manifest_data.get("best_model_selection", {})
    if not isinstance(best_model_selection, dict):
        best_model_selection = {}

    # =========================
    # AUTO-GENERAR GRÁFICOS EXPLORATORIOS
    # =========================
    try:
        existing_plots = manifest_data.get("plots", [])
        if not isinstance(existing_plots, list):
            existing_plots = []

        exploratory_tokens = [
            "_missing",
            "_box",
            "_corr",
            "_scatter",
            "_hist",
            "_bar",
        ]

        has_exploratory = any(
            isinstance(p, str) and any(tok in p.lower() for tok in exploratory_tokens)
            for p in existing_plots
        )

        if kind == "dataset" and df is not None and not has_exploratory:
            current_app.logger.warning(
                f"[dataset_detail:auto_eda] ds{dataset_id}: generando gráficos exploratorios automáticos"
            )

            auto_generated = generate_general_dataset_plots(
                df,
                dataset_id,
                dataset_type=kind,
                research_area=research_area,
            ) or []

            merged_plots = []
            seen = set()

            for p in existing_plots + auto_generated:
                if isinstance(p, str) and p not in seen:
                    merged_plots.append(p)
                    seen.add(p)

            manifest_data["plots"] = merged_plots
            write_manifest(dataset_id, manifest_data)

    except Exception as e:
        current_app.logger.warning(f"[dataset_detail:auto_eda] ds{dataset_id}: {e}")

    # =========================
    # AUTO MODELO SADI
    # =========================
    try:
        if df is not None:
            manifest_data = auto_run_sadi_model(
                df=df,
                dataset_id=dataset_id,
                plots_dir=PLOTS_DIR,
                analysis_meta=analysis_meta,
                manifest_data=manifest_data,
            )

            write_manifest(dataset_id, manifest_data)

    except Exception as e:
        current_app.logger.warning(f"[dataset_detail:auto_model] ds{dataset_id}: {e}")

    manifest_plots = manifest_data.get("generated") or manifest_data.get("plots") or []
    if not isinstance(manifest_plots, list):
        manifest_plots = []

    def _basename_png(p: str) -> str:
        p = str(p).replace("\\", "/")
        return p.split("/")[-1]

    allowed_names = [
        _basename_png(p)
        for p in manifest_plots
        if isinstance(p, str) and p.lower().endswith(".png")
    ]

    meta_plots = analysis_meta.get("plots", []) if isinstance(analysis_meta, dict) else []
    if isinstance(meta_plots, list):
        for p in meta_plots:
            if isinstance(p, str) and p.lower().endswith(".png"):
                base = _basename_png(p)
                if base not in allowed_names:
                    allowed_names.append(base)

    try:
        all_png = [p for p in os.listdir(PLOTS_DIR) if p.lower().endswith(".png")]
    except Exception:
        all_png = []

    ds_re = re.compile(rf"^ds{dataset_id}_.+\.png$", re.IGNORECASE)
    found = [n for n in all_png if ds_re.match(n)]

    for n in found:
        if n not in allowed_names:
            allowed_names.append(n)

    if kind == "dataset":
        def order_key_dataset(n: str):
            low = n.lower()
            if "_missing" in low:
                return (0, n)
            if "_box" in low:
                return (1, n)
            if "_corr" in low:
                return (2, n)
            if "_scatter" in low:
                return (3, n)
            if "_hist" in low:
                return (4, n)
            if "_bar" in low:
                return (5, n)
            if "regression_" in low:
                return (6, n)
            if "rf_" in low:
                return (7, n)
            if "logistic" in low:
                return (8, n)
            return (9, n)

        allowed_names.sort(key=order_key_dataset)

    seen = set()
    unique_names = []
    for n in allowed_names:
        if n not in seen:
            unique_names.append(n)
            seen.add(n)

    valid_names = []
    for n in unique_names:
        full_path = os.path.join(PLOTS_DIR, n)
        if os.path.exists(full_path):
            valid_names.append(n)

    plots = [os.path.join("plots", n).replace("\\", "/") for n in valid_names]

    # =========================
    # Separar gráficos exploratorios vs modelos
    # =========================
    def _is_model_plot(p: str) -> bool:
        low = str(p).lower()
        model_tokens = [
            "rf_",
            "regression",
            "logistic",
            "pred_vs_real",
            "feature_importance",
            "residuals",
            "coefficients",
            "roc",
            "confusion",
            "classification",
        ]
        return any(tok in low for tok in model_tokens)

    eda_plots = [p for p in plots if not _is_model_plot(p)]
    model_plots = [p for p in plots if _is_model_plot(p)]

    cols = manifest_data.get("cols") or []
    if not isinstance(cols, list):
        cols = []
    cols = [c for c in cols if c is not None]
    manifest_data["cols"] = cols

    manifest_data["research_area"] = research_area
    manifest_data["dataset_type"] = kind
    manifest_data["analysis_meta"] = analysis_meta if isinstance(analysis_meta, dict) else {}
    manifest_data["suggested_plan"] = suggested_plan if isinstance(suggested_plan, dict) else {}
    manifest_data["suggested_area"] = suggested_area

    model_results = manifest_data.get("model_results", {})
    current_app.logger.warning(
        f"[DEBUG dataset_detail model_results] ds{dataset_id}: keys={list(model_results.keys()) if isinstance(model_results, dict) else 'NO_DICT'}"
    )
    if not isinstance(model_results, dict):
        model_results = {}

    regression_result = model_results.get("regression_result")
    rf_result = model_results.get("rf_result")

    try:
        model_comparison = build_model_comparison_summary(
            regression_result=regression_result,
            rf_result=rf_result,
        )
    except Exception as e:
        current_app.logger.warning(f"[dataset_detail:model_comparison] ds{dataset_id}: {e}")
        model_comparison = {"available": False}

    try:
        if analysis_meta.get("target_type") == "classification":
            advanced_model_interpretation = build_advanced_classification_interpretation(
                rf_result=rf_result,
            )
        else:
            advanced_model_interpretation = build_advanced_model_interpretation(
                regression_result=regression_result,
                rf_result=rf_result,
            )
    except Exception as e:
        current_app.logger.warning(f"[dataset_detail:advanced_model_interpretation] ds{dataset_id}: {e}")
        advanced_model_interpretation = []

    try:
        next_step_source = analysis_meta if isinstance(analysis_meta, dict) else {}
        next_step_recommendation = build_next_step_recommendation(
            analysis_meta=next_step_source,
            dataset_kind=kind,
            model_plots=model_plots,
        )
    except Exception as e:
        current_app.logger.warning(f"[dataset_detail:next_step] ds{dataset_id}: {e}")
        next_step_recommendation = None

    quick_recommendations = analysis_meta.get("quick_recommendations", []) if isinstance(analysis_meta, dict) else []
    final_plot_summary = analysis_meta.get("plot_summary") if isinstance(analysis_meta, dict) else None
    if not final_plot_summary:
        final_plot_summary = summarize_plot_tags(plots) if plots else []

    return render_template(
        "dataset_detail.html",
        ds=ds,
        preview=preview,
        plots=plots,
        eda_plots=eda_plots,
        model_plots=model_plots,
        plot_summary=final_plot_summary,
        dataset_kind=kind,
        research_area=research_area,
        manifest_data=manifest_data,
        analysis_meta=analysis_meta,
        meta=analysis_meta,
        suggested_plan=suggested_plan,
        suggested_area=suggested_area,
        classify_plot_tag=classify_plot_tag,
        prettify_plot_title=prettify_plot_title,
        describe_plot=describe_plot,
        quick_recommendations=quick_recommendations,
        regression_result=regression_result,
        rf_result=rf_result,
        model_comparison=model_comparison,
        advanced_model_interpretation=advanced_model_interpretation,
        next_step_recommendation=next_step_recommendation,
        best_model_selection=best_model_selection,
    )


@dataset_bp.get("/datasets/<int:dataset_id>/stats")
@login_required
def dataset_stats(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

        filename = ds.filename
        delimiter = ds.delimiter

        dataset_type = getattr(ds, "dataset_type", "dataset")
        research_area = getattr(ds, "research_area", "general")

        kind = (dataset_type or "dataset").strip().lower()

    file_path = os.path.join(UPLOAD_DIR, filename)

    try:
        df = _read_dataset_with_auto_repair(file_path, delimiter)

        if df.shape[1] == 1:
            only_col = str(df.columns[0]) if len(df.columns) > 0 else ""
            if any(x in only_col for x in [",", ";", "\t", "|", ":"]) or df.shape[0] == 0:
                flash(
                    "El archivo quedó con una sola columna incluso después del intento de reparación.",
                    "danger"
                )
                return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    except Exception as e:
        current_app.logger.exception(f"[dataset_stats] ds{dataset_id}: error leyendo dataset: {e}")
        flash(f"No se pudo analizar: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    # =========================================================
    # Limpiar SOLO gráficos exploratorios del módulo dataset
    # NO borrar gráficos predictivos / multivariados
    # =========================================================
    try:
        exploratory_tokens = [
            "_missing",
            "_box",
            "_corr",
            "_scatter",
            "_hist",
            "_bar",
        ]

        for p in glob.glob(os.path.join(PLOTS_DIR, f"ds{dataset_id}_*.png")):
            low = os.path.basename(p).lower()
            if any(tok in low for tok in exploratory_tokens):
                os.remove(p)

    except Exception as e:
        current_app.logger.warning(f"[dataset_stats] ds{dataset_id}: no se pudieron limpiar plots exploratorios previos: {e}")

    generated: list[str] = []
    meta = {}

    manifest_now = read_manifest_data(dataset_id)
    if not isinstance(manifest_now, dict):
        manifest_now = {}

    if kind == "dataset":
        try:
            meta = analyze_dataset_with_recommendations(
                df,
                dataset_type=kind,
                research_area=research_area,
            )
            if not isinstance(meta, dict):
                meta = {}

            meta["dataset_type"] = kind
            meta["research_area"] = research_area
            meta["n_total"] = int(df.shape[0])

        except Exception as e:
            current_app.logger.exception(f"[dataset_stats] ds{dataset_id}: error en análisis general inteligente: {e}")
            flash(f"Análisis general: error: {e}", "warning")
            meta = {}

        try:
            saved = generate_general_dataset_plots(
                df,
                dataset_id,
                dataset_type=kind,
                research_area=research_area,
            )

            if saved:
                generated.extend(saved)
                flash(f"Generados {len(saved)} gráficos exploratorios (dataset).", "success")
            else:
                flash("No se generaron gráficos exploratorios (dataset). Revisa tipos/valores.", "warning")

        except Exception as e:
            current_app.logger.exception(f"[dataset_stats] ds{dataset_id}: error generando gráficos: {e}")
            flash(f"Gráficos exploratorios: error: {e}", "warning")

        generated = sorted(set(generated))

        # =========================================================
        # Fusionar plots nuevos + existentes
        # para NO perder gráficos predictivos ya generados
        # =========================================================
        existing_plots = manifest_now.get("plots", [])
        if not isinstance(existing_plots, list):
            existing_plots = []

        merged_plots = []
        seen = set()

        for p in existing_plots + generated:
            if isinstance(p, str) and p not in seen:
                merged_plots.append(p)
                seen.add(p)

        # Mantener otras claves del manifest, por ejemplo model_results
        manifest_now["plots"] = merged_plots
        manifest_now["meta"] = meta if isinstance(meta, dict) else {}

        write_manifest(dataset_id, manifest_now)

        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    flash("Tipo de dataset no reconocido para este módulo. Selecciona 'dataset'.", "warning")

    existing_plots = manifest_now.get("plots", [])
    if not isinstance(existing_plots, list):
        existing_plots = []

    merged_plots = []
    seen = set()

    for p in existing_plots + generated:
        if isinstance(p, str) and p not in seen:
            merged_plots.append(p)
            seen.add(p)

    manifest_now["plots"] = merged_plots
    manifest_now["meta"] = meta if isinstance(meta, dict) else {}

    write_manifest(dataset_id, manifest_now)

    return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))


# =========================================================
# Upload / manual / delete / set_type
# =========================================================

@dataset_bp.post("/datasets/upload")
@login_required
def upload_dataset():
    import os
    import re
    import time
    import uuid
    from datetime import datetime

    from flask import current_app, flash, redirect, request, url_for
    from werkzeug.utils import secure_filename

    from .analysis import read_dataframe, suggest_research_area
    from .constans import normalize_research_area

    current_app.logger.warning("[upload] POST upload_dataset ENTRÓ")

    MAX_TITLE_LEN = 200
    DEFAULT_DELIMITER = ","
    ALLOWED_EXTENSIONS = {".csv", ".xlsx", ".xls", ".tsv", ".json"}

    def _normalize_delimiter(value: str | None) -> str:
        raw = (value or DEFAULT_DELIMITER).strip()
        aliases = {
            r"\t": "\t",
            "\\t": "\t",
            "tab": "\t",
            "TAB": "\t",
        }
        raw = aliases.get(raw, raw)
        if not raw or len(raw) != 1:
            return DEFAULT_DELIMITER
        return raw

    def _normalize_dataset_type(value: str | None) -> str:
        allowed = {
            "survey_likert_5",
            "survey_likert_7",
            "survey_normal",
            "dataset",
            "multivariate",
        }
        v = (value or "dataset").strip()
        return v if v in allowed else "dataset"

    def _safe_title(value: str | None) -> str:
        title = (value or "").strip()
        if len(title) > MAX_TITLE_LEN:
            title = title[:MAX_TITLE_LEN]
        cleaned = re.sub(r"[^a-zA-Z0-9_-]+", "_", title).strip("_")
        return cleaned[:60] if cleaned else "dataset"

    file = request.files.get("file")
    title = (request.form.get("title") or "").strip()
    delimiter = _normalize_delimiter(request.form.get("delimiter"))
    dataset_type = _normalize_dataset_type(
        request.form.get("dataset_type") or request.form.get("dataset_kind")
    )

    # 🔥 NUEVO: research_area correcto
    research_area_raw = normalize_research_area(request.form.get("research_area"))

    sheet_name = (request.form.get("sheet_name") or "").strip()

    current_app.logger.warning(
        f"[upload] title={title!r} delimiter={delimiter!r} "
        f"type={dataset_type!r} sheet_name={sheet_name!r} "
        f"has_file={file is not None}"
    )

    if not file:
        flash("Debes seleccionar un archivo.", "warning")
        return redirect(url_for("dataset.dashboard"))

    if not file.filename or not file.filename.strip():
        flash("El archivo no tiene nombre válido.", "warning")
        return redirect(url_for("dataset.dashboard"))

    original_name = file.filename.strip()
    ext = os.path.splitext(original_name)[1].lower()

    if ext not in ALLOWED_EXTENSIONS:
        flash("Formato no permitido. Usa CSV, XLSX, XLS, TSV o JSON.", "warning")
        return redirect(url_for("dataset.dashboard"))

    safe_original = secure_filename(original_name)
    if not safe_original:
        safe_original = f"archivo{ext}"

    unique_suffix = f"{int(time.time())}_{uuid.uuid4().hex[:8]}"
    safe_title = _safe_title(title)
    stored_name = f"up_{current_user.id}_{unique_suffix}_{safe_title}{ext}"
    fpath = os.path.join(UPLOAD_DIR, stored_name)

    # Guardar archivo
    try:
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        file.save(fpath)
    except Exception as e:
        current_app.logger.exception(f"[upload] error guardando archivo: {e}")
        flash(f"No se pudo guardar el archivo: {e}", "danger")
        return redirect(url_for("dataset.dashboard"))

    # Validar archivo guardado
    try:
        if not os.path.exists(fpath) or os.path.getsize(fpath) == 0:
            os.remove(fpath)
            flash("El archivo está vacío o no se guardó correctamente.", "warning")
            return redirect(url_for("dataset.dashboard"))
    except Exception as e:
        current_app.logger.exception(f"[upload] error verificando archivo: {e}")
        flash(f"No se pudo verificar el archivo: {e}", "danger")
        return redirect(url_for("dataset.dashboard"))

    # =========================
    # Validación de lectura
    # =========================
    try:
        if ext in {".xlsx", ".xls"} and sheet_name:
            current_app.logger.warning(
                f"[upload] sheet_name recibido pero no utilizado: {sheet_name!r}"
            )

        df = read_dataframe(fpath, delimiter if ext in {".csv", ".tsv"} else None)

        if df is None:
            raise ValueError("No se pudo leer el archivo como tabla.")

        if df.empty and len(df.columns) == 0:
            raise ValueError("El archivo no contiene datos utilizables.")

        # 🔥 AUTO DETECCIÓN DE ÁREA
        if research_area_raw in {"", "auto", "auto_detect"}:
            research_area = suggest_research_area(df)
        else:
            research_area = research_area_raw

        n_rows = int(df.shape[0])
        n_cols = int(df.shape[1])

        if n_cols <= 0:
            raise ValueError("El archivo no contiene columnas válidas.")

    except Exception as e:
        current_app.logger.exception(f"[upload] error validando lectura: {e}")
        try:
            if os.path.exists(fpath):
                os.remove(fpath)
        except Exception:
            pass
        flash(f"No se pudo leer o validar el archivo: {e}", "danger")
        return redirect(url_for("dataset.dashboard"))

    # =========================
    # Guardar en BD
    # =========================
    ds_id = None
    try:
        with SessionLocal() as db:
            ds = Dataset(
                user_id=current_user.id,
                title=title or safe_title or safe_original,
                filename=stored_name,
                original_name=original_name,
                delimiter=delimiter if ext in {".csv", ".tsv"} else None,
                n_rows=n_rows,
                n_cols=n_cols,
                dataset_type=dataset_type,

                # 🔥 CORRECTO
                research_area=research_area,

                uploaded_at=datetime.utcnow(),
            )

            db.add(ds)
            db.commit()
            db.refresh(ds)
            ds_id = ds.id

        current_app.logger.warning(
            f"[upload] OK dataset id={ds_id} "
            f"type={dataset_type} research_area={research_area} "
            f"rows={n_rows} cols={n_cols}"
        )

    except Exception as e:
        current_app.logger.exception(f"[upload] error guardando en BD: {e}")
        try:
            if os.path.exists(fpath):
                os.remove(fpath)
        except Exception:
            pass
        flash(f"No se pudo guardar el dataset en la base de datos: {e}", "danger")
        return redirect(url_for("dataset.dashboard"))

    if not ds_id:
        flash("No se pudo obtener el ID del dataset.", "warning")
        return redirect(url_for("dataset.dashboard"))

    flash("Dataset cargado correctamente.", "success")
    return redirect(url_for("dataset.dataset_detail", dataset_id=ds_id))


@dataset_bp.post("/datasets/<int:dataset_id>/delete", endpoint="dataset_delete")
@login_required
def dataset_delete(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

        try:
            file_path = os.path.join(UPLOAD_DIR, ds.filename)
            if ds.filename and os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"[dataset_delete] no se pudo borrar archivo ds{dataset_id}: {e}")

        try:
            for p in glob.glob(os.path.join(PLOTS_DIR, f"ds{dataset_id}_*.png")):
                try:
                    os.remove(p)
                except Exception:
                    pass
        except Exception as e:
            print(f"[dataset_delete] no se pudo limpiar plots ds{dataset_id}: {e}")

        try:
            db.delete(ds)
            db.commit()
        except Exception as e:
            db.rollback()
            flash(f"No se pudo eliminar el dataset: {e}", "danger")
            return redirect(url_for("dataset.dashboard"))

    flash("Dataset eliminado.", "success")
    return redirect(url_for("dataset.dashboard"))


@dataset_bp.post("/datasets/new_manual", endpoint="dataset_new_manual_post")
@login_required
def dataset_new_manual_post():
    import csv
    import json
    import os
    import re
    import time
    import uuid
    from datetime import datetime

    from flask import current_app, flash, redirect, request, url_for
    from .constans import DATASET_TYPE_VALUES, RESEARCH_AREA_VALUES

    current_app.logger.warning("[manual] POST dataset_new_manual_post ENTRÓ")

    # =========================================================
    # Configuración / límites defensivos
    # =========================================================
    MAX_TITLE_LEN = 200
    MAX_COLS = 300
    MAX_ROWS = 50000
    MAX_CELL_LEN = 10000
    DEFAULT_DELIMITER = ","

    # =========================================================
    # Lectura de form
    # =========================================================
    title = (request.form.get("title") or "").strip()
    delimiter = DEFAULT_DELIMITER
    dataset_type = (request.form.get("dataset_type") or "dataset").strip()
    research_area = (request.form.get("research_area") or "general").strip()
    data_json = (request.form.get("data_json") or "").strip()

    if dataset_type not in DATASET_TYPE_VALUES:
        current_app.logger.warning(f"[manual] dataset_type inválido recibido: {dataset_type!r}")
        dataset_type = "dataset"

    if research_area not in RESEARCH_AREA_VALUES:
        current_app.logger.warning(f"[manual] research_area inválido recibido: {research_area!r}")
        research_area = "general"

    # =========================================================
    # Validaciones iniciales
    # =========================================================
    if not data_json:
        flash("No se recibieron datos de la planilla.", "warning")
        return redirect(url_for("dataset.dataset_new_manual"))

    if len(title) > MAX_TITLE_LEN:
        title = title[:MAX_TITLE_LEN]

    if dataset_type not in DATASET_TYPE_VALUES:
        current_app.logger.warning(f"[manual] dataset_type inválido recibido: {dataset_type!r}")
        dataset_type = "dataset"

    if research_area not in RESEARCH_AREA_VALUES:
        current_app.logger.warning(f"[manual] research_area inválido recibido: {research_area!r}")
        research_area = "general"

    # =========================================================
    # Helpers
    # =========================================================
    def _normalize_cell(value):
        if value is None:
            return ""

        try:
            s = str(value)
        except Exception:
            s = ""

        s = s.replace("\\r\\n", "\n")
        s = s.replace("\\n", "\n")
        s = s.replace("\\r", "\r")
        s = s.replace("\\t", "\t")
        s = s.replace("\x00", "")
        s = s.replace("\u00A0", " ")
        s = s.strip()

        if len(s) > MAX_CELL_LEN:
            s = s[:MAX_CELL_LEN]

        return s

    def _normalize_row(row, width):
        if not isinstance(row, list):
            return None

        rr = [_normalize_cell(x) for x in row]

        if len(rr) < width:
            rr += [""] * (width - len(rr))
        elif len(rr) > width:
            rr = rr[:width]

        return rr

    def _has_any_real_data(row):
        return any(str(x).strip() != "" for x in row)

    def _normalize_header_name(name, idx):
        h = _normalize_cell(name)
        if not h:
            h = f"col_{idx}"

        h = re.sub(r"\s+", " ", h).strip()
        return h[:255]

    def _deduplicate_headers(headers):
        normalized = []
        used = set()

        for i, h in enumerate(headers, start=1):
            col_name = _normalize_header_name(h, i)
            base = col_name
            n = 2

            while col_name in used:
                col_name = f"{base}_{n}"
                n += 1

            used.add(col_name)
            normalized.append(col_name)

        return normalized

    # =========================================================
    # Parseo de JSON
    # =========================================================
    try:
        payload = json.loads(data_json)
    except Exception as e:
        current_app.logger.exception(f"[manual] JSON inválido: {e}")
        flash(f"Los datos de la planilla no tienen un formato JSON válido: {e}", "danger")
        return redirect(url_for("dataset.dataset_new_manual"))

    header = []
    rows = []

    try:
        if isinstance(payload, dict):
            if isinstance(payload.get("header"), list) and isinstance(payload.get("rows"), list):
                header = payload.get("header") or []
                rows = payload.get("rows") or []

            elif isinstance(payload.get("columns"), list) and isinstance(payload.get("data"), list):
                header = payload.get("columns") or []
                rows = payload.get("data") or []

            elif isinstance(payload.get("grid"), list):
                grid = payload.get("grid") or []
                if grid and isinstance(grid[0], list):
                    header = grid[0]
                    rows = grid[1:]

        elif isinstance(payload, list) and payload and isinstance(payload[0], list):
            header = payload[0]
            rows = payload[1:]

    except Exception as e:
        current_app.logger.exception(f"[manual] error interpretando payload: {e}")
        flash(f"No se pudo interpretar la planilla: {e}", "danger")
        return redirect(url_for("dataset.dataset_new_manual"))

    if not isinstance(header, list):
        header = []
    if not isinstance(rows, list):
        rows = []

    header = [_normalize_cell(x) for x in header]

    if not header or all(h == "" for h in header):
        flash("La planilla no tiene encabezados válidos.", "warning")
        return redirect(url_for("dataset.dataset_new_manual"))

    if len(header) > MAX_COLS:
        flash(f"La planilla supera el máximo permitido de columnas ({MAX_COLS}).", "warning")
        return redirect(url_for("dataset.dataset_new_manual"))

    header = _deduplicate_headers(header)

    # =========================================================
    # Limpiar filas
    # =========================================================
    clean_rows = []
    invalid_row_count = 0

    for r in rows:
        rr = _normalize_row(r, len(header))
        if rr is None:
            invalid_row_count += 1
            continue

        if _has_any_real_data(rr):
            clean_rows.append(rr)

        if len(clean_rows) > MAX_ROWS:
            flash(f"La planilla supera el máximo permitido de filas ({MAX_ROWS}).", "warning")
            return redirect(url_for("dataset.dataset_new_manual"))

    current_app.logger.warning(
        f"[manual] header_cols={len(header)} clean_rows={len(clean_rows)} invalid_rows={invalid_row_count}"
    )

    if not clean_rows:
        flash("Agregá al menos una fila de datos válida.", "warning")
        return redirect(url_for("dataset.dataset_new_manual"))

    # =========================================================
    # Guardar CSV
    # =========================================================
    safe_title = re.sub(r"[^a-zA-Z0-9_-]+", "_", title).strip("_")[:60] if title else "manual"
    unique_suffix = f"{int(time.time())}_{uuid.uuid4().hex[:8]}"
    fname = f"manual_{current_user.id}_{unique_suffix}_{safe_title}.csv"
    fpath = os.path.join(UPLOAD_DIR, fname)

    try:
        os.makedirs(UPLOAD_DIR, exist_ok=True)

        with open(fpath, "w", encoding="utf-8", newline="") as f:
            writer = csv.writer(
                f,
                delimiter=delimiter,
                quotechar='"',
                quoting=csv.QUOTE_MINIMAL,
                lineterminator="\n",
            )
            writer.writerow(header)
            writer.writerows(clean_rows)

        current_app.logger.warning(
            f"[manual] CSV guardado OK path={fpath} rows={len(clean_rows)} cols={len(header)}"
        )

    except Exception as e:
        current_app.logger.exception(f"[manual] error guardando CSV: {e}")
        flash(f"No se pudo guardar el archivo: {e}", "danger")
        return redirect(url_for("dataset.dataset_new_manual"))

    # =========================================================
    # Validación inmediata leyendo lo recién guardado
    # =========================================================
    try:
        df_check = read_dataframe(fpath, delimiter)

        if df_check is None or (df_check.empty and len(clean_rows) > 0):
            current_app.logger.warning(
                f"[manual] validación: dataframe vacío tras guardar {fpath}"
            )

        if not df_check.empty and df_check.shape[1] == 1:
            only_col = str(df_check.columns[0]) if len(df_check.columns) > 0 else ""
            if any(x in only_col for x in [",", ";", "\t", "|", ":"]):
                current_app.logger.warning(
                    f"[manual] validación falló: archivo quedó como 1 columna. columns={list(df_check.columns)}"
                )
                try:
                    os.remove(fpath)
                except Exception:
                    pass

                flash(
                    "El archivo manual quedó con formato inválido después del guardado. Intenta nuevamente.",
                    "danger",
                )
                return redirect(url_for("dataset.dataset_new_manual"))

    except Exception as e:
        current_app.logger.exception(f"[manual] validación de lectura falló: {e}")
        try:
            if os.path.exists(fpath):
                os.remove(fpath)
        except Exception:
            pass

        flash(f"No se pudo validar el archivo guardado: {e}", "danger")
        return redirect(url_for("dataset.dataset_new_manual"))

    # =========================================================
    # Guardar en BD
    # =========================================================
    ds_id = None

    try:
        with SessionLocal() as db:
            ds = Dataset(
                user_id=current_user.id,
                title=title or safe_title or "Dataset manual",
                filename=fname,
                original_name=fname,
                delimiter=delimiter,
                n_rows=len(clean_rows),
                n_cols=len(header),
                dataset_type=dataset_type,
                research_area=research_area,
                uploaded_at=datetime.utcnow(),
            )

            db.add(ds)
            db.commit()
            db.refresh(ds)
            ds_id = ds.id

        current_app.logger.warning(
            f"[manual] OK creado dataset id={ds_id} user_id={current_user.id} "
            f"file={fname} type={dataset_type} research_area={research_area}"
        )

    except Exception as e:
        try:
            if os.path.exists(fpath):
                os.remove(fpath)
        except Exception:
            pass

        current_app.logger.exception(f"[manual] error guardando en BD: {e}")
        flash(f"No se pudo guardar en la base de datos: {e}", "danger")
        return redirect(url_for("dataset.dataset_new_manual"))

    if not ds_id:
        flash("Se guardó el archivo pero no se pudo obtener el ID del dataset.", "warning")
        return redirect(url_for("dataset.dashboard"))

    flash("Dataset manual creado correctamente.", "success")

    current_kind = (dataset_type or "").strip().lower()

    if current_kind == "multivariate":
        return redirect(url_for("multivariate.multivariate_detail", dataset_id=ds_id))
    elif current_kind in ("survey_likert_5", "survey_likert_7"):
        return redirect(url_for("dataset.dataset_detail", dataset_id=ds_id))
    else:
        return redirect(url_for("dataset.dataset_detail", dataset_id=ds_id))


@dataset_bp.post("/datasets/<int:dataset_id>/set_type", endpoint="dataset_set_type")
@login_required
def dataset_set_type(dataset_id: int):
    raw_kind_1 = request.form.get("dataset_kind")
    raw_kind_2 = request.form.get("dataset_type")
    new_kind = (raw_kind_1 or raw_kind_2 or "").strip()

    allowed = {
        "survey_likert_5",
        "survey_likert_7",
        "survey_normal",
        "dataset",
        "multivariate",
        "no_definido",
    }

    if new_kind not in allowed:
        msg = f"Tipo de dataset no reconocido: {new_kind!r}"
        if request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return jsonify(ok=False, message=msg, received=new_kind), 400
        flash(msg, "warning")
        return redirect(url_for("dataset.dashboard"))

    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            msg = "Dataset no encontrado."
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return jsonify(ok=False, message=msg), 404
            flash(msg, "warning")
            return redirect(url_for("dataset.dashboard"))

        ds.dataset_type = new_kind
        db.commit()

    def label_and_badge(k: str):
        if k == "survey_likert_5":
            return ("Encuesta Likert 5 puntos", "badge-ok")
        if k == "survey_likert_7":
            return ("Encuesta Likert 7 puntos", "badge-ok")
        if k == "survey_normal":
            return ("Encuesta normal", "badge-ok")
        if k == "dataset":
            return ("No encuesta", "badge-info")
        return ("No definido", "badge-warn")

    label, badge_class = label_and_badge(new_kind)

    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return jsonify(
            ok=True,
            kind=new_kind,
            label=label,
            badge_class=badge_class,
            saved_field="dataset_type",
        )

    flash("Tipo de dataset actualizado.", "success")
    return redirect(url_for("dataset.dashboard"))


# =========================================================
# Exports for dataset = no encuesta
# =========================================================

@dataset_bp.get("/datasets/<int:dataset_id>/pdf", endpoint="dataset_pdf")
@login_required
def dataset_pdf(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    kind = (getattr(ds, "dataset_type", None) or "no_definido").strip()
    if kind != "dataset":
        flash("Este informe PDF aplica solo para 'No encuesta (Dataset)'.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    manifest = read_manifest_data(dataset_id) or {}
    meta = (manifest.get("meta") or {}) if isinstance(manifest, dict) else {}
    plots = (manifest.get("generated") or manifest.get("plots") or []) if isinstance(manifest, dict) else []

    if not isinstance(meta, dict):
        meta = {}
    if not isinstance(plots, list):
        plots = []

    plots = normalize_plot_catalog(dataset_id, manifest_data, summary)
    exploratory_plots, model_plots = split_plots(plots)

    if not meta:
        flash("Primero genera el análisis exploratorio del dataset.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    reports_dir = os.path.join(PLOTS_DIR, "reports")
    os.makedirs(reports_dir, exist_ok=True)
    pdf_path = os.path.join(reports_dir, f"ds{dataset_id}_dataset_exploratorio.pdf")

    try:
        doc = SimpleDocTemplate(
            pdf_path,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()
        h1 = styles["Heading1"]
        h2 = styles["Heading2"]
        body = styles["BodyText"]
        small = ParagraphStyle("small", parent=body, fontSize=9, leading=11)
        caption = ParagraphStyle("caption", parent=body, fontSize=10, leading=12)

        elements = []

        elements.append(Paragraph(f"Informe exploratorio del dataset: {ds.title or ds.original_name}", styles["Title"]))
        elements.append(Spacer(1, 0.2 * inch))
        elements.append(Paragraph(f"<b>Archivo:</b> {ds.original_name}", body))
        elements.append(Paragraph(f"<b>Dimensiones:</b> {ds.n_rows} × {ds.n_cols}", body))
        elements.append(Spacer(1, 0.15 * inch))

        data = [
            ["Indicador", "Valor"],
            ["Filas", str(meta.get("n_rows", ds.n_rows))],
            ["Columnas", str(meta.get("n_cols", ds.n_cols))],
            ["Missing global", _fmt_pct(meta.get("missing_global_pct", 0.0))],
            ["Duplicados", str(meta.get("n_duplicates", "—"))],
            ["Numéricas", str(meta.get("n_num", 0))],
            ["Categóricas", str(meta.get("n_cat", 0))],
            ["Fechas", str(meta.get("n_dt", 0))],
        ]

        table = Table(data, colWidths=[7 * cm, 7 * cm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 0.2 * inch))

        if meta.get("insights_text"):
            elements.append(Paragraph("Insights automáticos", h1))
            elements.append(Paragraph(str(meta.get("insights_text")), body))
            elements.append(Spacer(1, 0.15 * inch))

        if meta.get("strong_correlations"):
            elements.append(Paragraph("Correlaciones fuertes", h1))
            for r in meta.get("strong_correlations", [])[:10]:
                elements.append(
                    Paragraph(
                        f"{r.get('a', '—')} vs {r.get('b', '—')} (r={_fmt_num(r.get('corr'), 3)})",
                        body
                    )
                )
            elements.append(Spacer(1, 0.15 * inch))

        if plots:
            elements.append(PageBreak())
            elements.append(Paragraph("Gráficos del análisis", h1))
            elements.append(Spacer(1, 0.1 * inch))

            for i, rel in enumerate(plots, start=1):
                fname = rel.split("/")[-1]
                img_path = os.path.join(PLOTS_DIR, fname)
                if not os.path.exists(img_path):
                    continue

                nice_title = prettify_plot_title(rel)
                nice_desc = describe_plot(rel)
                nice_tag = classify_plot_tag(rel)

                img = RLImage(img_path)
                max_img_w = 16 * cm
                max_img_h = 10 * cm
                try:
                    img._restrictSize(max_img_w, max_img_h)
                except Exception:
                    img.drawWidth = max_img_w
                    img.drawHeight = max_img_h

                block = [
                    Paragraph(f"Gráfico {i}. {nice_title}", h2),
                    Paragraph(f"<b>Tipo:</b> {nice_tag}", small),
                    Paragraph(nice_desc, caption),
                    Spacer(1, 0.05 * inch),
                    img,
                    Spacer(1, 0.18 * inch),
                ]
                elements.append(KeepTogether(block))

        elements.append(PageBreak())
        elements.append(Paragraph("Conclusión exploratoria", h1))
        cierre = (
            "El análisis exploratorio permitió identificar la estructura general del dataset, "
            "su nivel de completitud, las variables con mayor proporción de datos faltantes "
            "y las principales relaciones entre variables numéricas."
        )
        elements.append(Paragraph(cierre, body))

        doc.build(elements)

        return send_file(
            pdf_path,
            as_attachment=True,
            download_name=f"dataset_{dataset_id}_exploratorio.pdf"
        )

    except Exception as e:
        print(f"[dataset_pdf] ERROR: {e}")
        flash(f"No se pudo generar el PDF: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))


@dataset_bp.get("/datasets/<int:dataset_id>/word", endpoint="dataset_word")
@login_required
def dataset_word(dataset_id: int):
    db = SessionLocal()
    try:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

        current_kind = (
            getattr(ds, "dataset_kind", None)
            or getattr(ds, "dataset_type", None)
            or getattr(ds, "kind", None)
            or getattr(ds, "type", None)
            or "dataset"
        )
        current_kind = (str(current_kind).strip() or "dataset").lower()

        if current_kind != "dataset":
            flash("Este informe Word aplica solo para 'No encuesta (Dataset)'.", "warning")
            return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

        research_area = (
            getattr(ds, "research_area", None)
            or "general"
        )
        research_area = (str(research_area).strip() or "general")

        preview = None
        analysis_meta = {}
        suggested_plan = {}
        suggested_area = None
        advanced_model_interpretation = []
        next_step_recommendation = None
        regression_result = None
        rf_result = None
        model_comparison = {"available": False}
        best_model_selection = {}

        # =========================
        # Leer dataset igual que dataset_detail
        # =========================
        file_path = os.path.join(UPLOAD_DIR, ds.filename)
        df = _read_dataset_with_auto_repair(file_path, ds.delimiter)
        preview = df.head(25)

        ds.n_rows = int(df.shape[0])
        ds.n_cols = int(df.shape[1])

        # =========================
        # Análisis inteligente
        # =========================
        analysis_meta = analyze_dataset_with_recommendations(
            df,
            dataset_type=current_kind,
            research_area=research_area,
        ) or {}

        if not isinstance(analysis_meta, dict):
            analysis_meta = {}

        suggested_plan = analysis_meta.get("suggested_plan") or {}
        suggested_area = analysis_meta.get("research_area_suggested")

        # =========================
        # Leer manifest
        # =========================
        manifest_data = read_manifest_data(dataset_id) or {}
        if not isinstance(manifest_data, dict):
            manifest_data = {}

        best_model_selection = manifest_data.get("best_model_selection", {})
        if not isinstance(best_model_selection, dict):
            best_model_selection = {}

        # =========================
        # Auto-generar EDA si falta
        # =========================
        existing_plots = manifest_data.get("plots", [])
        if not isinstance(existing_plots, list):
            existing_plots = []

        exploratory_tokens = ["_missing", "_box", "_corr", "_scatter", "_hist", "_bar"]
        has_exploratory = any(
            isinstance(p, str) and any(tok in p.lower() for tok in exploratory_tokens)
            for p in existing_plots
        )

        if df is not None and not has_exploratory:
            auto_generated = generate_general_dataset_plots(
                df,
                dataset_id,
                dataset_type=current_kind,
                research_area=research_area,
            ) or []

            merged_plots = []
            seen = set()
            for p in existing_plots + auto_generated:
                if isinstance(p, str) and p not in seen:
                    merged_plots.append(p)
                    seen.add(p)

            manifest_data["plots"] = merged_plots
            write_manifest(dataset_id, manifest_data)

        # =========================
        # Auto modelo SADI
        # =========================
        manifest_data = auto_run_sadi_model(
            df=df,
            dataset_id=dataset_id,
            plots_dir=PLOTS_DIR,
            analysis_meta=analysis_meta,
            manifest_data=manifest_data,
        )
        write_manifest(dataset_id, manifest_data)

        # =========================
        # Reconstruir plots como en dataset_detail
        # =========================
        manifest_plots = manifest_data.get("generated") or manifest_data.get("plots") or []
        if not isinstance(manifest_plots, list):
            manifest_plots = []

        def _basename_png(p: str) -> str:
            p = str(p).replace("\\", "/")
            return p.split("/")[-1]

        allowed_names = [
            _basename_png(p)
            for p in manifest_plots
            if isinstance(p, str) and p.lower().endswith(".png")
        ]

        meta_plots = analysis_meta.get("plots", []) if isinstance(analysis_meta, dict) else []
        if isinstance(meta_plots, list):
            for p in meta_plots:
                if isinstance(p, str) and p.lower().endswith(".png"):
                    base = _basename_png(p)
                    if base not in allowed_names:
                        allowed_names.append(base)

        try:
            all_png = [p for p in os.listdir(PLOTS_DIR) if p.lower().endswith(".png")]
        except Exception:
            all_png = []

        ds_re = re.compile(rf"^ds{dataset_id}_.+\.png$", re.IGNORECASE)
        found = [n for n in all_png if ds_re.match(n)]

        for n in found:
            if n not in allowed_names:
                allowed_names.append(n)

        def order_key_dataset(n: str):
            low = n.lower()
            if "_missing" in low:
                return (0, n)
            if "_box" in low:
                return (1, n)
            if "_corr" in low:
                return (2, n)
            if "_scatter" in low:
                return (3, n)
            if "_hist" in low:
                return (4, n)
            if "_bar" in low:
                return (5, n)
            if "regression_" in low:
                return (6, n)
            if "rf_" in low:
                return (7, n)
            if "logistic" in low:
                return (8, n)
            return (9, n)

        allowed_names.sort(key=order_key_dataset)

        seen = set()
        unique_names = []
        for n in allowed_names:
            if n not in seen:
                unique_names.append(n)
                seen.add(n)

        valid_names = []
        for n in unique_names:
            full_path = os.path.join(PLOTS_DIR, n)
            if os.path.exists(full_path):
                valid_names.append(n)

        plots = [os.path.join("plots", n).replace("\\", "/") for n in valid_names]
        eda_plots = [p for p in plots if not _dataset_is_model_plot(p)]
        model_plots = [p for p in plots if _dataset_is_model_plot(p)]

        # =========================
        # Resultados de modelos como en dataset_detail
        # =========================
        model_results = manifest_data.get("model_results", {})
        if not isinstance(model_results, dict):
            model_results = {}

        regression_result = model_results.get("regression_result")
        rf_result = model_results.get("rf_result")

        try:
            model_comparison = build_model_comparison_summary(
                regression_result=regression_result,
                rf_result=rf_result,
            )
        except Exception:
            model_comparison = {"available": False}

        try:
            if analysis_meta.get("target_type") == "classification":
                advanced_model_interpretation = build_advanced_classification_interpretation(
                    rf_result=rf_result,
                )
            else:
                advanced_model_interpretation = build_advanced_model_interpretation(
                    regression_result=regression_result,
                    rf_result=rf_result,
                )
        except Exception:
            advanced_model_interpretation = []

        try:
            next_step_recommendation = build_next_step_recommendation(
                analysis_meta=analysis_meta,
                dataset_kind=current_kind,
                model_plots=model_plots,
            )
        except Exception:
            next_step_recommendation = None

        quick_recommendations = analysis_meta.get("quick_recommendations", []) if isinstance(analysis_meta, dict) else []
        insights = analysis_meta.get("insights", []) if isinstance(analysis_meta, dict) else []
        insights_text = analysis_meta.get("insights_text") if isinstance(analysis_meta, dict) else None

        final_plot_summary = analysis_meta.get("plot_summary") if isinstance(analysis_meta, dict) else None
        if not final_plot_summary:
            final_plot_summary = summarize_plot_tags(plots) if plots else []

        # =========================
        # Preparar Word
        # =========================
        reports_dir = os.path.join(PLOTS_DIR, "reports")
        os.makedirs(reports_dir, exist_ok=True)
        out_name = f"dataset_{dataset_id}_reporte.docx"
        out_path = os.path.join(reports_dir, out_name)

        doc = Document()

        # Título
        doc.add_heading(f"Informe exploratorio del dataset: {ds.title or ds.original_name}", 0)
        doc.add_paragraph(f"Archivo: {ds.original_name}")
        doc.add_paragraph(f"Tipo: {current_kind}")
        doc.add_paragraph(f"Área: {research_area}")

        # Abstract / Resumen
        abstract_text = _dataset_word_build_abstract(
            analysis_meta,
            ds.title or ds.original_name,
            research_area=research_area,
        )
        _safe_add_heading_paragraph(doc, "Resumen / Abstract", abstract_text)

        # Introducción
        introduction_text = _dataset_word_build_introduction(
            analysis_meta,
            ds.title or ds.original_name,
        )
        _safe_add_heading_paragraph(doc, "Introducción", introduction_text)

        # Resumen del perfil
        doc.add_heading("Resumen del perfil", level=1)
        doc.add_paragraph(f"Filas: {analysis_meta.get('n_rows', ds.n_rows)}")
        doc.add_paragraph(f"Columnas: {analysis_meta.get('n_cols', ds.n_cols)}")
        doc.add_paragraph(f"Numéricas: {analysis_meta.get('n_num', '—')}")
        doc.add_paragraph(f"Categóricas: {analysis_meta.get('n_cat', '—')}")
        doc.add_paragraph(f"Fecha: {analysis_meta.get('n_dt', '—')}")
        if analysis_meta.get("missing_global_pct") is not None:
            doc.add_paragraph(f"Missing global (%): {_fmt_num(analysis_meta.get('missing_global_pct'), 2)}")
        else:
            doc.add_paragraph("Missing global (%): —")

        # Insights automáticos
        _safe_add_bullets(doc, "Insights automáticos SADI", insights)

        # Conclusión automática
        _safe_add_heading_paragraph(doc, "Conclusión automática", insights_text)

        # Próximo paso recomendado
        if next_step_recommendation:
            doc.add_heading("Próximo paso recomendado por SADI", level=1)
            doc.add_paragraph(f"Título: {next_step_recommendation.get('title', '—')}")
            doc.add_paragraph(f"Por qué: {next_step_recommendation.get('reason', '—')}")
            doc.add_paragraph(f"Siguiente acción: {next_step_recommendation.get('action', '—')}")

        # Recomendaciones
        doc.add_heading("SADI recomienda", level=1)

        if quick_recommendations:
            doc.add_paragraph("Recomendaciones rápidas:")
            for r in quick_recommendations:
                doc.add_paragraph(f"• {r}")

        if suggested_plan:
            recommended_analysis = suggested_plan.get("recommended_analysis", []) or []
            recommended_plots = suggested_plan.get("recommended_plots", []) or []
            narrative_focus = suggested_plan.get("narrative_focus")
            warnings = suggested_plan.get("warnings", []) or []

            if recommended_analysis:
                doc.add_paragraph("Análisis sugeridos:")
                for a in recommended_analysis:
                    doc.add_paragraph(f"• {a}")

            if recommended_plots:
                doc.add_paragraph("Gráficos sugeridos:")
                for g in recommended_plots:
                    doc.add_paragraph(f"• {g}")

            if narrative_focus:
                doc.add_paragraph(f"Enfoque: {narrative_focus}")

            if warnings:
                doc.add_paragraph("Advertencias:")
                for w in warnings:
                    doc.add_paragraph(f"• {w}")

        # Detalle técnico
        doc.add_heading("Detalle técnico", level=1)
        doc.add_paragraph("Numéricas: " + ", ".join(map(str, analysis_meta.get("num_cols", []) or [])) if analysis_meta.get("num_cols") else "Numéricas: —")
        doc.add_paragraph("Categóricas: " + ", ".join(map(str, analysis_meta.get("cat_cols", []) or [])) if analysis_meta.get("cat_cols") else "Categóricas: —")
        doc.add_paragraph("Fecha: " + ", ".join(map(str, analysis_meta.get("dt_cols", []) or [])) if analysis_meta.get("dt_cols") else "Fecha: —")

        # Correlaciones importantes
        if analysis_meta.get("high_corr_pairs"):
            rows = []
            for row in analysis_meta.get("high_corr_pairs", []):
                rows.append([
                    row.get("col1", "—"),
                    row.get("col2", "—"),
                    row.get("corr", "—"),
                ])
            _safe_add_table(doc, "Correlaciones importantes", ["Variable 1", "Variable 2", "Correlación"], rows)

        # Variables más variables
        if analysis_meta.get("top_numeric_by_variability"):
            rows = []
            for row in analysis_meta.get("top_numeric_by_variability", []):
                rows.append([
                    row.get("column", "—"),
                    row.get("std", "—"),
                    row.get("range", "—"),
                ])
            _safe_add_table(doc, "Variables más variables", ["Variable", "Desv. estándar", "Rango"], rows)

        # Variables clave
        if analysis_meta.get("variable_importance"):
            rows = []
            for row in analysis_meta.get("variable_importance", []):
                rows.append([
                    row.get("column", "—"),
                    row.get("score", "—"),
                    row.get("std", "—"),
                    f"{row.get('missing_pct', '—')}%",
                ])
            _safe_add_table(doc, "Variables clave del dataset", ["Variable", "Score", "Desv. estándar", "% Missing"], rows)

        # Variable objetivo sugerida
        if analysis_meta.get("target_candidate"):
            doc.add_heading("Variable objetivo sugerida", level=1)
            doc.add_paragraph(f"Variable detectada: {analysis_meta.get('target_candidate', '—')}")
            if analysis_meta.get("target_type") == "regression":
                doc.add_paragraph("Tipo de problema: Regresión")
            elif analysis_meta.get("target_type") == "classification":
                doc.add_paragraph("Tipo de problema: Clasificación")
            else:
                doc.add_paragraph("Tipo de problema: No definido")

            if analysis_meta.get("target_reason"):
                doc.add_paragraph(str(analysis_meta.get("target_reason")))

        # Candidatos objetivo
        if analysis_meta.get("ranked_target_candidates"):
            rows = []
            for row in analysis_meta.get("ranked_target_candidates", []):
                rows.append([
                    row.get("column", "—"),
                    row.get("type", "—"),
                    row.get("score", "—"),
                    row.get("reason", "—"),
                ])
            _safe_add_table(doc, "Candidatos a variable objetivo", ["Variable", "Tipo", "Score", "Motivo"], rows)

        # Resumen clasificación
        if analysis_meta.get("target_type") == "classification" and isinstance(rf_result, dict) and rf_result:
            doc.add_heading("Resumen del modelo de clasificación", level=1)
            doc.add_paragraph(f"Accuracy: {rf_result.get('accuracy', '—')}")
            doc.add_paragraph(f"Precision: {rf_result.get('precision', '—')}")
            doc.add_paragraph(f"Recall: {rf_result.get('recall', '—')}")
            doc.add_paragraph(f"F1-score: {rf_result.get('f1_score', '—')}")

        # Modelo sugerido
        if analysis_meta.get("target_candidate") and analysis_meta.get("target_type") == "regression":
            doc.add_heading("Modelo sugerido por SADI", level=1)
            doc.add_paragraph(f"Variable objetivo detectada: {analysis_meta.get('target_candidate', '—')}")
            doc.add_paragraph("Tipo de problema: Regresión")
            if analysis_meta.get("model_suggestion"):
                doc.add_paragraph(f"Modelo sugerido: {analysis_meta.get('model_suggestion')}")

        if analysis_meta.get("target_candidate") and analysis_meta.get("target_type") == "classification":
            doc.add_heading("Modelo sugerido por SADI", level=1)
            doc.add_paragraph(f"Variable objetivo detectada: {analysis_meta.get('target_candidate', '—')}")
            doc.add_paragraph("Tipo de problema: Clasificación")
            if analysis_meta.get("model_suggestion"):
                doc.add_paragraph(f"Modelo sugerido: {analysis_meta.get('model_suggestion')}")

        # Outliers
        if analysis_meta.get("outlier_summary"):
            rows = []
            for row in analysis_meta.get("outlier_summary", []):
                rows.append([
                    row.get("column", "—"),
                    row.get("n_outliers", "—"),
                    f"{row.get('pct_outliers', '—')}%",
                ])
            _safe_add_table(doc, "Outliers detectados", ["Variable", "N outliers", "% outliers"], rows)

        # Resumen de gráficos
        if final_plot_summary:
            _safe_add_heading_paragraph(doc, "Resumen de gráficos", ", ".join(map(str, final_plot_summary)))

        # Interpretación de modelos
        if isinstance(manifest_data, dict):
            model_interpretation = manifest_data.get("model_interpretation", []) or []
        else:
            model_interpretation = []

        _safe_add_bullets(doc, "Interpretación automática del modelo", model_interpretation)
        _safe_add_bullets(doc, "Evaluación automática del modelo", advanced_model_interpretation)

        # Comparación de modelos
        if isinstance(model_comparison, dict):
            if not model_comparison.get("available") and analysis_meta.get("target_type") == "classification":
                doc.add_heading("Comparación automática de modelos", level=1)
                doc.add_paragraph(
                    "La comparación entre modelos de regresión no aplica en este caso, "
                    "ya que el problema detectado es de clasificación."
                )

            if model_comparison.get("available"):
                linear = model_comparison.get("linear") or {}
                random_forest = model_comparison.get("random_forest") or {}

                rows = [
                    ["Regresión lineal", linear.get("r2", "—"), linear.get("mae", "—"), linear.get("rmse", "—")],
                    ["Random Forest", random_forest.get("r2", "—"), random_forest.get("mae", "—"), random_forest.get("rmse", "—")],
                ]
                _safe_add_table(doc, "Comparación automática de modelos", ["Modelo", "R²", "MAE", "RMSE"], rows)

                winner = model_comparison.get("winner")
                if winner == "random_forest":
                    winner_txt = "Random Forest"
                elif winner == "linear_regression":
                    winner_txt = "Regresión lineal"
                elif winner == "tie":
                    winner_txt = "Rendimiento similar"
                else:
                    winner_txt = "No determinado"

                doc.add_paragraph(f"Modelo recomendado por SADI: {winner_txt}")

                if model_comparison.get("recommendation"):
                    doc.add_paragraph(str(model_comparison.get("recommendation")))

                if model_comparison.get("summary"):
                    for line in model_comparison.get("summary", []):
                        doc.add_paragraph(f"• {line}")

        # Selección mejor modelo
        if isinstance(best_model_selection, dict) and best_model_selection.get("best_model"):
            bm = best_model_selection.get("best_model")
            if bm == "random_forest":
                bm_txt = "Random Forest"
            elif bm == "linear_regression":
                bm_txt = "Regresión lineal"
            elif bm == "tie":
                bm_txt = "Rendimiento similar"
            else:
                bm_txt = "No determinado"

            doc.add_heading("Selección inteligente del mejor modelo", level=1)
            doc.add_paragraph(f"Variable objetivo: {best_model_selection.get('target_col', '—')}")
            doc.add_paragraph("Tipo de problema: Regresión")
            doc.add_paragraph(f"Modelo preferido por SADI: {bm_txt}")
            doc.add_paragraph(str(best_model_selection.get("reason", "—")))

            for line in best_model_selection.get("comparison_summary", []) or []:
                doc.add_paragraph(f"• {line}")

        # Gráficos exploratorios
        fig_idx = 1
        if eda_plots:
            doc.add_page_break()
            doc.add_heading("Gráficos exploratorios", level=1)
            doc.add_paragraph("Visualizaciones generadas por el análisis exploratorio del dataset.")
            for rel in eda_plots:
                if _dataset_add_figure_block(doc, fig_idx, rel, PLOTS_DIR):
                    fig_idx += 1

        # Gráficos de modelos
        if model_plots:
            doc.add_page_break()
            doc.add_heading("Gráficos de modelos predictivos", level=1)
            doc.add_paragraph("Visualizaciones generadas por modelos ejecutados y asociadas a este dataset.")
            for rel in model_plots:
                if _dataset_add_figure_block(doc, fig_idx, rel, PLOTS_DIR):
                    fig_idx += 1

        # Vista previa
        if preview is not None and len(preview.columns) > 0:
            doc.add_page_break()
            doc.add_heading("Vista previa", level=1)

            max_cols = min(len(preview.columns), 8)
            cols_to_use = list(preview.columns[:max_cols])

            table = doc.add_table(rows=1, cols=max_cols)
            table.style = "Table Grid"

            for i, col in enumerate(cols_to_use):
                table.rows[0].cells[i].text = str(col)

            for _, row_data in preview.iterrows():
                row = table.add_row().cells
                for i, col in enumerate(cols_to_use):
                    try:
                        row[i].text = str(row_data[col])
                    except Exception:
                        row[i].text = "—"

        doc.save(out_path)

        if not os.path.exists(out_path) or os.path.getsize(out_path) < 1000:
            flash("No se pudo generar correctamente el Word.", "danger")
            return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

        return send_file(
            out_path,
            as_attachment=True,
            download_name=out_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    except Exception as e:
        current_app.logger.exception(f"[dataset_word] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el Word: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))
    finally:
        db.close()


@dataset_bp.get("/datasets/<int:dataset_id>/latex", endpoint="dataset_latex")
@login_required
def dataset_latex(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    kind = (getattr(ds, "dataset_type", None) or "no_definido").strip()
    if kind != "dataset":
        flash("Este informe LaTeX aplica solo para 'No encuesta (Dataset)'.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    manifest = read_manifest_data(dataset_id) or {}
    meta = (manifest.get("meta") or {}) if isinstance(manifest, dict) else {}
    plots = (manifest.get("generated") or manifest.get("plots") or []) if isinstance(manifest, dict) else []

    if not isinstance(meta, dict):
        meta = {}
    if not isinstance(plots, list):
        plots = []

    plots = normalize_plot_catalog(dataset_id, manifest_data, summary)
    exploratory_plots, model_plots = split_plots(plots)

    if not meta:
        flash("Primero genera el análisis exploratorio del dataset.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    reports_dir = os.path.join(PLOTS_DIR, "reports")
    os.makedirs(reports_dir, exist_ok=True)

    pack_dir = os.path.join(reports_dir, f"ds{dataset_id}_latex_pack")
    images_dir = os.path.join(pack_dir, "images")
    tex_path = os.path.join(pack_dir, "main.tex")
    zip_path = os.path.join(reports_dir, f"ds{dataset_id}_dataset_latex.zip")

    try:
        if os.path.exists(pack_dir):
            shutil.rmtree(pack_dir)
        os.makedirs(images_dir, exist_ok=True)

        copied_images = []
        for rel in plots:
            rel = str(rel).replace("\\", "/")
            fname = rel.split("/")[-1]
            src = os.path.join(PLOTS_DIR, fname)
            if not os.path.exists(src):
                continue

            dst = os.path.join(images_dir, fname)
            shutil.copy2(src, dst)
            copied_images.append(fname)

        title_txt = _latex_escape(ds.title or ds.original_name)
        original_name_txt = _latex_escape(ds.original_name or "")
        insights_text = _latex_escape(meta.get("insights_text", "Sin insights automáticos disponibles."))

        latex_lines = [
            r"\documentclass[12pt]{article}",
            r"\usepackage[utf8]{inputenc}",
            r"\usepackage[T1]{fontenc}",
            r"\usepackage[spanish]{babel}",
            r"\usepackage{geometry}",
            r"\usepackage{graphicx}",
            r"\usepackage{float}",
            r"\usepackage{longtable}",
            r"\usepackage{booktabs}",
            r"\usepackage{array}",
            r"\usepackage{xcolor}",
            r"\geometry{margin=2.5cm}",
            r"\title{Informe exploratorio del dataset}",
            r"\author{SADI}",
            r"\date{\today}",
            "",
            r"\begin{document}",
            r"\maketitle",
            "",
            r"\section*{Identificación}",
            f"Título: {title_txt}",
            "",
            f"Archivo: {original_name_txt}",
            "",
            f"Dimensiones: {ds.n_rows} x {ds.n_cols}",
            "",
            r"\section*{Resumen general}",
            f"Filas: {meta.get('n_rows', ds.n_rows)}",
            "",
            f"Columnas: {meta.get('n_cols', ds.n_cols)}",
            "",
            f"Missing global (\\%): {_fmt_pct(meta.get('missing_global_pct', 0.0))}",
            "",
            f"Duplicados: {meta.get('n_duplicates', '—')}",
            "",
            f"Variables numéricas: {meta.get('n_num', 0)}",
            "",
            f"Variables categóricas: {meta.get('n_cat', 0)}",
            "",
            f"Variables de fecha: {meta.get('n_dt', 0)}",
            "",
            r"\section*{Insights automáticos}",
            insights_text,
            "",
        ]

        top_missing = meta.get("top_missing_cols") or []
        if top_missing:
            latex_lines.extend([
                r"\section*{Columnas con más faltantes}",
                r"\begin{longtable}{p{9cm}r}",
                r"\toprule",
                r"Columna & Missing (\%) \\",
                r"\midrule",
                r"\endhead",
            ])
            for row in top_missing[:15]:
                name = _latex_escape(row.get("name", "—"))
                miss = _fmt_pct(row.get("missing_pct", 0.0)).replace("%", r"\%")
                latex_lines.append(f"{name} & {miss} \\\\")
            latex_lines.extend([
                r"\bottomrule",
                r"\end{longtable}",
                "",
            ])

        strong = meta.get("strong_correlations") or []
        if strong:
            latex_lines.extend([
                r"\section*{Correlaciones fuertes}",
                r"\begin{longtable}{p{5.2cm}p{5.2cm}r}",
                r"\toprule",
                r"Variable A & Variable B & r \\",
                r"\midrule",
                r"\endhead",
            ])
            for row in strong[:20]:
                a = _latex_escape(row.get("a", "—"))
                b = _latex_escape(row.get("b", "—"))
                corr = _fmt_num(row.get("corr"), 3)
                latex_lines.append(f"{a} & {b} & {corr} \\\\")
            latex_lines.extend([
                r"\bottomrule",
                r"\end{longtable}",
                "",
            ])

        if copied_images:
            latex_lines.extend([
                r"\section*{Gráficos del análisis}",
                r"Los siguientes gráficos fueron generados automáticamente por SADI.",
                "",
            ])

            for i, fname in enumerate(copied_images, start=1):
                rel_plot = os.path.join("plots", fname).replace("\\", "/")
                nice_title = _latex_escape(prettify_plot_title(rel_plot))
                nice_desc = _latex_escape(describe_plot(rel_plot))
                nice_tag = _latex_escape(classify_plot_tag(rel_plot))

                latex_lines.extend([
                    r"\begin{figure}[H]",
                    r"\centering",
                    rf"\includegraphics[width=0.92\textwidth]{{images/{_latex_escape(fname)}}}",
                    rf"\caption{{{nice_title} ({nice_tag}). {nice_desc}}}",
                    rf"\label{{fig:{dataset_id}_{i}}}",
                    r"\end{figure}",
                    "",
                ])

        latex_lines.extend([
            r"\section*{Conclusión exploratoria}",
            _latex_escape(
                "El análisis exploratorio permitió identificar la estructura general del dataset, "
                "su nivel de completitud, las variables con mayor proporción de datos faltantes "
                "y las principales relaciones entre variables numéricas."
            ),
            "",
            r"\end{document}",
            "",
        ])

        with open(tex_path, "w", encoding="utf-8") as f:
            f.write("\n".join(latex_lines))

        if os.path.exists(zip_path):
            os.remove(zip_path)

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(tex_path, arcname="main.tex")
            for fname in copied_images:
                img_path = os.path.join(images_dir, fname)
                if os.path.exists(img_path):
                    zf.write(img_path, arcname=os.path.join("images", fname))

        return send_file(
            zip_path,
            as_attachment=True,
            download_name=f"dataset_{dataset_id}_latex.zip"
        )

    except Exception as e:
        print(f"[dataset_latex] ERROR: {e}")
        flash(f"No se pudo generar el paquete LaTeX: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))


@dataset_bp.get("/datasets/<int:dataset_id>/dataset_articulo_word", endpoint="dataset_articulo_word")
@login_required
def dataset_articulo_word(dataset_id: int):
    db = SessionLocal()
    try:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

        kind = (
            getattr(ds, "dataset_kind", None)
            or getattr(ds, "dataset_type", None)
            or getattr(ds, "kind", None)
            or getattr(ds, "type", None)
            or "dataset"
        )
        kind = (str(kind).strip() or "dataset").lower()

        if kind != "dataset":
            flash("Por ahora el artículo científico automático está disponible solo para 'No encuesta (Dataset)'.", "warning")
            return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

        research_area = (
            getattr(ds, "research_area", None)
            or "general"
        )
        research_area = (str(research_area).strip() or "general")

        # =========================
        # Helpers locales
        # =========================
        def _safe_dict(v):
            return v if isinstance(v, dict) else {}

        def _safe_list(v):
            if v is None:
                return []
            if isinstance(v, list):
                return v
            if isinstance(v, tuple):
                return list(v)
            return [v]

        def _safe_text(v, default="—"):
            if v is None:
                return default
            s = str(v).strip()
            return s if s else default

        def _basename_png(p: str) -> str:
            p = str(p).replace("\\", "/")
            return p.split("/")[-1]

        def _is_model_plot(p: str) -> bool:
            low = str(p).lower()
            model_tokens = [
                "rf_",
                "regression",
                "logistic",
                "pred_vs_real",
                "feature_importance",
                "residuals",
                "coefficients",
                "roc",
                "confusion",
                "classification",
            ]
            return any(tok in low for tok in model_tokens)

        def _winner_to_text(winner):
            if winner == "random_forest":
                return "Random Forest"
            if winner == "linear_regression":
                return "Regresión lineal"
            if winner == "tie":
                return "Rendimiento similar"
            return "No determinado"

        def _add_model_comparison_table(doc_obj, comparison):
            comparison = _safe_dict(comparison)
            linear = _safe_dict(comparison.get("linear"))
            random_forest = _safe_dict(comparison.get("random_forest"))

            table = doc_obj.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Modelo"
            hdr[1].text = "R²"
            hdr[2].text = "MAE"
            hdr[3].text = "RMSE"

            row = table.add_row().cells
            row[0].text = "Regresión lineal"
            row[1].text = _safe_text(linear.get("r2"))
            row[2].text = _safe_text(linear.get("mae"))
            row[3].text = _safe_text(linear.get("rmse"))

            row = table.add_row().cells
            row[0].text = "Random Forest"
            row[1].text = _safe_text(random_forest.get("r2"))
            row[2].text = _safe_text(random_forest.get("mae"))
            row[3].text = _safe_text(random_forest.get("rmse"))

        # =========================
        # Leer dataset
        # =========================
        file_path = os.path.join(UPLOAD_DIR, ds.filename)
        df = _read_dataset_with_auto_repair(file_path, ds.delimiter)

        ds.n_rows = int(df.shape[0])
        ds.n_cols = int(df.shape[1])

        # =========================
        # Recalcular análisis inteligente
        # =========================
        analysis_meta = analyze_dataset_with_recommendations(
            df,
            dataset_type=kind,
            research_area=research_area,
        ) or {}

        if not isinstance(analysis_meta, dict):
            analysis_meta = {}

        suggested_plan = analysis_meta.get("suggested_plan") or {}
        suggested_area = analysis_meta.get("research_area_suggested")

        # =========================
        # Leer manifest
        # =========================
        manifest_data = read_manifest_data(dataset_id) or {}
        if not isinstance(manifest_data, dict):
            manifest_data = {}

        # =========================
        # Auto EDA si falta
        # =========================
        existing_plots = manifest_data.get("plots", [])
        if not isinstance(existing_plots, list):
            existing_plots = []

        exploratory_tokens = ["_missing", "_box", "_corr", "_scatter", "_hist", "_bar"]
        has_exploratory = any(
            isinstance(p, str) and any(tok in p.lower() for tok in exploratory_tokens)
            for p in existing_plots
        )

        if df is not None and not has_exploratory:
            auto_generated = generate_general_dataset_plots(
                df,
                dataset_id,
                dataset_type=kind,
                research_area=research_area,
            ) or []

            merged_plots = []
            seen = set()
            for p in existing_plots + auto_generated:
                if isinstance(p, str) and p not in seen:
                    merged_plots.append(p)
                    seen.add(p)

            manifest_data["plots"] = merged_plots
            write_manifest(dataset_id, manifest_data)

        # =========================
        # Auto modelo SADI
        # =========================
        try:
            manifest_data = auto_run_sadi_model(
                df=df,
                dataset_id=dataset_id,
                plots_dir=PLOTS_DIR,
                analysis_meta=analysis_meta,
                manifest_data=manifest_data,
            ) or manifest_data
            write_manifest(dataset_id, manifest_data)
        except Exception as e:
            current_app.logger.warning(f"[dataset_articulo_word:auto_model] ds{dataset_id}: {e}")

        # REFRESCAR manifest después del auto modelo
        manifest_data = read_manifest_data(dataset_id) or manifest_data
        if not isinstance(manifest_data, dict):
            manifest_data = {}

        best_model_selection = manifest_data.get("best_model_selection", {})
        if not isinstance(best_model_selection, dict):
            best_model_selection = {}

        # =========================
        # Reconstruir lista completa de plots
        # =========================
        manifest_plots = manifest_data.get("generated") or manifest_data.get("plots") or []
        if not isinstance(manifest_plots, list):
            manifest_plots = []

        allowed_names = [
            _basename_png(p)
            for p in manifest_plots
            if isinstance(p, str) and p.lower().endswith(".png")
        ]

        meta_plots = analysis_meta.get("plots", []) if isinstance(analysis_meta, dict) else []
        if isinstance(meta_plots, list):
            for p in meta_plots:
                if isinstance(p, str) and p.lower().endswith(".png"):
                    base = _basename_png(p)
                    if base not in allowed_names:
                        allowed_names.append(base)

        try:
            all_png = [p for p in os.listdir(PLOTS_DIR) if p.lower().endswith(".png")]
        except Exception:
            all_png = []

        ds_re = re.compile(rf"^ds{dataset_id}_.+\.png$", re.IGNORECASE)
        found = [n for n in all_png if ds_re.match(n)]

        for n in found:
            if n not in allowed_names:
                allowed_names.append(n)

        def order_key_dataset(n: str):
            low = n.lower()
            if "_missing" in low:
                return (0, n)
            if "_box" in low:
                return (1, n)
            if "_corr" in low:
                return (2, n)
            if "_scatter" in low:
                return (3, n)
            if "_hist" in low:
                return (4, n)
            if "_bar" in low:
                return (5, n)
            if "regression_" in low:
                return (6, n)
            if "rf_" in low:
                return (7, n)
            if "logistic" in low:
                return (8, n)
            return (9, n)

        allowed_names.sort(key=order_key_dataset)

        seen = set()
        unique_names = []
        for n in allowed_names:
            if n not in seen:
                unique_names.append(n)
                seen.add(n)

        valid_names = []
        for n in unique_names:
            full_path = os.path.join(PLOTS_DIR, n)
            if os.path.exists(full_path):
                valid_names.append(n)

        plots = [os.path.join("plots", n).replace("\\", "/") for n in valid_names]
        eda_plots = [p for p in plots if not _is_model_plot(p)]
        model_plots = [p for p in plots if _is_model_plot(p)]

        # =========================
        # Resultados de modelos
        # =========================
        model_results = manifest_data.get("model_results", {})
        if not isinstance(model_results, dict):
            model_results = {}

        regression_result = _safe_dict(model_results.get("regression_result"))
        rf_result = _safe_dict(model_results.get("rf_result"))

        try:
            model_comparison = build_model_comparison_summary(
                regression_result=regression_result,
                rf_result=rf_result,
            ) or {"available": False}
        except Exception as e:
            current_app.logger.warning(f"[dataset_articulo_word:model_comparison] ds{dataset_id}: {e}")
            model_comparison = {"available": False}

        try:
            if analysis_meta.get("target_type") == "classification":
                advanced_model_interpretation = build_advanced_classification_interpretation(
                    rf_result=rf_result,
                ) or []
            else:
                advanced_model_interpretation = build_advanced_model_interpretation(
                    regression_result=regression_result,
                    rf_result=rf_result,
                ) or []
        except Exception as e:
            current_app.logger.warning(f"[dataset_articulo_word:advanced_model_interpretation] ds{dataset_id}: {e}")
            advanced_model_interpretation = []

        try:
            next_step_recommendation = build_next_step_recommendation(
                analysis_meta=analysis_meta,
                dataset_kind=kind,
                model_plots=model_plots,
            )
        except Exception as e:
            current_app.logger.warning(f"[dataset_articulo_word:next_step] ds{dataset_id}: {e}")
            next_step_recommendation = None

        quick_recommendations = analysis_meta.get("quick_recommendations", []) if isinstance(analysis_meta, dict) else []
        insights = analysis_meta.get("insights", []) if isinstance(analysis_meta, dict) else []
        insights_text = analysis_meta.get("insights_text") if isinstance(analysis_meta, dict) else None

        # =========================
        # Enriquecer meta para el texto del artículo
        # =========================
        article_meta = dict(analysis_meta or {})
        article_meta["quick_recommendations"] = quick_recommendations
        article_meta["suggested_plan"] = suggested_plan
        article_meta["next_step_recommendation"] = next_step_recommendation
        article_meta["advanced_model_interpretation"] = advanced_model_interpretation
        article_meta["model_comparison"] = model_comparison
        article_meta["best_model_selection"] = best_model_selection
        article_meta["regression_result"] = regression_result
        article_meta["rf_result"] = rf_result
        article_meta["research_area"] = research_area
        article_meta["research_area_suggested"] = suggested_area

        reports_dir = os.path.join(PLOTS_DIR, "reports")
        os.makedirs(reports_dir, exist_ok=True)
        docx_path = os.path.join(reports_dir, f"ds{dataset_id}_articulo_cientifico_dataset.docx")

        # =========================
        # Generar texto del artículo
        # =========================
        article_text = generate_scientific_article_dataset(
            article_meta,
            ds.title or ds.original_name
        )

        # =========================
        # Crear DOCX
        # =========================
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

        styles = doc.styles
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(11)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Análisis exploratorio automatizado del dataset '{ds.title or ds.original_name}' mediante SADI")
        run.bold = True
        run.font.size = Pt(17)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Documento preliminar para revisión y edición académica")
        run2.italic = True
        run2.font.size = Pt(11)

        doc.add_paragraph("")

        for block in article_text.split("\n\n"):
            block = block.strip()
            if block:
                upper = block.strip().upper()
                if upper in {
                    "TÍTULO",
                    "RESUMEN",
                    "1. INTRODUCCIÓN",
                    "2. METODOLOGÍA",
                    "3. RESULTADOS",
                    "4. INSIGHTS AUTOMÁTICOS",
                    "5. INTERPRETACIÓN AUTOMÁTICA",
                    "6. RECOMENDACIONES",
                    "7. PLAN DE ANÁLISIS SUGERIDO",
                    "8. DISCUSIÓN",
                    "9. CONCLUSIÓN",
                    "10. REFERENCIAS",
                }:
                    doc.add_heading(block, level=1)
                else:
                    doc.add_paragraph(block)

        # =========================
        # Anexo de gráficos exploratorios
        # =========================
        if eda_plots:
            doc.add_page_break()
            doc.add_heading("Anexo A. Gráficos exploratorios", level=1)

            for i, rel in enumerate(eda_plots, start=1):
                rel = str(rel).replace("\\", "/")
                fname = rel.split("/")[-1]
                img_path = os.path.join(PLOTS_DIR, fname)
                if not os.path.exists(img_path):
                    continue

                nice_title = prettify_plot_title(rel)
                nice_desc = describe_plot(rel)
                nice_tag = classify_plot_tag(rel)

                p_title = doc.add_paragraph()
                p_title.paragraph_format.keep_with_next = True
                r = p_title.add_run(f"Figura A{i}. {nice_title}")
                r.bold = True
                r.font.size = Pt(12)

                p_type = doc.add_paragraph()
                p_type.paragraph_format.keep_with_next = True
                r2 = p_type.add_run(f"Tipo: {nice_tag}")
                r2.italic = True

                if nice_desc:
                    p_desc = doc.add_paragraph(nice_desc)
                    p_desc.paragraph_format.keep_with_next = True

                inserted = safe_add_picture(doc, img_path, width_inches=6.0)
                if inserted:
                    doc.add_paragraph("")

        # =========================
        # Anexo B. Modelos predictivos + gráficos
        # =========================
        if model_plots:
            doc.add_page_break()
            doc.add_heading("Anexo B. Modelos predictivos", level=1)

            target_candidate = analysis_meta.get("target_candidate", "No detectado") if isinstance(analysis_meta, dict) else "No detectado"
            target_type = analysis_meta.get("target_type", "No definido") if isinstance(analysis_meta, dict) else "No definido"
            target_reason = analysis_meta.get("target_reason", "") if isinstance(analysis_meta, dict) else ""
            model_suggestion = analysis_meta.get("model_suggestion", "No sugerido") if isinstance(analysis_meta, dict) else "No sugerido"

            doc.add_heading("Contexto del modelado", level=2)
            doc.add_paragraph(f"Variable objetivo detectada: {target_candidate}")
            doc.add_paragraph(f"Tipo de problema: {target_type}")
            doc.add_paragraph(f"Modelo sugerido por SADI: {model_suggestion}")
            doc.add_paragraph(f"Justificación automática: {target_reason if target_reason else 'No disponible.'}")

            doc.add_heading("Comparación automática de modelos", level=2)
            if isinstance(model_comparison, dict) and model_comparison.get("available"):
                _add_model_comparison_table(doc, model_comparison)

                winner_txt = _winner_to_text(model_comparison.get("winner"))
                doc.add_paragraph(f"Modelo recomendado por SADI: {winner_txt}")
                doc.add_paragraph(
                    f"Interpretación: {_safe_text(model_comparison.get('recommendation'), 'No disponible.')}"
                )

                summary_lines = _safe_list(model_comparison.get("summary"))
                if summary_lines:
                    for line in summary_lines:
                        doc.add_paragraph(f"• {line}")
                else:
                    doc.add_paragraph("No se generó un resumen textual de la comparación.")
            else:
                doc.add_paragraph("No se pudo generar comparación automática de modelos.")

            doc.add_heading("Evaluación automática del modelo", level=2)
            if advanced_model_interpretation:
                for item in _safe_list(advanced_model_interpretation):
                    doc.add_paragraph(f"• {item}")
            else:
                doc.add_paragraph("• No se generó evaluación automática.")

            doc.add_heading("Selección del mejor modelo", level=2)
            if isinstance(best_model_selection, dict) and best_model_selection:
                best_model_txt = _winner_to_text(best_model_selection.get("best_model"))
                doc.add_paragraph(f"Modelo preferido: {best_model_txt}")
                doc.add_paragraph(f"Razón: {_safe_text(best_model_selection.get('reason'))}")

                for line in _safe_list(best_model_selection.get("comparison_summary")):
                    doc.add_paragraph(f"• {line}")
            else:
                doc.add_paragraph("No se pudo determinar un mejor modelo.")

            doc.add_paragraph("")

            for i, rel in enumerate(model_plots, start=1):
                rel = str(rel).replace("\\", "/")
                fname = rel.split("/")[-1]
                img_path = os.path.join(PLOTS_DIR, fname)
                if not os.path.exists(img_path):
                    continue

                nice_title = prettify_plot_title(rel)
                nice_desc = describe_plot(rel)

                p_title = doc.add_paragraph()
                p_title.paragraph_format.keep_with_next = True
                r = p_title.add_run(f"Figura B{i}. {nice_title}")
                r.bold = True
                r.font.size = Pt(12)

                p_type = doc.add_paragraph()
                p_type.paragraph_format.keep_with_next = True
                r2 = p_type.add_run("Tipo: Modelo predictivo")
                r2.italic = True

                if nice_desc:
                    p_desc = doc.add_paragraph(nice_desc)
                    p_desc.paragraph_format.keep_with_next = True

                inserted = safe_add_picture(doc, img_path, width_inches=6.0)
                if inserted:
                    doc.add_paragraph("")

        doc.save(docx_path)

        return send_file(
            docx_path,
            as_attachment=True,
            download_name=f"articulo_cientifico_dataset_{dataset_id}.docx"
        )

    except Exception as e:
        current_app.logger.exception(f"[dataset_articulo_word] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el artículo científico en Word: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    finally:
        db.close()

@dataset_bp.post("/datasets/<int:dataset_id>/set_research_area", endpoint="dataset_set_research_area")
@login_required
def dataset_set_research_area(dataset_id: int):
    from .constans import RESEARCH_AREA_LABELS, RESEARCH_AREA_VALUES

    research_area = (request.form.get("research_area") or "general").strip()
    if research_area not in RESEARCH_AREA_VALUES:
        research_area = "general"

    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            if request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return {"ok": False, "error": "Dataset no encontrado."}, 404
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

        ds.research_area = research_area
        db.commit()

    label = RESEARCH_AREA_LABELS.get(research_area, research_area)

    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        return {
            "ok": True,
            "research_area": research_area,
            "label": label,
            "badge_class": "badge-info",
        }

    flash("Área de investigación actualizada.", "success")
    return redirect(url_for("dataset.dashboard"))
@dataset_bp.get("/datasets/<int:dataset_id>/analyze_auto", endpoint="dataset_analyze_auto")
@login_required
def dataset_analyze_auto(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    kind = (
        getattr(ds, "dataset_type", None)
        or getattr(ds, "dataset_kind", None)
        or "dataset"
    )
    kind = str(kind).strip().lower()

    if kind in ("survey_likert_5", "survey_likert_7"):
        return redirect(url_for("likert.dataset_likert_analysis", dataset_id=dataset_id))

    if kind == "multivariate":
        return redirect(url_for("multivariate.multivariate_detail", dataset_id=dataset_id))

    if kind in ("dataset", "survey_normal"):
        return redirect(url_for("dataset.dataset_stats", dataset_id=dataset_id))

    flash(f"Tipo de dataset no reconocido: {kind}", "warning")
    return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))