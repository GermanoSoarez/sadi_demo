from __future__ import annotations

import glob
import os

from flask import (
    Blueprint,
    flash,
    redirect,
    send_file,
    url_for,
    current_app,
)
from flask_login import login_required, current_user

import pandas as pd

# 🔹 IMPORTAR lector de dataset
from blueprints.dataset.analysis import read_dataframe


from extensions import SessionLocal
from models import Dataset
from config import PLOTS_DIR, UPLOAD_DIR
from utils.manifest import read_manifest_data, write_manifest
from utils.plot_manager import normalize_plot_catalog, split_plots
# === IMPORTA AQUÍ LAS FUNCIONES QUE YA TIENES EN sadi.py ===
# Muévelas luego a blueprints/survey/analysis.py si quieres ordenar mejor.
from blueprints.dataset.analysis import analyze_dataset_with_recommendations, build_next_step_recommendation
from blueprints.survey.analysis import (
    read_dataframe,
    repair_broken_csv_file,
    detect_best_segment_column,
    summarize_column,
    numeric_group_tests,
    make_crosstab_tests,
    build_group_comparisons,
    build_group_comparison_summary,
    build_group_findings,
    cramers_v_strength,
    cramers_v_class,
    detect_key_variables,
    generate_results_text_survey_normal,
    build_survey_insights,
    plot_survey_normal_question,
    generate_group_comparison_plots,
    filter_plots_for_dataset,
    build_survey_key_findings,
    generate_survey_normal_report_pdf,
    build_academic_report_context,
    summarize_plot_tags,
    prettify_plot_title,
    describe_plot,
    classify_plot_tag,
    safe_text,
    safe_add_picture,
)
    
survey_bp = Blueprint("survey", __name__) 

def append_survey_sadi_sections(doc, report_ctx: dict, *, start_num: int = 14):
    n = start_num

    if report_ctx.get("sadi_insights"):
        doc.add_heading(f"{n}. Insights automáticos SADI", level=1)
        for line in str(report_ctx["sadi_insights"]).split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        n += 1

    if report_ctx.get("sadi_recommendations"):
        doc.add_heading(f"{n}. Recomendaciones", level=1)
        for rec in report_ctx["sadi_recommendations"]:
            doc.add_paragraph(f"• {rec}")
        n += 1

    if report_ctx.get("sadi_plan"):
        doc.add_heading(f"{n}. Plan de análisis sugerido", level=1)
        plan = report_ctx["sadi_plan"]
        if isinstance(plan, dict):
            for k, v in plan.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            doc.add_paragraph(str(plan))
        n += 1

    if report_ctx.get("sadi_priority"):
        doc.add_heading(f"{n}. Prioridad analítica sugerida", level=1)
        for step in report_ctx["sadi_priority"]:
            doc.add_paragraph(f"• {step}")
        n += 1

    if report_ctx.get("top_numeric_by_variability"):
        doc.add_heading(f"{n}. Variables con mayor variabilidad", level=1)
        for item in report_ctx["top_numeric_by_variability"][:10]:
            if isinstance(item, dict):
                doc.add_paragraph(
                    f"• {item.get('column', '—')} | std={item.get('std', '—')} | rango={item.get('range', '—')}"
                )
        n += 1

    if report_ctx.get("variable_importance"):
        doc.add_heading(f"{n}. Variables clave del dataset", level=1)
        for item in report_ctx["variable_importance"][:10]:
            if isinstance(item, dict):
                doc.add_paragraph(
                    f"• {item.get('column', '—')} | score={item.get('score', '—')} | missing={item.get('missing_pct', '—')}%"
                )
        n += 1

    if report_ctx.get("target_candidate"):
        doc.add_heading(f"{n}. Variable objetivo sugerida", level=1)
        doc.add_paragraph(f"Variable: {report_ctx.get('target_candidate')}")
        doc.add_paragraph(f"Tipo: {report_ctx.get('target_type')}")
        if report_ctx.get("target_reason"):
            doc.add_paragraph(f"Justificación: {report_ctx.get('target_reason')}")
        if report_ctx.get("model_suggestion"):
            doc.add_paragraph(f"Modelo sugerido: {report_ctx.get('model_suggestion')}")
        n += 1

    if report_ctx.get("ranked_target_candidates"):
        doc.add_heading(f"{n}. Candidatos a variable objetivo", level=1)
        for item in report_ctx["ranked_target_candidates"][:10]:
            if isinstance(item, dict):
                doc.add_paragraph(
                    f"• {item.get('column', '—')} | tipo={item.get('type', '—')} | motivo={item.get('reason', '—')}"
                )
        n += 1

    if report_ctx.get("next_step_recommendation"):
        doc.add_heading(f"{n}. Próximo paso recomendado", level=1)
        step = report_ctx["next_step_recommendation"]
        if isinstance(step, dict):
            for k, v in step.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            doc.add_paragraph(str(step))
        n += 1

    if report_ctx.get("warnings"):
        doc.add_heading(f"{n}. Advertencias", level=1)
        for w in report_ctx["warnings"]:
            doc.add_paragraph(f"⚠ {w}")

@survey_bp.post("/datasets/<int:dataset_id>/dataset_survey_normal_summary", endpoint="dataset_survey_normal_summary")
@login_required
def dataset_survey_normal_summary(dataset_id: int):
    lines: list[str] = []

    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

        filename = ds.filename
        delimiter = ds.delimiter
        file_path = os.path.join(UPLOAD_DIR, filename)

    # =========================
    # LECTURA / REPARACIÓN
    # =========================
    try:
        df = read_dataframe(file_path, delimiter)
    except Exception as e:
        flash(f"No se pudo leer dataset: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    if df.shape[1] == 1:
        only_col = str(df.columns[0]) if len(df.columns) > 0 else ""
        if any(x in only_col for x in [",", ";", "\t", "|", ":"]) or df.shape[0] == 0:
            repair_info = repair_broken_csv_file(file_path, delimiter)

            if repair_info.get("ok") and repair_info.get("repaired"):
                try:
                    df = read_dataframe(file_path, delimiter)
                    flash(
                        "El archivo CSV tenía un formato irregular y fue reparado automáticamente.",
                        "info"
                    )
                except Exception as e:
                    flash(f"Se intentó reparar el archivo, pero aún no pudo leerse: {e}", "danger")
                    return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))
            else:
                flash(
                    repair_info.get("message") or "El archivo se leyó como una sola columna. Revisa el delimitador del dataset.",
                    "danger"
                )
                return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    current_app.logger.warning(
        f"[survey_normal_summary] ds{dataset_id} shape={df.shape} columns={list(df.columns)} delimiter={delimiter!r}"
    )

    if df.shape[1] == 1:
        only_col = str(df.columns[0])
        if any(x in only_col for x in [",", ";", "\t", "|", ":"]):
            flash(
                "El archivo se leyó como una sola columna. Revisa el delimitador del dataset.",
                "danger"
            )
            return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    # =========================
    # META BASE
    # =========================
    meta = {}
    meta["n_total"] = int(df.shape[0])
    meta["n_rows"] = int(df.shape[0])
    meta["n_cols"] = int(df.shape[1])

    # tipos básicos
    num_cols = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c])]
    dt_cols = [c for c in df.columns if pd.api.types.is_datetime64_any_dtype(df[c])]
    cat_cols = [c for c in df.columns if c not in num_cols and c not in dt_cols]

    meta["num_cols"] = num_cols
    meta["cat_cols"] = cat_cols
    meta["dt_cols"] = dt_cols
    meta["n_num"] = len(num_cols)
    meta["n_cat"] = len(cat_cols)
    meta["n_dt"] = len(dt_cols)

    total_cells = int(df.shape[0] * df.shape[1]) if df.shape[0] and df.shape[1] else 0
    missing_total = int(df.isna().sum().sum()) if total_cells else 0
    meta["missing_global_pct"] = round((missing_total / total_cells) * 100.0, 2) if total_cells else 0.0

    # segmentación
    segment_col = detect_best_segment_column(df)
    meta["segment_column"] = segment_col

    # resumen por columna
    meta["columns_summary"] = [summarize_column(df[col]) for col in df.columns]

    # =========================
    # ANÁLISIS POR GRUPOS
    # =========================
    if segment_col:
        meta["numeric_tests"] = numeric_group_tests(df, segment_col)
        meta["crosstabs"] = make_crosstab_tests(df, segment_col)
        meta["group_comparisons"] = build_group_comparisons(df, segment_col)
        meta["group_comparison_summary"] = build_group_comparison_summary(meta["group_comparisons"])
        meta["group_findings"] = build_group_findings(meta["group_comparisons"])
    else:
        meta["numeric_tests"] = []
        meta["crosstabs"] = []
        meta["group_comparisons"] = []
        meta["group_comparison_summary"] = ""
        meta["group_findings"] = []

    # =========================
    # CROSSTABS ENRIQUECIDOS
    # =========================
    crosstabs = meta.get("crosstabs") or []
    if isinstance(crosstabs, list) and crosstabs:
        enriched = []
        for r in crosstabs:
            if not isinstance(r, dict):
                continue

            p = r.get("p_value")
            v = r.get("cramers_v")

            sig = bool(isinstance(p, (int, float)) and p < 0.05)
            strength = cramers_v_strength(v)
            css = cramers_v_class(v)

            rr = dict(r)
            rr["is_significant"] = sig
            rr["strength"] = strength
            rr["css_class"] = css
            enriched.append(rr)

        def sort_key(x):
            p = x.get("p_value")
            v = x.get("cramers_v")
            sig = 1 if x.get("is_significant") else 0
            v = float(v) if isinstance(v, (int, float)) else -1.0
            p = float(p) if isinstance(p, (int, float)) else 1e9
            return (-sig, -v, p)

        enriched.sort(key=sort_key)
        meta["crosstabs"] = enriched

        top = [
            x for x in enriched
            if x.get("is_significant") and isinstance(x.get("cramers_v"), (int, float))
        ][:8]

        if top:
            lines.append("Principales asociaciones detectadas (p < 0.05), ordenadas por magnitud:")
            for x in top:
                var = x.get("variable") or ""
                v = x.get("cramers_v")
                p = x.get("p_value")
                st = x.get("strength") or ""
                lines.append(f"• {var}: Cramér’s V = {v:.3f} ({st}), p = {p:.6f}")
            meta["crosstabs_insights"] = "\n".join(lines)
        else:
            meta["crosstabs_insights"] = (
                "No se detectaron asociaciones significativas (p < 0.05) "
                "en las combinaciones evaluadas."
            )
    else:
        meta["crosstabs_insights"] = ""

    # =========================
    # PERFIL SURVEY
    # =========================
    meta["survey_profile"] = {
        "n_total": meta.get("n_total"),
        "segment_column": meta.get("segment_column"),
        "n_variables": int(df.shape[1]),
        "n_categorical": len([
            c for c in meta.get("columns_summary", [])
            if isinstance(c, dict) and c.get("dtype") == "categorical"
        ]),
        "n_numeric": len([
            c for c in meta.get("columns_summary", [])
            if isinstance(c, dict) and c.get("dtype") == "numeric"
        ]),
    }

    # =========================
    # TEXTO BASE SURVEY
    # =========================
    meta["important_variables"] = detect_key_variables(meta)
    meta["results_text"] = generate_results_text_survey_normal(meta)
    meta["survey_insights"] = build_survey_insights(meta)
    meta["survey_key_findings"] = build_survey_key_findings(meta)

    # =========================
    # LIMPIAR GRÁFICOS ANTIGUOS
    # =========================
    try:
        for old_plot in glob.glob(os.path.join(PLOTS_DIR, f"ds{dataset_id}_*.png")):
            try:
                os.remove(old_plot)
            except Exception:
                pass
    except Exception:
        pass

    # =========================
    # GENERAR GRÁFICOS
    # =========================
    plot_files = []

    for col in df.columns:
        fname = plot_survey_normal_question(df[col], dataset_id, PLOTS_DIR)
        if fname:
            plot_files.append(fname)

    if segment_col and meta.get("group_comparisons"):
        comparison_plots = generate_group_comparison_plots(
            df=df,
            dataset_id=dataset_id,
            segment_col=segment_col,
            group_comparisons=meta.get("group_comparisons"),
            plots_dir=PLOTS_DIR,
            max_plots=8,
        )
        if comparison_plots:
            plot_files.extend(comparison_plots)

    plot_files = filter_plots_for_dataset(dataset_id, plot_files)

    valid_plot_files = []
    for p in plot_files:
        rel = str(p).replace("\\", "/").strip()
        fname = rel.split("/")[-1]
        full_path = os.path.join(PLOTS_DIR, fname)
        if os.path.exists(full_path):
            valid_plot_files.append(rel)

        # =========================
    # CAPA SADI FINAL
    # =========================

    # Base survey
    base_insights_text = meta.get("survey_insights") or meta.get("results_text") or ""
    base_insights_list = meta.get("group_findings") or meta.get("survey_key_findings") or []

    # Análisis SADI avanzado
    try:
        sadi_analysis = analyze_dataset_with_recommendations(
            df,
            dataset_type="survey_normal",
            research_area="general",
        ) or {}
    except Exception as e:
        current_app.logger.warning(f"[survey:sadi_analysis] ds{dataset_id}: {e}")
        sadi_analysis = {}

    if not isinstance(sadi_analysis, dict):
        sadi_analysis = {}

    # =========================
    # INSIGHTS
    # =========================
    # Estos son los que usa la pantalla como "avanzados"
    advanced_insights_text = sadi_analysis.get("insights_text") or ""
    advanced_insights_list = sadi_analysis.get("insights") or []

        # Guardar ambos niveles
    meta["survey_insights"] = base_insights_text
    meta["insights"] = advanced_insights_list or base_insights_list
    meta["insights_text"] = advanced_insights_text or base_insights_text

    # Construir insight avanzado más rico
    advanced_parts = []

    if advanced_insights_text:
        advanced_parts.append(advanced_insights_text)

    quick_recs = sadi_analysis.get("quick_recommendations", []) if isinstance(sadi_analysis, dict) else []
    if quick_recs:
        advanced_parts.append(
            "Recomendaciones principales: " + "; ".join(map(str, quick_recs[:5])) + "."
        )

    priority_order = sadi_analysis.get("priority_order", []) if isinstance(sadi_analysis, dict) else []
    if priority_order:
        advanced_parts.append(
            "Prioridad analítica sugerida: " + " > ".join(map(str, priority_order[:6])) + "."
        )

    suggested_plan = sadi_analysis.get("suggested_plan", {}) if isinstance(sadi_analysis, dict) else {}
    if isinstance(suggested_plan, dict) and suggested_plan:
        recommended_analysis = suggested_plan.get("recommended_analysis") or []
        recommended_plots = suggested_plan.get("recommended_plots") or []
        narrative_focus = suggested_plan.get("narrative_focus")

        if recommended_analysis:
            advanced_parts.append(
                "Análisis sugeridos: " + "; ".join(map(str, recommended_analysis[:6])) + "."
            )

        if recommended_plots:
            advanced_parts.append(
                "Gráficos sugeridos: " + "; ".join(map(str, recommended_plots[:6])) + "."
            )

        if narrative_focus:
            advanced_parts.append(f"Enfoque narrativo recomendado: {narrative_focus}.")

    # Se completa después con next_step si existe
    meta["sadi_insights"] = "\n\n".join([p for p in advanced_parts if p]) or (advanced_insights_text or base_insights_text)

    # =========================
    # RECOMENDACIONES / PLAN
    # =========================
    meta["quick_recommendations"] = sadi_analysis.get("quick_recommendations", []) or []
    meta["suggested_plan"] = sadi_analysis.get("suggested_plan", {}) or {}
    meta["priority_order"] = sadi_analysis.get("priority_order", []) or []
    meta["warnings"] = sadi_analysis.get("warnings", []) or []
    meta["research_area_suggested"] = sadi_analysis.get("research_area_suggested")

    meta["sadi_recommendations"] = meta.get("quick_recommendations", [])
    meta["sadi_plan"] = meta.get("suggested_plan", {})
    meta["sadi_priority"] = meta.get("priority_order", [])

    # =========================
    # PRÓXIMO PASO
    # =========================
    try:
        meta["next_step_recommendation"] = build_next_step_recommendation(
            analysis_meta=meta,
            dataset_kind="survey_normal",
            model_plots=valid_plot_files,
        )
    
        if meta.get("next_step_recommendation"):
            step = meta["next_step_recommendation"]
            if isinstance(step, dict):
                step_title = step.get("title")
                step_reason = step.get("reason")
                step_action = step.get("action")

                step_lines = []
                if step_title:
                    step_lines.append(f"Próximo paso recomendado: {step_title}.")
                if step_reason:
                    step_lines.append(f"Justificación: {step_reason}.")
                if step_action:
                    step_lines.append(f"Acción sugerida: {step_action}.")

                if step_lines:
                    meta["sadi_insights"] = (
                        (meta.get("sadi_insights") or "").strip() + "\n\n" + " ".join(step_lines)
                    ).strip()

    except Exception as e:
        current_app.logger.warning(f"[survey:next_step] ds{dataset_id}: {e}")
        meta["next_step_recommendation"] = None

    # =========================
    # VARIABLES MÁS VARIABLES
    # =========================
    try:
        top_var = []
        for col in num_cols:
            ser = pd.to_numeric(df[col], errors="coerce").dropna()
            if ser.empty:
                continue
            top_var.append({
                "column": col,
                "std": round(float(ser.std()), 4) if len(ser) > 1 else 0.0,
                "range": round(float(ser.max() - ser.min()), 4),
            })
        top_var.sort(key=lambda x: (x["std"], x["range"]), reverse=True)
        meta["top_numeric_by_variability"] = top_var[:10]
    except Exception as e:
        current_app.logger.warning(f"[survey:variability] ds{dataset_id}: {e}")
        meta["top_numeric_by_variability"] = []

    # =========================
    # VARIABLES CLAVE
    # =========================
    try:
        variable_importance = []
        for item in meta.get("columns_summary", []):
            if not isinstance(item, dict):
                continue
            numeric = item.get("numeric") or {}
            variable_importance.append({
                "column": item.get("name", "—"),
                "score": item.get("n_valid", 0),
                "std": numeric.get("std", "—") if isinstance(numeric, dict) else "—",
                "missing_pct": item.get("missing_pct", "—"),
            })

        variable_importance.sort(
            key=lambda x: (
                x["score"] if isinstance(x["score"], (int, float)) else -1,
                -(x["missing_pct"] if isinstance(x["missing_pct"], (int, float)) else 9999)
            ),
            reverse=True
        )
        meta["variable_importance"] = variable_importance[:10]
    except Exception as e:
        current_app.logger.warning(f"[survey:variable_importance] ds{dataset_id}: {e}")
        meta["variable_importance"] = []

    # =========================
    # VARIABLE OBJETIVO / MODELO SUGERIDO
    # =========================
    try:
        ranked = []
        for idx, var in enumerate(meta.get("important_variables", [])[:10], start=1):
            ranked.append({
                "column": var,
                "type": "classification" if segment_col and var != segment_col else "regression",
                "score": round(1.0 - ((idx - 1) * 0.05), 3),
                "reason": "Variable destacada automáticamente por SADI a partir del análisis descriptivo.",
            })

        meta["ranked_target_candidates"] = ranked

        if ranked:
            meta["target_candidate"] = ranked[0]["column"]
            meta["target_type"] = ranked[0]["type"]
            meta["target_reason"] = ranked[0]["reason"]
            meta["model_suggestion"] = (
                "Random Forest Classifier"
                if ranked[0]["type"] == "classification"
                else "Random Forest Regressor"
            )
        else:
            fallback_target = None
            if meta.get("important_variables"):
                fallback_target = meta["important_variables"][0]

            meta["target_candidate"] = fallback_target
            meta["target_type"] = "classification" if segment_col else "regression"
            meta["target_reason"] = (
                "Variable sugerida por fallback de SADI a partir de las variables más relevantes del análisis."
                if fallback_target else None
            )
            meta["model_suggestion"] = (
                "Random Forest Classifier"
                if segment_col else "Random Forest Regressor"
            ) if fallback_target else None
    except Exception as e:
        current_app.logger.warning(f"[survey:target_candidate] ds{dataset_id}: {e}")
        meta["ranked_target_candidates"] = []
        meta["target_candidate"] = None
        meta["target_type"] = None
        meta["target_reason"] = None
        meta["model_suggestion"] = None

    current_app.logger.warning(
        f"[SURVEY FINAL CHECK] ds{dataset_id} "
        f"survey_insights={bool(meta.get('survey_insights'))} "
        f"sadi_insights={bool(meta.get('sadi_insights'))} "
        f"quick_recommendations={len(meta.get('quick_recommendations') or [])} "
        f"sadi_plan={bool(meta.get('sadi_plan'))} "
        f"sadi_priority={len(meta.get('sadi_priority') or [])} "
        f"next_step={bool(meta.get('next_step_recommendation'))}"
    )

    # =========================
    # GUARDAR MANIFEST
    # =========================
    try:
        manifest_data = read_manifest_data(dataset_id) or {}
        if not isinstance(manifest_data, dict):
            manifest_data = {}

        prev_generated = manifest_data.get("generated", []) or []
        prev_plots = manifest_data.get("plots", []) or []

        merged_plots = []
        seen = set()

        for p in list(prev_generated) + list(prev_plots) + list(valid_plot_files):
            if isinstance(p, str):
                rel = str(p).replace("\\", "/").strip()
                if rel and rel not in seen:
                    merged_plots.append(rel)
                    seen.add(rel)

        manifest_data["generated"] = merged_plots
        manifest_data["plots"] = merged_plots
        manifest_data["meta"] = meta
        manifest_data["dataset_type"] = "survey_normal"
        manifest_data["cols"] = list(df.columns)

        write_manifest(dataset_id, manifest_data)

    except Exception as e:
        current_app.logger.warning(f"[survey:save_manifest] ds{dataset_id}: {e}")

    flash("Encuesta normal analizada correctamente.", "success")
    return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))


@survey_bp.get("/datasets/<int:dataset_id>/survey_normal_report.pdf", endpoint="dataset_survey_normal_report")
@login_required
def dataset_survey_normal_report(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    manifest_data = read_manifest_data(dataset_id)
    out_path = os.path.join(PLOTS_DIR, f"ds{dataset_id}_survey_normal_report.pdf")

    try:
        generate_survey_normal_report_pdf(
            dataset_id=dataset_id,
            dataset_title=(ds.title or ds.original_name or f"Dataset {dataset_id}"),
            manifest_data=manifest_data,
            plots_dir=PLOTS_DIR,
            output_path=out_path,
        )
    except Exception as e:
        current_app.logger.exception(f"[survey_normal_report] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el PDF de encuesta: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    if not os.path.exists(out_path) or os.path.getsize(out_path) < 1000:
        flash("Se intentó generar el PDF pero quedó vacío o no se creó.", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    return send_file(
        out_path,
        as_attachment=True,
        download_name=f"SADI_encuesta_ds{dataset_id}.pdf"
    )


@survey_bp.get("/datasets/<int:dataset_id>/survey_normal_word", endpoint="dataset_survey_normal_word")
@login_required
def dataset_survey_normal_word(dataset_id: int):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    kind = (getattr(ds, "dataset_type", None) or "no_definido").strip()
    if kind != "survey_normal":
        flash("Este informe Word aplica solo para 'Encuesta normal'.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    manifest_data = read_manifest_data(dataset_id) or {}
    meta = manifest_data.get("meta") or {}
    plots = manifest_data.get("generated") or manifest_data.get("plots") or []

    if not isinstance(meta, dict):
        meta = {}
    if not isinstance(plots, list):
        plots = []

    if not meta:
        flash("Primero genera el análisis descriptivo de encuesta.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    reports_dir = os.path.join(PLOTS_DIR, "reports")
    os.makedirs(reports_dir, exist_ok=True)
    docx_path = os.path.join(reports_dir, f"ds{dataset_id}_encuesta_normal.docx")

    try:
        plots = normalize_plot_catalog(dataset_id, manifest_data, meta)
        report_ctx = build_academic_report_context(ds, meta, plots)

        advanced_insights = (
            report_ctx.get("sadi_insights")
            or report_ctx.get("insights_text")
            or meta.get("sadi_insights")
            or meta.get("insights_text")
            or report_ctx.get("survey_insights")
            or meta.get("survey_insights")
            or ""
        )

        recommendations = (
            report_ctx.get("sadi_recommendations")
            or report_ctx.get("quick_recommendations")
            or meta.get("sadi_recommendations")
            or meta.get("quick_recommendations")
            or []
        )
        if not isinstance(recommendations, list):
            recommendations = [recommendations] if recommendations else []

        plan = (
            report_ctx.get("sadi_plan")
            or report_ctx.get("suggested_plan")
            or meta.get("sadi_plan")
            or meta.get("suggested_plan")
            or {}
        )
        if not isinstance(plan, dict):
            plan = {}

        priority = (
            report_ctx.get("sadi_priority")
            or report_ctx.get("priority_order")
            or meta.get("sadi_priority")
            or meta.get("priority_order")
            or []
        )
        if not isinstance(priority, list):
            priority = [priority] if priority else []

        next_step = (
            report_ctx.get("next_step_recommendation")
            or meta.get("next_step_recommendation")
            or {}
        )
        if not isinstance(next_step, dict):
            next_step = {}

        variable_importance = (
            report_ctx.get("variable_importance")
            or meta.get("variable_importance")
            or []
        )
        if not isinstance(variable_importance, list):
            variable_importance = []

        top_variability = (
            report_ctx.get("top_numeric_by_variability")
            or meta.get("top_numeric_by_variability")
            or []
        )
        if not isinstance(top_variability, list):
            top_variability = []

        target_candidate = (
            report_ctx.get("target_candidate")
            or meta.get("target_candidate")
        )
        target_type = (
            report_ctx.get("target_type")
            or meta.get("target_type")
        )
        target_reason = (
            report_ctx.get("target_reason")
            or meta.get("target_reason")
        )
        model_suggestion = (
            report_ctx.get("model_suggestion")
            or meta.get("model_suggestion")
        )

        ranked_target_candidates = (
            report_ctx.get("ranked_target_candidates")
            or meta.get("ranked_target_candidates")
            or []
        )
        if not isinstance(ranked_target_candidates, list):
            ranked_target_candidates = []

        warnings_list = (
            report_ctx.get("warnings")
            or meta.get("warnings")
            or []
        )
        if not isinstance(warnings_list, list):
            warnings_list = [warnings_list] if warnings_list else []

        survey_profile = report_ctx.get("survey_profile") or {}
        if not isinstance(survey_profile, dict):
            survey_profile = {}

        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        try:
            doc.styles["Normal"].font.name = "Calibri"
            doc.styles["Normal"].font.size = Pt(10.5)
        except Exception:
            pass

        def set_cell_text(cell, text, bold=False):
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run("" if text is None else str(text))
            r.bold = bold

        # =========================
        # PORTADA
        # =========================
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run(report_ctx.get("title") or "Informe académico — Encuesta normal")
        r.bold = True
        r.font.size = Pt(18)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = subtitle.add_run(report_ctx.get("dataset_name") or "")
        r2.italic = True
        r2.font.size = Pt(12)

        doc.add_paragraph("")

        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = "Table Grid"
        info_rows = [
            ("Dataset", report_ctx.get("dataset_name")),
            ("Archivo", report_ctx.get("file_name")),
            ("Fecha de generación", report_ctx.get("generated_at")),
            ("Tipo detectado", report_ctx.get("dataset_type")),
            ("Total de participantes/registros", report_ctx.get("n_total")),
        ]
        for i, (k, v) in enumerate(info_rows):
            set_cell_text(info_table.rows[i].cells[0], k, bold=True)
            set_cell_text(info_table.rows[i].cells[1], v)

        # =========================
        # 1. INTRODUCCIÓN
        # =========================
        doc.add_page_break()
        doc.add_heading("1. Introducción", 1)
        doc.add_paragraph(report_ctx.get("intro_text") or "")

        # =========================
        # 2. PERFIL DE LA ENCUESTA
        # =========================
        doc.add_heading("2. Perfil de la encuesta", 1)
        t_profile = doc.add_table(rows=1, cols=2)
        t_profile.style = "Table Grid"
        t_profile.rows[0].cells[0].text = "Indicador"
        t_profile.rows[0].cells[1].text = "Valor"

        profile_rows = [
            ("Total de participantes", survey_profile.get("n_total", meta.get("n_total", "—"))),
            ("Variables analizadas", survey_profile.get("n_variables", ds.n_cols)),
            ("Variables categóricas", survey_profile.get("n_categorical", "—")),
            ("Variables numéricas", survey_profile.get("n_numeric", "—")),
            ("Segmento principal", survey_profile.get("segment_column", meta.get("segment_column", "No detectado"))),
        ]
        for label, value in profile_rows:
            row = t_profile.add_row().cells
            row[0].text = str(label)
            row[1].text = str(value)

        # =========================
        # 3. INSIGHTS AUTOMÁTICOS SADI
        # =========================
        if advanced_insights:
            doc.add_heading("3. Insights automáticos SADI", 1)
            for line in str(advanced_insights).split("\n"):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)

        # =========================
        # 4. RESULTADOS
        # =========================
        if report_ctx.get("academic_results_text"):
            doc.add_heading("4. Resultados", 1)
            for block in str(report_ctx.get("academic_results_text")).split("\n\n"):
                block = block.strip()
                if block:
                    doc.add_paragraph(block)

        # =========================
        # 5. RECOMENDACIONES
        # =========================
        if recommendations:
            doc.add_heading("5. Recomendaciones", 1)
            for rec in recommendations:
                doc.add_paragraph(f"• {rec}")

        # =========================
        # 6. PLAN SUGERIDO
        # =========================
        if plan:
            doc.add_heading("6. Plan de análisis sugerido", 1)
            recommended_analysis = plan.get("recommended_analysis") or []
            recommended_plots = plan.get("recommended_plots") or []
            narrative_focus = plan.get("narrative_focus")
            plan_warnings = plan.get("warnings") or []

            if recommended_analysis:
                doc.add_paragraph("Análisis sugeridos:")
                for item in recommended_analysis:
                    doc.add_paragraph(f"• {item}")

            if recommended_plots:
                doc.add_paragraph("Gráficos sugeridos:")
                for item in recommended_plots:
                    doc.add_paragraph(f"• {item}")

            if narrative_focus:
                doc.add_paragraph(f"Enfoque narrativo recomendado: {narrative_focus}")

            if plan_warnings:
                doc.add_paragraph("Advertencias del plan:")
                for item in plan_warnings:
                    doc.add_paragraph(f"• {item}")

        # =========================
        # 7. PRIORIDAD ANALÍTICA
        # =========================
        if priority:
            doc.add_heading("7. Prioridad analítica sugerida", 1)
            for item in priority:
                doc.add_paragraph(f"• {item}")

        # =========================
        # 8. VARIABLES CLAVE
        # =========================
        if variable_importance:
            doc.add_heading("8. Variables clave", 1)
            t = doc.add_table(rows=1, cols=4)
            t.style = "Table Grid"
            t.rows[0].cells[0].text = "Variable"
            t.rows[0].cells[1].text = "Score"
            t.rows[0].cells[2].text = "Desv. estándar"
            t.rows[0].cells[3].text = "% Missing"

            for item in variable_importance[:10]:
                if not isinstance(item, dict):
                    continue
                row = t.add_row().cells
                row[0].text = str(item.get("column", "—"))
                row[1].text = str(item.get("score", "—"))
                row[2].text = str(item.get("std", "—"))
                row[3].text = f"{item.get('missing_pct', '—')}%"

        # =========================
        # 9. VARIABLES CON MAYOR VARIABILIDAD
        # =========================
        if top_variability:
            doc.add_heading("9. Variables con mayor variabilidad", 1)
            t = doc.add_table(rows=1, cols=3)
            t.style = "Table Grid"
            t.rows[0].cells[0].text = "Variable"
            t.rows[0].cells[1].text = "Desv. estándar"
            t.rows[0].cells[2].text = "Rango"

            for item in top_variability[:10]:
                if not isinstance(item, dict):
                    continue
                row = t.add_row().cells
                row[0].text = str(item.get("column", "—"))
                row[1].text = str(item.get("std", "—"))
                row[2].text = str(item.get("range", "—"))

        # =========================
        # 10. VARIABLE OBJETIVO / MODELO
        # =========================
        if target_candidate:
            doc.add_heading("10. Variable objetivo sugerida", 1)
            doc.add_paragraph(f"Variable: {target_candidate}")
            doc.add_paragraph(f"Tipo: {target_type or '—'}")
            if target_reason:
                doc.add_paragraph(f"Justificación: {target_reason}")
            if model_suggestion:
                doc.add_paragraph(f"Modelo sugerido: {model_suggestion}")

        # =========================
        # 11. CANDIDATOS
        # =========================
        if ranked_target_candidates:
            doc.add_heading("11. Candidatos a variable objetivo", 1)
            t = doc.add_table(rows=1, cols=4)
            t.style = "Table Grid"
            t.rows[0].cells[0].text = "Variable"
            t.rows[0].cells[1].text = "Tipo"
            t.rows[0].cells[2].text = "Score"
            t.rows[0].cells[3].text = "Motivo"

            for item in ranked_target_candidates[:10]:
                if not isinstance(item, dict):
                    continue
                row = t.add_row().cells
                row[0].text = str(item.get("column", "—"))
                row[1].text = str(item.get("type", "—"))
                row[2].text = str(item.get("score", "—"))
                row[3].text = str(item.get("reason", "—"))

        # =========================
        # 12. PRÓXIMO PASO
        # =========================
        if next_step:
            doc.add_heading("12. Próximo paso recomendado", 1)
            if next_step.get("title"):
                doc.add_paragraph(f"Título: {next_step.get('title')}")
            if next_step.get("reason"):
                doc.add_paragraph(f"Por qué: {next_step.get('reason')}")
            if next_step.get("action"):
                doc.add_paragraph(f"Acción sugerida: {next_step.get('action')}")

        # =========================
        # 13. CONCLUSIÓN
        # =========================
        doc.add_heading("13. Conclusión", 1)
        doc.add_paragraph(
            report_ctx.get("conclusion_text")
            or advanced_insights
            or report_ctx.get("survey_insights")
            or ""
        )

        # =========================
        # 14. ADVERTENCIAS
        # =========================
        if warnings_list:
            doc.add_heading("14. Advertencias", 1)
            for w in warnings_list:
                doc.add_paragraph(f"⚠ {w}")

        # =========================
        # 15. GRÁFICOS DEL ANÁLISIS
        # =========================
        if plots:
            doc.add_page_break()
            doc.add_heading("15. Gráficos del análisis", 1)

            for i, rel in enumerate(plots, start=1):
                rel = str(rel).replace("\\", "/")
                fname = rel.split("/")[-1]
                img_path = os.path.join(PLOTS_DIR, fname)
                if not os.path.exists(img_path):
                    continue

                nice_title = prettify_plot_title(rel)
                nice_desc = describe_plot(rel)
                nice_tag = classify_plot_tag(rel)

                p_title = doc.add_paragraph()
                p_title.paragraph_format.keep_with_next = True
                r = p_title.add_run(f"Gráfico {i}. {nice_title}")
                r.bold = True
                r.font.size = Pt(12)

                p_type = doc.add_paragraph()
                p_type.paragraph_format.keep_with_next = True
                r2 = p_type.add_run(f"Tipo: {nice_tag}")
                r2.italic = True

                if nice_desc:
                    p_desc = doc.add_paragraph(nice_desc)
                    p_desc.paragraph_format.keep_with_next = True

                inserted = safe_add_picture(doc, img_path, width_inches=6.2)
                if inserted:
                    doc.add_paragraph("")

        doc.save(docx_path)

    except Exception as e:
        current_app.logger.exception(f"[survey_normal_word] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el informe Word: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    return send_file(
        docx_path,
        as_attachment=True,
        download_name=f"encuesta_normal_{dataset_id}.docx"
    )

@survey_bp.get("/datasets/<int:dataset_id>/survey_articulo_word", endpoint="dataset_survey_articulo_word")
@login_required
def dataset_survey_articulo_word(dataset_id: int):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    kind = (getattr(ds, "dataset_type", None) or "no_definido").strip()
    if kind != "survey_normal":
        flash("Este artículo aplica solo para 'Encuesta normal'.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    manifest_data = read_manifest_data(dataset_id) or {}
    meta = manifest_data.get("meta") or {}
    plots = manifest_data.get("generated") or manifest_data.get("plots") or []

    if not isinstance(meta, dict):
        meta = {}
    if not isinstance(plots, list):
        plots = []

    if not meta:
        flash("Primero genera el análisis descriptivo de encuesta.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    reports_dir = os.path.join(PLOTS_DIR, "reports")
    os.makedirs(reports_dir, exist_ok=True)
    docx_path = os.path.join(reports_dir, f"ds{dataset_id}_articulo_encuesta.docx")

    try:
        plots = normalize_plot_catalog(dataset_id, manifest_data, meta)
        report_ctx = build_academic_report_context(ds, meta, plots)

        survey_profile = report_ctx.get("survey_profile") or {}
        if not isinstance(survey_profile, dict):
            survey_profile = {}

        academic_results = report_ctx.get("academic_results_text") or ""
        key_findings = report_ctx.get("key_findings") or []
        crosstabs = report_ctx.get("crosstabs") or []
        group_summary = report_ctx.get("group_comparison_summary") or ""

        advanced_insights = (
            report_ctx.get("sadi_insights")
            or report_ctx.get("insights_text")
            or meta.get("sadi_insights")
            or meta.get("insights_text")
            or report_ctx.get("survey_insights")
            or meta.get("survey_insights")
            or ""
        )

        recommendations = (
            report_ctx.get("sadi_recommendations")
            or report_ctx.get("quick_recommendations")
            or meta.get("sadi_recommendations")
            or meta.get("quick_recommendations")
            or []
        )
        if not isinstance(recommendations, list):
            recommendations = [recommendations] if recommendations else []

        plan = (
            report_ctx.get("sadi_plan")
            or report_ctx.get("suggested_plan")
            or meta.get("sadi_plan")
            or meta.get("suggested_plan")
            or {}
        )
        if not isinstance(plan, dict):
            plan = {}

        priority = (
            report_ctx.get("sadi_priority")
            or report_ctx.get("priority_order")
            or meta.get("sadi_priority")
            or meta.get("priority_order")
            or []
        )
        if not isinstance(priority, list):
            priority = [priority] if priority else []

        next_step = (
            report_ctx.get("next_step_recommendation")
            or meta.get("next_step_recommendation")
            or {}
        )
        if not isinstance(next_step, dict):
            next_step = {}

        variable_importance = (
            report_ctx.get("variable_importance")
            or meta.get("variable_importance")
            or []
        )
        if not isinstance(variable_importance, list):
            variable_importance = []

        target_candidate = (
            report_ctx.get("target_candidate")
            or meta.get("target_candidate")
        )
        target_type = (
            report_ctx.get("target_type")
            or meta.get("target_type")
        )
        target_reason = (
            report_ctx.get("target_reason")
            or meta.get("target_reason")
        )
        model_suggestion = (
            report_ctx.get("model_suggestion")
            or meta.get("model_suggestion")
        )

        warnings_list = (
            report_ctx.get("warnings")
            or meta.get("warnings")
            or []
        )
        if not isinstance(warnings_list, list):
            warnings_list = [warnings_list] if warnings_list else []

        current_app.logger.warning(
            f"[SURVEY ARTICLE CHECK] ds{dataset_id} "
            f"advanced_insights={bool(advanced_insights)} "
            f"recommendations={len(recommendations)} "
            f"plan={bool(plan)} "
            f"priority={len(priority)} "
            f"next_step={bool(next_step)} "
            f"variable_importance={len(variable_importance)} "
            f"target_candidate={target_candidate!r}"
        )

        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        try:
            doc.styles["Normal"].font.name = "Times New Roman"
            doc.styles["Normal"].font.size = Pt(11)
        except Exception:
            pass

        # =========================
        # TÍTULO
        # =========================
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run(
            f"Análisis descriptivo y comparativo de la encuesta: {safe_text(report_ctx.get('dataset_name'))}"
        )
        r.bold = True
        r.font.size = Pt(15)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = subtitle.add_run("Artículo científico generado automáticamente por SADI")
        r2.italic = True
        r2.font.size = Pt(11)

        # =========================
        # RESUMEN
        # =========================
        doc.add_heading("Resumen", level=1)
        resumen = advanced_insights or (
            f"Se analizaron {survey_profile.get('n_total', meta.get('n_total', '—'))} registros "
            f"correspondientes a una encuesta normal. El estudio incluyó análisis descriptivo, "
            f"asociaciones categóricas y comparaciones entre grupos cuando fue posible."
        )
        doc.add_paragraph(str(resumen))

        # =========================
        # 1. INTRODUCCIÓN
        # =========================
        doc.add_heading("1. Introducción", level=1)
        doc.add_paragraph(
            report_ctx.get("intro_text")
            or "Las encuestas constituyen una herramienta clave para describir percepciones, opiniones y patrones de comportamiento en contextos académicos e institucionales."
        )

        # =========================
        # 2. MÉTODO
        # =========================
        doc.add_heading("2. Método", level=1)
        doc.add_paragraph(
            f"El análisis se realizó sobre {survey_profile.get('n_total', meta.get('n_total', '—'))} registros. "
            f"Se procesaron {survey_profile.get('n_variables', getattr(ds, 'n_cols', '—'))} variables, "
            f"de las cuales {survey_profile.get('n_categorical', '—')} fueron categóricas y "
            f"{survey_profile.get('n_numeric', '—')} numéricas."
        )

        segment_column = survey_profile.get("segment_column", meta.get("segment_column"))
        if segment_column:
            doc.add_paragraph(
                f"Se identificó la variable '{segment_column}' como eje principal de segmentación, "
                "utilizada para explorar asociaciones y contrastes entre grupos."
            )
        else:
            doc.add_paragraph(
                "No se identificó una variable de segmentación suficientemente robusta, por lo que el estudio "
                "se concentró principalmente en la descripción general del conjunto de datos."
            )

        # =========================
        # 3. RESULTADOS
        # =========================
        doc.add_heading("3. Resultados", level=1)

        if academic_results:
            for block in str(academic_results).split("\n\n"):
                block = block.strip()
                if block:
                    doc.add_paragraph(block)

        if key_findings:
            doc.add_paragraph("")
            doc.add_paragraph("Hallazgos principales:")
            for item in key_findings[:8]:
                txt = str(item).strip()
                if txt:
                    doc.add_paragraph(f"• {txt}")

        if crosstabs:
            doc.add_paragraph("")
            doc.add_paragraph("Asociaciones categóricas relevantes:")
            for row in crosstabs[:6]:
                if not isinstance(row, dict):
                    continue
                var = row.get("variable", "—")
                p = row.get("p_value", None)
                v = row.get("cramers_v", None)
                st = row.get("strength", "")
                p_txt = f"{float(p):.6f}" if isinstance(p, (int, float)) else "—"
                v_txt = f"{float(v):.3f}" if isinstance(v, (int, float)) else "—"
                doc.add_paragraph(f"• {var}: Cramér’s V = {v_txt}, p = {p_txt}, intensidad {st}.")

        if group_summary:
            doc.add_paragraph("")
            doc.add_paragraph(group_summary)

        # =========================
        # 4. DISCUSIÓN
        # =========================
        doc.add_heading("4. Discusión", level=1)
        doc.add_paragraph(
            advanced_insights
            or "Los resultados muestran que el análisis automatizado permite identificar patrones relevantes en el dataset."
        )

        # =========================
        # 5. RECOMENDACIONES
        # =========================
        if recommendations:
            doc.add_heading("5. Recomendaciones", level=1)
            for r in recommendations:
                doc.add_paragraph(f"• {r}")

        # =========================
        # 6. PLAN SUGERIDO
        # =========================
        if plan:
            doc.add_heading("6. Plan sugerido", level=1)

            recommended_analysis = plan.get("recommended_analysis") or []
            recommended_plots = plan.get("recommended_plots") or []
            narrative_focus = plan.get("narrative_focus")
            plan_warnings = plan.get("warnings") or []

            if recommended_analysis:
                doc.add_paragraph("Análisis sugeridos:")
                for item in recommended_analysis:
                    doc.add_paragraph(f"• {item}")

            if recommended_plots:
                doc.add_paragraph("Gráficos sugeridos:")
                for item in recommended_plots:
                    doc.add_paragraph(f"• {item}")

            if narrative_focus:
                doc.add_paragraph(f"Enfoque narrativo recomendado: {narrative_focus}")

            if plan_warnings:
                doc.add_paragraph("Advertencias metodológicas:")
                for item in plan_warnings:
                    doc.add_paragraph(f"• {item}")

        # =========================
        # 7. PRIORIDAD ANALÍTICA
        # =========================
        if priority:
            doc.add_heading("7. Prioridad analítica sugerida", level=1)
            for item in priority:
                doc.add_paragraph(f"• {item}")

        # =========================
        # 8. VARIABLES CLAVE
        # =========================
        if variable_importance:
            doc.add_heading("8. Variables clave", level=1)
            for item in variable_importance[:10]:
                if isinstance(item, dict):
                    doc.add_paragraph(
                        f"• {item.get('column', '—')} | score={item.get('score', '—')} | missing={item.get('missing_pct', '—')}%"
                    )

        # =========================
        # 9. MODELO PREDICTIVO SUGERIDO
        # =========================
        if target_candidate:
            doc.add_heading("9. Modelo predictivo sugerido", level=1)
            doc.add_paragraph(f"Variable objetivo: {target_candidate}")
            doc.add_paragraph(f"Tipo: {target_type or '—'}")
            if target_reason:
                doc.add_paragraph(f"Justificación: {target_reason}")
            if model_suggestion:
                doc.add_paragraph(f"Modelo sugerido: {model_suggestion}")

        # =========================
        # 10. PRÓXIMO PASO
        # =========================
        if next_step:
            doc.add_heading("10. Próximo paso recomendado", level=1)
            if next_step.get("title"):
                doc.add_paragraph(f"Título: {next_step.get('title')}")
            if next_step.get("reason"):
                doc.add_paragraph(f"Por qué: {next_step.get('reason')}")
            if next_step.get("action"):
                doc.add_paragraph(f"Acción sugerida: {next_step.get('action')}")

        # =========================
        # 11. CONCLUSIÓN
        # =========================
        doc.add_heading("11. Conclusión", level=1)
        doc.add_paragraph(
            report_ctx.get("conclusion_text")
            or advanced_insights
            or "En conclusión, la encuesta analizada ofrece información relevante sobre la estructura descriptiva de los datos."
        )

        # =========================
        # 12. ADVERTENCIAS
        # =========================
        if warnings_list:
            doc.add_heading("12. Advertencias", level=1)
            for w in warnings_list:
                doc.add_paragraph(f"⚠ {w}")

        # =========================
        # FIGURAS
        # =========================
        if plots:
            doc.add_page_break()
            doc.add_heading("Anexo. Figuras del análisis", level=1)

            for i, rel in enumerate(plots, start=1):
                rel = str(rel).replace("\\", "/")
                fname = rel.split("/")[-1]
                img_path = os.path.join(PLOTS_DIR, fname)
                if not os.path.exists(img_path):
                    continue

                nice_title = prettify_plot_title(rel)
                nice_desc = describe_plot(rel)
                nice_tag = classify_plot_tag(rel)

                p_title = doc.add_paragraph()
                p_title.paragraph_format.keep_with_next = True
                rr = p_title.add_run(f"Figura {i}. {nice_title}")
                rr.bold = True
                rr.font.size = Pt(12)

                p_type = doc.add_paragraph()
                p_type.paragraph_format.keep_with_next = True
                rr2 = p_type.add_run(f"Tipo: {nice_tag}")
                rr2.italic = True

                if nice_desc:
                    p_desc = doc.add_paragraph(nice_desc)
                    p_desc.paragraph_format.keep_with_next = True

                safe_add_picture(doc, img_path, width_inches=6.0)
                doc.add_paragraph("")

        doc.save(docx_path)

    except Exception as e:
        current_app.logger.exception(f"[survey_articulo_word] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el artículo Word: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    return send_file(
        docx_path,
        as_attachment=True,
        download_name=f"articulo_encuesta_{dataset_id}.docx"
    )


@survey_bp.get("/datasets/<int:dataset_id>/survey_normal_latex", endpoint="dataset_survey_normal_latex")
@login_required
def dataset_survey_normal_latex(dataset_id: int):
    def tex_escape(value: str) -> str:
        text = str(value)
        replacements = {
            "\\": r"\textbackslash{}",
            "&": r"\&",
            "%": r"\%",
            "$": r"\$",
            "#": r"\#",
            "_": r"\_",
            "{": r"\{",
            "}": r"\}",
            "~": r"\textasciitilde{}",
            "^": r"\textasciicircum{}",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text

    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    kind = (getattr(ds, "dataset_type", None) or "no_definido").strip()
    if kind != "survey_normal":
        flash("Este informe LaTeX aplica solo para 'Encuesta normal'.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    manifest = read_manifest_data(dataset_id) or {}
    meta = (manifest.get("meta") or {}) if isinstance(manifest, dict) else {}
    plots = (manifest.get("generated") or manifest.get("plots") or []) if isinstance(manifest, dict) else []

    if not isinstance(meta, dict):
        meta = {}
    if not isinstance(plots, list):
        plots = []

    if not meta:
        flash("Primero genera el análisis descriptivo de encuesta.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    reports_dir = os.path.join(PLOTS_DIR, "reports")
    os.makedirs(reports_dir, exist_ok=True)
    out_path = os.path.join(reports_dir, f"ds{dataset_id}_encuesta_normal.tex")

    try:
        plots = normalize_plot_catalog(dataset_id, manifest_data, summary)
        exploratory_plots, model_plots = split_plots(plots)
        report_ctx = build_academic_report_context(ds, meta, plots)
        survey_profile = report_ctx.get("survey_profile") or meta.get("survey_profile") or {}
        key_findings = report_ctx.get("survey_key_findings") or meta.get("survey_key_findings") or []
        survey_insights = report_ctx.get("survey_insights") or meta.get("survey_insights") or ""
        academic_results = report_ctx.get("academic_results_text") or meta.get("results_text") or ""
        crosstabs = report_ctx.get("crosstabs") or []
        group_summary = report_ctx.get("group_comparison_summary") or ""

        findings_tex = ""
        if key_findings:
            findings_tex += "\\begin{itemize}\n"
            for item in key_findings[:8]:
                findings_tex += f"\\item {tex_escape(item)}\n"
            findings_tex += "\\end{itemize}\n"

        crosstabs_tex = ""
        if crosstabs:
            crosstabs_tex += "\\begin{itemize}\n"
            for row in crosstabs[:6]:
                if not isinstance(row, dict):
                    continue
                var = tex_escape(row.get("variable", "—"))
                p = row.get("p_value", None)
                v = row.get("cramers_v", None)
                st = tex_escape(row.get("strength", ""))
                p_txt = f"{float(p):.6f}" if isinstance(p, (int, float)) else "—"
                v_txt = f"{float(v):.3f}" if isinstance(v, (int, float)) else "—"
                crosstabs_tex += f"\\item {var}: Cramér’s V = {v_txt}, p = {p_txt}, intensidad {st}.\n"
            crosstabs_tex += "\\end{itemize}\n"

        latex = rf"""
\documentclass[12pt]{{article}}
\usepackage[utf8]{{inputenc}}
\usepackage[spanish]{{babel}}
\usepackage{{geometry}}
\geometry{{margin=2.5cm}}
\usepackage{{longtable}}
\usepackage{{array}}

\title{{Informe académico de encuesta normal}}
\author{{SADI}}
\date{{\today}}

\begin{{document}}
\maketitle

\section*{{Identificación}}
Título del dataset: {tex_escape(ds.title or ds.original_name)}

Archivo: {tex_escape(ds.original_name)}

Tipo: Encuesta normal

Total de registros: {tex_escape(survey_profile.get("n_total", meta.get("n_total", "—")))}

Variables analizadas: {tex_escape(survey_profile.get("n_variables", getattr(ds, "n_cols", "—")))}

Variables categóricas: {tex_escape(survey_profile.get("n_categorical", "—"))}

Variables numéricas: {tex_escape(survey_profile.get("n_numeric", "—"))}

Segmento principal: {tex_escape(survey_profile.get("segment_column", meta.get("segment_column", "No detectado")))}

\section*{{Resumen ejecutivo}}
{tex_escape(survey_insights or "No se generaron insights automáticos.")}

\section*{{Introducción}}
{tex_escape(report_ctx.get("intro_text") or "Este informe resume el análisis descriptivo y comparativo de una encuesta normal, identificando variables relevantes, asociaciones categóricas y diferencias entre grupos cuando fue posible.")}

\section*{{Hallazgos clave}}
{findings_tex if findings_tex else tex_escape("No se detectaron hallazgos clave automáticos.")}

\section*{{Resultados}}
{tex_escape(academic_results or "No se generó texto automático de resultados.")}

\section*{{Asociaciones categóricas}}
{crosstabs_tex if crosstabs_tex else tex_escape("No se detectaron asociaciones categóricas relevantes.")}

\section*{{Comparación entre grupos}}
{tex_escape(group_summary or "No se generaron comparaciones significativas entre grupos.")}

\section*{{Conclusión}}
{tex_escape(report_ctx.get("conclusion_text") or "La encuesta analizada ofrece un panorama descriptivo útil y puede servir como base para interpretaciones académicas posteriores.")}

\end{{document}}
""".strip()

        with open(out_path, "w", encoding="utf-8") as f:
            f.write(latex)

    except Exception as e:
        current_app.logger.exception(f"[survey_normal_latex] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el archivo LaTeX: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    return send_file(
        out_path,
        as_attachment=True,
        download_name=f"encuesta_normal_{dataset_id}.tex"
    )