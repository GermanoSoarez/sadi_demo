from __future__ import annotations

import os
import re
from datetime import datetime
from typing import Optional
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from flask import (
    Blueprint,
    flash,
    redirect,
    send_file,
    url_for,
    current_app,
)
from flask_login import login_required, current_user

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Image as RLImage,
    Table,
    TableStyle,
    PageBreak,
)
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

from factor_analyzer import FactorAnalyzer
from factor_analyzer.factor_analyzer import calculate_kmo, calculate_bartlett_sphericity

from extensions import SessionLocal, engine
from models import Dataset
from config import PLOTS_DIR, UPLOAD_DIR
from utils.manifest import read_manifest_data, write_manifest
from utils.plot_manager import normalize_plot_catalog, split_plots
from blueprints.survey.analysis import read_dataframe, repair_broken_csv_file
from blueprints.dataset.analysis import analyze_dataset_with_recommendations
likert_bp = Blueprint("likert", __name__)


# =========================================================
# HELPERS GENERALES
# =========================================================

def _reports_dir() -> str:
    out = os.path.join(PLOTS_DIR, "reports")
    os.makedirs(out, exist_ok=True)
    return out


def _static_root() -> str:
    # si PLOTS_DIR = .../static/plots
    return os.path.dirname(PLOTS_DIR)


def _fmt_num(x, decimals=4, default="—"):
    try:
        return f"{float(x):.{decimals}f}"
    except Exception:
        return default


def _fmt_pct(x, default="—"):
    try:
        return f"{float(x):.1f}%"
    except Exception:
        return default


def _resolve_plot(rel_or_name: str | None) -> str | None:
    if not rel_or_name:
        return None

    rel = str(rel_or_name).replace("\\", "/").strip()
    fname = rel.split("/")[-1]

    candidates = [
        os.path.join(PLOTS_DIR, fname),
        os.path.join(_static_root(), "plots", fname),
        rel,
    ]
    for path in candidates:
        if path and os.path.exists(path) and os.path.getsize(path) > 0:
            return path
    return None


def safe_add_picture(doc: Document, image_path: str, width_inches: float = 6.0) -> bool:
    try:
        if image_path and os.path.exists(image_path):
            doc.add_picture(image_path, width=Inches(width_inches))
            return True
    except Exception:
        pass
    return False

def _read_dataset_with_auto_repair(path: str, delimiter: str | None = None) -> pd.DataFrame:
    """
    Lee el dataset usando el lector robusto central.
    Si detecta una sola columna sospechosa, intenta reparación automática.
    """
    df = read_dataframe(path, delimiter)

    if df.shape[1] == 1:
        only_col = str(df.columns[0]) if len(df.columns) > 0 else ""
        if any(x in only_col for x in [",", ";", "\t", "|", ":"]) or df.shape[0] == 0:
            repair_info = repair_broken_csv_file(path, delimiter)

            if repair_info.get("ok") and repair_info.get("repaired"):
                current_app.logger.warning(
                    f"[AUTO-REPAIR] Archivo reparado: {path} backup={repair_info.get('backup_path')}"
                )
                df = read_dataframe(path, delimiter)
            elif not repair_info.get("ok"):
                current_app.logger.warning(
                    f"[AUTO-REPAIR] No se pudo reparar {path}: {repair_info.get('message')}"
                )

    return df

def get_likert_scale_from_dataset_type(dataset_type: str) -> int | None:
    dataset_type = (dataset_type or "").strip().lower()
    if dataset_type == "survey_likert_5":
        return 5
    if dataset_type == "survey_likert_7":
        return 7
    return None


def coerce_likert_cell_to_int(v, scale: int | None = None) -> int | None:
    if v is None:
        return None

    try:
        if pd.isna(v):
            return None
    except Exception:
        pass

    value = None

    if isinstance(v, (int, float)):
        try:
            value = int(round(float(v)))
        except Exception:
            value = None
    else:
        s = str(v).strip()
        if not s or s.lower() in ("nan", "none", "null"):
            return None

        s = s.replace(",", ".")

        try:
            value = int(round(float(s)))
        except Exception:
            value = None

        if value is None:
            m = re.search(r"\b([1-9])(?:\.0+)?\b", s)
            if not m:
                return None
            try:
                value = int(m.group(1))
            except Exception:
                return None

    if value is None:
        return None

    if scale is not None and not (1 <= value <= scale):
        return None

    return value


def detect_likert_columns_for_scale(df: pd.DataFrame, scale: int, min_valid_ratio: float = 0.80) -> list[str]:
    if scale not in (5, 7):
        return []

    excluded_names = {
        "id", "id_respuesta", "id_estudiante", "edad", "semestre", "anho", "ano",
        "anio", "year", "curso", "codigo", "cod", "nro", "numero", "fila", "index",
        "orden", "orden_item", "grupo", "seccion", "paralelo"
    }

    likert_cols: list[str] = []

    for col in df.columns:
        col_l = str(col).strip().lower()
        if col_l in excluded_names:
            current_app.logger.warning(f"[LIKERT DETECT] col={col!r} excluida por nombre")
            continue

        s = df[col]
        coerced = s.map(lambda x: coerce_likert_cell_to_int(x, scale=scale))

        total = len(coerced)
        if total == 0:
            continue

        valid_series = coerced.dropna()
        valid = len(valid_series)
        ratio = valid / total if total > 0 else 0.0

        if valid == 0:
            current_app.logger.warning(
                f"[LIKERT DETECT] col={col!r} valid=0 total={total} ratio=0.000 uniques=[]"
            )
            continue

        uniques = sorted(set(int(v) for v in valid_series.unique()))
        in_range = [u for u in uniques if 1 <= u <= scale]
        out_of_range = [u for u in uniques if u < 1 or u > scale]

        current_app.logger.warning(
            f"[LIKERT DETECT] col={col!r} total={total} valid={valid} "
            f"ratio={ratio:.3f} uniques={uniques} in_range={in_range} out_of_range={out_of_range}"
        )

        min_categories = 2 if scale == 5 else 3
        if ratio >= min_valid_ratio and len(in_range) >= min_categories and len(out_of_range) == 0:
            likert_cols.append(col)

    current_app.logger.warning(f"[LIKERT DETECT] scale={scale} -> likert_cols={likert_cols}")
    return likert_cols


def cronbach_alpha(df: pd.DataFrame) -> float | None:
    if df is None or df.empty or df.shape[1] < 2:
        return None

    x = df.dropna().copy()
    if x.empty or x.shape[1] < 2:
        return None

    try:
        item_vars = x.var(axis=0, ddof=1)
        total_var = x.sum(axis=1).var(ddof=1)
        if total_var == 0:
            return None
        k = x.shape[1]
        alpha = (k / (k - 1)) * (1 - item_vars.sum() / total_var)
        return float(alpha)
    except Exception:
        return None


def interpret_cronbach(alpha: float | None) -> str:
    if alpha is None:
        return "No aplicable"
    if alpha >= 0.90:
        return "Excelente"
    if alpha >= 0.80:
        return "Buena"
    if alpha >= 0.70:
        return "Aceptable"
    if alpha >= 0.60:
        return "Cuestionable"
    return "Baja"


def make_cronbach_alpha_plot(df: pd.DataFrame, dataset_id: int) -> tuple[float | None, list[str], str, str | None]:
    usable = [c for c in df.columns if df[c].nunique(dropna=True) >= 2]
    x = df[usable].dropna() if usable else df.dropna()

    alpha = cronbach_alpha(x)
    if alpha is None or x.shape[1] < 2:
        return None, [], "No aplica", None

    label = interpret_cronbach(alpha)

    fig, ax = plt.subplots(figsize=(9, 3.2))

    # Zonas de interpretación
    zones = [
        (0.00, 0.60, "Baja"),
        (0.60, 0.70, "Cuestionable"),
        (0.70, 0.80, "Aceptable"),
        (0.80, 0.90, "Buena"),
        (0.90, 1.00, "Excelente"),
    ]

    y = 0
    for start, end, zone_label in zones:
        ax.barh([y], [end - start], left=start, height=0.45, alpha=0.35)
        ax.text(
            (start + end) / 2,
            y - 0.32,
            zone_label,
            ha="center",
            va="center",
            fontsize=9
        )

    # Barra del valor observado
    ax.barh([y], [alpha], height=0.22)

    # Línea del alfa observado
    ax.axvline(alpha, linestyle="--", linewidth=1.8)

    # Texto principal
    ax.text(
        min(alpha + 0.015, 0.97),
        y + 0.02,
        f"α = {alpha:.3f} ({label})",
        va="center",
        fontsize=11,
        fontweight="bold"
    )

    # Configuración del eje
    ax.set_xlim(0, 1.0)
    ax.set_ylim(-0.6, 0.6)
    ax.set_yticks([])
    ax.set_xlabel("Valor del alfa de Cronbach")
    ax.set_title("Consistencia interna del instrumento")

    # Líneas guía
    for ref in [0.60, 0.70, 0.80, 0.90]:
        ax.axvline(ref, linestyle=":", linewidth=0.9, alpha=0.8)

    # Nota técnica
    ax.text(
        0.00,
        -0.48,
        f"Ítems usados: {x.shape[1]} | Casos válidos: {x.shape[0]}",
        fontsize=9,
        ha="left",
        va="center"
    )

    # Limpiar marco
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_visible(False)

    out_name = f"ds{dataset_id}_cronbach.png"
    out_path = os.path.join(PLOTS_DIR, out_name)

    fig.tight_layout()
    fig.savefig(out_path, dpi=180, bbox_inches="tight")
    plt.close(fig)

    return alpha, list(x.columns), label, out_name


def plot_likert_summary_scores(df: pd.DataFrame, dataset_id: int) -> str | None:
    try:
        means = df.mean().sort_values()
        fig, ax = plt.subplots(figsize=(10, max(4, len(means) * 0.35)))
        means.plot(kind="barh", ax=ax)
        ax.set_title("Promedio por ítem (Likert)")
        ax.set_xlabel("Promedio")
        ax.set_ylabel("Ítems")

        out_name = f"ds{dataset_id}_likert_summary_scores.png"
        out_path = os.path.join(PLOTS_DIR, out_name)
        fig.tight_layout()
        fig.savefig(out_path, dpi=160)
        plt.close(fig)

        return out_name
    except Exception as e:
        current_app.logger.warning(f"[plot_likert_summary_scores] ds{dataset_id}: {e}")
        return None


def plot_likert_divergent(df: pd.DataFrame, dataset_id: int, output_path: str, scale: int = 5) -> str:
    likert_cols = detect_likert_columns_for_scale(df, scale=scale, min_valid_ratio=0.80)
    if not likert_cols:
        raise ValueError(f"No se detectaron columnas Likert válidas para la escala {scale}.")

    labels = []
    data = []

    for col in likert_cols:
        vals = df[col].map(lambda x: coerce_likert_cell_to_int(x, scale=scale)).dropna().astype(int)
        counts = vals.value_counts(normalize=True).reindex(range(1, scale + 1), fill_value=0) * 100
        labels.append(str(col))
        data.append(counts.values)

    plot_df = pd.DataFrame(data, columns=list(range(1, scale + 1)), index=labels)

    fig, ax = plt.subplots(figsize=(12, max(5, len(labels) * 0.6)))

    if scale == 5:
        left_1 = -(plot_df[1] + plot_df[2] + plot_df[3] / 2)
        left_2 = -(plot_df[2] + plot_df[3] / 2)
        left_3 = -(plot_df[3] / 2)
        left_4 = plot_df[3] / 2
        left_5 = plot_df[3] / 2 + plot_df[4]

        ax.barh(labels, plot_df[1], left=left_1, label="Totalmente en desacuerdo")
        ax.barh(labels, plot_df[2], left=left_2, label="En desacuerdo")
        ax.barh(labels, plot_df[3], left=left_3, label="Neutral")
        ax.barh(labels, plot_df[4], left=left_4, label="De acuerdo")
        ax.barh(labels, plot_df[5], left=left_5, label="Totalmente de acuerdo")

    else:
        left_1 = -(plot_df[1] + plot_df[2] + plot_df[3] + plot_df[4] / 2)
        left_2 = -(plot_df[2] + plot_df[3] + plot_df[4] / 2)
        left_3 = -(plot_df[3] + plot_df[4] / 2)
        left_4 = -(plot_df[4] / 2)
        left_5 = plot_df[4] / 2
        left_6 = plot_df[4] / 2 + plot_df[5]
        left_7 = plot_df[4] / 2 + plot_df[5] + plot_df[6]

        ax.barh(labels, plot_df[1], left=left_1, label="Muy en desacuerdo")
        ax.barh(labels, plot_df[2], left=left_2, label="En desacuerdo")
        ax.barh(labels, plot_df[3], left=left_3, label="Algo en desacuerdo")
        ax.barh(labels, plot_df[4], left=left_4, label="Neutral")
        ax.barh(labels, plot_df[5], left=left_5, label="Algo de acuerdo")
        ax.barh(labels, plot_df[6], left=left_6, label="De acuerdo")
        ax.barh(labels, plot_df[7], left=left_7, label="Muy de acuerdo")

    ax.axvline(0, linewidth=1)
    ax.set_xlabel("Porcentaje")
    ax.set_title(f"Gráfico Likert divergente (escala 1 a {scale})")
    ax.legend(loc="lower right", fontsize=8)
    fig.tight_layout()
    fig.savefig(output_path, dpi=220, bbox_inches="tight")
    plt.close(fig)

    return output_path


def build_general_figure_catalog_likert(dataset_id: int, plots: list[str]) -> list[dict]:
    catalog = []
    for rel in plots or []:
        fname = os.path.basename(str(rel).replace("\\", "/"))
        full = os.path.join(PLOTS_DIR, fname)
        if not os.path.exists(full):
            continue

        catalog.append({
            "filename": f"plots/{fname}",
            "basename": fname,
            "title": prettify_plot_title(fname),
            "caption": describe_plot(fname),
            "section": classify_plot_tag(fname),
        })
    return catalog


def build_likert_sadi_payload(
    *,
    ds,
    likert_df: pd.DataFrame,
    summary: dict,
    manifest_meta: dict | None = None,
    plots: list[str] | None = None,
) -> dict:
    """
    Unifica el análisis psicométrico Likert con el motor central SADI.
    """
    from blueprints.dataset.analysis import analyze_dataset_with_recommendations

    manifest_meta = manifest_meta or {}
    plots = plots or []

    dataset_type = (getattr(ds, "dataset_type", None) or "survey_likert").strip()
    research_area = (getattr(ds, "research_area", None) or "general").strip()

    analysis_meta = analyze_dataset_with_recommendations(
        likert_df,
        dataset_type=dataset_type,
        research_area=research_area,
    ) or {}

    alpha = summary.get("cronbach_alpha")
    n = summary.get("n", len(likert_df))
    scale = summary.get("scale") or manifest_meta.get("likert_scale") or 5
    means_by_item = summary.get("means_by_item", {}) or {}

    # Insights Likert específicos
    likert_insights = []

    if scale:
        likert_insights.append(f"El instrumento fue procesado como una escala Likert de {scale} puntos.")

    if n:
        likert_insights.append(f"Se analizaron {n} casos válidos para la evaluación psicométrica.")

    if means_by_item:
        try:
            ordered = sorted(means_by_item.items(), key=lambda x: x[1])
            low_item, low_mean = ordered[0]
            high_item, high_mean = ordered[-1]
            likert_insights.append(
                f"El ítem con promedio más bajo fue '{low_item}' ({float(low_mean):.2f}), mientras que el más alto fue '{high_item}' ({float(high_mean):.2f})."
            )
        except Exception:
            pass

    if alpha is not None:
        label = interpret_cronbach(alpha)
        likert_insights.append(
            f"El alfa de Cronbach global fue {float(alpha):.4f}, lo que sugiere una consistencia interna {label.lower()}."
        )

    kmo_report = manifest_meta.get("kmo_report") or {}
    bart_report = manifest_meta.get("bartlett_report") or {}
    efa_report = manifest_meta.get("efa_report") or {}

    kmo_global = kmo_report.get("kmo_global")
    if kmo_global is not None:
        likert_insights.append(
            f"El índice KMO global alcanzó {float(kmo_global):.3f}, indicando una adecuación muestral {interpret_kmo(float(kmo_global))}."
        )

    p_bart = bart_report.get("p_value")
    if p_bart is not None:
        try:
            if float(p_bart) < 0.05:
                likert_insights.append(
                    "La prueba de Bartlett fue significativa, por lo que la matriz de correlaciones resulta adecuada para análisis factorial."
                )
            else:
                likert_insights.append(
                    "La prueba de Bartlett no fue significativa, por lo que la estructura factorial debe interpretarse con cautela."
                )
        except Exception:
            pass

    n_factors = efa_report.get("n_factors")
    if n_factors:
        likert_insights.append(
            f"El análisis factorial exploratorio sugirió una estructura de {n_factors} factor(es)."
        )

    factor_interp = manifest_meta.get("factor_interpretation")
    psych_interp = manifest_meta.get("psychometric_interpretation")

    # Combinar insights del motor central + Likert
    base_insights = analysis_meta.get("insights", []) or []
    analysis_meta["insights"] = base_insights + likert_insights

    # Narrativa enriquecida
    base_text = (analysis_meta.get("insights_text") or "").strip()
    extra_parts = []
    if psych_interp:
        extra_parts.append(str(psych_interp).strip())
    if factor_interp:
        extra_parts.append(str(factor_interp).strip())

    full_text = " ".join([x for x in [base_text] + extra_parts if x]).strip()
    if full_text:
        analysis_meta["insights_text"] = full_text

    # Recomendaciones rápidas adicionales Likert
    qr = analysis_meta.get("quick_recommendations", []) or []
    if alpha is not None and float(alpha) < 0.70:
        qr.append("Revisar ítems con baja consistencia interna y considerar depuración del instrumento.")
    if kmo_global is not None and float(kmo_global) >= 0.60:
        qr.append("La adecuación muestral permite profundizar en la interpretación de factores latentes.")
    if n_factors:
        qr.append("Interpretar las dimensiones emergentes y contrastarlas con la teoría del instrumento.")
    analysis_meta["quick_recommendations"] = list(dict.fromkeys(qr))

    # Plan sugerido enriquecido
    sp = analysis_meta.get("suggested_plan", {}) or {}
    ra = sp.get("recommended_analysis", []) or []
    rp = sp.get("recommended_plots", []) or []

    ra.extend([
        "Consistencia interna mediante alfa de Cronbach",
        "Prueba KMO y esfericidad de Bartlett",
        "Análisis factorial exploratorio",
        "Interpretación de dimensiones del instrumento",
    ])
    rp.extend([
        "Resumen Likert por ítem",
        "Gráfico Likert divergente",
        "Scree plot",
        "Mapa de cargas factoriales",
        "Modelo factorial",
    ])

    sp["recommended_analysis"] = list(dict.fromkeys(ra))
    sp["recommended_plots"] = list(dict.fromkeys(rp))
    if not sp.get("narrative_focus"):
        sp["narrative_focus"] = (
            "Priorizar la consistencia interna del instrumento, la validez de la estructura factorial "
            "y la interpretación sustantiva de las dimensiones detectadas."
        )
    analysis_meta["suggested_plan"] = sp

    plot_summary = summarize_plot_tags(plots)
    general_figure_catalog = build_general_figure_catalog_likert(dataset_id=getattr(ds, "id"), plots=plots)

    return {
        "analysis_meta": analysis_meta,
        "insights": analysis_meta.get("insights", []),
        "insights_text": analysis_meta.get("insights_text"),
        "quick_recommendations": analysis_meta.get("quick_recommendations", []),
        "suggested_plan": analysis_meta.get("suggested_plan", {}),
        "plot_summary": plot_summary,
        "general_figure_catalog": general_figure_catalog,
    }

def classify_plot_tag(plot_path: str) -> str:
    name = os.path.basename(str(plot_path)).lower()
    clean = re.sub(r"^ds\d+_", "", name)
    clean = re.sub(r"\.png$", "", clean)

    if clean == "cronbach":
        return "Psicométrico"
    if clean in ("likert_summary_scores", "dimension_radar"):
        return "Comparativo"
    if clean.startswith("likert_q"):
        return "Psicométrico"
    if clean.startswith("scree") or clean.startswith("factor_"):
        return "Factorial"
    if clean == "likert_divergent":
        return "Comparativo"
    if clean == "corr_heatmap":
        return "Relacional"
    return "Analítico"


def describe_plot(plot_path: str) -> str:
    name = os.path.basename(str(plot_path)).lower()
    clean = re.sub(r"^ds\d+_", "", name)
    clean = re.sub(r"\.png$", "", clean)

    if clean == "cronbach":
        return "Resume la consistencia interna del instrumento mediante el alfa de Cronbach."
    if clean == "likert_summary_scores":
        return "Compara los promedios obtenidos entre los ítems del instrumento Likert."
    if clean == "likert_divergent":
        return "Resume la distribución porcentual de respuestas por ítem en formato divergente."
    if clean == "factor_model":
        return "Representa la estructura general del modelo factorial obtenido."
    if clean == "factor_loadings":
        return "Presenta las cargas factoriales de los ítems en cada factor extraído."
    if clean == "scree_plot":
        return "Ayuda a identificar el número adecuado de factores mediante los autovalores observados."
    return "Gráfico generado automáticamente por SADI para apoyar la interpretación del instrumento."


def prettify_plot_title(plot_path: str) -> str:
    name = os.path.basename(str(plot_path)).lower()
    clean = re.sub(r"^ds\d+_", "", name)
    clean = re.sub(r"\.png$", "", clean)

    mapping = {
        "cronbach": "Consistencia interna (Alfa de Cronbach)",
        "likert_summary_scores": "Resumen Likert por ítem",
        "likert_divergent": "Gráfico Likert divergente",
        "scree_plot": "Gráfico de sedimentación (Scree Plot)",
        "factor_loadings": "Cargas factoriales",
        "factor_model": "Modelo factorial",
    }
    return mapping.get(clean, clean.replace("_", " ").title())


def summarize_plot_tags(plots: list[str]) -> list[dict]:
    from collections import Counter

    if not plots:
        return []

    counts = Counter()
    for p in plots:
        counts[classify_plot_tag(p)] += 1

    preferred_order = [
        "Psicométrico",
        "Comparativo",
        "Factorial",
        "Relacional",
        "Analítico",
    ]

    result = []
    used = set()
    for tag in preferred_order:
        if tag in counts:
            result.append({"tag": tag, "count": int(counts[tag])})
            used.add(tag)

    for tag in sorted(counts.keys()):
        if tag not in used:
            result.append({"tag": tag, "count": int(counts[tag])})

    return result


def filter_plots_for_dataset(dataset_id: int, plots) -> list[str]:
    if not isinstance(plots, list):
        return []

    prefix = f"ds{dataset_id}_"
    clean = []

    for p in plots:
        if not p:
            continue
        rel = str(p).replace("\\", "/").strip()
        fname = rel.split("/")[-1]
        if fname.startswith(prefix):
            clean.append(rel)

    return clean


# =========================================================
# PSICOMETRÍA AVANZADA
# =========================================================

def get_likert_analysis_df(dataset_id: int) -> tuple[pd.DataFrame, int, list[str]]:
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds:
            raise ValueError("Dataset no encontrado.")

        dataset_type = (getattr(ds, "dataset_type", None) or "").strip().lower()
        filename = ds.filename
        delimiter = ds.delimiter

    scale = get_likert_scale_from_dataset_type(dataset_type)
    if scale not in (5, 7):
        raise ValueError("El dataset no está configurado como Likert 5 o 7.")

    path = os.path.join(UPLOAD_DIR, filename)
    df = _read_dataset_with_auto_repair(path, delimiter)

    likert_cols = detect_likert_columns_for_scale(df, scale=scale, min_valid_ratio=0.80)
    if not likert_cols:
        raise ValueError(f"No se detectaron columnas Likert válidas para escala 1 a {scale}.")

    likert_df = df[likert_cols].copy()
    for col in likert_cols:
        likert_df[col] = likert_df[col].map(lambda x: coerce_likert_cell_to_int(x, scale=scale))

    likert_df = likert_df.dropna()
    if likert_df.empty:
        raise ValueError("No hay suficientes datos válidos para análisis psicométrico.")

    return likert_df, scale, likert_cols


def compute_kmo_report(df: pd.DataFrame) -> dict:
    kmo_per_item, kmo_model = calculate_kmo(df)
    return {
        "kmo_global": float(kmo_model),
        "kmo_per_item": {
            col: float(val) for col, val in zip(df.columns, kmo_per_item)
        }
    }


def compute_bartlett_report(df: pd.DataFrame) -> dict:
    chi2, p_value = calculate_bartlett_sphericity(df)
    return {
        "chi2": float(chi2),
        "p_value": float(p_value),
    }


def interpret_kmo(kmo_global: float) -> str:
    if kmo_global >= 0.90:
        return "excelente"
    elif kmo_global >= 0.80:
        return "muy buena"
    elif kmo_global >= 0.70:
        return "buena"
    elif kmo_global >= 0.60:
        return "aceptable"
    elif kmo_global >= 0.50:
        return "débil"
    return "insuficiente"


def interpret_bartlett(p_value: float) -> str:
    if p_value < 0.05:
        return "La prueba de esfericidad de Bartlett fue significativa, lo que indica que la matriz de correlaciones es adecuada para análisis factorial."
    return "La prueba de esfericidad de Bartlett no fue significativa, por lo que no se recomienda realizar análisis factorial con estos datos."


def compute_efa_report(df: pd.DataFrame, n_factors: int | None = None, rotation: str = "varimax") -> dict:
    fa_test = FactorAnalyzer(rotation=None)
    fa_test.fit(df)

    ev, _ = fa_test.get_eigenvalues()
    eigenvalues = [float(x) for x in ev]

    if n_factors is None:
        n_factors = sum(1 for x in ev if x > 1.0)
        if n_factors < 1:
            n_factors = 1

    fa = FactorAnalyzer(n_factors=n_factors, rotation=rotation)
    fa.fit(df)

    loadings = fa.loadings_
    communalities = fa.get_communalities()
    variance = fa.get_factor_variance()

    loadings_df = pd.DataFrame(
        loadings,
        index=df.columns,
        columns=[f"Factor_{i+1}" for i in range(n_factors)]
    )

    communalities_df = pd.DataFrame({
        "Ítem": df.columns,
        "Comunalidad": communalities
    })

    variance_df = pd.DataFrame({
        "Factor": [f"Factor_{i+1}" for i in range(n_factors)],
        "SS_Loadings": variance[0],
        "Proportion_Var": variance[1],
        "Cumulative_Var": variance[2],
    })

    return {
        "n_factors": int(n_factors),
        "rotation": rotation,
        "eigenvalues": eigenvalues,
        "loadings_df": loadings_df,
        "communalities_df": communalities_df,
        "variance_df": variance_df,
    }


def make_scree_plot(df: pd.DataFrame, dataset_id: int) -> str | None:
    try:
        fa = FactorAnalyzer(rotation=None)
        fa.fit(df)
        ev, _ = fa.get_eigenvalues()

        fig = plt.figure()
        ax = fig.add_subplot(111)
        ax.plot(range(1, len(ev) + 1), ev, marker="o")
        ax.axhline(y=1.0, linestyle="--")
        ax.set_title("Scree Plot")
        ax.set_xlabel("Número de factores")
        ax.set_ylabel("Eigenvalue")

        out_name = f"ds{dataset_id}_scree_plot.png"
        out_path = os.path.join(PLOTS_DIR, out_name)
        fig.tight_layout()
        fig.savefig(out_path, dpi=150)
        plt.close(fig)

        return out_name
    except Exception as e:
        current_app.logger.warning(f"[scree_plot] ds{dataset_id}: {e}")
        return None


def make_factor_loadings_heatmap(loadings_df: pd.DataFrame, dataset_id: int) -> str | None:
    try:
        fig = plt.figure(figsize=(8, max(4, len(loadings_df) * 0.4)))
        ax = fig.add_subplot(111)

        im = ax.imshow(loadings_df.values, aspect="auto")
        ax.set_xticks(range(loadings_df.shape[1]))
        ax.set_xticklabels(loadings_df.columns, rotation=45, ha="right")
        ax.set_yticks(range(loadings_df.shape[0]))
        ax.set_yticklabels(loadings_df.index)
        ax.set_title("Cargas factoriales")
        fig.colorbar(im, ax=ax)

        out_name = f"ds{dataset_id}_factor_loadings.png"
        out_path = os.path.join(PLOTS_DIR, out_name)
        fig.tight_layout()
        fig.savefig(out_path, dpi=150)
        plt.close(fig)

        return out_name
    except Exception as e:
        current_app.logger.warning(f"[factor_loadings_heatmap] ds{dataset_id}: {e}")
        return None


def infer_factor_label_from_items(items: list[str]) -> str:
    if not items:
        return "Factor sin etiqueta"

    joined = " ".join(str(x).lower() for x in items)
    rules = [
        (["docente", "clase", "clases", "profesor", "claridad", "disponibilidad"], "Calidad docente"),
        (["organizacion", "programa", "contenido", "contenidos", "curricular"], "Organización académica"),
        (["tecnologia", "plataforma", "virtual", "digital"], "Recursos tecnológicos"),
        (["infraestructura", "aula", "aulas", "laboratorio", "instalacion"], "Infraestructura académica"),
        (["evaluacion", "aprendizaje", "rendimiento"], "Evaluación del aprendizaje"),
        (["satisfaccion", "percepcion", "valoracion"], "Satisfacción general"),
    ]

    for keywords, label in rules:
        if any(k in joined for k in keywords):
            return label

    return "Dimensión emergente"


def build_factor_structure_summary(efa_report: dict, threshold: float = 0.40) -> list[dict]:
    loadings = efa_report.get("loadings") or {}
    if not isinstance(loadings, dict) or not loadings:
        return []

    factor_names = []
    for _, factor_dict in loadings.items():
        factor_names = list((factor_dict or {}).keys())
        break

    grouped = {fname: [] for fname in factor_names}

    for item_name, factor_dict in loadings.items():
        if not isinstance(factor_dict, dict):
            continue

        best_factor = None
        best_value = None

        for fname, val in factor_dict.items():
            if isinstance(val, (int, float)):
                abs_val = abs(val)
                if best_value is None or abs_val > best_value:
                    best_value = abs_val
                    best_factor = fname

        if best_factor is not None and best_value is not None and best_value >= threshold:
            grouped[best_factor].append({
                "item": item_name,
                "loading": float(best_value),
            })

    result = []
    for fname, items in grouped.items():
        item_names = [x["item"] for x in items]
        label = infer_factor_label_from_items(item_names)

        result.append({
            "factor": fname,
            "label": label,
            "n_items": len(items),
            "items": items,
        })

    return result


def make_factor_model_diagram(dataset_id: int, factor_structure: list[dict]) -> str | None:
    from matplotlib.patches import FancyBboxPatch

    if not factor_structure:
        return None

    fig_w = 14
    fig_h = max(6, len(factor_structure) * 3.2)
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    n_factors = len(factor_structure)
    top_margin = 0.92
    bottom_margin = 0.08
    usable_h = top_margin - bottom_margin
    block_h = usable_h / max(1, n_factors)

    factor_x = 0.12
    item_x = 0.58
    factor_box_w = 0.24
    factor_box_h = 0.08
    item_box_w = 0.28
    item_box_h = 0.055

    for i, factor in enumerate(factor_structure):
        center_y = top_margin - (i + 0.5) * block_h

        factor_name = str(factor.get("factor", f"Factor_{i+1}"))
        factor_label = str(factor.get("label", "Dimensión"))
        factor_title = f"{factor_name}\n{factor_label}"

        factor_box = FancyBboxPatch(
            (factor_x, center_y - factor_box_h / 2),
            factor_box_w,
            factor_box_h,
            boxstyle="round,pad=0.02",
            linewidth=1.5,
            edgecolor="black",
            facecolor="#dbeafe",
        )
        ax.add_patch(factor_box)
        ax.text(
            factor_x + factor_box_w / 2,
            center_y,
            factor_title,
            ha="center",
            va="center",
            fontsize=10,
            fontweight="bold",
        )

        items = factor.get("items", []) or []
        if not items:
            continue

        n_items = len(items)
        item_total_h = n_items * item_box_h + max(0, n_items - 1) * 0.018
        start_y = center_y + item_total_h / 2 - item_box_h

        for j, item in enumerate(items):
            item_y = start_y - j * (item_box_h + 0.018)
            item_name = str(item.get("item", "Ítem"))
            loading = item.get("loading")
            item_text = item_name
            if isinstance(loading, (int, float)):
                item_text += f"\nλ={loading:.2f}"

            item_box = FancyBboxPatch(
                (item_x, item_y),
                item_box_w,
                item_box_h,
                boxstyle="round,pad=0.015",
                linewidth=1.0,
                edgecolor="black",
                facecolor="#fef3c7",
            )
            ax.add_patch(item_box)
            ax.text(
                item_x + item_box_w / 2,
                item_y + item_box_h / 2,
                item_text,
                ha="center",
                va="center",
                fontsize=8.5,
            )

            ax.annotate(
                "",
                xy=(item_x, item_y + item_box_h / 2),
                xytext=(factor_x + factor_box_w, center_y),
                arrowprops=dict(arrowstyle="->", lw=1.2),
            )

    ax.set_title("Modelo factorial inferido", fontsize=14, fontweight="bold", pad=12)

    out_name = f"ds{dataset_id}_factor_model.png"
    out_path = os.path.join(PLOTS_DIR, out_name)
    fig.tight_layout()
    fig.savefig(out_path, dpi=180, bbox_inches="tight")
    plt.close(fig)

    return out_name


def generate_factor_interpretation_text(efa_report: dict, threshold: float = 0.40) -> str:
    factors = build_factor_structure_summary(efa_report, threshold=threshold)
    if not factors:
        return "No fue posible identificar una estructura factorial interpretable a partir de las cargas obtenidas."

    parts = [
        "La interpretación de la estructura factorial sugiere la presencia de dimensiones conceptualmente diferenciadas dentro del instrumento."
    ]

    for f in factors:
        if not f["items"]:
            continue
        item_names = [x["item"] for x in f["items"]]
        item_text = ", ".join(item_names[:6])
        if len(item_names) > 6:
            item_text += ", entre otros"

        parts.append(
            f"{f['factor']} puede interpretarse como una dimensión de {f['label'].lower()}, "
            f"al agrupar principalmente los ítems: {item_text}."
        )

    parts.append(
        "En conjunto, esta organización respalda la existencia de componentes internos con significado teórico, "
        "lo cual fortalece la interpretación del instrumento en contextos de investigación."
    )
    return " ".join(parts)


def generate_psychometric_interpretation(kmo_report: dict, bartlett_report: dict, efa_report: dict) -> str:
    parts = []

    kmo_global = kmo_report.get("kmo_global")
    p_value = bartlett_report.get("p_value")
    n_factors = efa_report.get("n_factors")
    variance_df = efa_report.get("variance_df")

    if kmo_global is not None:
        parts.append(
            f"El índice KMO global fue de {kmo_global:.3f}, lo que indica una adecuación muestral {interpret_kmo(kmo_global)}."
        )

    if p_value is not None:
        parts.append(interpret_bartlett(p_value))

    if n_factors is not None:
        parts.append(
            f"El análisis factorial exploratorio sugirió una estructura de {n_factors} factor(es)."
        )

    if variance_df is not None and not variance_df.empty:
        last_cum = float(variance_df["Cumulative_Var"].iloc[-1]) * 100
        parts.append(
            f"La varianza acumulada explicada por la solución factorial fue de {last_cum:.2f}%."
        )

    parts.append(
        "En conjunto, estos resultados aportan evidencia sobre la estructura interna del instrumento y respaldan su utilidad para fines de investigación."
    )

    return " ".join(parts)


def ensure_advanced_psychometrics(dataset_id: int) -> dict:
    manifest_data = read_manifest_data(dataset_id) or {}
    meta = (manifest_data.get("meta") or {})
    if not isinstance(meta, dict):
        meta = {}

    has_all = (
        meta.get("kmo_report") and
        meta.get("bartlett_report") and
        meta.get("efa_report") and
        meta.get("psychometric_interpretation")
    )
    if has_all:
        return meta

    likert_df, scale, likert_cols = get_likert_analysis_df(dataset_id)

    kmo_report = compute_kmo_report(likert_df)
    bartlett_report = compute_bartlett_report(likert_df)
    efa_report = compute_efa_report(likert_df)

    scree_plot = make_scree_plot(likert_df, dataset_id)
    factor_heatmap = make_factor_loadings_heatmap(efa_report["loadings_df"], dataset_id)

    interpretation = generate_psychometric_interpretation(
        kmo_report, bartlett_report, efa_report
    )
    factor_interpretation = generate_factor_interpretation_text(efa_report)
    factor_structure = build_factor_structure_summary(efa_report)
    factor_model_plot = make_factor_model_diagram(dataset_id, factor_structure)

    meta.update({
        "likert_scale": scale,
        "likert_cols": likert_cols,
        "kmo_report": {
            "kmo_global": kmo_report["kmo_global"],
            "kmo_per_item": kmo_report["kmo_per_item"],
        },
        "bartlett_report": bartlett_report,
        "efa_report": {
            "n_factors": efa_report["n_factors"],
            "rotation": efa_report["rotation"],
            "eigenvalues": efa_report["eigenvalues"],
            "loadings": efa_report["loadings_df"].round(4).to_dict(),
            "communalities": efa_report["communalities_df"].round(4).to_dict(orient="records"),
            "variance": efa_report["variance_df"].round(4).to_dict(orient="records"),
            "factor_interpretation": factor_interpretation,
            "factor_structure": factor_structure,
            "factor_model_plot": factor_model_plot,
        },
        "factor_interpretation": factor_interpretation,
        "factor_structure": factor_structure,
        "factor_model_plot": factor_model_plot,
        "psychometric_interpretation": interpretation,
    })

    generated = manifest_data.get("generated", []) if isinstance(manifest_data, dict) else []
    for g in (scree_plot, factor_heatmap, factor_model_plot):
        if g:
            generated.append(g)

    generated = sorted(set(generated))
    write_manifest(dataset_id, generated, meta=meta)
    return meta


# =========================================================
# WORD / PDF / ARTÍCULO
# =========================================================

def set_paragraph_keep_with_next(paragraph):
    try:
        paragraph.paragraph_format.keep_with_next = True
    except Exception:
        pass


def add_logo_if_exists(doc: Document, logo_path: str | None, width_inches: float = 1.2):
    if not logo_path or not os.path.exists(logo_path):
        return

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(width_inches))
    except Exception:
        pass


def add_heading_keep(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    set_paragraph_keep_with_next(p)
    return p


def set_doc_base_styles(doc: Document):
    styles = doc.styles

    try:
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(11)
    except Exception:
        pass

    for style_name, size, bold in [
        ("Title", 18, True),
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
    ]:
        try:
            st = styles[style_name]
            st.font.name = "Times New Roman"
            st.font.size = Pt(size)
            st.font.bold = bold
        except Exception:
            pass


def add_cover_title(doc: Document, dataset, meta: dict, logo_path: str | None = None):
    add_logo_if_exists(doc, logo_path, width_inches=1.3)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME PSICOMÉTRICO")
    r.bold = True
    r.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(str(dataset.title or dataset.original_name))
    r2.font.size = Pt(13)

    doc.add_paragraph("")

    info = [
        f"Fecha: {datetime.now().strftime('%Y-%m-%d')}",
        f"Participantes: {meta.get('n_total', meta.get('n', '—'))}",
        f"Escala Likert: {meta.get('likert_scale', '—')} puntos",
        f"Número de ítems: {len(meta.get('likert_columns', []) or meta.get('likert_cols', []))}",
    ]

    for line in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)

    doc.add_paragraph("")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Generado automáticamente por SADI")
    r3.italic = True
    r3.font.size = Pt(10)

    doc.add_page_break()


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    if row_idx == 0:
                        run.bold = True


def build_psychometric_conclusion(meta: dict) -> str:
    alpha = meta.get("cronbach_alpha")
    if alpha is None:
        return (
            "No fue posible establecer una conclusión psicométrica robusta, "
            "debido a que el coeficiente alfa de Cronbach no pudo ser calculado."
        )

    quality = interpret_cronbach(alpha).lower()
    return (
        f"En términos generales, el instrumento presenta una consistencia interna {quality} "
        f"(α = {alpha:.4f}). Se recomienda interpretar estos resultados junto con el contenido "
        "conceptual de los ítems y los objetivos teóricos del instrumento."
    )


def build_psychometric_key_findings(meta: dict) -> list[str]:
    findings = []

    scale = meta.get("likert_scale")
    n_total = meta.get("n_total", meta.get("n"))
    alpha = meta.get("cronbach_alpha")

    if n_total is not None:
        findings.append(f"La muestra analizada estuvo conformada por {n_total} participantes.")
    if scale is not None:
        findings.append(f"El instrumento fue procesado como una escala Likert de {scale} puntos.")
    if alpha is not None:
        findings.append(f"El alfa de Cronbach global alcanzó un valor de {alpha:.4f}.")

    return findings[:6]


def generate_scientific_article(
    *,
    meta: dict,
    scale: int,
    n_items: int,
    n_cases: int,
    dataset_title: str | None = None,
) -> str:
    title = dataset_title or "Estudio Psicométrico"

    meta = meta or {}

    kmo = meta.get("kmo_report", {}) or {}
    bart = meta.get("bartlett_report", {}) or {}
    efa = meta.get("efa_report", {}) or {}

    psych_interp = (meta.get("psychometric_interpretation") or "").strip()
    factor_interp = (
        meta.get("factor_interpretation")
        or efa.get("factor_interpretation")
        or ""
    ).strip()

    cronbach_alpha = meta.get("cronbach_alpha")
    insights_text = (meta.get("insights_text") or "").strip()
    quick_recommendations = meta.get("quick_recommendations", []) or []
    suggested_plan = meta.get("suggested_plan", {}) or {}

    def fmt_num(v, decimals=4, default="—"):
        try:
            return f"{float(v):.{decimals}f}"
        except Exception:
            return default

    def interpret_alpha(alpha) -> str:
        try:
            a = float(alpha)
        except Exception:
            return "no concluyente"

        if a >= 0.90:
            return "excelente"
        if a >= 0.80:
            return "buena"
        if a >= 0.70:
            return "aceptable"
        if a >= 0.60:
            return "cuestionable"
        return "baja"

    def interpret_kmo_level(kmo_value) -> str:
        try:
            k = float(kmo_value)
        except Exception:
            return "no concluyente"

        if k >= 0.90:
            return "excelente"
        if k >= 0.80:
            return "muy buena"
        if k >= 0.70:
            return "buena"
        if k >= 0.60:
            return "aceptable"
        if k >= 0.50:
            return "débil"
        return "insuficiente"

    alpha_text = ""
    if cronbach_alpha is not None:
        alpha_text = (
            f"El instrumento presentó un alfa de Cronbach global de {fmt_num(cronbach_alpha)}, "
            f"lo que sugiere una consistencia interna {interpret_alpha(cronbach_alpha)}. "
        )

    kmo_text = ""
    if kmo.get("kmo_global") is not None:
        kmo_text = (
            f"El índice KMO global fue de {fmt_num(kmo.get('kmo_global'))}, "
            f"valor que puede interpretarse como una adecuación muestral {interpret_kmo_level(kmo.get('kmo_global'))}. "
        )

    bartlett_text = ""
    if bart.get("chi2") is not None and bart.get("p_value") is not None:
        bartlett_text = (
            f"La prueba de esfericidad de Bartlett reportó χ² = {fmt_num(bart.get('chi2'))} "
            f"y p = {fmt_num(bart.get('p_value'), decimals=6)}. "
        )

    efa_text = ""
    if efa.get("n_factors") is not None:
        efa_text = (
            f"El análisis factorial exploratorio sugirió una solución de {efa.get('n_factors')} factor(es), "
            "lo que aporta evidencia sobre la estructura interna del instrumento. "
        )

    narrative_focus = suggested_plan.get("narrative_focus")
    recommended_analysis = suggested_plan.get("recommended_analysis", []) or []

    text = f"{title}\n\n"

    text += "1. Introducción\n\n"
    text += (
        "El presente estudio tuvo como propósito analizar las propiedades psicométricas "
        "de un instrumento estructurado en formato Likert, con énfasis en su consistencia interna, "
        "adecuación muestral y estructura factorial. Este tipo de evaluación resulta fundamental "
        "para determinar la calidad métrica del instrumento y su pertinencia en contextos de investigación académica.\n\n"
    )

    text += "2. Método\n\n"
    text += (
        f"La muestra analizada estuvo compuesta por {n_cases} participantes. "
        f"El instrumento evaluado estuvo conformado por {n_items} ítems con una escala Likert de {scale} puntos. "
        "Se aplicaron procedimientos de análisis descriptivo y psicométrico, incluyendo el cálculo del alfa de Cronbach, "
        "el índice Kaiser-Meyer-Olkin (KMO), la prueba de esfericidad de Bartlett "
        "y el análisis factorial exploratorio (EFA), con el fin de examinar la calidad interna del instrumento.\n\n"
    )

    text += "3. Resultados\n\n"
    resultados = alpha_text + kmo_text + bartlett_text + efa_text
    if resultados.strip():
        text += resultados.strip() + "\n\n"

    if psych_interp:
        text += psych_interp + "\n\n"

    if factor_interp:
        text += factor_interp + "\n\n"

    if insights_text:
        text += (
            "Desde la perspectiva analítica automatizada de SADI, se obtuvo la siguiente síntesis interpretativa: "
            f"{insights_text}\n\n"
        )

    text += "4. Discusión\n\n"

    discussion_parts = []

    if cronbach_alpha is not None:
        discussion_parts.append(
            f"En primer lugar, la consistencia interna observada puede calificarse como {interpret_alpha(cronbach_alpha)}, "
            "lo que constituye un indicador relevante de coherencia entre los ítems del instrumento."
        )

    if kmo.get("kmo_global") is not None:
        discussion_parts.append(
            f"Asimismo, la adecuación muestral fue {interpret_kmo_level(kmo.get('kmo_global'))}, "
            "lo que respalda la pertinencia del análisis factorial sobre la matriz de datos."
        )

    if bart.get("p_value") is not None:
        try:
            pval = float(bart.get("p_value"))
            if pval < 0.05:
                discussion_parts.append(
                    "La significancia de la prueba de Bartlett refuerza la idea de que las correlaciones entre ítems "
                    "no son aleatorias, sino suficientemente estructuradas como para justificar un estudio factorial."
                )
            else:
                discussion_parts.append(
                    "La ausencia de significancia en la prueba de Bartlett sugiere cautela al interpretar la estructura factorial, "
                    "pues las asociaciones entre ítems podrían no ser suficientemente robustas."
                )
        except Exception:
            pass

    if efa.get("n_factors") is not None:
        discussion_parts.append(
            f"La solución factorial de {efa.get('n_factors')} factor(es) sugiere que el instrumento podría organizarse "
            "en dimensiones internas con significado teórico, lo que incrementa su valor interpretativo."
        )

    if narrative_focus:
        discussion_parts.append(
            f"Además, SADI sugiere enfocar la interpretación en el siguiente eje analítico: {narrative_focus}."
        )

    if not discussion_parts:
        discussion_parts.append(
            "Los resultados obtenidos ofrecen evidencia preliminar favorable sobre el comportamiento psicométrico del instrumento."
        )

    text += " ".join(discussion_parts) + "\n\n"

    if quick_recommendations:
        text += "5. Recomendaciones metodológicas\n\n"
        for rec in quick_recommendations:
            text += f"- {rec}\n"
        text += "\n"
        conclusion_number = "6"
    else:
        conclusion_number = "5"

    text += f"{conclusion_number}. Conclusión\n\n"

    conclusion_parts = []
    if cronbach_alpha is not None:
        conclusion_parts.append(
            f"En conjunto, el instrumento evidenció una consistencia interna {interpret_alpha(cronbach_alpha)}."
        )

    if efa.get("n_factors") is not None:
        conclusion_parts.append(
            f"Adicionalmente, la estructura factorial identificada ({efa.get('n_factors')} factor(es)) "
            "aporta evidencia a favor de la organización interna del instrumento."
        )

    if recommended_analysis:
        conclusion_parts.append(
            "Estos resultados justifican la continuidad de análisis complementarios orientados a profundizar "
            "la interpretación sustantiva del fenómeno estudiado."
        )

    if not conclusion_parts:
        conclusion_parts.append(
            "En conjunto, el instrumento presenta propiedades psicométricas útiles para su empleo en investigación académica."
        )

    text += " ".join(conclusion_parts).strip() + "\n"

    return text


def generate_likert_article_docx(dataset_id: int, summary: dict, output_path: str):
    summary = summary or {}
    if not isinstance(summary, dict):
        summary = {}

    manifest_data = read_manifest_data(dataset_id) or {}
    if not isinstance(manifest_data, dict):
        manifest_data = {}

    manifest_meta = manifest_data.get("meta") or {}
    if not isinstance(manifest_meta, dict):
        manifest_meta = {}

    analysis_meta = summary.get("analysis_meta", {}) or {}
    if not isinstance(analysis_meta, dict):
        analysis_meta = {}

    meta = dict(manifest_meta)
    meta.update(summary)
    meta.update(analysis_meta)

    try:
        if not (
            meta.get("kmo_report")
            and meta.get("bartlett_report")
            and meta.get("efa_report")
            and meta.get("psychometric_interpretation")
        ):
            advanced = ensure_advanced_psychometrics(dataset_id)
            if isinstance(advanced, dict):
                meta.update(advanced)
            manifest_data = read_manifest_data(dataset_id) or manifest_data
            meta2 = manifest_data.get("meta") or {}
            if isinstance(meta2, dict):
                meta.update(meta2)
    except Exception as e:
        current_app.logger.warning(f"[generate_likert_article_docx] ds{dataset_id}: {e}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    try:
        section = doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    except Exception:
        pass

    def _safe_list(v):
        if v is None:
            return []
        if isinstance(v, list):
            return v
        if isinstance(v, tuple):
            return list(v)
        return [v]

    def _safe_dict(v):
        return v if isinstance(v, dict) else {}

    def _safe_text(v, default="—"):
        if v is None:
            return default
        s = str(v).strip()
        return s if s else default

    def _winner_to_text(winner):
        if winner == "random_forest":
            return "Random Forest"
        if winner == "linear_regression":
            return "Regresión lineal"
        if winner == "tie":
            return "Rendimiento similar"
        return "No determinado"

    def _is_model_plot(name: str) -> bool:
        low = str(name).lower()
        model_tokens = [
            "rf_", "regression", "logistic", "pred_vs_real",
            "feature_importance", "residuals", "coefficients",
            "roc", "confusion", "classification"
        ]
        return any(tok in low for tok in model_tokens)

    def _add_section(title: str, text: str | None):
        if text is None:
            return
        text = str(text).strip()
        if not text:
            return
        doc.add_heading(title, level=1)
        doc.add_paragraph(text)

    def _add_bullets(title: str, items):
        items = _safe_list(items)
        if not items:
            return
        doc.add_heading(title, level=1)
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")

    def _add_table(title: str, headers, rows):
        rows = rows or []
        if not rows:
            return
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = str(h)
        for row_data in rows:
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = "" if val is None else str(val)

    def _add_plot_block(fig_label: str, fname: str, model_plot: bool = False):
        img_path = _resolve_plot(fname) if "_resolve_plot" in globals() else None
        if not img_path:
            return False

        rel = f"plots/{fname}"
        nice_title = prettify_plot_title(rel)
        nice_desc = describe_plot(rel)

        p_title = doc.add_paragraph()
        p_title.paragraph_format.keep_with_next = True
        r = p_title.add_run(f"{fig_label}. {nice_title}")
        r.bold = True
        r.font.size = Pt(12)

        p_type = doc.add_paragraph()
        p_type.paragraph_format.keep_with_next = True
        r2 = p_type.add_run("Tipo: Modelo predictivo" if model_plot else f"Tipo: {classify_plot_tag(rel)}")
        r2.italic = True

        if nice_desc:
            p_desc = doc.add_paragraph(nice_desc)
            p_desc.paragraph_format.keep_with_next = True

        inserted = safe_add_picture(doc, img_path, width_inches=6.0)
        if inserted:
            doc.add_paragraph("")
            return True
        return False

    # =========================
    # PLOTS UNIFICADOS
    # =========================
    try:
        plots = normalize_plot_catalog(dataset_id, manifest_data, summary)
        exploratory_plots, model_plots = split_plots(plots)
    except Exception as e:
        current_app.logger.warning(f"[plots normalization] ds{dataset_id}: {e}")
        plots = []
        exploratory_plots = []
        model_plots = []

    current_app.logger.warning(
        f"[ARTICLE PLOTS] ds{dataset_id}: total={len(plots)} | eda={len(exploratory_plots)} | model={len(model_plots)}"
    )

    # =========================
    # DATOS BASE
    # =========================
    dataset_title = (
        summary.get("dataset_title")
        or meta.get("dataset_name")
        or meta.get("title")
        or f"Dataset {dataset_id}"
    )

    scale = summary.get("scale") or meta.get("likert_scale") or 5
    n_items = meta.get("n_items", len(meta.get("likert_columns") or meta.get("likert_cols") or meta.get("columns") or []))
    n_cases = meta.get("n", summary.get("n", 0))
    n_total = meta.get("n_total", n_cases)
    cronbach_alpha = meta.get("cronbach_alpha")
    reliability_level = meta.get("reliability_level")
    global_mean = meta.get("global_mean")
    trend = meta.get("trend")
    insights = _safe_list(meta.get("insights"))
    quick_recommendations = _safe_list(meta.get("quick_recommendations"))
    suggested_plan = _safe_dict(meta.get("suggested_plan"))
    next_step = _safe_dict(meta.get("next_step_recommendation"))
    warnings = _safe_list(meta.get("warnings"))
    strong_dimensions = _safe_list(meta.get("strong_dimensions"))
    weak_dimensions = _safe_list(meta.get("weak_dimensions"))

    # =========================
    # PORTADA
    # =========================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Artículo científico automático")
    r.bold = True
    r.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Análisis de encuesta Likert")
    r2.italic = True
    r2.font.size = Pt(12)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(dataset_title)
    r3.font.size = Pt(12)

    doc.add_paragraph("")

    # =========================
    # RESUMEN
    # =========================
    doc.add_heading("RESUMEN", level=1)
    abstract_parts = [
        f"Se presenta un análisis automatizado del instrumento Likert '{dataset_title}'.",
        f"El cuestionario fue analizado con una escala de {scale} puntos, {n_items} ítems y {n_cases} respuestas válidas."
    ]
    if cronbach_alpha is not None:
        abstract_parts.append(
            f"La consistencia interna alcanzó un alfa de Cronbach de {cronbach_alpha:.4f}, interpretado como {reliability_level}."
        )
    if global_mean is not None:
        abstract_parts.append(
            f"La media global observada fue {global_mean:.3f}, lo que refleja una tendencia {trend} en las respuestas."
        )
    if strong_dimensions:
        abstract_parts.append(f"Las dimensiones más fuertes fueron: {', '.join(map(str, strong_dimensions))}.")
    if weak_dimensions:
        abstract_parts.append(f"Las dimensiones con menor desempeño fueron: {', '.join(map(str, weak_dimensions))}.")
    doc.add_paragraph(" ".join(abstract_parts))

    # =========================
    # INTRODUCCIÓN
    # =========================
    doc.add_heading("1. INTRODUCCIÓN", level=1)
    intro = (
        f"El análisis de instrumentos tipo Likert constituye una fase esencial en estudios académicos basados en percepción, "
        f"satisfacción y evaluación de procesos. En este trabajo se analiza el instrumento '{dataset_title}' mediante SADI, "
        f"integrando lectura psicométrica, hallazgos automáticos, interpretación metodológica y, cuando aplica, modelado predictivo."
    )
    doc.add_paragraph(intro)

    # =========================
    # METODOLOGÍA
    # =========================
    doc.add_heading("2. METODOLOGÍA", level=1)
    methodology = (
        "Se aplicó un análisis automatizado que incluyó detección de ítems Likert, cálculo de medias por ítem y dimensión, "
        "consistencia interna mediante alfa de Cronbach, visualización exploratoria, identificación de dimensiones fuertes y débiles, "
        "y reconstrucción de resultados predictivos disponibles."
    )
    doc.add_paragraph(methodology)

    # =========================
    # RESULTADOS
    # =========================
    doc.add_heading("3. RESULTADOS", level=1)
    result_text = []
    if cronbach_alpha is not None:
        result_text.append(
            f"El instrumento presentó un alfa de Cronbach de {cronbach_alpha:.4f}, con una confiabilidad interpretada como {reliability_level}."
        )
    if global_mean is not None:
        result_text.append(
            f"La media global del instrumento fue {global_mean:.3f}, indicando una tendencia {trend} en las respuestas."
        )
    if strong_dimensions:
        result_text.append(f"Entre las dimensiones mejor valoradas se identificaron: {', '.join(map(str, strong_dimensions))}.")
    if weak_dimensions:
        result_text.append(f"Las dimensiones con mayor necesidad de revisión fueron: {', '.join(map(str, weak_dimensions))}.")
    doc.add_paragraph(" ".join(result_text) if result_text else "No se identificaron resultados narrativos suficientes.")

    _add_bullets("4. INSIGHTS AUTOMÁTICOS SADI", insights)
    _add_section("5. CONCLUSIÓN AUTOMÁTICA", meta.get("insights_text"))

    if next_step:
        doc.add_heading("6. PRÓXIMO PASO RECOMENDADO POR SADI", level=1)
        doc.add_paragraph(f"Título: {_safe_text(next_step.get('title'))}")
        doc.add_paragraph(f"Por qué: {_safe_text(next_step.get('reason'))}")
        doc.add_paragraph(f"Siguiente acción: {_safe_text(next_step.get('action'))}")

    doc.add_heading("7. LECTURA METODOLÓGICA DEL INSTRUMENTO", level=1)
    if strong_dimensions:
        doc.add_paragraph(f"Dimensiones fuertes: {', '.join(map(str, strong_dimensions))}.")
    if weak_dimensions:
        doc.add_paragraph(f"Dimensiones débiles: {', '.join(map(str, weak_dimensions))}.")
    if warnings:
        doc.add_paragraph("Advertencias metodológicas:")
        for w in warnings:
            doc.add_paragraph(str(w), style="List Bullet")

    _add_bullets("8. RECOMENDACIONES RÁPIDAS", quick_recommendations)

    if suggested_plan:
        doc.add_heading("9. PLAN DE ANÁLISIS SUGERIDO", level=1)

        recommended_analysis = _safe_list(suggested_plan.get("recommended_analysis"))
        recommended_plots = _safe_list(suggested_plan.get("recommended_plots"))
        narrative_focus = suggested_plan.get("narrative_focus")
        plan_warnings = _safe_list(suggested_plan.get("warnings"))

        if recommended_analysis:
            doc.add_paragraph("Análisis sugeridos:")
            for x in recommended_analysis:
                doc.add_paragraph(str(x), style="List Bullet")

        if recommended_plots:
            doc.add_paragraph("Gráficos sugeridos:")
            for x in recommended_plots:
                doc.add_paragraph(str(x), style="List Bullet")

        if narrative_focus:
            doc.add_paragraph(f"Enfoque narrativo recomendado: {narrative_focus}")

        if plan_warnings:
            doc.add_paragraph("Advertencias:")
            for x in plan_warnings:
                doc.add_paragraph(str(x), style="List Bullet")

    dimension_ranking = _safe_list(meta.get("dimension_ranking"))
    if dimension_ranking:
        rows = []
        for row_data in dimension_ranking:
            if isinstance(row_data, dict):
                rows.append([
                    row_data.get("dimension", "—"),
                    row_data.get("mean", "—")
                ])
        _add_table("10. RANKING DE DIMENSIONES", ["Dimensión", "Media"], rows)

    top_numeric = _safe_list(meta.get("top_numeric_by_variability"))
    if top_numeric:
        rows = []
        for row in top_numeric:
            if isinstance(row, dict):
                rows.append([
                    row.get("column", "—"),
                    row.get("std", "—"),
                    row.get("range", "—")
                ])
        _add_table("11. VARIABLES MÁS VARIABLES", ["Variable", "Desv. estándar", "Rango"], rows)

    variable_importance = _safe_list(meta.get("variable_importance"))
    if variable_importance:
        rows = []
        for row in variable_importance:
            if isinstance(row, dict):
                rows.append([
                    row.get("column", "—"),
                    row.get("score", "—"),
                    row.get("std", "—"),
                    f"{row.get('missing_pct', '—')}%"
                ])
        _add_table("12. VARIABLES CLAVE DEL DATASET", ["Variable", "Score", "Desv. estándar", "% Missing"], rows)

    if meta.get("target_candidate"):
        doc.add_heading("13. VARIABLE OBJETIVO SUGERIDA", level=1)
        doc.add_paragraph(f"Variable detectada: {_safe_text(meta.get('target_candidate'))}")
        doc.add_paragraph(f"Tipo de problema: {_safe_text(meta.get('target_type'), 'No definido')}")
        if meta.get("target_reason"):
            doc.add_paragraph(str(meta.get("target_reason")))

    ranked_target_candidates = _safe_list(meta.get("ranked_target_candidates"))
    if ranked_target_candidates:
        rows = []
        for row in ranked_target_candidates:
            if isinstance(row, dict):
                rows.append([
                    row.get("column", "—"),
                    row.get("type", "—"),
                    row.get("score", "—"),
                    row.get("reason", "—"),
                ])
        _add_table("14. CANDIDATOS A VARIABLE OBJETIVO", ["Variable", "Tipo", "Score", "Motivo"], rows)

    if meta.get("target_candidate") and meta.get("model_suggestion"):
        doc.add_heading("15. MODELO SUGERIDO POR SADI", level=1)
        doc.add_paragraph(f"Variable objetivo detectada: {_safe_text(meta.get('target_candidate'))}")
        doc.add_paragraph(f"Tipo de problema: {_safe_text(meta.get('target_type'))}")
        doc.add_paragraph(f"Modelo sugerido: {_safe_text(meta.get('model_suggestion'))}")

    # =========================
    # GRÁFICOS EXPLORATORIOS SOLO UNA VEZ
    # =========================
    if exploratory_plots:
        doc.add_heading("Anexo A. Gráficos exploratorios", level=1)
        fig_num = 1
        for rel in exploratory_plots:
            fname = rel.split("/")[-1]
            if _add_plot_block(f"Figura A{fig_num}", fname, model_plot=False):
                fig_num += 1

    # =========================
    # CAPÍTULO 16
    # =========================
    regression_result = _safe_dict(meta.get("regression_result"))
    rf_result = _safe_dict(meta.get("rf_result"))
    model_comparison = _safe_dict(meta.get("model_comparison"))
    best_model_selection = _safe_dict(meta.get("best_model_selection"))
    model_interpretation = _safe_list(meta.get("model_interpretation"))
    advanced_model_interpretation = _safe_list(meta.get("advanced_model_interpretation"))

    has_predictive = any([
        regression_result,
        rf_result,
        model_comparison,
        best_model_selection,
        model_interpretation,
        advanced_model_interpretation,
        meta.get("target_candidate"),
        meta.get("model_suggestion"),
    ])

    if has_predictive:
        doc.add_heading("16. RESULTADOS DEL MODELO PREDICTIVO", level=1)

        doc.add_heading("16.1 Contexto del modelado", level=2)
        doc.add_paragraph(f"Variable objetivo detectada: {_safe_text(meta.get('target_candidate'), 'No detectado')}")
        doc.add_paragraph(f"Tipo de problema: {_safe_text(meta.get('target_type'), 'No definido')}")
        doc.add_paragraph(f"Modelo sugerido por SADI: {_safe_text(meta.get('model_suggestion'), 'No sugerido')}")
        doc.add_paragraph(f"Justificación automática: {_safe_text(meta.get('target_reason'), 'No disponible.')}")

        doc.add_heading("16.2 Evaluación automática del modelo", level=2)

        evaluation_items = []
        if advanced_model_interpretation:
            evaluation_items.extend(_safe_list(advanced_model_interpretation))
        elif model_interpretation:
            evaluation_items.extend(_safe_list(model_interpretation))
        elif best_model_selection.get("reason"):
            evaluation_items.append(best_model_selection.get("reason"))
        elif model_comparison.get("recommendation"):
            evaluation_items.append(model_comparison.get("recommendation"))

        if evaluation_items:
            for item in evaluation_items:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph("No se generó evaluación automática detallada del modelo.")

        doc.add_heading("16.3 Comparación automática de modelos", level=2)

        linear = _safe_dict(model_comparison.get("linear"))
        random_forest = _safe_dict(model_comparison.get("random_forest"))
        comparison_available = bool(model_comparison.get("available"))

        if not comparison_available and (regression_result or rf_result):
            linear = {
                "r2": regression_result.get("r2"),
                "mae": regression_result.get("mae"),
                "rmse": regression_result.get("rmse"),
            }
            random_forest = {
                "r2": rf_result.get("r2"),
                "mae": rf_result.get("mae"),
                "rmse": rf_result.get("rmse"),
            }
            comparison_available = True

        if comparison_available:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Modelo"
            hdr[1].text = "R²"
            hdr[2].text = "MAE"
            hdr[3].text = "RMSE"

            row = table.add_row().cells
            row[0].text = "Regresión lineal"
            row[1].text = _safe_text(linear.get("r2"))
            row[2].text = _safe_text(linear.get("mae"))
            row[3].text = _safe_text(linear.get("rmse"))

            row = table.add_row().cells
            row[0].text = "Random Forest"
            row[1].text = _safe_text(random_forest.get("r2"))
            row[2].text = _safe_text(random_forest.get("mae"))
            row[3].text = _safe_text(random_forest.get("rmse"))

            winner_value = best_model_selection.get("best_model") or model_comparison.get("winner")
            winner_txt = _winner_to_text(winner_value)

            recommendation_text = (
                model_comparison.get("recommendation")
                or best_model_selection.get("reason")
                or "No disponible."
            )

            doc.add_paragraph(f"Modelo recomendado por SADI: {winner_txt}")
            doc.add_paragraph(f"Interpretación comparativa: {_safe_text(recommendation_text, 'No disponible.')}")

            summary_lines = _safe_list(model_comparison.get("summary"))
            if not summary_lines and best_model_selection.get("comparison_summary"):
                summary_lines = _safe_list(best_model_selection.get("comparison_summary"))

            if summary_lines:
                for line in summary_lines:
                    doc.add_paragraph(str(line), style="List Bullet")
        else:
            if str(meta.get("target_type")).lower() == "classification":
                doc.add_paragraph(
                    "La comparación automática entre modelos de regresión no aplica en este caso, "
                    "ya que el problema detectado es de clasificación."
                )
            else:
                doc.add_paragraph(
                    "No se dispuso de resultados suficientes para construir una comparación automática completa entre modelos."
                )

        doc.add_heading("16.4 Selección inteligente del mejor modelo", level=2)

        if best_model_selection:
            winner_txt = _winner_to_text(best_model_selection.get('best_model'))
            doc.add_paragraph(f"Variable objetivo: {_safe_text(best_model_selection.get('target_col'), _safe_text(meta.get('target_candidate')))}")
            doc.add_paragraph(f"Tipo de problema: {_safe_text(best_model_selection.get('problem_type'), _safe_text(meta.get('target_type')))}")
            doc.add_paragraph(f"Modelo preferido por SADI: {winner_txt}")
            doc.add_paragraph(f"Razón: {_safe_text(best_model_selection.get('reason'), 'No disponible')}")

            comparison_summary = _safe_list(best_model_selection.get("comparison_summary"))
            if comparison_summary:
                for line in comparison_summary:
                    doc.add_paragraph(str(line), style="List Bullet")
        else:
            winner_txt = _winner_to_text(model_comparison.get('winner'))
            if winner_txt != "No determinado":
                doc.add_paragraph(f"Modelo recomendado por SADI: {winner_txt}")
                doc.add_paragraph(f"Razón: {_safe_text(model_comparison.get('recommendation'), 'No disponible')}")
            else:
                doc.add_paragraph("No se pudo determinar automáticamente un mejor modelo.")

    # =========================
    # GRÁFICOS DE MODELO SOLO UNA VEZ
    # =========================
    if model_plots:
        doc.add_heading("Anexo B. Gráficos del modelo predictivo", level=1)
        fig_num = 1
        for rel in model_plots:
            fname = rel.split("/")[-1]
            if _add_plot_block(f"Figura B{fig_num}", fname, model_plot=True):
                fig_num += 1

    doc.save(output_path)

def generate_psychometric_word(*, dataset, meta, plots):
    meta = meta or {}
    if not isinstance(meta, dict):
        meta = {}

    analysis_meta = meta.get("analysis_meta", {}) or {}
    if not isinstance(analysis_meta, dict):
        analysis_meta = {}

    # fusionar analysis_meta sobre meta solo para lectura del Word
    merged = dict(meta)
    merged.update(analysis_meta)

    try:
        if not (
            merged.get("kmo_report")
            and merged.get("bartlett_report")
            and merged.get("efa_report")
            and merged.get("psychometric_interpretation")
        ):
            try:
                advanced = ensure_advanced_psychometrics(dataset.id)
                if isinstance(advanced, dict):
                    merged.update(advanced)
            except Exception as e:
                current_app.logger.warning(f"[generate_psychometric_word] ds{dataset.id}: {e}")
    except Exception as e:
        current_app.logger.warning(f"[generate_psychometric_word] ds{dataset.id}: {e}")

    doc = Document()
    set_doc_base_styles(doc)

    try:
        section = doc.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    except Exception:
        pass

    static_root = _static_root()
    logo_path = os.path.join(static_root, "logo2.ico")
    if not os.path.exists(logo_path):
        alt_logo = os.path.join(static_root, "logo2.png")
        logo_path = alt_logo if os.path.exists(alt_logo) else None

    def _safe_list(v):
        if v is None:
            return []
        if isinstance(v, list):
            return v
        if isinstance(v, tuple):
            return list(v)
        return [v]

    def _safe_dict(v):
        return v if isinstance(v, dict) else {}

    def _safe_text(v, default="—"):
        if v is None:
            return default
        s = str(v).strip()
        return s if s else default

    def _winner_to_text(winner):
        if winner == "random_forest":
            return "Random Forest"
        if winner == "linear_regression":
            return "Regresión lineal"
        if winner == "tie":
            return "Rendimiento similar"
        return "No determinado"

    def _add_kv_table(title, rows):
        rows = rows or []
        if not rows:
            return
        add_heading_keep(doc, title, level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Indicador"
        hdr[1].text = "Valor"
        for a, b in rows:
            row = table.add_row().cells
            row[0].text = str(a)
            row[1].text = str(b)

    def _add_table(title, headers, rows):
        rows = rows or []
        if not rows:
            return
        if title:
            add_heading_keep(doc, title, level=1)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = str(h)
        for row_data in rows:
            row = table.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = "" if val is None else str(val)

    def _add_bullets(title, items):
        items = _safe_list(items)
        if not items:
            return
        add_heading_keep(doc, title, level=1)
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")

    try:
        plots = filter_plots_for_dataset(dataset.id, plots)
    except Exception:
        plots = []

    all_plot_candidates = list(plots or [])
    advanced_plot_names = [
        f"ds{dataset.id}_scree_plot.png",
        f"ds{dataset.id}_factor_loadings.png",
        f"ds{dataset.id}_factor_model.png",
        f"ds{dataset.id}_likert_divergent.png",
        f"ds{dataset.id}_cronbach.png",
        f"ds{dataset.id}_corr_heatmap.png",
        f"ds{dataset.id}_likert_summary_scores.png",
        f"ds{dataset.id}_dimension_summary_scores.png",
        f"ds{dataset.id}_dimension_radar.png",
        f"ds{dataset.id}_multivariate_corr.png",
        f"ds{dataset.id}_regression_coefficients.png",
        f"ds{dataset.id}_regression_pred_vs_real.png",
        f"ds{dataset.id}_regression_residuals.png",
        f"ds{dataset.id}_rf_feature_importance.png",
        f"ds{dataset.id}_rf_pred_vs_real.png",
        f"ds{dataset.id}_roc_curve.png",
        f"ds{dataset.id}_confusion_matrix.png",
    ]
    for ap in advanced_plot_names:
        if ap not in all_plot_candidates and _resolve_plot(ap):
            all_plot_candidates.append(ap)

    valid_plots = []
    used = set()
    for p in all_plot_candidates:
        rel = str(p).replace("\\", "/").strip()
        fname = rel.split("/")[-1]
        if fname in used:
            continue
        if _resolve_plot(fname):
            valid_plots.append(fname)
            used.add(fname)

    plot_summary = summarize_plot_tags([f"plots/{x}" for x in valid_plots]) if valid_plots else []

    def _is_model_plot(name: str) -> bool:
        low = str(name).lower()
        model_tokens = [
            "rf_", "regression", "logistic", "pred_vs_real",
            "feature_importance", "residuals", "coefficients",
            "roc", "confusion", "classification"
        ]
        return any(tok in low for tok in model_tokens)

    exploratory_plots = [x for x in valid_plots if not _is_model_plot(x)]
    model_plots = [x for x in valid_plots if _is_model_plot(x)]

    add_cover_title(doc, dataset, merged, logo_path=logo_path)

    # Resumen ejecutivo
    survey_insights = (
        merged.get("survey_insights")
        or merged.get("results_text")
        or ""
    )
    insights = _safe_list(merged.get("insights"))
    if str(survey_insights).strip() or insights:
        add_heading_keep(doc, "Resumen ejecutivo", level=1)
        if str(survey_insights).strip():
            for para in str(survey_insights).split("\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)
        if insights:
            doc.add_paragraph("Insights avanzados SADI:")
            for item in insights:
                doc.add_paragraph(str(item), style="List Bullet")

    # Conclusión automática
    if merged.get("insights_text"):
        add_heading_keep(doc, "Conclusión automática", level=1)
        doc.add_paragraph(str(merged.get("insights_text")))

    # Próximo paso recomendado
    next_step = _safe_dict(merged.get("next_step_recommendation"))
    if next_step:
        add_heading_keep(doc, "Próximo paso recomendado por SADI", level=1)
        doc.add_paragraph(f"Título: {_safe_text(next_step.get('title'))}")
        doc.add_paragraph(f"Por qué: {_safe_text(next_step.get('reason'))}")
        doc.add_paragraph(f"Siguiente acción: {_safe_text(next_step.get('action'))}")

    # 1. Perfil del instrumento
    likert_cols = merged.get("likert_columns", []) or merged.get("likert_cols", []) or merged.get("columns", []) or []
    likert_scale = merged.get("likert_scale", merged.get("scale", "—"))
    n_items = merged.get("n_items", len(likert_cols))
    n_dimensions = merged.get("n_dimensions", "—")
    n_total = merged.get("n_total", merged.get("n", "—"))
    n_valid = merged.get("n_valid", merged.get("n", "—"))
    response_rate = merged.get("response_rate")
    data_quality = merged.get("data_quality", "—")
    reliability_level = merged.get("reliability_level", "—")
    alpha = merged.get("cronbach_alpha")

    profile_rows = [
        ("Tipo de instrumento", "Likert"),
        ("Escala", likert_scale),
        ("Número de ítems", n_items),
        ("Número de dimensiones", n_dimensions),
        ("Participantes totales", n_total),
        ("Casos válidos", n_valid),
        ("Calidad de datos", data_quality),
        ("Nivel de confiabilidad", reliability_level),
    ]
    if response_rate is not None:
        profile_rows.append(("Tasa de respuesta válida", f"{response_rate * 100:.2f}%"))
    if alpha is not None:
        profile_rows.append(("Alfa de Cronbach", f"{alpha:.4f}"))
    if merged.get("trend") is not None:
        profile_rows.append(("Tendencia general", merged.get("trend")))
    _add_kv_table("1. Perfil del instrumento", profile_rows)

    # 2. Detalle técnico
    num_cols = merged.get("num_cols", []) or []
    cat_cols = merged.get("cat_cols", []) or []
    dt_cols = merged.get("dt_cols", []) or []
    detail_rows = [
        ("Variables numéricas", ", ".join(map(str, num_cols)) if num_cols else "—"),
        ("Variables categóricas", ", ".join(map(str, cat_cols)) if cat_cols else "—"),
        ("Variables de fecha", ", ".join(map(str, dt_cols)) if dt_cols else "—"),
    ]
    _add_kv_table("2. Detalle técnico", detail_rows)

    # 3. Hallazgos clave
    findings = build_psychometric_key_findings(merged)
    extra_findings = []

    strong_dimensions = _safe_list(merged.get("strong_dimensions"))
    weak_dimensions = _safe_list(merged.get("weak_dimensions"))
    if strong_dimensions:
        extra_findings.append(f"Dimensiones fuertes: {', '.join(map(str, strong_dimensions))}.")
    if weak_dimensions:
        extra_findings.append(f"Dimensiones débiles: {', '.join(map(str, weak_dimensions))}.")

    for i in insights[:6]:
        extra_findings.append(str(i))

    merged_findings = []
    seen = set()
    for f in list(findings or []) + extra_findings:
        key = str(f).strip()
        if key and key not in seen:
            merged_findings.append(key)
            seen.add(key)
    _add_bullets("3. Hallazgos clave", merged_findings)

    # 4. Consistencia interna
    add_heading_keep(doc, "4. Consistencia interna", level=1)
    if alpha is not None:
        doc.add_paragraph(f"Alfa de Cronbach global: {alpha:.4f}")
    if merged.get("cronbach_text"):
        doc.add_paragraph(str(merged.get("cronbach_text")))
    if reliability_level:
        doc.add_paragraph(f"Nivel de confiabilidad interpretado por SADI: {reliability_level}.")

    # 5. Ranking de dimensiones
    dimension_ranking = _safe_list(merged.get("dimension_ranking"))
    if dimension_ranking:
        rows = []
        for d in dimension_ranking:
            if isinstance(d, dict):
                rows.append([d.get("dimension", "—"), d.get("mean", "—")])
        _add_table("5. Ranking de dimensiones", ["Dimensión", "Media"], rows)

    # 6. Variables más variables
    top_numeric = _safe_list(merged.get("top_numeric_by_variability"))
    if top_numeric:
        rows = []
        for row in top_numeric:
            if isinstance(row, dict):
                rows.append([row.get("column", "—"), row.get("std", "—"), row.get("range", "—")])
        _add_table("6. Variables más variables", ["Variable", "Desv. estándar", "Rango"], rows)

    # 7. Variables clave
    variable_importance = _safe_list(merged.get("variable_importance"))
    if variable_importance:
        rows = []
        for row in variable_importance:
            if isinstance(row, dict):
                rows.append([
                    row.get("column", "—"),
                    row.get("score", "—"),
                    row.get("std", "—"),
                    f"{row.get('missing_pct', '—')}%"
                ])
        _add_table("7. Variables clave del dataset", ["Variable", "Score", "Desv. estándar", "% Missing"], rows)

    # 8. Variable objetivo sugerida
    if merged.get("target_candidate"):
        add_heading_keep(doc, "8. Variable objetivo sugerida", level=1)
        doc.add_paragraph(f"Variable detectada: {_safe_text(merged.get('target_candidate'))}")
        target_type = merged.get("target_type")
        if target_type == "regression":
            doc.add_paragraph("Tipo de problema: Regresión")
        elif target_type == "classification":
            doc.add_paragraph("Tipo de problema: Clasificación")
        else:
            doc.add_paragraph(f"Tipo de problema: {_safe_text(target_type, 'No definido')}")
        if merged.get("target_reason"):
            doc.add_paragraph(str(merged.get("target_reason")))

    # 9. Candidatos a variable objetivo
    ranked_target_candidates = _safe_list(merged.get("ranked_target_candidates"))
    if ranked_target_candidates:
        rows = []
        for row in ranked_target_candidates:
            if isinstance(row, dict):
                rows.append([
                    row.get("column", "—"),
                    row.get("type", "—"),
                    row.get("score", "—"),
                    row.get("reason", "—"),
                ])
        _add_table("9. Candidatos a variable objetivo", ["Variable", "Tipo", "Score", "Motivo"], rows)

    # 10. Modelo sugerido
    if merged.get("target_candidate") and merged.get("model_suggestion"):
        add_heading_keep(doc, "10. Modelo sugerido por SADI", level=1)
        doc.add_paragraph(f"Variable objetivo detectada: {_safe_text(merged.get('target_candidate'))}")
        if merged.get("target_type") == "regression":
            doc.add_paragraph("Tipo de problema: Regresión")
        elif merged.get("target_type") == "classification":
            doc.add_paragraph("Tipo de problema: Clasificación")
        doc.add_paragraph(f"Modelo sugerido: {_safe_text(merged.get('model_suggestion'))}")

    # 11. Recomendaciones automáticas SADI
    quick_recommendations = _safe_list(merged.get("quick_recommendations"))
    suggested_plan = _safe_dict(merged.get("suggested_plan"))
    warnings = _safe_list(merged.get("warnings"))
    if quick_recommendations or suggested_plan or warnings:
        add_heading_keep(doc, "11. Recomendaciones automáticas SADI", level=1)

        if quick_recommendations:
            doc.add_paragraph("Recomendaciones rápidas:")
            for r in quick_recommendations:
                doc.add_paragraph(str(r), style="List Bullet")

        if suggested_plan:
            ra = _safe_list(suggested_plan.get("recommended_analysis"))
            rp = _safe_list(suggested_plan.get("recommended_plots"))
            nf = suggested_plan.get("narrative_focus")
            sw = _safe_list(suggested_plan.get("warnings"))

            if ra:
                doc.add_paragraph("Análisis sugeridos:")
                for x in ra:
                    doc.add_paragraph(str(x), style="List Bullet")

            if rp:
                doc.add_paragraph("Gráficos sugeridos:")
                for x in rp:
                    doc.add_paragraph(str(x), style="List Bullet")

            if nf:
                doc.add_paragraph(f"Enfoque recomendado: {nf}")

            if sw:
                doc.add_paragraph("Advertencias interpretativas:")
                for x in sw:
                    doc.add_paragraph(str(x), style="List Bullet")

        if warnings:
            doc.add_paragraph("Advertencias metodológicas adicionales:")
            for w in warnings:
                doc.add_paragraph(str(w), style="List Bullet")

    # 12. Interpretación psicométrica
    if merged.get("psychometric_interpretation"):
        add_heading_keep(doc, "12. Interpretación psicométrica", level=1)
        doc.add_paragraph(str(merged.get("psychometric_interpretation")))

    # 13. Resumen exploratorio visual
    if plot_summary:
        add_heading_keep(doc, "13. Resumen exploratorio visual", level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Tipo"
        hdr[1].text = "Cantidad"
        for item in plot_summary:
            row = table.add_row().cells
            row[0].text = str(item.get("tag", "Analítico"))
            row[1].text = str(item.get("count", 0))
        style_table(table)

    # 14. Gráficos exploratorios
    if exploratory_plots:
        add_heading_keep(doc, "14. Gráficos exploratorios", level=1)
        figure_num = 1
        for fname in exploratory_plots:
            img_path = _resolve_plot(fname)
            if not img_path:
                continue

            rel = f"plots/{fname}"
            nice_title = prettify_plot_title(rel)
            nice_desc = describe_plot(rel)
            nice_tag = classify_plot_tag(rel)

            p_title = doc.add_paragraph()
            p_title.paragraph_format.keep_with_next = True
            r = p_title.add_run(f"Gráfico {figure_num}. {nice_title}")
            r.bold = True
            r.font.size = Pt(12)

            p_type = doc.add_paragraph()
            p_type.paragraph_format.keep_with_next = True
            r2 = p_type.add_run(f"Tipo: {nice_tag}")
            r2.italic = True

            p_desc = doc.add_paragraph(nice_desc)
            p_desc.paragraph_format.keep_with_next = True

            inserted = safe_add_picture(doc, img_path, width_inches=6.0)
            if inserted:
                doc.add_paragraph("")
                figure_num += 1

    # 15. Resultados del modelo predictivo
    regression_result = _safe_dict(merged.get("regression_result"))
    rf_result = _safe_dict(merged.get("rf_result"))
    model_comparison = _safe_dict(merged.get("model_comparison"))
    best_model_selection = _safe_dict(merged.get("best_model_selection"))
    model_interpretation = _safe_list(merged.get("model_interpretation"))
    advanced_model_interpretation = _safe_list(merged.get("advanced_model_interpretation"))

    has_predictive = any([
        regression_result,
        rf_result,
        model_comparison,
        best_model_selection,
        model_interpretation,
        advanced_model_interpretation,
        merged.get("target_candidate"),
        merged.get("model_suggestion"),
    ])

    if has_predictive:
        add_heading_keep(doc, "15. Resultados del modelo predictivo", level=1)

        add_heading_keep(doc, "15.1 Contexto del modelado", level=2)
        doc.add_paragraph(f"Variable objetivo detectada: {_safe_text(merged.get('target_candidate'), 'No detectado')}")
        doc.add_paragraph(f"Tipo de problema: {_safe_text(merged.get('target_type'), 'No definido')}")
        doc.add_paragraph(f"Modelo sugerido por SADI: {_safe_text(merged.get('model_suggestion'), 'No sugerido')}")
        doc.add_paragraph(f"Justificación automática: {_safe_text(merged.get('target_reason'), 'No disponible.')}")

        add_heading_keep(doc, "15.2 Resumen de regresión lineal", level=2)
        if regression_result:
            doc.add_paragraph(f"R²: {_safe_text(regression_result.get('r2'))}")
            doc.add_paragraph(f"MAE: {_safe_text(regression_result.get('mae'))}")
            doc.add_paragraph(f"RMSE: {_safe_text(regression_result.get('rmse'))}")
        else:
            doc.add_paragraph("No se encontraron resultados completos de regresión lineal.")

        add_heading_keep(doc, "15.3 Resumen de Random Forest", level=2)
        if rf_result:
            if str(merged.get("target_type")).lower() == "classification":
                doc.add_paragraph(f"Accuracy: {_safe_text(rf_result.get('accuracy'))}")
                doc.add_paragraph(f"Precision: {_safe_text(rf_result.get('precision'))}")
                doc.add_paragraph(f"Recall: {_safe_text(rf_result.get('recall'))}")
                doc.add_paragraph(f"F1-score: {_safe_text(rf_result.get('f1_score'))}")
            else:
                doc.add_paragraph(f"R²: {_safe_text(rf_result.get('r2'))}")
                doc.add_paragraph(f"MAE: {_safe_text(rf_result.get('mae'))}")
                doc.add_paragraph(f"RMSE: {_safe_text(rf_result.get('rmse'))}")
        else:
            doc.add_paragraph("No se encontraron resultados completos de Random Forest.")

        add_heading_keep(doc, "15.4 Interpretación automática del modelo", level=2)
        if model_interpretation:
            for item in model_interpretation:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            doc.add_paragraph("No se generó interpretación automática básica del modelo.")

        add_heading_keep(doc, "15.5 Evaluación automática del modelo", level=2)
        if advanced_model_interpretation:
            for item in advanced_model_interpretation:
                doc.add_paragraph(str(item), style="List Bullet")
        else:
            fallback_eval = None
            if best_model_selection.get("reason"):
                fallback_eval = best_model_selection.get("reason")
            elif model_comparison.get("recommendation"):
                fallback_eval = model_comparison.get("recommendation")

            if fallback_eval:
                doc.add_paragraph(str(fallback_eval), style="List Bullet")
            else:
                doc.add_paragraph("No se generó evaluación automática detallada del modelo.")

        add_heading_keep(doc, "15.6 Comparación automática de modelos", level=2)
        if model_comparison.get("available"):
            linear = _safe_dict(model_comparison.get("linear"))
            random_forest = _safe_dict(model_comparison.get("random_forest"))
            rows = [
                ["Regresión lineal", _safe_text(linear.get("r2")), _safe_text(linear.get("mae")), _safe_text(linear.get("rmse"))],
                ["Random Forest", _safe_text(random_forest.get("r2")), _safe_text(random_forest.get("mae")), _safe_text(random_forest.get("rmse"))],
            ]
            _add_table(None, ["Modelo", "R²", "MAE", "RMSE"], rows)

            winner_txt = _winner_to_text(model_comparison.get("winner"))
            doc.add_paragraph(f"Modelo recomendado por SADI: {winner_txt}")
            doc.add_paragraph(f"Interpretación comparativa: {_safe_text(model_comparison.get('recommendation'), 'No disponible.')}")

            summary_lines = _safe_list(model_comparison.get("summary"))
            if summary_lines:
                for line in summary_lines:
                    doc.add_paragraph(str(line), style="List Bullet")
            else:
                doc.add_paragraph("No se generó un resumen textual adicional de la comparación.")
        else:
            if str(merged.get("target_type")).lower() == "classification":
                doc.add_paragraph(
                    "La comparación automática entre modelos de regresión no aplica en este caso, "
                    "ya que el problema detectado es de clasificación."
                )
            else:
                doc.add_paragraph(
                    "No se dispuso de resultados suficientes para construir una comparación automática completa entre modelos."
                )

        add_heading_keep(doc, "15.7 Selección inteligente del mejor modelo", level=2)
        if best_model_selection:
            winner_txt = _winner_to_text(best_model_selection.get("best_model"))
            doc.add_paragraph(f"Variable objetivo: {_safe_text(best_model_selection.get('target_col'), _safe_text(merged.get('target_candidate')))}")
            doc.add_paragraph(f"Tipo de problema: {_safe_text(best_model_selection.get('problem_type'), _safe_text(merged.get('target_type')))}")
            doc.add_paragraph(f"Modelo preferido por SADI: {winner_txt}")
            doc.add_paragraph(f"Razón: {_safe_text(best_model_selection.get('reason'), 'No disponible')}")

            comparison_summary = _safe_list(best_model_selection.get("comparison_summary"))
            if comparison_summary:
                for line in comparison_summary:
                    doc.add_paragraph(str(line), style="List Bullet")
            else:
                doc.add_paragraph("No se generó un resumen adicional de selección.")
        else:
            winner_txt = _winner_to_text(model_comparison.get("winner"))
            if winner_txt != "No determinado":
                doc.add_paragraph(f"Modelo recomendado por SADI: {winner_txt}")
                doc.add_paragraph(f"Razón: {_safe_text(model_comparison.get('recommendation'), 'No disponible')}")
            else:
                doc.add_paragraph("No se pudo determinar automáticamente un mejor modelo.")

    # 16. Gráficos de modelos predictivos
    if model_plots:
        add_heading_keep(doc, "16. Gráficos de modelos predictivos", level=1)
        figure_num = 1
        for fname in model_plots:
            img_path = _resolve_plot(fname)
            if not img_path:
                continue

            rel = f"plots/{fname}"
            nice_title = prettify_plot_title(rel)
            nice_desc = describe_plot(rel)

            p_title = doc.add_paragraph()
            p_title.paragraph_format.keep_with_next = True
            r = p_title.add_run(f"Gráfico M{figure_num}. {nice_title}")
            r.bold = True
            r.font.size = Pt(12)

            p_type = doc.add_paragraph()
            p_type.paragraph_format.keep_with_next = True
            r2 = p_type.add_run("Tipo: Modelo predictivo")
            r2.italic = True

            if nice_desc:
                p_desc = doc.add_paragraph(nice_desc)
                p_desc.paragraph_format.keep_with_next = True

            inserted = safe_add_picture(doc, img_path, width_inches=6.0)
            if inserted:
                doc.add_paragraph("")
                figure_num += 1

    # 17. Conclusión final
    add_heading_keep(doc, "17. Conclusión final", level=1)
    final_conclusion = merged.get("insights_text") or build_psychometric_conclusion(merged)
    doc.add_paragraph(str(final_conclusion))

    reports_dir = _reports_dir()
    out_path = os.path.join(reports_dir, f"informe_psicometrico_{dataset.id}.docx")
    doc.save(out_path)

    current_app.logger.warning(f"[generate_psychometric_word] ds{dataset.id} out_path={out_path}")
    return out_path


def generate_psychometric_report_pdf(*, dataset_id: int, manifest_data: dict, plots_dir: str, output_path: str) -> None:
    if not isinstance(manifest_data, dict):
        manifest_data = {}

    plots = (manifest_data.get("generated") or manifest_data.get("plots") or [])
    meta = manifest_data.get("meta") or {}

    if not isinstance(plots, list):
        plots = []
    if not isinstance(meta, dict):
        meta = {}

    try:
        if not (
            meta.get("kmo_report")
            and meta.get("bartlett_report")
            and meta.get("efa_report")
            and meta.get("psychometric_interpretation")
        ):
            meta = ensure_advanced_psychometrics(dataset_id)
            manifest_data = read_manifest_data(dataset_id) or {}
            plots = manifest_data.get("generated") or manifest_data.get("plots") or plots
            meta = manifest_data.get("meta") or meta
    except Exception as e:
        current_app.logger.warning(f"[generate_psychometric_report_pdf] ds{dataset_id}: {e}")

    plots = normalize_plot_catalog(dataset_id, manifest_data, summary)
    exploratory_plots, model_plots = split_plots(plots)

    styles = getSampleStyleSheet()
    body = styles["BodyText"]
    title_style = styles["Title"]
    h1 = styles["Heading1"]
    small = ParagraphStyle("small", parent=body, fontSize=9, leading=11)

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    elements = []
    elements.append(Paragraph("Informe psicométrico", title_style))
    elements.append(Spacer(1, 0.2 * inch))
    elements.append(Paragraph(f"Dataset ID: {dataset_id}", body))
    elements.append(Paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body))
    elements.append(Spacer(1, 0.2 * inch))

    alpha = meta.get("cronbach_alpha")
    data = [
        ["Indicador", "Valor"],
        ["Escala Likert", str(meta.get("likert_scale", "—"))],
        ["N participantes", str(meta.get("n_total", meta.get("n", "—")))],
        ["Alfa de Cronbach", _fmt_num(alpha)],
    ]
    table = Table(data, colWidths=[7 * cm, 7 * cm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
    ]))
    elements.append(table)
    elements.append(Spacer(1, 0.2 * inch))

    if meta.get("psychometric_interpretation"):
        elements.append(Paragraph("Interpretación psicométrica", h1))
        elements.append(Paragraph(str(meta.get("psychometric_interpretation")), body))
        elements.append(Spacer(1, 0.15 * inch))

    if plots:
        elements.append(PageBreak())
        elements.append(Paragraph("Gráficos", h1))
        for rel in plots:
            fname = str(rel).replace("\\", "/").split("/")[-1]
            img_path = _resolve_plot(fname)
            if not img_path:
                continue

            elements.append(Paragraph(prettify_plot_title(rel), styles["Heading2"]))
            elements.append(Paragraph(describe_plot(rel), small))
            img = RLImage(img_path)
            img._restrictSize(16 * cm, 10 * cm)
            elements.append(img)
            elements.append(Spacer(1, 0.2 * inch))

    doc.build(elements)

def build_likert_dimensions(df: pd.DataFrame) -> dict[str, list[str]]:
    """
    Agrupa ítems Likert por dimensión usando el prefijo antes de '_' .
    Ejemplos:
      D1_P1, D1_P2 -> dimensión D1
      Satisfaccion_1, Satisfaccion_2 -> dimensión Satisfaccion
    Si no encuentra '_', usa una dimensión general.
    """
    dimensions: dict[str, list[str]] = {}

    for col in df.columns:
        name = str(col).strip()
        if "_" in name:
            dim = name.split("_")[0].strip()
        else:
            dim = "General"

        dimensions.setdefault(dim, []).append(col)

    return {k: v for k, v in dimensions.items() if v}


def compute_dimension_summary(df: pd.DataFrame) -> list[dict]:
    """
    Calcula promedio, desviación y número de ítems por dimensión.
    """
    dims = build_likert_dimensions(df)
    rows: list[dict] = []

    for dim, cols in dims.items():
        sub = df[cols].copy()
        item_means = sub.mean(axis=0, numeric_only=True)
        dim_score = sub.mean(axis=1, numeric_only=True)

        rows.append({
            "dimension": dim,
            "n_items": len(cols),
            "items": cols,
            "mean": float(dim_score.mean()) if not dim_score.empty else None,
            "std": float(dim_score.std()) if not dim_score.empty else None,
            "item_means": {c: float(item_means[c]) for c in item_means.index},
        })

    rows.sort(key=lambda x: (x["mean"] is None, -(x["mean"] or 0)))
    return rows


def plot_dimension_summary_scores(df: pd.DataFrame, dataset_id: int) -> str | None:
    """
    Gráfico de barras horizontales con promedio por dimensión.
    """
    try:
        summary = compute_dimension_summary(df)
        if not summary:
            return None

        labels = [row["dimension"] for row in summary]
        values = [row["mean"] for row in summary]

        fig, ax = plt.subplots(figsize=(9, max(4, len(labels) * 0.55)))
        y = np.arange(len(labels))

        ax.barh(y, values)
        ax.set_yticks(y)
        ax.set_yticklabels(labels)
        ax.invert_yaxis()
        ax.set_xlabel("Promedio")
        ax.set_title("Promedio por dimensión")

        for i, v in enumerate(values):
            if v is not None:
                ax.text(v + 0.03, i, f"{v:.2f}", va="center", fontsize=9)

        out_name = f"ds{dataset_id}_dimension_summary_scores.png"
        out_path = os.path.join(PLOTS_DIR, out_name)
        fig.tight_layout()
        fig.savefig(out_path, dpi=180, bbox_inches="tight")
        plt.close(fig)

        return out_name
    except Exception as e:
        current_app.logger.warning(f"[plot_dimension_summary_scores] ds{dataset_id}: {e}")
        return None


def plot_dimension_radar(df: pd.DataFrame, dataset_id: int, scale: int = 5) -> str | None:
    """
    Gráfico radar con promedio por dimensión.
    """
    try:
        summary = compute_dimension_summary(df)
        if not summary or len(summary) < 2:
            return None

        labels = [row["dimension"] for row in summary]
        values = [float(row["mean"] or 0) for row in summary]

        angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
        values += values[:1]
        angles += angles[:1]

        fig = plt.figure(figsize=(7, 7))
        ax = plt.subplot(111, polar=True)

        ax.plot(angles, values, linewidth=2)
        ax.fill(angles, values, alpha=0.20)

        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(labels)
        ax.set_ylim(0, scale)
        ax.set_yticks(range(1, scale + 1))
        ax.set_title("Radar de dimensiones", pad=20)

        out_name = f"ds{dataset_id}_dimension_radar.png"
        out_path = os.path.join(PLOTS_DIR, out_name)
        fig.tight_layout()
        fig.savefig(out_path, dpi=180, bbox_inches="tight")
        plt.close(fig)

        return out_name
    except Exception as e:
        current_app.logger.warning(f"[plot_dimension_radar] ds{dataset_id}: {e}")
        return None
# =========================================================
# RESUMEN LIKERT BASE
# =========================================================

@likert_bp.post("/datasets/<int:dataset_id>/likert_summary", endpoint="dataset_likert_summary")
@login_required
def dataset_likert_summary(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)

        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

        kind = (getattr(ds, "dataset_type", None) or "").strip().lower()
        if kind not in {"survey_likert_5", "survey_likert_7"}:
            flash("Este dataset no es Likert.", "warning")
            return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

        filename = ds.filename
        delimiter = ds.delimiter

    try:
        file_path = os.path.join(UPLOAD_DIR, filename)
        current_app.logger.warning(f"[LIKERT] file_path={file_path} exists={os.path.exists(file_path)}")

        df = _read_dataset_with_auto_repair(file_path, delimiter)

        if df.shape[1] == 1:
            only_col = str(df.columns[0]) if len(df.columns) > 0 else ""
            if any(x in only_col for x in [",", ";", "\t", "|", ":"]) or df.shape[0] == 0:
                flash("El archivo quedó con una sola columna incluso después del intento de reparación.", "danger")
                return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    except Exception as e:
        flash(f"No se pudo leer dataset: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    try:
        scale = get_likert_scale_from_dataset_type(kind)
        likert_cols = detect_likert_columns_for_scale(df, scale=scale, min_valid_ratio=0.80)
        if not likert_cols:
            raise ValueError("No se detectaron columnas Likert válidas.")

        items_df = df[likert_cols].copy()
        for col in likert_cols:
            items_df[col] = items_df[col].map(lambda x: coerce_likert_cell_to_int(x, scale=scale))

        items_df = items_df.dropna()

        summary = build_likert_summary(dataset_id)
        if not isinstance(summary, dict):
            summary = {}

        generated = []
        summary_plots = summary.get("plots", []) or []

        for p in summary_plots:
            if isinstance(p, str) and p:
                generated.append(p)

        summary_plot = summary.get("summary_plot")
        if summary_plot:
            generated.append(summary_plot)

        cronbach_plot = summary.get("cronbach_plot")
        if cronbach_plot:
            generated.append(cronbach_plot)

        dimension_summary_plot = summary.get("dimension_summary_plot")
        if dimension_summary_plot:
            generated.append(dimension_summary_plot)

        dimension_radar_plot = summary.get("dimension_radar_plot")
        if dimension_radar_plot:
            generated.append(dimension_radar_plot)

        divergent_plot = summary.get("divergent_plot")
        if divergent_plot:
            generated.append(divergent_plot)

        manifest = read_manifest_data(dataset_id) or {}
        if not isinstance(manifest, dict):
            manifest = {}

        prev_generated = manifest.get("generated", []) or []
        prev_plots = manifest.get("plots", []) or []

        merged = []
        seen = set()
        for p in list(prev_generated) + list(prev_plots) + list(generated):
            if isinstance(p, str) and p not in seen:
                merged.append(p)
                seen.add(p)

        manifest["generated"] = merged
        manifest["plots"] = merged
        manifest["meta"] = summary
        manifest["analysis_meta"] = summary.get("analysis_meta", {})
        manifest["dataset_type"] = kind
        manifest["research_area"] = summary.get("research_area")
        manifest["suggested_plan"] = summary.get("suggested_plan", {})
        manifest["cols"] = likert_cols

        write_manifest(dataset_id, manifest)
        flash("Análisis Likert generado correctamente.", "success")

    except Exception as e:
        current_app.logger.exception(f"[dataset_likert_summary] ds{dataset_id}: {e}")
        flash(f"Error en análisis Likert: {e}", "danger")

    return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))


@likert_bp.get("/datasets/<int:dataset_id>/psychometric_analysis", endpoint="dataset_psychometric_analysis")
@login_required
def dataset_psychometric_analysis(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    try:
        meta = ensure_advanced_psychometrics(dataset_id)
        manifest_data = read_manifest_data(dataset_id) or {}
        generated = manifest_data.get("generated", []) if isinstance(manifest_data, dict) else []
        write_manifest(dataset_id, generated, meta=meta)
        flash("Análisis psicométrico generado correctamente.", "success")
    except Exception as e:
        current_app.logger.exception(f"[dataset_psychometric_analysis] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el análisis psicométrico: {e}", "danger")

    return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))


@likert_bp.get("/datasets/<int:dataset_id>/psychometric_report.pdf", endpoint="dataset_psychometric_report")
@login_required
def dataset_psychometric_report(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    manifest_data = read_manifest_data(dataset_id) or {}
    out_path = os.path.join(_reports_dir(), f"ds{dataset_id}_psychometric_report.pdf")

    try:
        generate_psychometric_report_pdf(
            dataset_id=dataset_id,
            manifest_data=manifest_data,
            plots_dir=PLOTS_DIR,
            output_path=out_path,
        )
    except Exception as e:
        current_app.logger.exception(f"[psychometric_report] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el PDF psicométrico: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    if not os.path.exists(out_path) or os.path.getsize(out_path) < 1000:
        flash("Se intentó generar el PDF pero quedó vacío o no se creó.", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    return send_file(out_path, as_attachment=True, download_name=f"SADI_psicometrico_ds{dataset_id}.pdf")


@likert_bp.get("/datasets/<int:dataset_id>/likert_word", endpoint="dataset_likert_word")
@login_required
def dataset_likert_word(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    kind = (
        getattr(ds, "dataset_kind", None)
        or getattr(ds, "dataset_type", None)
        or getattr(ds, "kind", None)
        or getattr(ds, "type", None)
        or ""
    )
    kind = str(kind).strip().lower()

    allowed_likert = {
        "survey_likert_5", "survey_likert_7",
        "likert_5", "likert_7",
        "survey_likert", "likert"
    }
    if kind not in allowed_likert:
        flash(f"Este informe Word aplica solo para datasets Likert. Tipo detectado: {kind}", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    try:
        manifest_data = read_manifest_data(dataset_id) or {}
        if not isinstance(manifest_data, dict):
            manifest_data = {}

        manifest_meta = manifest_data.get("meta") or {}
        if not isinstance(manifest_meta, dict):
            manifest_meta = {}

        manifest_plots = manifest_data.get("generated") or manifest_data.get("plots") or []
        if not isinstance(manifest_plots, list):
            manifest_plots = []

        # =========================
        # Reconstruir summary Likert inteligente
        # =========================
        summary = build_likert_summary(dataset_id)
        if not isinstance(summary, dict):
            summary = {}

        summary_meta = summary.get("analysis_meta") or {}
        if not isinstance(summary_meta, dict):
            summary_meta = {}

        # fusionar meta general + meta summary
        meta = dict(manifest_meta)
        meta.update(summary)
        meta.update(summary_meta)

        # =========================
        # Completar psicometría avanzada si existe
        # =========================
        try:
            advanced_meta = ensure_advanced_psychometrics(dataset_id)
            if isinstance(advanced_meta, dict):
                meta.update(advanced_meta)
        except Exception as e:
            current_app.logger.warning(
                f"[dataset_likert_word] ds{dataset_id}: no se pudo completar psicometría avanzada: {e}"
            )

        # =========================
        # RECONSTRUIR BLOQUE PREDICTIVO
        # =========================
        model_results = manifest_data.get("model_results", {})
        if not isinstance(model_results, dict):
            model_results = {}

        regression_result = model_results.get("regression_result") or {}
        if not isinstance(regression_result, dict):
            regression_result = {}

        rf_result = model_results.get("rf_result") or {}
        if not isinstance(rf_result, dict):
            rf_result = {}

        try:
            model_comparison = build_model_comparison_summary(
                regression_result=regression_result,
                rf_result=rf_result,
            ) or {"available": False}
        except Exception as e:
            current_app.logger.warning(f"[dataset_likert_word:model_comparison] ds{dataset_id}: {e}")
            model_comparison = {"available": False}

        best_model_selection = manifest_data.get("best_model_selection", {}) or {}
        if not isinstance(best_model_selection, dict):
            best_model_selection = {}

        model_interpretation = manifest_data.get("model_interpretation", []) or []
        if not isinstance(model_interpretation, list):
            model_interpretation = []

        try:
            if meta.get("target_type") == "classification":
                advanced_model_interpretation = build_advanced_classification_interpretation(
                    rf_result=rf_result,
                ) or []
            else:
                advanced_model_interpretation = build_advanced_model_interpretation(
                    regression_result=regression_result,
                    rf_result=rf_result,
                ) or []
        except Exception as e:
            current_app.logger.warning(f"[dataset_likert_word:advanced_model_interpretation] ds{dataset_id}: {e}")
            advanced_model_interpretation = []

        try:
            model_plot_candidates = manifest_plots + (summary.get("plots", []) or [])
            next_step_recommendation = build_next_step_recommendation(
                analysis_meta=meta,
                dataset_kind=kind,
                model_plots=model_plot_candidates,
            )
        except Exception as e:
            current_app.logger.warning(f"[dataset_likert_word:next_step] ds{dataset_id}: {e}")
            next_step_recommendation = None

        # inyectar al meta final para el Word
        meta["regression_result"] = regression_result
        meta["rf_result"] = rf_result
        meta["model_comparison"] = model_comparison
        meta["best_model_selection"] = best_model_selection
        meta["model_interpretation"] = model_interpretation
        meta["advanced_model_interpretation"] = advanced_model_interpretation
        meta["next_step_recommendation"] = next_step_recommendation

        # mantener análisis interno alineado también
        if "analysis_meta" not in meta or not isinstance(meta.get("analysis_meta"), dict):
            meta["analysis_meta"] = {}
        meta["analysis_meta"].update({
            "regression_result": regression_result,
            "rf_result": rf_result,
            "model_comparison": model_comparison,
            "best_model_selection": best_model_selection,
            "model_interpretation": model_interpretation,
            "advanced_model_interpretation": advanced_model_interpretation,
            "next_step_recommendation": next_step_recommendation,
        })

        # =========================
        # Reconstruir plots
        # =========================
        plots = []
        seen = set()

        for p in manifest_plots + (summary.get("plots", []) or []):
            if isinstance(p, str) and p not in seen:
                plots.append(p)
                seen.add(p)

        extra_plot_keys = [
            "cronbach_plot",
            "summary_plot",
            "divergent_plot",
            "dimension_summary_plot",
            "dimension_radar_plot",
            "scree_plot",
            "factor_loadings_plot",
            "corr_heatmap_plot",
            "factor_model_plot",
        ]
        for key in extra_plot_keys:
            p = summary.get(key)
            if isinstance(p, str) and p:
                rel = os.path.join("plots", p).replace("\\", "/") if not p.startswith("plots/") else p
                if rel not in seen:
                    plots.append(rel)
                    seen.add(rel)

        # agregar plots predictivos conocidos si existen
        predictive_plot_names = [
            f"plots/ds{dataset_id}_regression_coefficients.png",
            f"plots/ds{dataset_id}_regression_pred_vs_real.png",
            f"plots/ds{dataset_id}_regression_residuals.png",
            f"plots/ds{dataset_id}_rf_feature_importance.png",
            f"plots/ds{dataset_id}_rf_pred_vs_real.png",
            f"plots/ds{dataset_id}_roc_curve.png",
            f"plots/ds{dataset_id}_confusion_matrix.png",
        ]
        for rel in predictive_plot_names:
            if rel not in seen:
                fname = rel.split("/")[-1]
                abs_path = _resolve_plot(fname) if "_resolve_plot" in globals() else None
                if abs_path:
                    plots.append(rel)
                    seen.add(rel)

        plots = normalize_plot_catalog(dataset_id, manifest_data, summary)
        exploratory_plots, model_plots = split_plots(plots)

        file_path = generate_psychometric_word(
            dataset=ds,
            meta=meta,
            plots=plots,
        )

    except Exception as e:
        current_app.logger.exception(f"[likert_word] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el informe psicométrico Word: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    if not file_path or not os.path.exists(file_path) or os.path.getsize(file_path) < 1000:
        flash("Se intentó generar el Word psicométrico pero quedó vacío o no se creó.", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    return send_file(
        file_path,
        as_attachment=True,
        download_name=f"informe_psicometrico_ds{dataset_id}.docx"
    )


@likert_bp.get("/datasets/<int:dataset_id>/likert_article", endpoint="dataset_likert_article")
@login_required
def dataset_likert_article(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    kind = (
        getattr(ds, "dataset_kind", None)
        or getattr(ds, "dataset_type", None)
        or getattr(ds, "kind", None)
        or getattr(ds, "type", None)
        or ""
    )
    kind = str(kind).strip().lower()

    allowed_likert = {
        "survey_likert_5", "survey_likert_7",
        "likert_5", "likert_7",
        "survey_likert", "likert"
    }
    if kind not in allowed_likert:
        flash(f"Este artículo aplica solo para datasets Likert. Tipo detectado: {kind}", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    try:
        manifest_data = read_manifest_data(dataset_id) or {}
        if not isinstance(manifest_data, dict):
            manifest_data = {}

        # Asegurar psicometría avanzada
        try:
            ensure_advanced_psychometrics(dataset_id)
            manifest_data = read_manifest_data(dataset_id) or manifest_data
        except Exception as e:
            current_app.logger.warning(
                f"[dataset_likert_article] ds{dataset_id}: no se pudo completar psicometría avanzada: {e}"
            )

        summary = build_likert_summary(dataset_id)
        if not isinstance(summary, dict):
            summary = {}

        # reconstruir bloque predictivo igual que en Word
        model_results = manifest_data.get("model_results", {})
        if not isinstance(model_results, dict):
            model_results = {}

        regression_result = model_results.get("regression_result") or {}
        if not isinstance(regression_result, dict):
            regression_result = {}

        rf_result = model_results.get("rf_result") or {}
        if not isinstance(rf_result, dict):
            rf_result = {}

        try:
            model_comparison = build_model_comparison_summary(
                regression_result=regression_result,
                rf_result=rf_result,
            ) or {"available": False}
        except Exception as e:
            current_app.logger.warning(f"[dataset_likert_article:model_comparison] ds{dataset_id}: {e}")
            model_comparison = {"available": False}

        best_model_selection = manifest_data.get("best_model_selection", {}) or {}
        if not isinstance(best_model_selection, dict):
            best_model_selection = {}

        model_interpretation = manifest_data.get("model_interpretation", []) or []
        if not isinstance(model_interpretation, list):
            model_interpretation = []

        try:
            if summary.get("target_type") == "classification":
                advanced_model_interpretation = build_advanced_classification_interpretation(
                    rf_result=rf_result,
                ) or []
            else:
                advanced_model_interpretation = build_advanced_model_interpretation(
                    regression_result=regression_result,
                    rf_result=rf_result,
                ) or []
        except Exception as e:
            current_app.logger.warning(f"[dataset_likert_article:advanced_model_interpretation] ds{dataset_id}: {e}")
            advanced_model_interpretation = []

        manifest_plots = manifest_data.get("generated") or manifest_data.get("plots") or []
        if not isinstance(manifest_plots, list):
            manifest_plots = []

        try:
            next_step_recommendation = build_next_step_recommendation(
                analysis_meta=summary.get("analysis_meta", {}) or summary,
                dataset_kind=kind,
                model_plots=manifest_plots + (summary.get("plots", []) or []),
            )
        except Exception as e:
            current_app.logger.warning(f"[dataset_likert_article:next_step] ds{dataset_id}: {e}")
            next_step_recommendation = None

        # inyectar para el artículo
        summary["regression_result"] = regression_result
        summary["rf_result"] = rf_result
        summary["model_comparison"] = model_comparison
        summary["best_model_selection"] = best_model_selection
        summary["model_interpretation"] = model_interpretation
        summary["advanced_model_interpretation"] = advanced_model_interpretation
        summary["next_step_recommendation"] = next_step_recommendation

        if "analysis_meta" not in summary or not isinstance(summary.get("analysis_meta"), dict):
            summary["analysis_meta"] = {}
        summary["analysis_meta"].update({
            "regression_result": regression_result,
            "rf_result": rf_result,
            "model_comparison": model_comparison,
            "best_model_selection": best_model_selection,
            "model_interpretation": model_interpretation,
            "advanced_model_interpretation": advanced_model_interpretation,
            "next_step_recommendation": next_step_recommendation,
        })

        output_path = os.path.join(_reports_dir(), f"ds{dataset_id}_likert_article.docx")

        generate_likert_article_docx(
            dataset_id=dataset_id,
            summary=summary,
            output_path=output_path
        )

    except Exception as e:
        current_app.logger.exception(f"[likert_article] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el artículo científico: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    if not os.path.exists(output_path) or os.path.getsize(output_path) < 1000:
        flash("Se intentó generar el artículo Likert, pero el archivo no se creó correctamente.", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    return send_file(
        output_path,
        as_attachment=True,
        download_name=f"likert_article_{dataset_id}.docx"
    )


@likert_bp.get("/datasets/<int:dataset_id>/scientific_paper.docx", endpoint="dataset_scientific_paper")
@login_required
def dataset_scientific_paper(dataset_id: int):
    return dataset_likert_article(dataset_id)


@likert_bp.get("/datasets/<int:dataset_id>/likert_divergent_plot", endpoint="dataset_likert_divergent_plot")
@login_required
def dataset_likert_divergent_plot(dataset_id: int):
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds or ds.user_id != current_user.id:
            flash("Dataset no encontrado.", "warning")
            return redirect(url_for("dataset.dashboard"))

    scale = get_likert_scale_from_dataset_type((getattr(ds, "dataset_type", None) or "").strip())
    if scale not in (5, 7):
        flash("El dataset no está configurado como Likert 5 o Likert 7.", "warning")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    path = os.path.join(UPLOAD_DIR, ds.filename)

    try:
        df = _read_dataset_with_auto_repair(path, ds.delimiter)

        if df.shape[1] == 1:
            only_col = str(df.columns[0]) if len(df.columns) > 0 else ""
            if any(x in only_col for x in [",", ";", "\t", "|", ":"]) or df.shape[0] == 0:
                flash("El archivo quedó con una sola columna incluso después del intento de reparación.", "danger")
                return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    except Exception as e:
        flash(f"No se pudo leer dataset: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    output_path = os.path.join(PLOTS_DIR, f"ds{dataset_id}_likert_divergent.png")

    try:
        plot_likert_divergent(
            df=df,
            dataset_id=dataset_id,
            output_path=output_path,
            scale=scale
        )
    except Exception as e:
        current_app.logger.exception(f"[likert_divergent] ds{dataset_id}: {e}")
        flash(f"No se pudo generar el gráfico Likert: {e}", "danger")
        return redirect(url_for("dataset.dataset_detail", dataset_id=dataset_id))

    return send_file(output_path, mimetype="image/png")

def build_likert_summary(dataset_id: int) -> dict:
    with SessionLocal() as db:
        ds = db.get(Dataset, dataset_id)
        if not ds:
            raise ValueError("Dataset no encontrado.")

        dataset_type = (getattr(ds, "dataset_type", None) or "").strip().lower()
        research_area = (getattr(ds, "research_area", None) or "general").strip().lower()
        filename = ds.filename
        delimiter = ds.delimiter
        dataset_title = getattr(ds, "title", None) or f"Dataset {dataset_id}"

    scale = get_likert_scale_from_dataset_type(dataset_type)
    if scale not in (5, 7):
        raise ValueError("El dataset no está marcado como Likert 5 o Likert 7.")

    path = os.path.join(UPLOAD_DIR, filename)
    df = _read_dataset_with_auto_repair(path, delimiter)

    likert_cols = detect_likert_columns_for_scale(df, scale=scale, min_valid_ratio=0.80)
    if not likert_cols:
        raise ValueError(f"No se detectaron columnas Likert válidas (escala 1 a {scale}).")

    likert_df = df[likert_cols].copy()
    for col in likert_cols:
        likert_df[col] = likert_df[col].map(lambda x: coerce_likert_cell_to_int(x, scale=scale))

    # Mantener versión con posibles NA para descriptivos
    likert_df_stats = likert_df.copy()

    # Para psicometría robusta
    likert_df_clean = likert_df.dropna()
    if likert_df_clean.empty:
        raise ValueError("No hay suficientes datos válidos para análisis Likert.")

    means = likert_df_stats.mean(numeric_only=True).sort_values(ascending=False)
    global_mean = float(likert_df_stats.mean(numeric_only=True).mean()) if not likert_df_stats.empty else None

    alpha, usable_items, alpha_label, cronbach_plot = make_cronbach_alpha_plot(likert_df_clean, dataset_id)

    top_item = str(means.index[0]) if not means.empty else None
    top_mean = float(means.iloc[0]) if not means.empty else None
    low_item = str(means.index[-1]) if not means.empty else None
    low_mean = float(means.iloc[-1]) if not means.empty else None

    item_stats = []
    for col in likert_df_stats.columns:
        s = likert_df_stats[col].dropna()
        item_stats.append({
            "item": str(col),
            "mean": float(s.mean()) if not s.empty else None,
            "median": float(s.median()) if not s.empty else None,
            "std": float(s.std()) if not s.empty else None,
            "n_valid": int(s.shape[0]),
        })

    # =========================
    # Perfil del instrumento
    # =========================

    # Nivel de completitud
    missing_pct = 0.0
    if len(likert_df_stats):
        missing_pct = 100.0 * (1 - (len(likert_df_clean) / len(likert_df_stats)))

    # Clasificación de calidad de datos
    if missing_pct < 5:
        data_quality = "alta"
    elif missing_pct < 15:
        data_quality = "moderada"
    else:
        data_quality = "baja"

    n_items = len(likert_cols)
    response_rate = (len(likert_df_clean) / len(likert_df_stats)) if len(likert_df_stats) else None

    def interpret_alpha_quality(alpha_value):
        if alpha_value is None:
            return "no disponible"
        if alpha_value >= 0.90:
            return "excelente"
        if alpha_value >= 0.80:
            return "buena"
        if alpha_value >= 0.70:
            return "aceptable"
        if alpha_value >= 0.60:
            return "débil"
        return "deficiente"

    alpha_quality = interpret_alpha_quality(alpha)

    # =========================
    # Motor central SADI
    # =========================
    try:
        analysis_meta = analyze_dataset_with_recommendations(
            likert_df_clean,
            dataset_type=dataset_type,
            research_area=research_area,
        ) or {}
        if not isinstance(analysis_meta, dict):
            analysis_meta = {}
    except Exception as e:
        current_app.logger.warning(f"[build_likert_summary:sadi_core] ds{dataset_id}: {e}")
        analysis_meta = {}

    # =========================
    # Gráficos y resumen por dimensión
    # =========================
    summary_plot = plot_likert_summary_scores(likert_df_stats, dataset_id)
    dimension_summary = compute_dimension_summary(likert_df_stats)
    dimension_summary_plot = plot_dimension_summary_scores(likert_df_stats, dataset_id)
    dimension_radar_plot = plot_dimension_radar(likert_df_stats, dataset_id, scale=scale)

    divergent_plot = None
    try:
        divergent_name = f"ds{dataset_id}_likert_divergent.png"
        divergent_path = os.path.join(PLOTS_DIR, divergent_name)
        plot_likert_divergent(likert_df_stats, dataset_id, divergent_path, scale=scale)
        if os.path.exists(divergent_path):
            divergent_plot = divergent_name
    except Exception as e:
        current_app.logger.warning(f"[build_likert_summary:divergent] ds{dataset_id}: {e}")

    # =========================
    # Análisis de dimensiones
    # =========================
    valid_dims = [d for d in (dimension_summary or []) if d.get("mean") is not None]
    n_dimensions = len(valid_dims)

    dimension_ranking = sorted(
        valid_dims,
        key=lambda x: x.get("mean", 0),
        reverse=True
    ) if valid_dims else []

    strong_dimensions = [d.get("dimension") for d in dimension_ranking[:2] if d.get("dimension")]
    weak_dimensions = [d.get("dimension") for d in dimension_ranking[-2:] if d.get("dimension")] if len(dimension_ranking) >= 2 else []

    best_dim = max(valid_dims, key=lambda x: x["mean"], default=None)
    low_dim = min(valid_dims, key=lambda x: x["mean"], default=None)

    # =========================
    # Tendencia global
    # =========================
    trend = "moderada"
    if global_mean is not None:
        if scale == 5:
            if global_mean >= 4.2:
                trend = "muy positiva"
            elif global_mean >= 3.5:
                trend = "positiva"
            elif global_mean >= 2.8:
                trend = "intermedia"
            elif global_mean >= 2.0:
                trend = "negativa"
            else:
                trend = "muy negativa"
        elif scale == 7:
            if global_mean >= 6.2:
                trend = "muy positiva"
            elif global_mean >= 5.2:
                trend = "positiva"
            elif global_mean >= 3.8:
                trend = "intermedia"
            elif global_mean >= 2.5:
                trend = "negativa"
            else:
                trend = "muy negativa"

    # =========================
    # Insights Likert propios
    # =========================
    likert_insights = []

    if alpha is not None:
        likert_insights.append(
            f"El instrumento presentó un alfa de Cronbach de {_fmt_num(alpha)}, lo que indica una consistencia interna {alpha_label.lower()}."
        )
        likert_insights.append(
            f"En términos metodológicos, la confiabilidad global del instrumento puede calificarse como {alpha_quality}."
        )

    if global_mean is not None:
        likert_insights.append(
            f"La media global del instrumento fue de {_fmt_num(global_mean, 3)}, lo que resume una tendencia general {trend} en las respuestas."
        )

        if global_mean >= scale * 0.75:
            likert_insights.append(
                "Se observa una tendencia alta en las respuestas, lo que sugiere una percepción general favorable."
            )
        elif global_mean <= scale * 0.40:
            likert_insights.append(
                "Se observa una tendencia baja en las respuestas, lo que sugiere percepciones críticas o desfavorables."
            )
        else:
            likert_insights.append(
                "Las respuestas muestran un comportamiento relativamente equilibrado, sin polarización extrema."
            )

    if top_item and top_mean is not None:
        likert_insights.append(
            f"El ítem con valoración media más alta fue '{top_item}' con un promedio de {_fmt_num(top_mean, 3)}."
        )

    if low_item and low_mean is not None:
        likert_insights.append(
            f"El ítem con valoración media más baja fue '{low_item}' con un promedio de {_fmt_num(low_mean, 3)}."
        )

    if len(likert_cols) >= 8:
        likert_insights.append(
            f"Se identificaron {len(likert_cols)} ítems Likert válidos, lo que brinda una base suficiente para análisis psicométrico avanzado."
        )
    else:
        likert_insights.append(
            f"Se identificaron {len(likert_cols)} ítems Likert válidos, adecuados para un análisis descriptivo y de consistencia interna."
        )

    if best_dim:
        likert_insights.append(
            f"La dimensión con promedio más alto fue '{best_dim['dimension']}' con una media de {_fmt_num(best_dim['mean'], 3)}."
        )

    if low_dim:
        likert_insights.append(
            f"La dimensión con promedio más bajo fue '{low_dim['dimension']}' con una media de {_fmt_num(low_dim['mean'], 3)}."
        )

    if strong_dimensions:
        likert_insights.append(
            f"Las dimensiones más fuertes del instrumento son: {', '.join(strong_dimensions)}."
        )

    if weak_dimensions:
        likert_insights.append(
            f"Las dimensiones con mayor necesidad de revisión o refuerzo son: {', '.join(weak_dimensions)}."
        )

    if alpha is not None and alpha < 0.70:
        likert_insights.append(
            "Se recomienda revisar o depurar ítems, ya que la consistencia interna del instrumento es limitada."
        )

    if response_rate is not None and response_rate < 0.80:
        likert_insights.append(
            "Se detectó una proporción relevante de casos incompletos, lo que puede afectar la estabilidad del análisis."
        )
    # =========================
    # INSIGHTS SADI AVANZADOS
    # =========================

    # Calidad del instrumento
    likert_insights.append(
        f"La calidad general de los datos puede considerarse {data_quality}, con un {round(missing_pct,2)}% de valores faltantes."
    )

    # Evaluación metodológica global
    if alpha_quality in ["excelente", "buena"]:
        likert_insights.append(
            "El instrumento presenta una estructura consistente y adecuada para análisis interpretativos confiables."
        )
    elif alpha_quality == "aceptable":
        likert_insights.append(
            "El instrumento es usable, aunque se recomienda revisar algunos ítems para mejorar la consistencia."
        )
    else:
        likert_insights.append(
            "El instrumento presenta problemas de consistencia interna y debería ser revisado antes de análisis avanzados."
        )

    # Balance de dimensiones
    if n_dimensions >= 3:
        likert_insights.append(
            f"El instrumento presenta {n_dimensions} dimensiones, lo que sugiere una estructura conceptual multidimensional."
        )

    # Sesgo de respuestas
    if global_mean:
        if global_mean >= scale * 0.8:
            likert_insights.append("Se detecta posible sesgo de techo en las respuestas (tendencia a valores altos).")
        elif global_mean <= scale * 0.3:
            likert_insights.append("Se detecta posible sesgo de piso en las respuestas (tendencia a valores bajos).")
    # =========================
    # Recomendaciones Likert propias
    # =========================
    likert_recommendations = []

    if alpha is not None and alpha < 0.70:
        likert_recommendations.append(
            "Revisar la redacción, coherencia y redundancia de los ítems para mejorar la confiabilidad del instrumento."
        )
    else:
        likert_recommendations.append(
            "La consistencia interna del instrumento permite continuar con interpretación por ítems y dimensiones."
        )

    if weak_dimensions:
        likert_recommendations.append(
            f"Profundizar el análisis de las dimensiones más débiles: {', '.join(weak_dimensions)}."
        )

    if strong_dimensions:
        likert_recommendations.append(
            f"Considerar las dimensiones fuertes ({', '.join(strong_dimensions)}) como ejes interpretativos principales."
        )

    if response_rate is not None and response_rate < 0.80:
        likert_recommendations.append(
            "Revisar la completitud de respuestas y considerar tratamiento de datos faltantes antes de análisis avanzados."
        )

    likert_recommendations.append(
        "Complementar la interpretación con gráficos por dimensión, ranking de ítems y distribución de respuestas."
    )
    # Advertencias metodológicas
    likert_warnings = []

    if alpha is not None and alpha < 0.70:
        likert_warnings.append(
            "La baja confiabilidad puede afectar la validez de las conclusiones."
        )

    if missing_pct > 15:
        likert_warnings.append(
            "El alto porcentaje de datos faltantes puede sesgar los resultados."
        )

    if global_mean and (global_mean >= scale * 0.8 or global_mean <= scale * 0.3):
        likert_warnings.append(
            "Existe posible sesgo en las respuestas que debería considerarse en la interpretación."
        )
    # =========================
    # Plan sugerido específico Likert
    # =========================
    likert_suggested_plan = {
        "recommended_analysis": [
            "Análisis descriptivo por ítem",
            "Análisis descriptivo por dimensión",
            "Consistencia interna mediante alfa de Cronbach",
            "Comparación entre dimensiones del instrumento",
        ],
        "recommended_plots": [
            "Resumen gráfico de puntajes Likert",
            "Radar de dimensiones",
            "Gráfico divergente Likert",
            "Gráfico resumen por dimensión",
        ],
        "narrative_focus": (
            "Interpretar la consistencia del instrumento, la tendencia general de respuesta "
            "y las diferencias entre dimensiones o bloques temáticos."
        ),
        "warnings": (
            ["La confiabilidad del instrumento es baja y conviene revisar ítems antes de sacar conclusiones fuertes."]
            if alpha is not None and alpha < 0.70 else
            ["La interpretación debe considerar tanto medias como dispersión y consistencia interna."]
        ),
    }

    # =========================
    # Fusionar insights del motor central + Likert
    # =========================
    merged_insights = []
    core_insights = analysis_meta.get("insights", []) or []
    if isinstance(core_insights, list):
        merged_insights.extend(core_insights)
    merged_insights.extend(likert_insights)

    seen = set()
    deduped_insights = []
    for x in merged_insights:
        key = str(x).strip()
        if key and key not in seen:
            deduped_insights.append(x)
            seen.add(key)

    # =========================
    # Quick recommendations fusionadas
    # =========================
    merged_quick_recommendations = []
    core_recs = analysis_meta.get("quick_recommendations", []) or []
    if isinstance(core_recs, list):
        merged_quick_recommendations.extend(core_recs)
    merged_quick_recommendations.extend(likert_recommendations)

    seen_recs = set()
    deduped_recommendations = []
    for x in merged_quick_recommendations:
        key = str(x).strip()
        if key and key not in seen_recs:
            deduped_recommendations.append(x)
            seen_recs.add(key)

    # =========================
    # Texto narrativo
    # =========================
    insights_text = analysis_meta.get("insights_text") or ""
    if not insights_text:
        narrative_parts = []

        if global_mean is not None:
            narrative_parts.append(
                f"El análisis del instrumento Likert mostró una media global de {_fmt_num(global_mean, 3)}, reflejando una tendencia {trend} en las respuestas."
            )

        if alpha is not None:
            narrative_parts.append(
                f"Asimismo, el alfa de Cronbach alcanzó un valor de {_fmt_num(alpha)}, lo que sugiere una consistencia interna {alpha_label.lower()} y una calidad metodológica {alpha_quality}."
            )

        if top_item and low_item:
            narrative_parts.append(
                f"Los ítems extremos observados fueron '{top_item}' como el de mayor valoración media y '{low_item}' como el de menor promedio."
            )

        if best_dim and low_dim:
            narrative_parts.append(
                f"A nivel dimensional, '{best_dim['dimension']}' mostró el promedio más alto, mientras que '{low_dim['dimension']}' presentó el promedio más bajo."
            )

        if weak_dimensions:
            narrative_parts.append(
                f"Las dimensiones que requieren mayor atención interpretativa son: {', '.join(weak_dimensions)}."
            )

        insights_text = " ".join(narrative_parts).strip()

    # =========================
    # Plots consolidados
    # =========================
    plots = []
    for name in [
        cronbach_plot,
        summary_plot,
        divergent_plot,
        dimension_summary_plot,
        dimension_radar_plot,
    ]:
        if name:
            plots.append(os.path.join("plots", name).replace("\\", "/"))

    plot_summary = summarize_plot_tags(plots)

    # =========================
    # Inyectar al analysis_meta para dataset_detail
    # =========================
    analysis_meta["instrument_type"] = "Likert"
    analysis_meta["missing_pct"] = missing_pct
    analysis_meta["data_quality"] = data_quality
    analysis_meta["warnings"] = likert_warnings
    analysis_meta["n_total"] = len(likert_df_stats)
    analysis_meta["n_valid"] = len(likert_df_clean)
    analysis_meta["likert_scale"] = scale
    analysis_meta["n_items"] = n_items
    analysis_meta["n_dimensions"] = n_dimensions
    analysis_meta["response_rate"] = response_rate
    analysis_meta["cronbach_alpha"] = alpha
    analysis_meta["reliability_level"] = alpha_quality
    analysis_meta["cronbach_label"] = alpha_label
    analysis_meta["global_mean"] = global_mean
    analysis_meta["trend"] = trend
    analysis_meta["dimension_summary"] = dimension_summary
    analysis_meta["dimension_ranking"] = dimension_ranking
    analysis_meta["strong_dimensions"] = strong_dimensions
    analysis_meta["weak_dimensions"] = weak_dimensions
    analysis_meta["dimension_summary_plot"] = dimension_summary_plot
    analysis_meta["dimension_radar_plot"] = dimension_radar_plot
    analysis_meta["plots"] = plots
    analysis_meta["plot_summary"] = plot_summary
    analysis_meta["insights"] = deduped_insights[:12]
    analysis_meta["insights_text"] = insights_text
    analysis_meta["quick_recommendations"] = deduped_recommendations[:10]
    analysis_meta["suggested_plan"] = likert_suggested_plan
    analysis_meta["priority_order"] = [
        "Verificar confiabilidad global",
        "Interpretar dimensiones más fuertes y más débiles",
        "Analizar ítems extremos",
        "Revisar posibles dimensiones débiles o ítems problemáticos",
    ]
    analysis_meta["research_area_suggested"] = analysis_meta.get("research_area_suggested") or research_area

    analysis_meta["instrument_profile"] = {
        "instrument_type": "Likert",
        "likert_scale": scale,
        "n_items": n_items,
        "n_dimensions": n_dimensions,
        "n_total": len(likert_df_stats),
        "n_valid": len(likert_df_clean),
        "response_rate": response_rate,
        "missing_pct": missing_pct,
        "data_quality": data_quality,
        "reliability_level": alpha_quality,
    }

    analysis_meta["methodological_warnings"] = likert_warnings
    analysis_meta["top_strengths"] = strong_dimensions
    analysis_meta["top_weaknesses"] = weak_dimensions

    return {
        "dataset_id": dataset_id,
        "dataset_title": dataset_title,
        "research_area": research_area,
        "dataset_type": dataset_type,

        "n": int(len(likert_df_clean)),
        "n_total": int(len(likert_df_stats)),
        "scale": scale,
        "likert_scale": scale,

        "columns": likert_cols,
        "likert_cols": likert_cols,
        "likert_columns": likert_cols,
        "usable_items": usable_items,

        "n_items": n_items,
        "n_dimensions": n_dimensions,
        "response_rate": response_rate,
        "global_mean": global_mean,
        "trend": trend,

        "means_by_item": {k: float(v) for k, v in means.to_dict().items()},
        "item_stats": item_stats,

        "cronbach_alpha": alpha,
        "cronbach_label": alpha_label,
        "cronbach_text": f"El alfa de Cronbach global fue {_fmt_num(alpha)} ({alpha_label})." if alpha is not None else "",
        "reliability_level": alpha_quality,

        "top_item": top_item,
        "top_mean": top_mean,
        "low_item": low_item,
        "low_mean": low_mean,

        "alpha_if_deleted": [],
        "dimension_alpha": [],
        "dimension_summary": dimension_summary,
        "dimension_ranking": dimension_ranking,
        "strong_dimensions": strong_dimensions,
        "weak_dimensions": weak_dimensions,

        "results_text": insights_text,
        "survey_insights": "\n".join(str(x) for x in deduped_insights),

        # SADI central
        "analysis_meta": analysis_meta,
        "insights": deduped_insights[:12],
        "insights_text": insights_text,
        "quick_recommendations": deduped_recommendations[:10],
        "suggested_plan": likert_suggested_plan,
        "priority_order": analysis_meta.get("priority_order", []) or [],
        "research_area_suggested": analysis_meta.get("research_area_suggested"),

        # gráficos
        "cronbach_plot": cronbach_plot,
        "summary_plot": summary_plot,
        "divergent_plot": divergent_plot,
        "plots": plots,
        "plot_summary": plot_summary,
        "dimension_summary_plot": dimension_summary_plot,
        "dimension_radar_plot": dimension_radar_plot,
    }